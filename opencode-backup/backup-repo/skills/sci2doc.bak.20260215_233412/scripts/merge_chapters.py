#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合并章节文档并生成完整论文

功能：
1. 按顺序合并多个 docx 章节文件
2. 保持原有样式和格式
3. 自动插入分页符
4. 生成目录
5. 添加页眉页脚

作者：Sci2Doc Team
日期：2024-03-15
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os
import re
import json


def add_page_break(doc):
    """在文档中添加分页符"""
    doc.add_page_break()


def extract_chapter_number(filename):
    """从文件名中提取章节号"""
    match = re.search(r'第(\d+)章', filename)
    if match:
        return int(match.group(1))
    return 999  # 未识别的放到最后


def merge_docx_files(file_list, output_path):
    """
    合并多个 docx 文件
    
    Args:
        file_list: docx 文件路径列表
        output_path: 输出文件路径
    
    Returns:
        dict: 合并结果
    """
    if not file_list:
        return {
            'success': False,
            'error': '文件列表为空'
        }
    
    # 按章节号排序
    file_list_sorted = sorted(file_list, key=lambda x: extract_chapter_number(os.path.basename(x)))
    
    # 创建主文档（基于第一个文件）
    try:
        master_doc = Document(file_list_sorted[0])
    except Exception as e:
        return {
            'success': False,
            'error': f'无法打开第一个文件：{str(e)}'
        }
    
    # 添加后续文件
    for file_path in file_list_sorted[1:]:
        try:
            # 添加分页符
            add_page_break(master_doc)
            
            # 读取章节文档
            chapter_doc = Document(file_path)
            
            # 复制段落
            for para in chapter_doc.paragraphs:
                # 创建新段落
                new_para = master_doc.add_paragraph()
                new_para.style = para.style
                new_para.alignment = para.alignment
                
                # 复制段落格式
                new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
                new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                new_para.paragraph_format.space_before = para.paragraph_format.space_before
                new_para.paragraph_format.space_after = para.paragraph_format.space_after
                new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
                
                # 复制文本和格式
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                    if run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
            
            # 复制表格
            for table in chapter_doc.tables:
                # 简化：仅复制文本内容
                # 完整实现需要复制表格样式、单元格合并等
                rows = len(table.rows)
                cols = len(table.columns)
                new_table = master_doc.add_table(rows=rows, cols=cols)
                
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.rows[i].cells[j].text = cell.text
            
            print(f"✅ 已合并：{os.path.basename(file_path)}")
            
        except Exception as e:
            print(f"⚠️  合并失败：{os.path.basename(file_path)} - {str(e)}")
            continue
    
    # 保存合并后的文档
    try:
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        master_doc.save(output_path)
        return {
            'success': True,
            'output_path': output_path,
            'merged_files': len(file_list_sorted),
            'file_list': [os.path.basename(f) for f in file_list_sorted]
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'保存文件失败：{str(e)}'
        }


def generate_toc(doc):
    """
    生成目录（简化版）
    注意：完整的 Word 目录需要 VBA 或 COM 自动化
    
    Args:
        doc: Document 对象
    """
    toc_entries = []
    
    # 扫描所有标题
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            toc_entries.append({
                'text': para.text,
                'level': level
            })
    
    # 在文档开头插入目录
    toc_para = doc.paragraphs[0].insert_paragraph_before()
    toc_para.add_run('目  录').bold = True
    toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_para.paragraph_format.space_after = Pt(24)
    
    for entry in toc_entries:
        toc_line = doc.paragraphs[1].insert_paragraph_before()
        indent = (entry['level'] - 1) * 0.74  # 每级缩进 0.74cm
        toc_line.paragraph_format.left_indent = Cm(indent)
        toc_line.add_run(entry['text'])
        toc_line.add_run('.' * 30)  # 占位符
        toc_line.add_run('[页码待更新]')
    
    print("✅ 目录已生成（页码需要在 Word 中手动更新字段）")


def add_header_footer(doc, thesis_title):
    """
    添加页眉页脚
    
    Args:
        doc: Document 对象
        thesis_title: 论文标题
    """
    section = doc.sections[0]
    
    # 页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"中南大学博士学位论文                                {thesis_title}"
    header_para.style = doc.styles['Header']
    
    # 页脚（页码）
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加页码字段
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    print("✅ 页眉页脚已添加")


def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='合并 Word 章节文档')
    parser.add_argument('--input-dir', required=True, help='章节文件所在目录')
    parser.add_argument('--output', required=True, help='输出文件路径')
    parser.add_argument('--title', default='论文标题', help='论文标题（用于页眉）')
    parser.add_argument('--add-toc', action='store_true', help='生成目录')
    parser.add_argument('--add-header', action='store_true', help='添加页眉页脚')
    
    args = parser.parse_args()
    
    # 获取所有 docx 文件
    if not os.path.exists(args.input_dir):
        print(f"❌ 目录不存在：{args.input_dir}")
        sys.exit(1)
    
    file_list = []
    for filename in os.listdir(args.input_dir):
        if filename.endswith('.docx') and not filename.startswith('~'):
            file_path = os.path.join(args.input_dir, filename)
            file_list.append(file_path)
    
    if not file_list:
        print(f"❌ 目录中没有找到 docx 文件：{args.input_dir}")
        sys.exit(1)
    
    print(f"📁 找到 {len(file_list)} 个章节文件")
    print(f"🔄 开始合并...")
    
    # 执行合并
    result = merge_docx_files(file_list, args.output)
    
    if not result['success']:
        print(f"❌ 合并失败：{result['error']}")
        sys.exit(1)
    
    print(f"\n✅ 合并完成！")
    print(f"📄 输出文件：{result['output_path']}")
    print(f"📊 合并章节数：{result['merged_files']}")
    print(f"📝 章节列表：")
    for filename in result['file_list']:
        print(f"   - {filename}")
    
    # 可选：生成目录
    if args.add_toc:
        print(f"\n📑 正在生成目录...")
        doc = Document(args.output)
        generate_toc(doc)
        doc.save(args.output)
    
    # 可选：添加页眉页脚
    if args.add_header:
        print(f"\n📋 正在添加页眉页脚...")
        doc = Document(args.output)
        add_header_footer(doc, args.title)
        doc.save(args.output)
    
    print(f"\n🎉 所有操作完成！")


if __name__ == '__main__':
    main()
