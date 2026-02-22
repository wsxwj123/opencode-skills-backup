#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 转 Word 工具（应用中南大学样式）

功能：
1. 解析 Markdown 文本
2. 转换为 Word 文档
3. 自动应用中南大学博士论文样式
4. 处理图表占位符

作者：Sci2Doc Team
日期：2024-03-15
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys
import os


def parse_markdown_line(line):
    """
    解析 Markdown 行，识别类型
    
    Returns:
        tuple: (type, content, level)
    """
    line = line.rstrip()
    
    # 空行
    if not line.strip():
        return ('empty', '', 0)
    
    # 一级标题 # Title
    if line.startswith('# ') and not line.startswith('## '):
        return ('heading1', line[2:].strip(), 1)
    
    # 二级标题 ## Title
    elif line.startswith('## ') and not line.startswith('### '):
        return ('heading2', line[3:].strip(), 2)
    
    # 三级标题 ### Title
    elif line.startswith('### '):
        return ('heading3', line[4:].strip(), 3)
    
    # 图片占位符 [图 1-1：标题]
    elif re.match(r'\[图\s*\d+-\d+[：:].+\]', line):
        return ('figure', line.strip(), 0)
    
    # 表格占位符 [表 1-1：标题]
    elif re.match(r'\[表\s*\d+-\d+[：:].+\]', line):
        return ('table', line.strip(), 0)
    
    # 正文段落
    else:
        return ('paragraph', line.strip(), 0)


def set_run_font(run, latin, east_asia, size_pt, bold=None):
    """
    同时设置拉丁字体和东亚字体，避免 Word 回退到意外字体。
    """
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), east_asia)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def apply_csu_heading1_style(paragraph):
    """应用一级标题样式"""
    paragraph.style = "Heading 1"
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    
    for run in paragraph.runs:
        set_run_font(run, latin='Times New Roman', east_asia='SimHei', size_pt=16, bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)


def apply_csu_heading2_style(paragraph):
    """应用二级标题样式"""
    paragraph.style = "Heading 2"
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    
    for run in paragraph.runs:
        set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=14, bold=False)


def apply_csu_heading3_style(paragraph):
    """应用三级标题样式"""
    paragraph.style = "Heading 3"
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    
    for run in paragraph.runs:
        set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=12, bold=False)


def apply_csu_normal_style(paragraph):
    """应用正文样式"""
    paragraph.style = "Normal"
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)  # 2字符
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    
    for run in paragraph.runs:
        set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=12, bold=False)


def apply_csu_caption_style(paragraph):
    """应用图表题注样式"""
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    
    for run in paragraph.runs:
        set_run_font(run, latin='Times New Roman', east_asia='KaiTi', size_pt=10.5, bold=False)


def _set_cell_border(cell, **kwargs):
    """
    设置单元格边框。
    
    用法: _set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                                  bottom={"sz": 12, "val": "single", "color": "000000"})
    sz 单位为 1/8 pt，所以 12 = 1.5pt, 4 = 0.5pt
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge, attrs in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tcBorders.append(element)
        for attr_name, attr_val in attrs.items():
            element.set(qn(f'w:{attr_name}'), str(attr_val))


def _clear_table_borders(table):
    """清除表格级别的所有边框（后续由单元格级别精确控制）"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{border_name}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _parse_table_row(line):
    """解析 Markdown 表格行，返回单元格列表，如果不是表格行返回 None"""
    line = line.strip()
    if not line.startswith('|') or not line.endswith('|'):
        return None
    cells = [c.strip() for c in line.strip('|').split('|')]
    return cells


def _is_separator_row(cells):
    """判断是否为 Markdown 表格分隔行（如 | --- | --- |）"""
    if not cells:
        return False
    return all(re.match(r'^:?-{1,}:?$', c.strip()) for c in cells)


def create_three_line_table(doc, headers, data_rows):
    """
    在文档中创建三线表。
    
    Args:
        doc: Document 对象
        headers: 表头列表 ['col1', 'col2', ...]
        data_rows: 数据行列表 [['val1', 'val2', ...], ...]
    
    Returns:
        table: 创建的表格对象
    """
    num_cols = len(headers)
    num_rows = 1 + len(data_rows)  # 表头 + 数据行
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 清除默认边框
    _clear_table_borders(table)
    
    # 定义边框样式
    thick = {"sz": "12", "val": "single", "color": "000000", "space": "0"}   # 1.5pt
    thin  = {"sz": "4",  "val": "single", "color": "000000", "space": "0"}   # 0.5pt
    none  = {"sz": "0",  "val": "none",   "color": "auto",   "space": "0"}
    
    total_rows = len(table.rows)
    
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            # 填充内容
            if row_idx == 0:
                cell.text = headers[col_idx] if col_idx < len(headers) else ''
            else:
                data_row = data_rows[row_idx - 1]
                cell.text = data_row[col_idx] if col_idx < len(data_row) else ''
            
            # 设置字体
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.first_line_indent = Cm(0)
                for run in paragraph.runs:
                    if row_idx == 0:
                        # 表头：黑体加粗
                        set_run_font(run, latin='Times New Roman', east_asia='SimHei', size_pt=10.5, bold=True)
                    else:
                        # 表体：宋体
                        set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=10.5, bold=False)
            
            # 设置边框：精确控制每个单元格
            # 左右边框始终为 none
            borders = {"left": none, "right": none}
            
            if row_idx == 0:
                # 表头行：顶部粗线，底部细线
                borders["top"] = thick
                borders["bottom"] = thin
            elif row_idx == total_rows - 1:
                # 最后一行：顶部无线，底部粗线
                borders["top"] = none
                borders["bottom"] = thick
            else:
                # 中间数据行：无边框
                borders["top"] = none
                borders["bottom"] = none
            
            _set_cell_border(cell, **borders)
    
    return table


def markdown_to_docx(md_content, output_path, chapter_num=None):
    """
    将 Markdown 内容转换为 Word 文档
    
    Args:
        md_content: Markdown 文本内容
        output_path: 输出文件路径
        chapter_num: 章节号（可选）
    
    Returns:
        bool: 转换是否成功
    """
    try:
        # 创建文档
        doc = Document()
        
        # 设置页面
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        
        # 逐行解析（支持多行表格累积）
        lines = md_content.split('\n')
        table_buffer = []       # 累积表格行
        table_caption = None    # 表格题注（表 X-X：标题）
        pending_caption = None  # 上一行是表格题注，等待判断下一行是否为表格
        
        def flush_table(doc, table_buffer, table_caption):
            """将累积的表格行转换为三线表并写入文档"""
            if not table_buffer:
                return
            
            headers = None
            data_rows = []
            
            for row_cells in table_buffer:
                if headers is None:
                    headers = row_cells
                elif _is_separator_row(row_cells):
                    continue  # 跳过分隔行
                else:
                    data_rows.append(row_cells)
            
            if headers:
                # 先写题注（表格上方）
                if table_caption:
                    caption_text = table_caption.strip('[]')
                    para = doc.add_paragraph(caption_text)
                    apply_csu_caption_style(para)
                
                # 创建三线表
                create_three_line_table(doc, headers, data_rows)
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # 尝试解析为表格行
            row_cells = _parse_table_row(line)
            
            if row_cells is not None:
                # 当前行是表格行
                if pending_caption is not None:
                    # 前一行是表格题注，确认关联
                    table_caption = pending_caption
                    pending_caption = None
                table_buffer.append(row_cells)
                i += 1
                continue
            
            # 当前行不是表格行 → 先 flush 之前累积的表格
            if table_buffer:
                flush_table(doc, table_buffer, table_caption)
                table_buffer = []
                table_caption = None
            
            # 如果有未消费的 pending_caption（后面没跟表格），作为普通题注输出
            if pending_caption is not None:
                para = doc.add_paragraph(pending_caption.strip('[]'))
                apply_csu_caption_style(para)
                pending_caption = None
            
            # 正常解析当前行
            line_type, content, level = parse_markdown_line(line)
            
            if line_type == 'empty':
                i += 1
                continue
            
            elif line_type == 'heading1':
                para = doc.add_heading(content, level=1)
                apply_csu_heading1_style(para)
            
            elif line_type == 'heading2':
                para = doc.add_heading(content, level=2)
                apply_csu_heading2_style(para)
            
            elif line_type == 'heading3':
                para = doc.add_heading(content, level=3)
                apply_csu_heading3_style(para)
            
            elif line_type == 'figure':
                para = doc.add_paragraph(content)
                apply_csu_caption_style(para)
            
            elif line_type == 'table':
                # 表格题注：暂存，等下一行判断是否跟着 Markdown 表格
                pending_caption = content
            
            elif line_type == 'paragraph':
                # 检查是否为独立的表格题注行（不带方括号的 表 X-X：格式）
                if re.match(r'^表\s*\d+-\d+[：:]', content):
                    pending_caption = content
                elif content:
                    para = doc.add_paragraph(content)
                    apply_csu_normal_style(para)
            
            i += 1
        
        # 文件结束时 flush 残留的表格
        if table_buffer:
            flush_table(doc, table_buffer, table_caption)
        elif pending_caption is not None:
            para = doc.add_paragraph(pending_caption.strip('[]'))
            apply_csu_caption_style(para)
        
        # 保存文档
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        doc.save(output_path)
        print(f"✅ 转换成功：{output_path}")
        return True
    
    except Exception as e:
        print(f"❌ 转换失败：{str(e)}")
        return False


def convert_markdown_file(md_file_path, output_path=None):
    """
    转换 Markdown 文件
    
    Args:
        md_file_path: Markdown 文件路径
        output_path: 输出文件路径（可选）
    """
    if not os.path.exists(md_file_path):
        print(f"❌ 文件不存在：{md_file_path}")
        return False
    
    # 读取 Markdown 文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 确定输出路径
    if output_path is None:
        output_path = md_file_path.replace('.md', '.docx')
    
    # 提取章节号（如果文件名包含）
    chapter_num = None
    match = re.search(r'第(\d+)章', os.path.basename(md_file_path))
    if match:
        chapter_num = int(match.group(1))
    
    # 执行转换
    return markdown_to_docx(md_content, output_path, chapter_num)


def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Markdown 转 Word（中南大学样式）')
    parser.add_argument('input', help='输入 Markdown 文件路径')
    parser.add_argument('-o', '--output', help='输出 Word 文件路径（可选）')
    parser.add_argument('-c', '--chapter', type=int, help='章节号（可选）')
    
    args = parser.parse_args()
    
    # 执行转换
    success = convert_markdown_file(args.input, args.output)
    
    if success:
        print("\n📋 样式说明：")
        print("  - 一级标题：三号黑体加粗，居中")
        print("  - 二级标题：四号宋体，顶格")
        print("  - 三级标题：小四号宋体，顶格")
        print("  - 正文：小四号宋体/Times New Roman，首行缩进")
        print("  - 图表题注：五号楷体，居中")
        sys.exit(0)
    else:
        sys.exit(1)


if __name__ == '__main__':
    main()
