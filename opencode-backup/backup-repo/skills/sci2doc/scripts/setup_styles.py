#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
初始化 Word 文档样式（中南大学博士学位论文标准）

功能：
1. 创建基础样式模板 docx 文件
2. 设置中南大学要求的所有样式
3. 可被其他脚本引用以应用统一样式

作者：Sci2Doc Team
日期：2024-03-15
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import sys
import os


def create_csu_styles(output_path):
    """
    创建符合中南大学标准的 Word 样式模板
    
    Args:
        output_path: 输出文件路径
    """
    doc = Document()
    
    # ========== 页面设置 ==========
    section = doc.sections[0]
    
    # 页边距（单位：厘米）
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    
    # 页眉页脚距离
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    
    # ========== 样式定义 ==========
    styles = doc.styles
    
    # 1. 正文样式
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)  # 小四号
    
    normal_para = normal_style.paragraph_format
    normal_para.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal_para.line_spacing = Pt(20)
    normal_para.first_line_indent = Cm(0.74)  # 首行缩进 2 字符
    normal_para.space_before = Pt(0)
    normal_para.space_after = Pt(0)
    
    # 2. 一级标题样式
    heading1_style = styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'SimHei'  # 黑体
    heading1_font.size = Pt(16)  # 三号
    heading1_font.bold = True
    heading1_font.color.rgb = RGBColor(0, 0, 0)
    
    heading1_para = heading1_style.paragraph_format
    heading1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1_para.space_before = Pt(18)
    heading1_para.space_after = Pt(12)
    heading1_para.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading1_para.line_spacing = Pt(20)
    heading1_para.page_break_before = True
    
    # 3. 二级标题样式
    heading2_style = styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'SimSun'  # 宋体
    heading2_font.size = Pt(14)  # 四号
    heading2_font.bold = False
    
    heading2_para = heading2_style.paragraph_format
    heading2_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading2_para.space_before = Pt(10)
    heading2_para.space_after = Pt(8)
    heading2_para.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading2_para.line_spacing = Pt(20)
    heading2_para.first_line_indent = Cm(0)
    
    # 4. 三级标题样式
    heading3_style = styles['Heading 3']
    heading3_font = heading3_style.font
    heading3_font.name = 'SimSun'  # 宋体
    heading3_font.size = Pt(12)  # 小四号
    heading3_font.bold = False
    
    heading3_para = heading3_style.paragraph_format
    heading3_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading3_para.space_before = Pt(10)
    heading3_para.space_after = Pt(8)
    heading3_para.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading3_para.line_spacing = Pt(20)
    heading3_para.first_line_indent = Cm(0)
    
    # ========== 添加示例内容 ==========
    doc.add_heading('第一章 示例章节', level=1)
    doc.add_heading('1.1 示例节标题', level=2)
    doc.add_heading('1.1.1 示例小节标题', level=3)
    
    para = doc.add_paragraph(
        '这是正文段落示例。中文使用小四号宋体，英文使用小四号 Times New Roman，'
        '行距固定值 20 磅，首行缩进 2 字符。This is an example paragraph with '
        'mixed Chinese and English text following CSU thesis standards.'
    )
    
    # ========== 保存文件 ==========
    doc.save(output_path)
    print(f"✅ 样式模板已创建：{output_path}")
    return output_path


def apply_csu_style_to_paragraph(para, style_type='normal'):
    """
    应用中南大学样式到段落
    
    Args:
        para: 段落对象
        style_type: 样式类型 ('normal', 'heading1', 'heading2', 'heading3')
    """
    if style_type == 'normal':
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para.paragraph_format.line_spacing = Pt(20)
        para.paragraph_format.first_line_indent = Cm(0.74)
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    elif style_type == 'heading1':
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
        for run in para.runs:
            run.font.name = 'SimHei'
            run.font.size = Pt(16)
            run.font.bold = True
    
    elif style_type == 'heading2':
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(8)
        for run in para.runs:
            run.font.name = 'SimSun'
            run.font.size = Pt(14)
    
    elif style_type == 'heading3':
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(8)
        for run in para.runs:
            run.font.name = 'SimSun'
            run.font.size = Pt(12)


def main():
    """命令行入口"""
    if len(sys.argv) < 2:
        output_path = "csu_styles_template.docx"
    else:
        output_path = sys.argv[1]
    
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    create_csu_styles(output_path)
    
    print("\n📋 样式说明：")
    print("  - 正文：小四号宋体/Times New Roman，行距 20 磅，首行缩进 2 字符")
    print("  - 一级标题：三号黑体加粗，居中，段前 18 磅，段后 12 磅")
    print("  - 二级标题：四号宋体，顶格，段前 10 磅，段后 8 磅")
    print("  - 三级标题：小四号宋体，顶格，段前 10 磅，段后 8 磅")
    print("\n💡 使用方法：")
    print(f"  from docx import Document")
    print(f"  doc = Document('{output_path}')")
    print(f"  # 继续添加内容...")


if __name__ == '__main__':
    main()
