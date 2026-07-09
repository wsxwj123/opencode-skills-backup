#!/usr/bin/env python3
from __future__ import annotations

import argparse
import itertools
import json
import re
import sys
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from common import read_json, read_docx_paragraphs, strip_inline_format_markers


PROFILE_PRESETS: dict[str, dict[str, Any]] = {
    "journal-manuscript": {
        "body_font": "Times New Roman",
        "body_size": 11,
        "title_font": "Arial",
        "title_size": 16,
        "heading_font": "Arial",
        "heading1_size": 13,
        "heading2_size": 12,
        "heading3_size": 11.5,
        "heading4_size": 10.5,
        "label_font": "Arial",
        "body_line_spacing": 1.15,
        "body_space_after": 6,
        "header_text_align": WD_ALIGN_PARAGRAPH.CENTER,
        "ref_size": 9.5,
        "abstract_size": 10.5,
        "table_text_size": 9.5,
        "table_header_fill": "D9E2F3",
        "table_header_font": "Arial",
    },
    "nature-review": {
        "body_font": "Garamond",
        "body_size": 11,
        "title_font": "Arial",
        "title_size": 17,
        "heading_font": "Arial",
        "heading1_size": 13.5,
        "heading2_size": 12.5,
        "heading3_size": 11.5,
        "heading4_size": 10.5,
        "label_font": "Arial",
        "body_line_spacing": 1.12,
        "body_space_after": 5,
        "header_text_align": WD_ALIGN_PARAGRAPH.CENTER,
        "ref_size": 9.5,
        "abstract_size": 10.5,
        "table_text_size": 9.5,
        "table_header_fill": "EAF2F8",
        "table_header_font": "Arial",
    },
    "cell-press": {
        "body_font": "Times New Roman",
        "body_size": 11,
        "title_font": "Georgia",
        "title_size": 16.5,
        "heading_font": "Georgia",
        "heading1_size": 13,
        "heading2_size": 12,
        "heading3_size": 11,
        "heading4_size": 10.5,
        "label_font": "Georgia",
        "body_line_spacing": 1.15,
        "body_space_after": 6,
        "header_text_align": WD_ALIGN_PARAGRAPH.CENTER,
        "ref_size": 9.5,
        "abstract_size": 10.5,
        "table_text_size": 9.5,
        "table_header_fill": "F2E5D7",
        "table_header_font": "Georgia",
    },
    "lancet-review": {
        "body_font": "Times New Roman",
        "body_size": 10.5,
        "title_font": "Arial",
        "title_size": 16,
        "heading_font": "Arial",
        "heading1_size": 12.5,
        "heading2_size": 11.5,
        "heading3_size": 11,
        "heading4_size": 10.5,
        "label_font": "Arial",
        "body_line_spacing": 1.15,
        "body_space_after": 5,
        "header_text_align": WD_ALIGN_PARAGRAPH.CENTER,
        "ref_size": 9,
        "abstract_size": 10,
        "table_text_size": 9,
        "table_header_fill": "F3F6F9",
        "table_header_font": "Arial",
    },
    "review-response": {
        "body_font": "Times New Roman",
        "body_size": 11,
        "title_font": "Arial",
        "title_size": 15.5,
        "heading_font": "Arial",
        "heading1_size": 13.5,
        "heading2_size": 12,
        "heading3_size": 11.5,
        "heading4_size": 10.5,
        "label_font": "Arial",
        "body_line_spacing": 1.15,
        "body_space_after": 5,
        "header_text_align": WD_ALIGN_PARAGRAPH.CENTER,
        "ref_size": 9.5,
        "abstract_size": 10.5,
        "table_text_size": 9.5,
        "table_header_fill": "DDEBF7",
        "table_header_font": "Arial",
    },
}


def doc_for_template(template_path: str) -> Document:
    if template_path:
        return Document(template_path)
    return Document()


INLINE_TOKEN_RE = re.compile(
    r"(\*\*.+?\*\*)"          # bold (matched before italic)
    r"|(\*.+?\*)"            # italic
    r"|(<sup>.*?</sup>)"     # superscript
    r"|(<sub>.*?</sub>)",    # subscript
    re.DOTALL,
)


def add_runs_with_bold(paragraph, text: str) -> None:
    """Render inline markers (**bold**, *italic*, <sup>, <sub>) as real Word runs.

    `**bold**` is matched before `*italic*` so a bold span is not mis-split by the
    single-asterisk branch. Markers are non-nesting (flat); unmatched text is plain.
    """
    pos = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if match.group(1) is not None:
            paragraph.add_run(token[2:-2]).bold = True
        elif match.group(2) is not None:
            paragraph.add_run(token[1:-1]).italic = True
        elif match.group(3) is not None:
            paragraph.add_run(token[5:-6]).font.superscript = True
        elif match.group(4) is not None:
            paragraph.add_run(token[5:-6]).font.subscript = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def profile_settings(profile_name: str) -> dict[str, Any]:
    return PROFILE_PRESETS.get(profile_name, PROFILE_PRESETS["journal-manuscript"])


DEFAULT_CJK_FONT = "宋体"


def set_style_font(style, name: str, size: float, bold: bool = False, east_asian: str = DEFAULT_CJK_FONT) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    # Set w:eastAsia so CJK glyphs fall back to a real Chinese face instead of the
    # Latin body font (which renders Chinese as tofu / wrong metrics).
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east_asian)


def ensure_custom_styles(doc: Document, profile_name: str, doc_kind: str) -> None:
    profile = profile_settings(profile_name)
    style_names = {style.name for style in doc.styles}
    if "ReviseSciTitle" not in style_names:
        style = doc.styles.add_style("ReviseSciTitle", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Title"]
    if "ReviseSciAbstractLabel" not in style_names:
        style = doc.styles.add_style("ReviseSciAbstractLabel", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
    if "ReviseSciAbstractBody" not in style_names:
        style = doc.styles.add_style("ReviseSciAbstractBody", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
    if "ReviseSciBody" not in style_names:
        style = doc.styles.add_style("ReviseSciBody", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
    if "ReviseSciRefEntry" not in style_names:
        style = doc.styles.add_style("ReviseSciRefEntry", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
    if "ReviseSciLabel" not in style_names:
        style = doc.styles.add_style("ReviseSciLabel", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
    if "ReviseSciReviewerHeading" not in style_names:
        style = doc.styles.add_style("ReviseSciReviewerHeading", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Heading 1"]
    if "ReviseSciCommentHeading" not in style_names:
        style = doc.styles.add_style("ReviseSciCommentHeading", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Heading 3"]
    normal_style = doc.styles["Normal"]
    set_style_font(normal_style, profile["body_font"], profile["body_size"])

    set_style_font(doc.styles["ReviseSciTitle"], profile["title_font"], profile["title_size"], bold=True)
    set_style_font(doc.styles["ReviseSciAbstractLabel"], profile["label_font"], profile["body_size"], bold=True)
    set_style_font(doc.styles["ReviseSciAbstractBody"], profile["body_font"], profile["abstract_size"])
    set_style_font(doc.styles["ReviseSciBody"], profile["body_font"], profile["body_size"])
    set_style_font(doc.styles["ReviseSciRefEntry"], profile["body_font"], profile["ref_size"])
    set_style_font(doc.styles["ReviseSciLabel"], profile["label_font"], profile["body_size"], bold=True)
    set_style_font(doc.styles["ReviseSciReviewerHeading"], profile["heading_font"], profile["heading1_size"], bold=True)
    set_style_font(doc.styles["ReviseSciCommentHeading"], profile["heading_font"], profile["heading3_size"], bold=True)

    title_style = doc.styles["ReviseSciTitle"]
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_before = Pt(6)
    title_style.paragraph_format.space_after = Pt(12)

    abstract_label_style = doc.styles["ReviseSciAbstractLabel"]
    abstract_label_style.paragraph_format.space_before = Pt(8)
    abstract_label_style.paragraph_format.space_after = Pt(0)

    abstract_body_style = doc.styles["ReviseSciAbstractBody"]
    abstract_body_style.paragraph_format.line_spacing = profile["body_line_spacing"]
    abstract_body_style.paragraph_format.space_after = Pt(6)
    abstract_body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    normal_style = doc.styles["Normal"]
    body_style = doc.styles["ReviseSciBody"]
    body_style.paragraph_format.line_spacing = profile["body_line_spacing"]
    body_style.paragraph_format.space_after = Pt(profile["body_space_after"])
    body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ref_style = doc.styles["ReviseSciRefEntry"]
    ref_style.paragraph_format.line_spacing = 1.0
    ref_style.paragraph_format.space_after = Pt(2)
    ref_style.paragraph_format.left_indent = Inches(0.25)
    ref_style.paragraph_format.first_line_indent = Inches(-0.25)
    ref_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    label_style = doc.styles["ReviseSciLabel"]
    label_style.paragraph_format.space_before = Pt(6)
    label_style.paragraph_format.space_after = Pt(0)
    reviewer_style = doc.styles["ReviseSciReviewerHeading"]
    reviewer_style.paragraph_format.space_before = Pt(12)
    reviewer_style.paragraph_format.space_after = Pt(6)
    reviewer_style.paragraph_format.keep_with_next = True
    comment_style = doc.styles["ReviseSciCommentHeading"]
    comment_style.paragraph_format.space_before = Pt(10)
    comment_style.paragraph_format.space_after = Pt(4)
    comment_style.paragraph_format.keep_with_next = True
    for heading_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = doc.styles[heading_name]
        set_style_font(style, profile["heading_font"], profile[f"{heading_name.lower().replace(' ', '')}_size"], bold=True)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(6)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(10)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(4)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(3)
    doc.styles["Heading 4"].paragraph_format.space_before = Pt(6)
    doc.styles["Heading 4"].paragraph_format.space_after = Pt(2)
    if doc_kind == "response":
        reviewer_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        comment_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_markdown_paragraph(doc: Document, line: str, in_references: bool = False, in_abstract: bool = False) -> None:
    label_match = re.fullmatch(r"\*\*(.+?)\*\*", line.strip())
    if label_match:
        paragraph = doc.add_paragraph(style="ReviseSciLabel")
        paragraph.add_run(label_match.group(1)).bold = True
        return
    bullet_match = re.match(r"^- (.+)$", line)
    if bullet_match:
        paragraph = doc.add_paragraph(style="List Bullet")
        add_runs_with_bold(paragraph, bullet_match.group(1))
        return
    number_match = re.match(r"^\d+\. (.+)$", line)
    if number_match:
        paragraph = doc.add_paragraph(style="ReviseSciRefEntry" if in_references else "List Number")
        add_runs_with_bold(paragraph, line if in_references else number_match.group(1))
        return
    paragraph = doc.add_paragraph(style="ReviseSciRefEntry" if in_references else ("ReviseSciAbstractBody" if in_abstract else "ReviseSciBody"))
    add_runs_with_bold(paragraph, line)


def parse_markdown_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def is_markdown_table_separator(line: str) -> bool:
    cells = parse_markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_markdown_table(doc: Document, rows: list[list[str]], profile_name: str) -> None:
    if not rows:
        return
    profile = profile_settings(profile_name)
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(column_count):
            text = row[col_idx] if col_idx < len(row) else ""
            cell_paragraph = table.cell(row_idx, col_idx).paragraphs[0]
            add_runs_with_bold(cell_paragraph, text)
            for run in cell_paragraph.runs:
                run.font.name = profile["body_font"]
                run.font.size = Pt(profile["table_text_size"])
            if row_idx == 0:
                for run in cell_paragraph.runs:
                    run.bold = True
                    run.font.name = profile["table_header_font"]
                cell = table.cell(row_idx, col_idx)
                tc_pr = cell._tc.get_or_add_tcPr()
                shade = OxmlElement("w:shd")
                shade.set(qn("w:fill"), profile["table_header_fill"])
                tc_pr.append(shade)
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def append_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def configure_document(doc: Document, header_text: str = "", profile_name: str = "journal-manuscript", doc_kind: str = "manuscript") -> None:
    profile = profile_settings(profile_name)
    ensure_custom_styles(doc, profile_name, doc_kind)
    normal_style = doc.styles["Normal"]
    normal_style.font.name = profile["body_font"]
    normal_style.font.size = Pt(profile["body_size"])
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)
        if header_text:
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = header_text
            header_para.alignment = profile["header_text_align"]
            if header_para.runs:
                header_para.runs[0].bold = True
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        append_field(footer_para, "PAGE", "1")


def add_toc_block(doc: Document) -> None:
    toc_title = doc.add_paragraph()
    toc_title.add_run("Table of Contents").bold = True
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_para = doc.add_paragraph()
    append_field(toc_para, 'TOC \\o "1-4" \\h \\z \\u', "Right-click to update field.")
    doc.add_paragraph("")


def markdown_to_docx(
    md_path: Path,
    docx_path: Path,
    template_path: str = "",
    page_break_before_comment: bool = False,
    header_text: str = "",
    include_toc: bool = False,
    profile_name: str = "journal-manuscript",
    doc_kind: str = "manuscript",
) -> None:
    doc = doc_for_template(template_path)
    configure_document(doc, header_text, profile_name=profile_name, doc_kind=doc_kind)
    lines = md_path.read_text(encoding="utf-8").splitlines()
    seen_comment = False
    toc_inserted = False
    seen_title = False
    in_references = False
    in_abstract = False
    idx = 0
    while idx < len(lines):
        raw_line = lines[idx]
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            idx += 1
            continue
        if "|" in line and idx + 1 < len(lines) and is_markdown_table_separator(lines[idx + 1]):
            table_rows = [parse_markdown_table_row(line)]
            idx += 2
            while idx < len(lines):
                next_line = lines[idx].rstrip()
                if not next_line or "|" not in next_line:
                    break
                table_rows.append(parse_markdown_table_row(next_line))
                idx += 1
            add_markdown_table(doc, table_rows, profile_name)
            continue
        if line.startswith("### Comment") and page_break_before_comment:
            if seen_comment:
                doc.add_page_break()
            seen_comment = True
        if line.startswith("#### "):
            in_abstract = False
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            in_abstract = False
            paragraph = doc.add_paragraph(style="ReviseSciCommentHeading")
            paragraph.add_run(line[4:]).bold = True
        elif line.startswith("## "):
            heading_text = line[3:]
            in_references = heading_text.lower() in {"references", "reference", "参考文献"}
            in_abstract = heading_text.lower() in {"abstract", "摘要"}
            if in_abstract:
                paragraph = doc.add_paragraph(style="ReviseSciAbstractLabel")
                paragraph.add_run(heading_text).bold = True
            else:
                doc.add_heading(heading_text, level=2)
        elif line.startswith("# "):
            heading_text = line[2:]
            in_references = heading_text.lower() in {"references", "reference", "参考文献"}
            in_abstract = heading_text.lower() in {"abstract", "摘要"}
            if doc_kind == "manuscript" and not seen_title:
                paragraph = doc.add_paragraph(style="ReviseSciTitle")
                paragraph.add_run(heading_text).bold = True
                seen_title = True
            elif page_break_before_comment and (heading_text.startswith("Reviewer #") or heading_text.startswith("Editor")):
                paragraph = doc.add_paragraph(style="ReviseSciReviewerHeading")
                paragraph.add_run(heading_text).bold = True
            else:
                doc.add_heading(heading_text, level=1)
            if include_toc and not toc_inserted:
                add_toc_block(doc)
                toc_inserted = True
        else:
            add_markdown_paragraph(doc, line, in_references=in_references, in_abstract=in_abstract)
        idx += 1
    doc.save(str(docx_path))


# ---------------------------------------------------------------------------
# In-place export: edit the ORIGINAL manuscript docx, replacing only the text of
# paragraphs that revise actually changed. Untouched paragraphs, tables, images
# and styles are preserved byte-for-byte (we never recreate the document).
# ---------------------------------------------------------------------------


def _paragraph_base_font(paragraph) -> dict[str, Any]:
    """Capture the paragraph's base run font so rebuilt runs inherit name/size/
    eastAsia instead of falling back to Normal. Prefers the first run with text;
    falls back to the paragraph style's font."""
    name = None
    size = None
    east_asia = None
    for run in paragraph.runs:
        if run.text and run.text.strip():
            name = run.font.name
            size = run.font.size
            try:
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is not None:
                    east_asia = rfonts.get(qn("w:eastAsia"))
            except Exception:
                east_asia = None
            break
    style = getattr(paragraph, "style", None)
    if name is None and style is not None:
        name = getattr(getattr(style, "font", None), "name", None)
    if size is None and style is not None:
        size = getattr(getattr(style, "font", None), "size", None)
    return {"name": name, "size": size, "east_asia": east_asia}


def _paragraph_has_embedded_image(paragraph) -> bool:
    """True if the paragraph contains an inline/floating image (<w:drawing>) or a
    legacy embedded object (<w:object>). Clearing such a paragraph's runs would
    delete the picture and leave an orphan image relationship, so callers must
    skip in-place run rebuild for these paragraphs."""
    p = paragraph._p
    return bool(
        p.findall(".//" + qn("w:drawing"))
        or p.findall(".//" + qn("w:object"))
    )


def _clear_paragraph_runs(paragraph) -> None:
    """Remove every run element from a paragraph, keeping its pPr (alignment,
    style, indentation) intact so paragraph-level formatting survives."""
    p = paragraph._p
    for run in list(paragraph.runs):
        p.remove(run._r)


def _apply_base_font(run, base: dict[str, Any]) -> None:
    if base.get("name"):
        run.font.name = base["name"]
    if base.get("size") is not None:
        run.font.size = base["size"]
    if base.get("east_asia"):
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), base["east_asia"])


def rebuild_paragraph_runs(paragraph, text: str, base: dict[str, Any]) -> None:
    """Replace paragraph content with runs built from inline markers, each run
    inheriting the captured base font and layering italic/bold/sup/sub on top."""
    _clear_paragraph_runs(paragraph)
    pos = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            _apply_base_font(run, base)
        if match.group(1) is not None:
            run = paragraph.add_run(match.group(0)[2:-2]); run.bold = True
        elif match.group(2) is not None:
            run = paragraph.add_run(match.group(0)[1:-1]); run.italic = True
        elif match.group(3) is not None:
            run = paragraph.add_run(match.group(0)[5:-6]); run.font.superscript = True
        elif match.group(4) is not None:
            run = paragraph.add_run(match.group(0)[5:-6]); run.font.subscript = True
        else:
            run = paragraph.add_run(match.group(0))
        _apply_base_font(run, base)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _apply_base_font(run, base)


# ---------------------------------------------------------------------------
# Word-level track-changes rebuild (opt-in via --track-changes). Instead of
# replacing a changed paragraph wholesale, diff original vs current at the word
# level and emit real Word revisions: <w:ins> for added spans, <w:del> for
# removed spans. Inline format (italic/bold/sup/sub) is carried on each token so
# emphasis survives the word split and lands on the right side of every change.
# ---------------------------------------------------------------------------

# Tokenizer: English words (with internal apostrophe/hyphen), numbers, whitespace
# runs, single CJK characters (Chinese has no spaces -> split per char), and any
# other single punctuation char. Splitting whitespace as its own token keeps
# equal spans equal so only truly-changed words fall inside <w:ins>/<w:del>.
WORD_SPLIT_RE = re.compile(
    r"[A-Za-z]+(?:[’'\-][A-Za-z]+)*"   # english word
    r"|\d+(?:[.,]\d+)*%?"              # number (optional decimal/percent)
    r"|\s+"                            # whitespace run
    r"|[一-鿿]"               # single CJK char
    r"|[^\sA-Za-z0-9一-鿿]"   # any other single char (punctuation)
)

# fmt tuple layout: (bold, italic, superscript, subscript). Hashable so tokens can
# feed difflib.SequenceMatcher directly as (text, fmt) pairs.
_PLAIN_FMT = (False, False, False, False)


def _split_words(segment: str) -> list[str]:
    return WORD_SPLIT_RE.findall(segment)


def tokenize_with_format(text: str) -> list[tuple[str, tuple[bool, bool, bool, bool]]]:
    """Split text carrying inline markers into (word, fmt) tokens.

    Reuses INLINE_TOKEN_RE (the same **bold** / *italic* / <sup> / <sub> vocabulary
    rebuild_paragraph_runs parses) so a formatted span's words each inherit that
    span's format, and plain runs stay plain. Format lives on the token, so when the
    diff slices a span the emphasis rides along instead of being dropped or shifted."""
    tokens: list[tuple[str, tuple[bool, bool, bool, bool]]] = []

    def emit(segment: str, fmt: tuple[bool, bool, bool, bool]) -> None:
        for word in _split_words(segment):
            tokens.append((word, fmt))

    pos = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > pos:
            emit(text[pos:match.start()], _PLAIN_FMT)
        if match.group(1) is not None:
            emit(match.group(0)[2:-2], (True, False, False, False))
        elif match.group(2) is not None:
            emit(match.group(0)[1:-1], (False, True, False, False))
        elif match.group(3) is not None:
            emit(match.group(0)[5:-6], (False, False, True, False))
        elif match.group(4) is not None:
            emit(match.group(0)[5:-6], (False, False, False, True))
        pos = match.end()
    if pos < len(text):
        emit(text[pos:], _PLAIN_FMT)
    return tokens


def _add_formatted_run(paragraph, text: str, base: dict[str, Any], fmt: tuple[bool, bool, bool, bool]):
    """Append a run carrying the base font plus this token's italic/bold/sup/sub.
    Built via python-docx so <w:rPr> child ordering stays schema-valid."""
    run = paragraph.add_run(text)
    _apply_base_font(run, base)
    bold, italic, sup, sub = fmt
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if sup:
        run.font.superscript = True
    if sub:
        run.font.subscript = True
    return run


def _add_change_block(
    paragraph,
    tokens: list[tuple[str, tuple[bool, bool, bool, bool]]],
    base: dict[str, Any],
    change_tag: str,
    id_gen,
    author: str,
    date: str,
) -> None:
    """Wrap `tokens` in a single <w:ins> or <w:del> revision. Runs are created via
    python-docx (valid rPr) then moved under the change wrapper; for deletions their
    <w:t> is retagged <w:delText> per the OOXML tracked-changes schema."""
    if not tokens:
        return
    is_del = change_tag == "del"
    runs = [_add_formatted_run(paragraph, token_str, base, fmt) for token_str, fmt in tokens]
    change = OxmlElement(f"w:{change_tag}")
    change.set(qn("w:id"), str(next(id_gen)))
    change.set(qn("w:author"), author)
    change.set(qn("w:date"), date)
    for run in runs:
        r = run._r
        r.getparent().remove(r)
        if is_del:
            for t in r.findall(qn("w:t")):
                t.tag = qn("w:delText")
        change.append(r)
    paragraph._p.append(change)


def rebuild_paragraph_runs_tracked(
    paragraph,
    original: str,
    current: str,
    base: dict[str, Any],
    id_gen,
    author: str,
    date: str,
) -> None:
    """Replace a paragraph's content with a word-level track-changes rendering of
    original -> current: equal spans as plain runs, deletions in <w:del>, insertions
    in <w:ins>, replacements as del-then-ins. Every run inherits the base font and
    keeps its token's inline emphasis."""
    _clear_paragraph_runs(paragraph)
    a = tokenize_with_format(original)
    b = tokenize_with_format(current)
    matcher = SequenceMatcher(None, a, b, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for token_str, fmt in a[i1:i2]:
                _add_formatted_run(paragraph, token_str, base, fmt)
        elif tag == "delete":
            _add_change_block(paragraph, a[i1:i2], base, "del", id_gen, author, date)
        elif tag == "insert":
            _add_change_block(paragraph, b[j1:j2], base, "ins", id_gen, author, date)
        elif tag == "replace":
            _add_change_block(paragraph, a[i1:i2], base, "del", id_gen, author, date)
            _add_change_block(paragraph, b[j1:j2], base, "ins", id_gen, author, date)


def collect_changed_paragraphs(project_root: Path) -> list[dict[str, Any]]:
    """From manuscript_section_index.json, return paragraphs whose current_text
    diverged from the original text (i.e. the ones revise actually edited)."""
    index = read_json(project_root / "manuscript_section_index.json", {"sections": []})
    changed: list[dict[str, Any]] = []
    for section in index.get("sections", []):
        for para in section.get("paragraphs", []):
            pidx = para.get("paragraph_index")
            original = para.get("text", "")
            current = para.get("current_text", original)
            if pidx is None:
                continue
            if current != original:
                changed.append({"paragraph_index": pidx, "original": original, "current": current})
    return changed


def export_inplace(
    manuscript_docx: Path,
    project_root: Path,
    output_docx: Path,
    track_changes: bool = False,
    author: str = "revise-sci",
    date: str | None = None,
) -> dict[str, Any]:
    """Edit the original manuscript docx in place: replace only changed paragraphs'
    text, preserve everything else. fail-closed on any location/identity mismatch.

    track_changes=False (default): changed paragraphs are rebuilt to their current
    text (clean, no revision marks) — unchanged legacy behavior.
    track_changes=True: changed paragraphs are re-rendered as word-level Word
    revisions (<w:ins>/<w:del>) diffing original vs current."""
    if date is None:
        date = datetime.now().isoformat()
    id_gen = itertools.count(1)
    doc = Document(str(manuscript_docx))
    paragraphs = doc.paragraphs
    changed = collect_changed_paragraphs(project_root)
    errors: list[str] = []
    applied: list[int] = []
    for entry in changed:
        pidx = entry["paragraph_index"]
        if pidx < 0 or pidx >= len(paragraphs):
            errors.append(f"paragraph_index {pidx} out of range (doc has {len(paragraphs)} paragraphs)")
            continue
        paragraph = paragraphs[pidx]
        # Identity check: the live paragraph (with the same inline serialization the
        # index was built from) must match the recorded original. A mismatch means
        # the docx drifted from the atomized snapshot -> fail-closed, write nothing.
        live = strip_inline_format_markers(paragraph.text)
        recorded = strip_inline_format_markers(entry["original"])
        if normalize_text_for_match(live) != normalize_text_for_match(recorded):
            errors.append(
                f"identity mismatch at paragraph_index {pidx}: "
                f"docx={live[:60]!r} != index={recorded[:60]!r}"
            )
            continue
    if errors:
        return {"ok": False, "rejected": True, "errors": errors, "mode": "in-place"}
    # second pass: all entries validated, now mutate.
    skipped_images: list[int] = []
    for entry in changed:
        pidx = entry["paragraph_index"]
        paragraph = paragraphs[pidx]
        # fail-safe: a revised paragraph that embeds an image cannot be run-rebuilt
        # in place — _clear_paragraph_runs would delete the <w:drawing>/<w:object>
        # and leave an orphan image rel. Skip its text rewrite, keep runs intact.
        if _paragraph_has_embedded_image(paragraph):
            skipped_images.append(pidx)
            print(
                f"[export_inplace] warning: paragraph {pidx} contains an embedded "
                f"image; skipping in-place rewrite to preserve the picture — "
                f"please revise this paragraph's text manually.",
                file=sys.stderr,
            )
            continue
        base = _paragraph_base_font(paragraph)
        if track_changes:
            rebuild_paragraph_runs_tracked(paragraph, entry["original"], entry["current"], base, id_gen, author, date)
        else:
            rebuild_paragraph_runs(paragraph, entry["current"], base)
        applied.append(pidx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return {
        "ok": True,
        "mode": "in-place-tracked" if track_changes else "in-place",
        "track_changes": track_changes,
        "paragraphs_changed": applied,
        "paragraphs_skipped_images": skipped_images,
        "paragraphs_total": len(paragraphs),
    }


def normalize_text_for_match(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def main() -> int:
    parser = argparse.ArgumentParser(description="Export markdown artifacts to docx")
    parser.add_argument("--project-root", required=True)
    parser.add_argument("--output-md", required=True)
    parser.add_argument("--output-docx", required=True)
    parser.add_argument("--reference-docx", default="")
    parser.add_argument("--manuscript-docx", default="", help="Original manuscript docx; when present, the revised manuscript is exported in-place to preserve its formatting/tables/images.")
    parser.add_argument("--no-inplace", action="store_true", help="Force the legacy full-rebuild-from-md path even when an original manuscript docx is available.")
    parser.add_argument("--allow-rebuild-fallback", action="store_true", help="When in-place export is rejected (location/identity mismatch), fall back to md full-rebuild instead of hard-stopping. Off by default so the user must explicitly accept a reformatted rebuild.")
    parser.add_argument("--track-changes", action="store_true", help="In-place export as word-level Word tracked changes (<w:ins>/<w:del>) instead of a clean rewrite. Requires an in-place export (original --manuscript-docx).")
    parser.add_argument("--author", default="revise-sci", help="w:author stamped on each tracked change (only with --track-changes).")
    parser.add_argument("--date", default="", help="w:date (ISO 8601) stamped on each tracked change; defaults to datetime.now().isoformat() (only with --track-changes).")
    parser.add_argument("--journal-style", choices=tuple(PROFILE_PRESETS.keys()), default="journal-manuscript")
    args = parser.parse_args()

    project_root = Path(args.project_root)
    output_md = Path(args.output_md)
    output_docx = Path(args.output_docx)
    response_md = project_root / "response_to_reviewers.md"
    response_docx = project_root / "response_to_reviewers.docx"

    # 改后稿默认 in-place:基于原始 manuscript docx 编辑,保原稿格式/表格/图片,
    # 只替换被改段落。无原始 docx 或 in-place 失败时,回退到 md 全量重建(legacy)。
    manuscript_docx = Path(args.manuscript_docx) if args.manuscript_docx else None
    inplace_result: dict[str, Any] | None = None
    use_inplace = (
        not args.no_inplace
        and manuscript_docx is not None
        and manuscript_docx.exists()
        and manuscript_docx.suffix.lower() == ".docx"
    )
    if args.track_changes and not use_inplace:
        print(
            "[export_docx] warning: --track-changes needs an in-place export "
            "(original --manuscript-docx, no --no-inplace); falling back to clean "
            "md-rebuild without revision marks.",
            file=sys.stderr,
        )
    if use_inplace:
        inplace_result = export_inplace(
            manuscript_docx,
            project_root,
            output_docx,
            track_changes=args.track_changes,
            author=args.author,
            date=args.date or None,
        )
        if not inplace_result.get("ok"):
            print(json.dumps({"inplace_rejected": inplace_result}, ensure_ascii=False), file=sys.stderr)
            if not args.allow_rebuild_fallback:
                # 🚪 fail-closed + 硬停:不静默降级为 md 全量重建(那会丢原稿表格/图位/对齐并按
                # 期刊模板重排),把失败原因和两个选项显式呈现给用户,由用户决定。
                errors = inplace_result.get("errors", [])
                lines = [
                    "=" * 68,
                    "[export_docx] 原地保格式导出被拒绝(in-place export rejected)。",
                    "改后稿无法基于原始 docx 原地修改:定位/身份与原子化快照对不上。",
                    "拒绝原因:",
                ]
                lines.extend(f"  - {e}" for e in errors)
                lines.extend([
                    "",
                    "已阻止自动降级为 md 全量重建(会丢失原稿表格/图片位置/对齐,并按期刊模板重排版)。",
                    "请二选一后重跑:",
                    "  1) 接受重排版本 —— 重跑并加 --allow-rebuild-fallback,",
                    "     产出 md 重建稿(表格/图/排版需人工核对,图号绑定见 merge 的 figure_map)。",
                    "  2) 修锚点后重跑 —— 定位不匹配通常因原稿与原子化快照漂移(段落被移动/改写)。",
                    "     用未改动的原稿重跑,或修正 manuscript_section_index.json 的段落定位使 in-place 命中。",
                    "=" * 68,
                ])
                print("\n".join(lines), file=sys.stderr)
                print(json.dumps({
                    "ok": False,
                    "error": "inplace_export_rejected",
                    "allow_rebuild_fallback": False,
                    "inplace_rejected": inplace_result,
                }, ensure_ascii=False))
                return 3
            # 用户已显式接受重排:回退到 md 全量重建。
            print("[export_docx] --allow-rebuild-fallback 已启用:回退到 md 全量重建(排版将按期刊模板重排)。", file=sys.stderr)
            inplace_result = None

    if inplace_result is None:
        markdown_to_docx(
            output_md,
            output_docx,
            args.reference_docx,
            page_break_before_comment=False,
            header_text="Revised Manuscript",
            profile_name=args.journal_style,
            doc_kind="manuscript",
        )
    markdown_to_docx(
        response_md,
        response_docx,
        args.reference_docx,
        page_break_before_comment=True,
        header_text="Response to Reviewers",
        include_toc=True,
        profile_name="review-response",
        doc_kind="response",
    )

    state = read_json(project_root / "project_state.json", {})
    state.setdefault("outputs", {})
    state["outputs"]["response_md"] = str(response_md.resolve())
    state["outputs"]["response_docx"] = str(response_docx.resolve())
    state["outputs"]["output_md"] = str(output_md.resolve())
    state["outputs"]["output_docx"] = str(output_docx.resolve())
    state.setdefault("inputs", {})
    state["inputs"]["journal_style"] = args.journal_style
    export_mode = inplace_result.get("mode", "in-place") if inplace_result is not None else "md-rebuild"
    state["outputs"]["manuscript_export_mode"] = export_mode
    (project_root / "project_state.json").write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")

    print(json.dumps({
        "ok": True,
        "response_docx": str(response_docx.resolve()),
        "output_docx": str(output_docx.resolve()),
        "journal_style": args.journal_style,
        "manuscript_export_mode": export_mode,
        "inplace": inplace_result,
    }, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
