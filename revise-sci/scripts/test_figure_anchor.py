#!/usr/bin/env python3
"""Round-trip test for the figure<->section anchor + fail-closed image merge.

Run directly:  python3 test_figure_anchor.py   (or via pytest)

Builds a minimal docx with 3 captioned figures (Results/Discussion) plus one
uncaptioned logo image, then asserts:
  1. atomize_manuscript.py writes a `figures` field on each section entry,
     anchoring each captioned figure to the section its caption lives in.
  2. merge_manuscript.py places all 3 captioned figures into their sections and
     routes the uncaptioned image to an "Unplaced figures" block — never drops a
     single image (phantom-deletion defense).
"""
from __future__ import annotations

import json
import struct
import subprocess
import sys
import tempfile
import zlib
from pathlib import Path

from docx import Document

SCRIPTS_DIR = Path(__file__).resolve().parent
ATOMIZE = SCRIPTS_DIR / "atomize_manuscript.py"
EXTRACT = SCRIPTS_DIR / "extract_docx_images.py"
MERGE = SCRIPTS_DIR / "merge_manuscript.py"


def _png_1x1(rgb: tuple[int, int, int]) -> bytes:
    """Distinct 1x1 RGB PNG per color so python-docx does not dedupe blobs."""
    def chunk(typ: bytes, data: bytes) -> bytes:
        body = typ + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)  # 1x1, 8-bit, RGB
    idat = zlib.compress(b"\x00" + bytes(rgb))
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _build_docx(path: Path, img_dir: Path) -> None:
    from docx.shared import Inches

    imgs = []
    for i, color in enumerate([(255, 0, 0), (0, 255, 0), (0, 0, 255), (128, 128, 128)], start=1):
        p = img_dir / f"src_{i}.png"
        p.write_bytes(_png_1x1(color))
        imgs.append(p)

    doc = Document()
    doc.add_paragraph("A Minimal Study Title")
    doc.add_heading("Results", level=1)
    doc.add_paragraph("Body text referring to Figure 1 and Figure 2 in the results.")
    doc.add_picture(str(imgs[0]), width=Inches(1))
    doc.add_paragraph("Figure 1. First results panel showing the effect.")
    doc.add_picture(str(imgs[1]), width=Inches(1))
    doc.add_paragraph("Figure 2. Second results panel with controls.")
    doc.add_heading("Discussion", level=1)
    doc.add_paragraph("Discussion text pointing to Figure 3 for the model.")
    doc.add_picture(str(imgs[2]), width=Inches(1))
    doc.add_paragraph("Figure 3. Proposed mechanistic model.")
    # An uncaptioned decorative image (e.g. a logo): in the docx media but not a
    # numbered figure, so no section can claim it -> must land in "Unplaced".
    doc.add_picture(str(imgs[3]), width=Inches(1))
    doc.save(str(path))


def _run(cmd: list[str]) -> subprocess.CompletedProcess:
    result = subprocess.run(cmd, cwd=str(SCRIPTS_DIR), capture_output=True, text=True)
    assert result.returncode == 0, f"{cmd[1]} failed: {result.stderr or result.stdout}"
    return result


def test_figure_section_anchor_and_failclosed_merge() -> None:
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        manuscript = root / "manuscript.docx"
        _build_docx(manuscript, root)

        # Real pipeline order: atomize (manifest absent) -> extract -> merge.
        _run([sys.executable, str(ATOMIZE), "--manuscript", str(manuscript), "--project-root", str(root)])

        index = json.loads((root / "manuscript_section_index.json").read_text(encoding="utf-8"))
        by_heading = {s["heading"]: s for s in index["sections"]}
        assert "figures" in index["sections"][0], "atomize did not emit a `figures` field"

        results_figs = {f["figure_id"] for f in by_heading["Results"].get("figures", [])}
        discussion_figs = {f["figure_id"] for f in by_heading["Discussion"].get("figures", [])}
        assert results_figs == {"Figure 1", "Figure 2"}, results_figs
        assert discussion_figs == {"Figure 3"}, discussion_figs
        # Caption text is carried, source records the anchoring basis.
        fig1 = next(f for f in by_heading["Results"]["figures"] if f["figure_id"] == "Figure 1")
        assert "First results panel" in fig1["caption"]
        assert fig1["source"] == "caption"

        _run([sys.executable, str(EXTRACT), "--manuscript", str(manuscript), "--project-root", str(root)])
        manifest = json.loads((root / "figures" / "image_manifest.json").read_text(encoding="utf-8"))
        assert manifest["count"] == 4, f"expected 4 extracted images, got {manifest['count']}"

        merge = _run([sys.executable, str(MERGE), "--project-root", str(root), "--output-md", str(root / "out.md")])
        report = json.loads(merge.stdout.strip().splitlines()[-1])

        assert report["images_total"] == 4, report
        assert report["images_placed"] == 3, report
        assert report["images_unplaced"] == ["figure_04.png"], report
        assert "could not be anchored" in merge.stderr, "fail-closed warning missing from stderr"

        out_md = (root / "out.md").read_text(encoding="utf-8")
        # Zero image loss: every extracted binary appears exactly once in the md.
        for fname in [img["filename"] for img in manifest["images"]]:
            assert out_md.count(f"figures/{fname}") == 1, f"{fname} not present exactly once in merged md"
        assert "# Unplaced figures" in out_md


if __name__ == "__main__":
    test_figure_section_anchor_and_failclosed_merge()
    print("OK")
