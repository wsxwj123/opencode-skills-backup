#!/usr/bin/env python3
"""Tests for export_docx.export_inplace — the function that produces the user's
revised manuscript docx.

Its contract: rebuild ONLY changed paragraphs, leave everything else byte-identical,
fail-closed (write nothing) on any location/identity mismatch, and never destroy an
embedded image. Uses python-docx to build fixtures; if unavailable the whole module
skips with exit 0.

Self-contained, standalone: `python3 test_export_docx_inplace.py`.
"""
from __future__ import annotations

import json
import os
import struct
import sys
import tempfile
import zlib
from pathlib import Path

_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)

try:
    from docx import Document
    from docx.shared import Inches
except Exception:  # python-docx missing -> skip, do not fail the suite
    print("SKIP: python-docx not installed")
    raise SystemExit(0)

import export_docx as e  # noqa: E402


def _png_1x1() -> bytes:
    def chunk(typ: bytes, data: bytes) -> bytes:
        body = typ + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + bytes((255, 0, 0)))
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _write_index(root: Path, paragraphs: list[dict]) -> None:
    (root / "manuscript_section_index.json").write_text(
        json.dumps({"sections": [{"heading": "S", "paragraphs": paragraphs}]}, ensure_ascii=False),
        encoding="utf-8",
    )


def test_inplace_rebuilds_only_changed_paragraph():
    with tempfile.TemporaryDirectory() as d:
        root = Path(d)
        doc = Document()
        doc.add_paragraph("First paragraph original text.")
        doc.add_paragraph("Second paragraph unchanged text.")
        doc.add_paragraph("Third paragraph original text.")
        ms = root / "m.docx"
        doc.save(str(ms))
        _write_index(root, [
            {"paragraph_index": 0, "text": "First paragraph original text.", "current_text": "First paragraph REVISED text."},
            {"paragraph_index": 1, "text": "Second paragraph unchanged text.", "current_text": "Second paragraph unchanged text."},
        ])
        out = root / "out.docx"
        report = e.export_inplace(ms, root, out)
        assert report["ok"] is True, report
        assert report["paragraphs_changed"] == [0], report
        texts = [p.text for p in Document(str(out)).paragraphs]
        assert texts == [
            "First paragraph REVISED text.",
            "Second paragraph unchanged text.",
            "Third paragraph original text.",
        ], texts


def test_inplace_identity_mismatch_rejected_no_save():
    with tempfile.TemporaryDirectory() as d:
        root = Path(d)
        doc = Document()
        doc.add_paragraph("Live text here.")
        ms = root / "m.docx"
        doc.save(str(ms))
        # Recorded original diverges from the live docx paragraph -> drift.
        _write_index(root, [{"paragraph_index": 0, "text": "DIFFERENT recorded text.", "current_text": "new"}])
        out = root / "out.docx"
        report = e.export_inplace(ms, root, out)
        assert report["ok"] is False and report.get("rejected") is True, report
        assert report.get("errors"), report
        assert not out.exists(), "fail-closed: no output on identity mismatch"


def test_inplace_paragraph_index_out_of_range_records_error():
    with tempfile.TemporaryDirectory() as d:
        root = Path(d)
        doc = Document()
        doc.add_paragraph("Only paragraph.")
        ms = root / "m.docx"
        doc.save(str(ms))
        _write_index(root, [{"paragraph_index": 9, "text": "x", "current_text": "y"}])
        out = root / "out.docx"
        report = e.export_inplace(ms, root, out)
        assert report["ok"] is False, report
        assert any("out of range" in err for err in report["errors"]), report
        assert not out.exists()


def test_inplace_embedded_image_paragraph_skipped_image_preserved():
    with tempfile.TemporaryDirectory() as d:
        root = Path(d)
        img = root / "i.png"
        img.write_bytes(_png_1x1())
        doc = Document()
        doc.add_picture(str(img), width=Inches(1))
        ms = root / "m.docx"
        doc.save(str(ms))
        # add_picture makes a trailing empty paragraph carrying the <w:drawing>.
        img_idx = len(Document(str(ms)).paragraphs) - 1
        _write_index(root, [{"paragraph_index": img_idx, "text": "", "current_text": "now has text"}])
        out = root / "out.docx"
        report = e.export_inplace(ms, root, out)
        assert report["ok"] is True, report
        assert report["paragraphs_skipped_images"] == [img_idx], report
        assert report["paragraphs_changed"] == [], report
        # The picture rel survives the export.
        od = Document(str(out))
        has_image = any(p._p.findall(".//" + e.qn("w:drawing")) for p in od.paragraphs)
        assert has_image, "embedded image must be preserved when its paragraph is skipped"


if __name__ == "__main__":
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            fn()
            print(f"ok {name}")
    print("OK: export_inplace rebuilds only changed paras, fails closed on drift/OOR, preserves images")
