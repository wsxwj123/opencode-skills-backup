import json
import unittest

from docx import Document

from helpers import TempProject, run_script


class StrictGateTests(unittest.TestCase):
    def setUp(self):
        self.project = TempProject()
        self.root = self.project.root
        self.project_root = self.root / "run"
        (self.project_root / "units").mkdir(parents=True)
        (self.project_root / "manuscript_sections").mkdir(parents=True)
        (self.project_root / "comment_records").mkdir(parents=True)

    def tearDown(self):
        self.project.cleanup()

    def _write_valid_base(self):
        unit = {
            "comment_id": "R1-Major-01",
            "reviewer": "Reviewer #1",
            "severity": "major",
            "reviewer_comment_en": "Please clarify the mechanism statement.",
            "reviewer_comment_zh_literal": "请澄清机制表述。",
            "intent_zh": "审稿人要求收紧表述边界。",
            "response_en": "We revised the sentence to limit the claim to the present dataset.",
            "response_zh": "我们已修订该句，将结论限定在当前数据范围内。",
            "original_excerpt_en": "Quercetin showed a protective effect in TAC-treated cells.",
            "revised_excerpt_en": "In the present dataset, Quercetin showed a protective effect in TAC-treated cells.",
            "revised_excerpt_zh": "在当前数据范围内，槲皮素在TAC处理细胞中表现出保护作用。",
            "modification_actions": [{"action": "修改", "reason": "收紧结论边界。"}],
            "notes_core_zh": ["已按评论修改正文。"],
            "notes_support_zh": ["已保留证据来源。"],
            "evidence_sources": [{"provider_family": "user-provided", "source": "manuscript_sections/01-results.md"}],
            "target_document": "manuscript",
            "editorial_intent": "clarify",
            "revision_plan": {
                "scope": "sentence_replace",
                "target_sentence_index": 0,
                "original_fragment": "Quercetin showed a protective effect in TAC-treated cells.",
                "raw_fragment": "In the present dataset, Quercetin showed a protective effect in TAC-treated cells.",
                "polished_fragment": "In the present dataset, Quercetin showed a protective effect in TAC-treated cells.",
                "paragraph_before": "Quercetin showed a protective effect in TAC-treated cells.",
                "paragraph_after_raw": "In the present dataset, Quercetin showed a protective effect in TAC-treated cells.",
                "paragraph_after_polished": "In the present dataset, Quercetin showed a protective effect in TAC-treated cells.",
                "changed_fragment_type": "modified_sentence",
            },
            "polish_applied": True,
            "polish_driver_mode": "local-heuristic",
            "polish_guard_ok": True,
            "status": "completed",
            "author_confirmation_reason": "",
            "atomic_location": {
                "manuscript_section_id": "manuscript-001",
                "si_section_id": "",
                "section_file": "manuscript_sections/01-results.md",
                "section_heading": "Results",
                "paragraph_index": 1,
                "matched_sentence": "Quercetin showed a protective effect in TAC-treated cells.",
                "matched_sentence_index": 0,
            },
        }
        (self.project_root / "units" / "001_R1-Major-01.json").write_text(json.dumps(unit), encoding="utf-8")
        (self.project_root / "comment_records" / "R1-Major-01.md").write_text("# R1-Major-01\n", encoding="utf-8")
        (self.project_root / "response_to_reviewers.md").write_text(
            "\n".join(
                [
                    "# 回复审稿人的邮件",
                    "",
                    "# Reviewer #1",
                    "",
                    "## Major",
                    "",
                    "### Comment 1",
                    "",
                    "Please clarify the mechanism statement.",
                    "",
                    "#### 2) Response to Reviewer（中英对照）",
                    "",
                    "We revised the sentence to limit the claim to the present dataset.",
                    "",
                    "In the present dataset, Quercetin showed a protective effect in TAC-treated cells.",
                    "",
                    "#### 5) Evidence Attachments",
                    "",
                    "**Text**",
                    "- user-provided: manuscript_sections/01-results.md",
                    "",
                    "**Image**",
                    "- Not provided by user",
                    "",
                    "**Table**",
                    "- Not provided by user",
                    "",
                ]
            ),
            encoding="utf-8",
        )
        response_doc = Document()
        response_doc.add_heading("回复审稿人的邮件", level=1)
        response_doc.add_heading("Reviewer #1", level=1)
        response_doc.add_heading("Major", level=2)
        response_doc.add_heading("Comment 1", level=3)
        response_doc.add_heading("2) Response to Reviewer（中英对照）", level=4)
        response_doc.add_paragraph("We revised the sentence to limit the claim to the present dataset.")
        response_doc.add_heading("5) Evidence Attachments", level=4)
        response_doc.add_paragraph("Text")
        response_doc.add_paragraph("Image")
        response_doc.add_paragraph("Table")
        response_doc.save(self.project_root / "response_to_reviewers.docx")

        revised_doc = Document()
        revised_doc.add_heading("Results", level=1)
        revised_doc.add_paragraph("In the present dataset, Quercetin showed a protective effect in TAC-treated cells.")
        revised_doc.save(self.project_root / "revised.docx")
        (self.project_root / "revised.md").write_text("# Results\n\nIn the present dataset, Quercetin showed a protective effect in TAC-treated cells.\n", encoding="utf-8")
        (self.project_root / "manuscript_edit_plan.md").write_text(
            "\n".join(
                [
                    "# manuscript_edit_plan",
                    "",
                    "| comment_id | 目标文档 | 段落索引 | 待替换片段 | 替换后文本 | 动作类型 |",
                    "|---|---|---|---|---|---|",
                    "| R1-Major-01 | manuscript | 1 | Quercetin showed a protective effect in TAC-treated cells. | In the present dataset, Quercetin showed a protective effect in TAC-treated cells. | 修改 |",
                    "",
                ]
            ),
            encoding="utf-8",
        )
        (self.project_root / "manuscript_section_index.json").write_text(
            json.dumps(
                {
                    "sections": [
                        {
                            "section_id": "manuscript-001",
                            "heading": "Results",
                            "file": "manuscript_sections/01-results.md",
                            "paragraphs": [
                                {
                                    "paragraph_index": 1,
                                    "text": "Quercetin showed a protective effect in TAC-treated cells.",
                                    "current_text": "In the present dataset, Quercetin showed a protective effect in TAC-treated cells.",
                                }
                            ],
                        }
                    ]
                }
            ),
            encoding="utf-8",
        )
        (self.project_root / "manuscript_sections" / "01-results.md").write_text(
            "# Results\n\nIn the present dataset, Quercetin showed a protective effect in TAC-treated cells.\n",
            encoding="utf-8",
        )
        (self.project_root / "project_state.json").write_text(
            json.dumps(
                {
                    "delivery_status": "ready_to_submit",
                    "counts": {"comment_units": 1},
                    "inputs": {"reference_search_decision": "ask"},
                    "outputs": {
                        "output_md": str((self.project_root / "revised.md").resolve()),
                        "output_docx": str((self.project_root / "revised.docx").resolve()),
                    },
                }
            ),
            encoding="utf-8",
        )
        (self.project_root / "revision_polish_manifest.json").write_text(
            json.dumps({"workflow": "revise-sci-polish", "candidates": [{"comment_id": "R1-Major-01"}]}),
            encoding="utf-8",
        )
        (self.project_root / "revision_polish_execution.json").write_text(
            json.dumps({"ok": True, "driver_mode": "local-heuristic", "candidate_count": 1, "polished_comment_ids": ["R1-Major-01"]}),
            encoding="utf-8",
        )

    def test_gate_rejects_placeholders(self):
        self._write_valid_base()
        unit_path = self.project_root / "units" / "001_R1-Major-01.json"
        unit = json.loads(unit_path.read_text(encoding="utf-8"))
        unit["response_en"] = "AI_FILL_REQUIRED"
        unit["response_zh"] = "待AI"
        unit_path.write_text(json.dumps(unit), encoding="utf-8")
        result = run_script("strict_gate.py", ["--project-root", str(self.project_root)], cwd=self.root)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("placeholder", result.stdout + result.stderr)

    def test_gate_rejects_missing_response_comment_mapping(self):
        self._write_valid_base()
        (self.project_root / "response_to_reviewers.md").write_text("# 回复审稿人的邮件\n", encoding="utf-8")
        result = run_script("strict_gate.py", ["--project-root", str(self.project_root)], cwd=self.root)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("missing comment mapping", result.stdout + result.stderr)

    def test_gate_rejects_completed_unit_when_manuscript_and_plan_do_not_match(self):
        self._write_valid_base()
        (self.project_root / "manuscript_sections" / "01-results.md").write_text(
            "# Results\n\nQuercetin showed a protective effect in TAC-treated cells.\n",
            encoding="utf-8",
        )
        (self.project_root / "manuscript_edit_plan.md").write_text("# manuscript_edit_plan\n", encoding="utf-8")
        result = run_script("strict_gate.py", ["--project-root", str(self.project_root)], cwd=self.root)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("completed excerpt not found in manuscript section", result.stdout + result.stderr)
        self.assertIn("edit plan missing comment_id", result.stdout + result.stderr)

    def test_gate_requires_literature_index_and_matrix_for_completed_citation_unit(self):
        self._write_valid_base()
        unit_path = self.project_root / "units" / "001_R1-Major-01.json"
        unit = json.loads(unit_path.read_text(encoding="utf-8"))
        unit["editorial_intent"] = "citation"
        unit["evidence_sources"] = [{"provider_family": "paper-search", "source": "PMID:123456"}]
        unit_path.write_text(json.dumps(unit), encoding="utf-8")
        (self.project_root / "reference_sync_report.json").write_text(
            json.dumps({"covered_comment_ids": ["R1-Major-01"], "references_added": 1}),
            encoding="utf-8",
        )
        result = run_script("strict_gate.py", ["--project-root", str(self.project_root)], cwd=self.root)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("literature_index", result.stdout + result.stderr)

    def test_gate_rejects_response_docx_without_comment_structure(self):
        self._write_valid_base()
        bad_doc = Document()
        bad_doc.add_heading("回复审稿人的邮件", level=1)
        bad_doc.add_paragraph("Only a title")
        bad_doc.save(self.project_root / "response_to_reviewers.docx")
        result = run_script("strict_gate.py", ["--project-root", str(self.project_root)], cwd=self.root)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("response_to_reviewers.docx", result.stdout + result.stderr)

    def test_gate_requires_user_decision_before_searching_for_new_references(self):
        self._write_valid_base()
        (self.project_root / "data").mkdir(exist_ok=True)
        (self.project_root / "data" / "reference_registry.json").write_text("[]", encoding="utf-8")
        (self.project_root / "data" / "reference_coverage_audit.json").write_text(
            json.dumps(
                {
                    "ok": False,
                    "citation_style": "numeric",
                    "reference_entries": 0,
                    "cited_numbers": [1, 2, 3],
                    "missing_reference_numbers": [1, 2, 3],
                    "author_year_citations": [],
                    "missing_author_year_citations": [],
                    "reference_source": "",
                    "reference_search_required": True,
                    "reference_search_decision": "ask",
                }
            ),
            encoding="utf-8",
        )
        state = json.loads((self.project_root / "project_state.json").read_text(encoding="utf-8"))
        state["inputs"] = {"references_source_path": "", "reference_search_decision": "ask"}
        (self.project_root / "project_state.json").write_text(json.dumps(state), encoding="utf-8")
        result = run_script("strict_gate.py", ["--project-root", str(self.project_root)], cwd=self.root)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("reference search decision required", result.stdout + result.stderr)

    def test_gate_requires_strategy_and_status_artifacts_when_search_is_approved(self):
        self._write_valid_base()
        (self.project_root / "data").mkdir(exist_ok=True)
        (self.project_root / "data" / "reference_registry.json").write_text("[]", encoding="utf-8")
        (self.project_root / "data" / "reference_coverage_audit.json").write_text(
            json.dumps(
                {
                    "ok": False,
                    "citation_style": "numeric",
                    "reference_entries": 0,
                    "cited_numbers": [1, 2, 3],
                    "missing_reference_numbers": [1, 2, 3],
                    "author_year_citations": [],
                    "missing_author_year_citations": [],
                    "reference_source": "",
                    "reference_search_required": True,
                    "reference_search_decision": "approved",
                }
            ),
            encoding="utf-8",
        )
        (self.project_root / "reference_search_manifest.json").write_text("{}", encoding="utf-8")
        (self.project_root / "reference_search_task.md").write_text("# task\n", encoding="utf-8")
        result = run_script("strict_gate.py", ["--project-root", str(self.project_root)], cwd=self.root)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("reference_search_strategy.json", result.stdout + result.stderr)
        self.assertIn("reference_search_status.json", result.stdout + result.stderr)

    def test_gate_rejects_non_review_writing_search_governance(self):
        self._write_valid_base()
        (self.project_root / "data").mkdir(exist_ok=True)
        (self.project_root / "data" / "reference_registry.json").write_text("[]", encoding="utf-8")
        (self.project_root / "data" / "reference_coverage_audit.json").write_text(
            json.dumps(
                {
                    "ok": True,
                    "citation_style": "numeric",
                    "reference_entries": 1,
                    "cited_numbers": [1],
                    "missing_reference_numbers": [],
                    "author_year_citations": [],
                    "missing_author_year_citations": [],
                    "reference_source": "",
                    "reference_search_required": False,
                    "reference_search_decision": "approved",
                }
            ),
            encoding="utf-8",
        )
        state = json.loads((self.project_root / "project_state.json").read_text(encoding="utf-8"))
        state["inputs"] = {"references_source_path": "", "reference_search_decision": "approved"}
        (self.project_root / "project_state.json").write_text(json.dumps(state), encoding="utf-8")
        (self.project_root / "paper_search_results.json").write_text(json.dumps({"results": []}), encoding="utf-8")
        (self.project_root / "paper_search_validated.json").write_text(json.dumps({"results": []}), encoding="utf-8")
        (self.project_root / "paper_search_guard_report.json").write_text(
            json.dumps({"summary": {"all_rows_guard_verified": True}}),
            encoding="utf-8",
        )
        (self.project_root / "reference_search_manifest.json").write_text(
            json.dumps(
                {
                    "reference_search_decision": "approved",
                    "governance_active": True,
                    "allowed_provider_families": ["paper-search"],
                    "forbidden_provider_families": ["websearch"],
                    "workflow_rules": {"rounds": [1, 2, 3]},
                }
            ),
            encoding="utf-8",
        )
        (self.project_root / "reference_search_task.md").write_text("# task\n", encoding="utf-8")
        (self.project_root / "reference_search_rounds.json").write_text(
            json.dumps({"workflow": "review-writing", "rounds": [{"round": 1, "queries": ["q1"]}, {"round": 2, "queries": []}, {"round": 3, "queries": []}]}),
            encoding="utf-8",
        )
        (self.project_root / "reference_search_strategy.json").write_text(
            json.dumps(
                {
                    "workflow": "not-review-writing",
                    "provider_policy": {"primary": ["paper-search"], "forbidden": ["websearch"]},
                    "round_model": [1, 2, 3],
                }
            ),
            encoding="utf-8",
        )
        (self.project_root / "reference_search_status.json").write_text(
            json.dumps(
                {
                    "reference_search_decision": "approved",
                    "governance_active": True,
                    "steps": {
                        "paper_search_batch_imported": True,
                        "validated_batch_present": True,
                        "citation_guard_passed": True,
                        "literature_index_built": True,
                        "synthesis_matrix_audited": True,
                        "reference_sync_completed": True,
                    },
                }
            ),
            encoding="utf-8",
        )
        (self.project_root / "data" / "literature_index.json").write_text("[]", encoding="utf-8")
        (self.project_root / "data" / "synthesis_matrix.json").write_text("[]", encoding="utf-8")
        (self.project_root / "data" / "synthesis_matrix_audit.json").write_text(
            json.dumps({"missing_claim": 0, "missing_key_fields": 0}),
            encoding="utf-8",
        )
        (self.project_root / "reference_sync_report.json").write_text(
            json.dumps({"covered_comment_ids": [], "references_added": 0}),
            encoding="utf-8",
        )
        result = run_script("strict_gate.py", ["--project-root", str(self.project_root)], cwd=self.root)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("review-writing", result.stdout + result.stderr)

    def test_gate_rejects_failed_reference_search_execution_report(self):
        self._write_valid_base()
        (self.project_root / "data").mkdir(exist_ok=True)
        (self.project_root / "data" / "reference_registry.json").write_text("[]", encoding="utf-8")
        (self.project_root / "data" / "reference_coverage_audit.json").write_text(
            json.dumps(
                {
                    "ok": True,
                    "citation_style": "numeric",
                    "reference_entries": 1,
                    "cited_numbers": [1],
                    "missing_reference_numbers": [],
                    "author_year_citations": [],
                    "missing_author_year_citations": [],
                    "reference_source": "",
                    "reference_search_required": False,
                    "reference_search_decision": "approved",
                }
            ),
            encoding="utf-8",
        )
        state = json.loads((self.project_root / "project_state.json").read_text(encoding="utf-8"))
        state["inputs"] = {"references_source_path": "", "reference_search_decision": "approved"}
        (self.project_root / "project_state.json").write_text(json.dumps(state), encoding="utf-8")
        for name, content in {
            "paper_search_results.json": {"results": []},
            "paper_search_validated.json": {"results": []},
            "paper_search_guard_report.json": {"summary": {"all_rows_guard_verified": True}},
            "reference_search_manifest.json": {
                "workflow": "review-writing",
                "reference_search_decision": "approved",
                "governance_active": True,
                "allowed_provider_families": ["paper-search"],
                "forbidden_provider_families": ["websearch"],
                "verification_policy": {"dual_verification_required": True, "allow_unverified": False, "guard_command": "python scripts/citation_guard.py"},
                "workflow_rules": {"rounds": [1, 2, 3]},
            },
            "reference_search_strategy.json": {
                "workflow": "review-writing",
                "provider_policy": {"primary": ["paper-search"], "forbidden": ["websearch"]},
                "mandatory_guard_command": "python scripts/citation_guard.py",
                "round_model": [1, 2, 3],
                "required_outputs": ["data/literature_index.json", "data/synthesis_matrix.json", "data/synthesis_matrix_audit.json"],
            },
            "reference_search_status.json": {
                "reference_search_decision": "approved",
                "governance_active": True,
                "steps": {
                    "search_round_plan_generated": True,
                    "paper_search_batch_imported": True,
                    "validated_batch_present": True,
                    "citation_guard_passed": True,
                    "literature_index_built": True,
                    "synthesis_matrix_audited": True,
                    "reference_sync_completed": True,
                },
            },
            "reference_search_rounds.json": {
                "workflow": "review-writing",
                "rounds": [
                    {"round": 1, "provider_family": "paper-search", "queries": ["q1"]},
                    {"round": 2, "provider_family": "paper-search", "queries": ["q2"]},
                    {"round": 3, "provider_family": "paper-search", "queries": ["q3"]},
                ],
            },
            "reference_search_execution.json": {"ok": False, "driver_mode": "opencode-driver"},
            "reference_sync_report.json": {"covered_comment_ids": [], "references_added": 0},
        }.items():
            (self.project_root / name).write_text(json.dumps(content), encoding="utf-8") if name.endswith(".json") else (self.project_root / name).write_text(content, encoding="utf-8")
        (self.project_root / "reference_search_task.md").write_text("# task\n", encoding="utf-8")
        (self.project_root / "data" / "literature_index.json").write_text("[]", encoding="utf-8")
        (self.project_root / "data" / "synthesis_matrix.json").write_text("[]", encoding="utf-8")
        (self.project_root / "data" / "synthesis_matrix_audit.json").write_text(
            json.dumps({"missing_claim": 0, "missing_key_fields": 0}),
            encoding="utf-8",
        )
        result = run_script("strict_gate.py", ["--project-root", str(self.project_root)], cwd=self.root)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("reference_search_execution.json exists but reports ok=false", result.stdout + result.stderr)
