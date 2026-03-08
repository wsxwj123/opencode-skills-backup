import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = SKILL_ROOT / "scripts"


def create_docx(path: Path, rows: list[tuple[str, str]]) -> Path:
    doc = Document()
    for style, text in rows:
        if style == "heading1":
            doc.add_heading(text, level=1)
        elif style == "heading2":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)
    doc.save(path)
    return path


def run_script(script_name: str, args: list[str], cwd: Path) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(SCRIPTS_DIR / script_name)] + args,
        cwd=cwd,
        text=True,
        capture_output=True,
    )


def create_fake_paper_search_runner(path: Path, output_payload: str) -> Path:
    path.write_text(
        "\n".join(
            [
                "#!/usr/bin/env python3",
                "import argparse",
                "from pathlib import Path",
                "",
                "parser = argparse.ArgumentParser()",
                "parser.add_argument('--rounds-json', required=True)",
                "parser.add_argument('--output', required=True)",
                "parser.add_argument('--project-root', required=True)",
                "args = parser.parse_args()",
                "Path(args.output).write_text(" + repr(output_payload) + ", encoding='utf-8')",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    path.chmod(0o755)
    return path


class TempProject:
    def __init__(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.root = Path(self._tmp.name)

    def cleanup(self):
        self._tmp.cleanup()
