import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from helpers import run_script


class ExportDocxTests(unittest.TestCase):
    def test_export_docx_converts_bold_markers_and_bullets_to_word_styles(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            project_root = root / "project"
            project_root.mkdir()
            md_path = project_root / "sample.md"
            md_path.write_text(
                "# Title\n\n**Text**\n\n- first item\n\n1. first number\n\nplain **bold** text\n",
                encoding="utf-8",
            )
            response_md = project_root / "response_to_reviewers.md"
            response_md.write_text("# 回复审稿人的邮件\n\n### Comment 1\n\n**Text**\n", encoding="utf-8")

            result = run_script(
                "export_docx.py",
                [
                    "--project-root",
                    str(project_root),
                    "--output-md",
                    str(md_path),
                    "--output-docx",
                    str(project_root / "sample.docx"),
                ],
                cwd=root,
            )
            self.assertEqual(result.returncode, 0, msg=result.stdout + result.stderr)

            doc = Document(str(project_root / "sample.docx"))
            text_styles = [(p.text, p.style.name) for p in doc.paragraphs if p.text]
            self.assertIn(("first item", "List Bullet"), text_styles)
            self.assertIn(("first number", "List Number"), text_styles)

            label_para = next(p for p in doc.paragraphs if p.text == "Text")
            self.assertTrue(any(run.bold for run in label_para.runs if run.text == "Text"))

            mixed_para = next(p for p in doc.paragraphs if p.text == "plain bold text")
            bold_runs = [run for run in mixed_para.runs if run.text == "bold"]
            self.assertTrue(bold_runs and bold_runs[0].bold)
            self.assertNotIn("**", "\n".join(p.text for p in doc.paragraphs))

    def test_export_docx_adds_header_footer_and_toc_field(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            project_root = root / "project"
            project_root.mkdir()
            md_path = project_root / "sample.md"
            md_path.write_text("# Title\n\n## Section\n\nBody\n", encoding="utf-8")
            response_md = project_root / "response_to_reviewers.md"
            response_md.write_text(
                "# 回复审稿人的邮件\n\n# Reviewer #1\n\n## Major\n\n### Comment 1\n\n#### 2) Response to Reviewer（中英对照）\n\nBody\n\n#### 5) Evidence Attachments\n\n**Text**\n\n**Image**\n\n**Table**\n",
                encoding="utf-8",
            )

            result = run_script(
                "export_docx.py",
                [
                    "--project-root",
                    str(project_root),
                    "--output-md",
                    str(md_path),
                    "--output-docx",
                    str(project_root / "sample.docx"),
                ],
                cwd=root,
            )
            self.assertEqual(result.returncode, 0, msg=result.stdout + result.stderr)
            with zipfile.ZipFile(project_root / "response_to_reviewers.docx") as zf:
                document_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
                footer_xml = zf.read("word/footer1.xml").decode("utf-8", errors="ignore")
                header_xml = zf.read("word/header1.xml").decode("utf-8", errors="ignore")
            self.assertIn("TOC", document_xml)
            self.assertIn("PAGE", footer_xml)
            self.assertIn("Response to Reviewers", header_xml)


if __name__ == "__main__":
    unittest.main()
