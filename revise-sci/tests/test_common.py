import importlib.util
import sys
import unittest
from pathlib import Path

from docx import Document


SCRIPTS_DIR = Path(__file__).resolve().parents[1] / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))

spec = importlib.util.spec_from_file_location("common", SCRIPTS_DIR / "common.py")
common = importlib.util.module_from_spec(spec)
assert spec and spec.loader
spec.loader.exec_module(common)


class CommonTests(unittest.TestCase):
    def test_comment_nature_prioritizes_methodology_over_table_anchor_noise(self):
        text = (
            "稿件没有交代检索数据库、关键词、纳入排除标准，也没有说明证据等级差异。"
            "Table 2 未提供系统纳入逻辑。"
        )
        self.assertEqual(common.comment_nature(text), "需要实质性解释、结构重构或方法学澄清")

    def test_autodiscover_reference_source_prefers_same_title_sibling_docx_with_references(self):
        from tempfile import TemporaryDirectory

        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            comments = root / "comments.html"
            comments.write_text("<html></html>", encoding="utf-8")
            manuscript = root / "Manuscript .docx"
            sibling = root / "Manuscript (1).docx"

            current = Document()
            current.add_paragraph("Advancement of Extracellular Vesicles Applications in Pulmonary Diseases")
            current.add_paragraph("Introduction")
            current.save(manuscript)

            legacy = Document()
            legacy.add_paragraph("Advancement of Extracellular Vesicles Applications in Pulmonary Diseases")
            legacy.add_paragraph("References")
            legacy.add_paragraph("1. Smith J. Study A. 2023.")
            legacy.add_paragraph("2. Lee K. Study B. 2024.")
            legacy.add_paragraph("3. Wang M. Study C. 2025.")
            legacy.save(sibling)

            discovered = common.autodiscover_reference_source(comments, None, root, manuscript)
            self.assertEqual(discovered, sibling.resolve())

    def test_autodiscover_reference_source_rejects_unrelated_sibling_docx(self):
        from tempfile import TemporaryDirectory

        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            comments = root / "comments.html"
            comments.write_text("<html></html>", encoding="utf-8")
            manuscript = root / "Manuscript .docx"
            unrelated = root / "Other Paper (1).docx"

            current = Document()
            current.add_paragraph("Advancement of Extracellular Vesicles Applications in Pulmonary Diseases")
            current.add_paragraph("Introduction")
            current.save(manuscript)

            legacy = Document()
            legacy.add_paragraph("Completely Different Topic")
            legacy.add_paragraph("References")
            legacy.add_paragraph("1. Smith J. Study A. 2023.")
            legacy.add_paragraph("2. Lee K. Study B. 2024.")
            legacy.add_paragraph("3. Wang M. Study C. 2025.")
            legacy.save(unrelated)

            discovered = common.autodiscover_reference_source(comments, None, root, manuscript)
            self.assertIsNone(discovered)

    def test_autodiscover_reference_source_ignores_word_lock_files(self):
        from tempfile import TemporaryDirectory

        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            comments = root / "comments.html"
            comments.write_text("<html></html>", encoding="utf-8")
            manuscript = root / "Manuscript .docx"
            sibling = root / "Manuscript (1).docx"
            lockfile = root / "~$文初稿3.7.docx"

            current = Document()
            current.add_paragraph("Advancement of Extracellular Vesicles Applications in Pulmonary Diseases")
            current.add_paragraph("Introduction")
            current.save(manuscript)

            legacy = Document()
            legacy.add_paragraph("Advancement of Extracellular Vesicles Applications in Pulmonary Diseases")
            legacy.add_paragraph("References")
            legacy.add_paragraph("1. Smith J. Study A. 2023.")
            legacy.add_paragraph("2. Lee K. Study B. 2024.")
            legacy.add_paragraph("3. Wang M. Study C. 2025.")
            legacy.save(sibling)

            lockfile.write_text("not-a-real-docx", encoding="utf-8")

            discovered = common.autodiscover_reference_source(comments, None, root, manuscript)
            self.assertEqual(discovered, sibling.resolve())

    def test_autodiscover_reference_source_finds_same_title_docx_in_subdirectory(self):
        from tempfile import TemporaryDirectory

        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            comments = root / "comments.html"
            comments.write_text("<html></html>", encoding="utf-8")
            manuscript = root / "Manuscript .docx"
            legacy_dir = root / "revision"
            legacy_dir.mkdir()
            sibling = legacy_dir / "revised manuscript.docx"

            current = Document()
            current.add_paragraph("Advancement of Extracellular Vesicles Applications in Pulmonary Diseases")
            current.add_paragraph("Introduction")
            current.save(manuscript)

            legacy = Document()
            legacy.add_paragraph("Advancement of Extracellular Vesicles Applications in Pulmonary Diseases")
            legacy.add_paragraph("References")
            legacy.add_paragraph("1. Smith J. Study A. 2023.")
            legacy.add_paragraph("2. Lee K. Study B. 2024.")
            legacy.add_paragraph("3. Wang M. Study C. 2025.")
            legacy.save(sibling)

            discovered = common.autodiscover_reference_source(comments, None, root, manuscript)
            self.assertEqual(discovered, sibling.resolve())


if __name__ == "__main__":
    unittest.main()
