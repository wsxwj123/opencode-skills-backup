#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
博士论文质量自检工具

功能：
1. 检查字数是否达标
2. 检查格式是否规范
3. 检查图表编号是否连续
4. 检查是否有列表项（应为段落）
5. 生成质量报告

作者：Sci2Doc Team
日期：2024-03-15
"""

from docx import Document
import sys
import os
import re
import json
from datetime import datetime


def check_word_count(doc):
    """检查字数是否达标"""
    issues = []
    
    # 使用 count_words_docx 的逻辑
    total_chinese = 0
    review_chinese = 0
    in_review = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if '综述' in text and para.style.name.startswith('Heading'):
            in_review = True
            continue
        
        if in_review:
            review_chinese += sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
        else:
            total_chinese += sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
    
    # 检查正文字数
    if total_chinese < 50000:
        issues.append({
            'level': 'error',
            'category': '字数',
            'message': f'正文字数不足：{total_chinese} / 50,000 字',
            'suggestion': '需要扩展内容以达到 50,000 字要求'
        })
    
    # 检查综述字数
    if review_chinese > 0 and review_chinese < 5000:
        issues.append({
            'level': 'warning',
            'category': '字数',
            'message': f'综述字数不足：{review_chinese} / 5,000 字',
            'suggestion': '建议扩展综述部分至 5,000 字以上'
        })
    
    return issues, {
        'body_words': total_chinese,
        'review_words': review_chinese
    }


def check_heading_levels(doc):
    """检查标题层级是否规范"""
    issues = []
    prev_level = 0
    
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            
            # 检查是否超过三级
            if level > 3:
                issues.append({
                    'level': 'error',
                    'category': '标题层级',
                    'location': f'第 {i+1} 段',
                    'message': f'标题层级过深：{level} 级（{para.text[:30]}...）',
                    'suggestion': '中南大学要求标题最多三级'
                })
            
            # 检查是否跳级（如 1 → 3）
            if level - prev_level > 1:
                issues.append({
                    'level': 'warning',
                    'category': '标题层级',
                    'location': f'第 {i+1} 段',
                    'message': f'标题层级跳跃：从 {prev_level} 级跳到 {level} 级',
                    'suggestion': '建议按顺序设置标题层级'
                })
            
            prev_level = level
    
    return issues


def check_figure_numbering(doc):
    """检查图表编号是否连续规范"""
    issues = []
    
    # 匹配图编号：图 1-1, 图 2-3 等
    figure_pattern = re.compile(r'图\s*(\d+)-(\d+)')
    table_pattern = re.compile(r'表\s*(\d+)-(\d+)')
    
    figures = []
    tables = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # 查找图编号
        for match in figure_pattern.finditer(text):
            chapter = int(match.group(1))
            number = int(match.group(2))
            figures.append({
                'chapter': chapter,
                'number': number,
                'location': i + 1,
                'text': text[:50]
            })
        
        # 查找表编号
        for match in table_pattern.finditer(text):
            chapter = int(match.group(1))
            number = int(match.group(2))
            tables.append({
                'chapter': chapter,
                'number': number,
                'location': i + 1,
                'text': text[:50]
            })
    
    # 检查图编号连续性
    figures_by_chapter = {}
    for fig in figures:
        chapter = fig['chapter']
        if chapter not in figures_by_chapter:
            figures_by_chapter[chapter] = []
        figures_by_chapter[chapter].append(fig)
    
    for chapter, figs in figures_by_chapter.items():
        figs_sorted = sorted(figs, key=lambda x: x['number'])
        for i, fig in enumerate(figs_sorted):
            expected = i + 1
            if fig['number'] != expected:
                issues.append({
                    'level': 'warning',
                    'category': '图表编号',
                    'location': f'第 {fig["location"]} 段',
                    'message': f'第 {chapter} 章图编号不连续：图 {chapter}-{fig["number"]}',
                    'suggestion': f'应为 图 {chapter}-{expected}'
                })
    
    # 检查表编号连续性（同理）
    tables_by_chapter = {}
    for tab in tables:
        chapter = tab['chapter']
        if chapter not in tables_by_chapter:
            tables_by_chapter[chapter] = []
        tables_by_chapter[chapter].append(tab)
    
    for chapter, tabs in tables_by_chapter.items():
        tabs_sorted = sorted(tabs, key=lambda x: x['number'])
        for i, tab in enumerate(tabs_sorted):
            expected = i + 1
            if tab['number'] != expected:
                issues.append({
                    'level': 'warning',
                    'category': '图表编号',
                    'location': f'第 {tab["location"]} 段',
                    'message': f'第 {chapter} 章表编号不连续：表 {chapter}-{tab["number"]}',
                    'suggestion': f'应为 表 {chapter}-{expected}'
                })
    
    return issues


def check_bullet_points(doc):
    """检查是否有列表项（应全部为段落）"""
    issues = []
    
    for i, para in enumerate(doc.paragraphs):
        # 检查是否为列表样式
        if 'List' in para.style.name:
            issues.append({
                'level': 'error',
                'category': '格式',
                'location': f'第 {i+1} 段',
                'message': f'发现列表项：{para.text[:30]}...',
                'suggestion': '论文要求段落式写作，不允许使用列表项'
            })
        
        # 检查是否有列表标记（简单检测）
        text = para.text.strip()
        if text and (
            text.startswith('• ') or
            text.startswith('- ') or
            text.startswith('* ') or
            re.match(r'^\d+\.\s', text) or
            re.match(r'^\(\d+\)\s', text)
        ):
            issues.append({
                'level': 'warning',
                'category': '格式',
                'location': f'第 {i+1} 段',
                'message': f'疑似列表项：{text[:30]}...',
                'suggestion': '请确认是否为列表项，如是请改为段落形式'
            })
    
    return issues


def check_reference_count(doc):
    """检查参考文献数量"""
    issues = []
    
    # 查找参考文献章节
    in_references = False
    ref_count = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if '参考文献' in text and para.style.name.startswith('Heading'):
            in_references = True
            continue
        
        if in_references:
            # 匹配 [1], [2] 等编号
            if re.match(r'^\[\d+\]', text):
                ref_count += 1
    
    if ref_count < 80:
        issues.append({
            'level': 'error',
            'category': '参考文献',
            'message': f'参考文献数量不足：{ref_count} / 80 篇',
            'suggestion': '医学博士论文要求参考文献不少于 80 篇'
        })
    
    return issues, ref_count


def check_paragraph_formatting(doc):
    """检查段落格式是否规范"""
    issues = []
    
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        
        # 跳过标题
        if para.style.name.startswith('Heading'):
            continue
        
        # 检查首行缩进（正文应有首行缩进）
        if para.style.name == 'Normal':
            if not para.paragraph_format.first_line_indent:
                issues.append({
                    'level': 'info',
                    'category': '格式',
                    'location': f'第 {i+1} 段',
                    'message': '正文段落缺少首行缩进',
                    'suggestion': '正文段落应设置首行缩进 2 字符'
                })
        
        # 检查行距（应为 20 磅）
        if para.paragraph_format.line_spacing:
            if para.paragraph_format.line_spacing != 20:
                issues.append({
                    'level': 'info',
                    'category': '格式',
                    'location': f'第 {i+1} 段',
                    'message': f'行距不符合要求：{para.paragraph_format.line_spacing} 磅',
                    'suggestion': '正文行距应设置为固定值 20 磅'
                })
    
    return issues


def generate_quality_report(docx_path):
    """生成完整质量报告"""
    try:
        doc = Document(docx_path)
    except Exception as e:
        return {
            'success': False,
            'error': f'无法打开文件：{str(e)}'
        }
    
    all_issues = []
    
    # 1. 检查字数
    print("🔍 检查字数...")
    word_issues, word_stats = check_word_count(doc)
    all_issues.extend(word_issues)
    
    # 2. 检查标题层级
    print("🔍 检查标题层级...")
    heading_issues = check_heading_levels(doc)
    all_issues.extend(heading_issues)
    
    # 3. 检查图表编号
    print("🔍 检查图表编号...")
    numbering_issues = check_figure_numbering(doc)
    all_issues.extend(numbering_issues)
    
    # 4. 检查列表项
    print("🔍 检查列表项...")
    bullet_issues = check_bullet_points(doc)
    all_issues.extend(bullet_issues)
    
    # 5. 检查参考文献
    print("🔍 检查参考文献...")
    ref_issues, ref_count = check_reference_count(doc)
    all_issues.extend(ref_issues)
    
    # 6. 检查段落格式
    print("🔍 检查段落格式...")
    format_issues = check_paragraph_formatting(doc)
    all_issues.extend(format_issues)
    
    # 计算总分（简化评分）
    error_count = len([i for i in all_issues if i['level'] == 'error'])
    warning_count = len([i for i in all_issues if i['level'] == 'warning'])
    info_count = len([i for i in all_issues if i['level'] == 'info'])
    
    total_score = 100 - error_count * 10 - warning_count * 3 - info_count * 1
    total_score = max(0, min(100, total_score))
    
    report = {
        'success': True,
        'file_path': docx_path,
        'check_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'overall_score': total_score,
        'statistics': {
            'body_words': word_stats['body_words'],
            'review_words': word_stats['review_words'],
            'reference_count': ref_count,
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables)
        },
        'issue_summary': {
            'total': len(all_issues),
            'error': error_count,
            'warning': warning_count,
            'info': info_count
        },
        'issues': all_issues,
        'recommendations': generate_recommendations(all_issues, word_stats)
    }
    
    return report


def generate_recommendations(issues, word_stats):
    """根据检查结果生成建议"""
    recommendations = []
    
    # 字数建议
    if word_stats['body_words'] < 50000:
        gap = 50000 - word_stats['body_words']
        recommendations.append(
            f"正文字数不足 {gap} 字，建议扩展以下内容：\n"
            "  - 增加机制讨论（从分子、细胞、体内三个层面）\n"
            "  - 补充文献对比分析\n"
            "  - 增加局限性讨论和未来展望"
        )
    
    # 错误处理建议
    errors = [i for i in issues if i['level'] == 'error']
    if errors:
        recommendations.append(
            f"发现 {len(errors)} 个严重错误，必须修正：\n" +
            "\n".join(f"  - {e['message']}" for e in errors[:3])
        )
    
    # 格式建议
    format_issues = [i for i in issues if i['category'] == '格式']
    if len(format_issues) > 5:
        recommendations.append(
            "格式问题较多，建议：\n"
            "  - 使用样式模板统一格式\n"
            "  - 检查所有段落的首行缩进和行距"
        )
    
    if not recommendations:
        recommendations.append("论文质量良好，符合基本要求！")
    
    return recommendations


def format_report_text(report):
    """格式化报告为可读文本"""
    lines = []
    lines.append("=" * 70)
    lines.append("📋 博士论文质量检查报告")
    lines.append("=" * 70)
    lines.append(f"📄 文件：{os.path.basename(report['file_path'])}")
    lines.append(f"📅 检查时间：{report['check_date']}")
    lines.append(f"⭐ 总体评分：{report['overall_score']} / 100")
    lines.append("")
    
    # 统计信息
    stats = report['statistics']
    lines.append("📊 统计信息：")
    lines.append(f"   正文字数：{stats['body_words']:,} / 50,000")
    lines.append(f"   综述字数：{stats['review_words']:,} / 5,000")
    lines.append(f"   参考文献：{stats['reference_count']} / 80")
    lines.append(f"   总段落数：{stats['total_paragraphs']:,}")
    lines.append(f"   总表格数：{stats['total_tables']}")
    lines.append("")
    
    # 问题汇总
    summary = report['issue_summary']
    lines.append("⚠️  问题汇总：")
    lines.append(f"   严重错误：{summary['error']} 个")
    lines.append(f"   警告：{summary['warning']} 个")
    lines.append(f"   提示：{summary['info']} 个")
    lines.append("")
    
    # 详细问题列表
    if report['issues']:
        lines.append("📝 详细问题：")
        for i, issue in enumerate(report['issues'][:20], 1):  # 只显示前 20 个
            level_icon = {
                'error': '❌',
                'warning': '⚠️ ',
                'info': 'ℹ️ '
            }.get(issue['level'], '•')
            
            location = issue.get('location', '')
            location_str = f" [{location}]" if location else ""
            
            lines.append(f"{i}. {level_icon} {issue['message']}{location_str}")
            lines.append(f"   💡 {issue['suggestion']}")
            lines.append("")
        
        if len(report['issues']) > 20:
            lines.append(f"   ... 还有 {len(report['issues']) - 20} 个问题未显示")
            lines.append("")
    
    # 建议
    lines.append("💡 改进建议：")
    for i, rec in enumerate(report['recommendations'], 1):
        lines.append(f"{i}. {rec}")
        lines.append("")
    
    lines.append("=" * 70)
    
    return "\n".join(lines)


def main():
    """命令行入口"""
    if len(sys.argv) < 2:
        print("用法: python3 check_quality.py <文件路径> [--output json|text]")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_format = 'text'
    
    if '--output' in sys.argv:
        idx = sys.argv.index('--output')
        if idx + 1 < len(sys.argv):
            output_format = sys.argv[idx + 1]
    
    if not os.path.exists(docx_path):
        print(f"❌ 文件不存在：{docx_path}")
        sys.exit(1)
    
    print("🔍 开始质量检查...\n")
    
    # 生成报告
    report = generate_quality_report(docx_path)
    
    if not report.get('success'):
        print(f"❌ 检查失败：{report.get('error')}")
        sys.exit(1)
    
    # 输出报告
    if output_format == 'json':
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        print(format_report_text(report))
    
    # 保存报告
    report_path = docx_path.replace('.docx', '_质量报告.json')
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    print(f"\n💾 报告已保存：{report_path}")
    
    # 返回退出码
    if report['overall_score'] >= 80:
        sys.exit(0)
    else:
        sys.exit(1)


if __name__ == '__main__':
    main()
