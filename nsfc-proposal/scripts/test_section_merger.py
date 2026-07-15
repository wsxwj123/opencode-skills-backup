#!/usr/bin/env python3
"""回归测试：section_merger.py 合并/校验/角标上标（基准=当前代码真实行为）。

自包含、纯 assert、tempfile 造临时 sections/，跑完自动清理，不改被测脚本。
覆盖：
  validate_order —— 缺必需节报缺 / 有P2拆分文件不误报P2缺 / 完全无P2报P2缺
  _p2_children  —— P2_1/P2_2/P2_10 数字排序（非字典序），无编号排末尾
  merge_selected —— --only 别名 P2 展开为拆分文件、裸名补 .md、未知别名丢弃、空选择 CLI rc=2
  _superscript_citations —— 正文 [1]/[2,3]/[4-6] 转上标并计数；进入"参考文献"标题后 [N] 不动；
                             [图1]/[表2] 不匹配；已上标 run 跳过（需 python-docx，缺则 skip）
"""
from __future__ import annotations

import subprocess
import sys
import tempfile
from pathlib import Path

_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))

import section_merger as smg  # noqa: E402


def _mk_sections(root: Path, names: list[str]) -> Path:
    sec = root / "sections"
    sec.mkdir(parents=True, exist_ok=True)
    for n in names:
        (sec / n).write_text("正文内容", encoding="utf-8")
    return sec


def test_validate_order_missing_required() -> None:
    # 缺 P1（必需节）但有 P2 拆分文件 → 应报缺 P1，且不误报 P2 缺
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        names = [n for n in smg.ORDER if n != "P2_研究内容.md"]
        names.remove("P1_立项依据.md")
        names.append("P2_1_子内容.md")  # 提供 P2 拆分文件
        sec = _mk_sections(root, names)
        vo = smg.validate_order(sec)
        assert vo["ok"] is False, "缺 P1 应 ok=False"
        assert "P1_立项依据.md" in vo["missing"], f"missing 应含 P1: {vo['missing']}"
        assert not any("P2" in m for m in vo["missing"]), f"有P2拆分不应报P2缺: {vo['missing']}"
    print("validate_order 缺P1报缺、P2拆分不误报：OK")


def test_validate_order_p2_split_ok() -> None:
    # 全部必需节齐 + 无 P2_研究内容.md 但有拆分文件 → ok=True，missing 空
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        names = [n for n in smg.ORDER if n != "P2_研究内容.md"] + ["P2_1_子内容.md"]
        sec = _mk_sections(root, names)
        vo = smg.validate_order(sec)
        assert vo["ok"] is True, f"P2拆分齐应过: {vo}"
        assert vo["missing"] == [], vo["missing"]
    print("validate_order P2拆分文件替代整节：OK")


def test_validate_order_no_p2_reports_missing() -> None:
    # 有全部必需节但完全没有任何 P2（整节+拆分都无）→ 报 P2 缺
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        names = [n for n in smg.ORDER if n != "P2_研究内容.md"]
        sec = _mk_sections(root, names)
        vo = smg.validate_order(sec)
        assert vo["ok"] is False, "完全无P2应 ok=False"
        assert any("P2" in m for m in vo["missing"]), f"应报P2缺: {vo['missing']}"
    print("validate_order 完全无P2报缺（不静默跳过）：OK")


def test_p2_children_numeric_sort() -> None:
    with tempfile.TemporaryDirectory() as td:
        sec = _mk_sections(
            Path(td),
            ["P2_10_c.md", "P2_2_b.md", "P2_1_a.md", "P2_研究内容.md", "P2_无编号.md"],
        )
        order = [p.name for p in smg._p2_children(sec)]
        # 数字排序 1,2,10（非字典序 1,10,2）；无编号排末尾；整节 P2_研究内容.md 被排除
        assert order == ["P2_1_a.md", "P2_2_b.md", "P2_10_c.md", "P2_无编号.md"], order
    print("_p2_children 数字排序 + 无编号末尾：OK")


def test_merge_selected_aliases() -> None:
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        sec = _mk_sections(root, ["P1_立项依据.md", "P2_1_a.md", "P2_2_b.md", "REF_参考文献.md"])
        out = root / "o.md"
        # P1_立项依据(裸名补.md)、P2(别名展开为拆分文件)、UNKNOWNX(未知别名丢弃)、REF(全名)
        used = smg.merge_selected(sec, ["P1_立项依据", "P2", "UNKNOWNX", "REF_参考文献.md"], out)
        assert used == ["P1_立项依据.md", "P2_1_a.md", "P2_2_b.md", "REF_参考文献.md"], used
        assert out.exists() and out.read_text(encoding="utf-8").strip(), "合并输出应非空"
    print("merge_selected 别名展开/裸名补/未知丢弃：OK")


def test_merge_only_empty_rc2() -> None:
    # 空选择（--only 只有分隔符）→ 无节被选中 → CLI rc=2
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        _mk_sections(root, ["P1_立项依据.md"])
        proc = subprocess.run(
            [sys.executable, str(_HERE / "section_merger.py"), "merge",
             "--sections-dir", str(root / "sections"),
             "--output", str(root / "o.md"), "--only", ",, "],
            capture_output=True, text=True)
        assert proc.returncode == 2, f"空选择应 rc=2，实际 {proc.returncode}: {proc.stdout}"
    print("merge --only 空选择 rc=2：OK")


def test_superscript_citations() -> None:
    try:
        from docx import Document
    except ImportError:
        print("test_superscript_citations: SKIP (python-docx 未安装)")
        return
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "t.docx"
        doc = Document()
        doc.add_paragraph("背景[1]研究[2,3]方法[4-6]。")   # 3 组角标
        doc.add_paragraph("见[图1]和[表2]示意。")           # 含中文，不匹配
        para = doc.add_paragraph("已上标")
        r = para.add_run("[8]")
        r.font.superscript = True                           # 已上标，应跳过
        doc.add_paragraph("参考文献")                        # 标题：其后停止处理
        doc.add_paragraph("[7] Author. Title. 2020.")       # 参考文献条目编号，不动
        doc.save(str(p))

        changed = smg._superscript_citations(p)
        assert changed == 3, f"正文 3 组角标应计 3，实际 {changed}"

        doc2 = Document(str(p))
        body = doc2.paragraphs[0]
        sup_texts = [run.text for run in body.runs if run.font.superscript]
        assert sup_texts == ["[1]", "[2,3]", "[4-6]"], sup_texts
        fig = doc2.paragraphs[1]
        assert not any(run.font.superscript for run in fig.runs), "[图1]/[表2] 不应上标"
        for para in doc2.paragraphs:
            if para.text.strip().startswith("[7]"):
                assert not any(run.font.superscript for run in para.runs), "参考文献列表 [7] 不应上标"
    print("_superscript_citations 正文转/图表跳/参考文献不动/已上标跳：OK")


if __name__ == "__main__":
    test_validate_order_missing_required()
    test_validate_order_p2_split_ok()
    test_validate_order_no_p2_reports_missing()
    test_p2_children_numeric_sort()
    test_merge_selected_aliases()
    test_merge_only_empty_rc2()
    test_superscript_citations()
    print("ALL PASS")
