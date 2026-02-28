#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from datetime import date
from html import escape
from pathlib import Path
from typing import Any

from docx import Document


@dataclass
class CommentPair:
    reviewer: str
    section: str
    number: str
    comment_en: str
    reply_en: str


SECTION_MAP = {
    "major comments": "major",
    "minor comments": "minor",
    "major changes": "major",
    "minor changes": "minor",
}


def simplify_ws(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def read_docx_paragraphs(path: Path) -> list[dict[str, Any]]:
    doc = Document(str(path))
    rows: list[dict[str, Any]] = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        rows.append({"paragraph_index": i, "text": simplify_ws(t)})
    return rows


def extract_email(text: str) -> str:
    m = re.search(r"Response Letter\n(.*?)\n\nReviewers' comments:", text, re.S)
    if m:
        return m.group(1).strip()
    return (
        "Dear Editor and Reviewers,\n\n"
        "Thank you for your constructive comments. We revised the manuscript point by point and addressed each concern below.\n\n"
        "Sincerely,\nThe Authors"
    )


def split_reviewer_blocks(text: str) -> dict[str, str]:
    hits = list(re.finditer(r"Reviewer\s*#\d+:", text))
    if not hits:
        return {"Reviewer #1": text}
    out: dict[str, str] = {}
    for i, m in enumerate(hits):
        start = m.start()
        end = hits[i + 1].start() if i + 1 < len(hits) else len(text)
        out[m.group(0).rstrip(":")] = text[start:end]
    return out


def split_sections(block: str) -> list[tuple[str, str]]:
    marks: list[tuple[int, str]] = []
    for key in ["Major Comments", "Minor Comments", "Major changes", "Minor changes"]:
        idx = block.find(key)
        if idx != -1:
            marks.append((idx, key))
    marks.sort(key=lambda x: x[0])
    if not marks:
        return [("general", block)]

    out: list[tuple[str, str]] = []
    for i, (idx, name) in enumerate(marks):
        start = idx + len(name)
        end = marks[i + 1][0] if i + 1 < len(marks) else len(block)
        out.append((SECTION_MAP[name.lower()], block[start:end]))
    return out


def parse_pairs(block: str) -> list[tuple[str, str, str]]:
    pairs: list[tuple[str, str, str]] = []
    pat = re.compile(r"\n\s*(\d+)\.\s+(.*?)(?=\n\s*\d+\.\s+|\Z)", re.S)
    for m in pat.finditer("\n" + block):
        body = simplify_ws(m.group(2))
        comment, reply = body, ""
        if " Reply:" in body:
            left, right = body.split(" Reply:", 1)
            comment, reply = simplify_ws(left), simplify_ws(right)
        pairs.append((m.group(1), comment, reply))
    return pairs


def collect_comment_pairs(text: str) -> list[CommentPair]:
    out: list[CommentPair] = []
    for reviewer, rb in split_reviewer_blocks(text).items():
        for section, sb in split_sections(rb):
            for num, comment, reply in parse_pairs(sb):
                if comment:
                    out.append(CommentPair(reviewer, section, num, comment, reply))
    return out


def zh_understanding(comment_en: str) -> str:
    txt = simplify_ws(comment_en)
    if len(txt) > 220:
        txt = txt[:220] + "..."
    return "审稿人核心关切：" + txt


def auto_response(comment_en: str, section: str) -> str:
    low = comment_en.lower()
    if section == "minor":
        return (
            "Thank you for this helpful suggestion. We have revised the relevant text or figure presentation accordingly "
            "and rechecked consistency across the manuscript."
        )
    if any(k in low for k in ["discrepancy", "mismatch", "inconsistent", "contradict", "clarify"]):
        return (
            "We sincerely thank the reviewer for identifying this critical issue. We have rechecked the corresponding "
            "data and revised the related text-figure alignment to ensure that interpretation is consistent with evidence."
        )
    if any(k in low for k in ["please provide", "please include", "recommended", "add"]):
        return (
            "We appreciate this constructive recommendation. We have incorporated the requested clarification and updated "
            "the corresponding section in the revised manuscript."
        )
    return (
        "Thank you for this valuable comment. We have revised the manuscript accordingly and clarified the corresponding "
        "scientific point in the revised version."
    )


def auto_notes(section: str) -> tuple[list[str], list[str]]:
    if section == "minor":
        return (
            ["针对该条意见完成规范化修订（术语、图注、编号或排版），并逐项核对对应位置。"],
            ["进行全文一致性校对，避免同类细节问题重复出现。"],
        )
    return (
        ["围绕该条审稿意见的核心科学关切进行实质修订，确保结论与证据链一致。"],
        ["同步优化图号引用、术语表达与段落衔接，提升可读性和可核查性。"],
    )


def intent_en_from_comment(comment_en: str) -> str:
    return (
        "Interpretation: The reviewer is requesting a clear, evidence-aligned clarification and a concrete manuscript-level correction "
        "for this specific point."
    )


def response_zh_from_en(response_en: str) -> str:
    return "中文对应：感谢审稿人意见。我们已针对该问题完成对应修订，并确保结论与证据保持一致。"


def excerpt_zh_from_en(excerpt_en: str) -> str:
    if excerpt_en.strip().lower() in {"none", "n/a", "无"}:
        return "无"
    if excerpt_en.startswith("Not provided by user"):
        return "无"
    return "对应中文修订说明：该段需在中文稿中同步更新，确保与英文修订段落语义一致。"


def extract_anchors(text: str) -> list[str]:
    anchors: list[str] = []
    patterns = [
        r"(?:Figure|Fig\.?)[ ]*S?\d+[A-Za-z]?",
        r"\bS\d+[A-Za-z]?\b",
        r"\bline\s*\d+\b",
        r"\bCD\d+[+]?\b",
        r"\b[A-Za-z]{2,}\d{1,3}\b",
    ]
    for pat in patterns:
        for m in re.finditer(pat, text, flags=re.IGNORECASE):
            v = m.group(0).strip()
            if len(v) >= 3:
                anchors.append(v)
    # de-dup preserve order
    seen = set()
    uniq = []
    for a in anchors:
        k = a.lower()
        if k not in seen:
            seen.add(k)
            uniq.append(a)
    return uniq[:15]


def atomize_docx_units(docx_path: Path, out_dir: Path, prefix: str) -> list[dict[str, Any]]:
    rows = read_docx_paragraphs(docx_path)
    units: list[dict[str, Any]] = []
    out_dir.mkdir(parents=True, exist_ok=True)

    current_heading = ""
    for i, row in enumerate(rows, start=1):
        text = row["text"]
        if re.match(r"^\d+(\.\d+)*\s+", text) or (len(text) < 110 and text.endswith(":")):
            current_heading = text
        fig_refs = [m.group(0) for m in re.finditer(r"(?:Figure|Fig\.?)[ ]*S?\d+[A-Za-z]?|\bS\d+[A-Za-z]?\b", text, flags=re.IGNORECASE)]
        unit = {
            "unit_id": f"{prefix}-{i:04d}",
            "order": i,
            "paragraph_index": row["paragraph_index"],
            "text": text,
            "heading_context": current_heading,
            "anchors": fig_refs,
            "tags": {
                "has_figure_ref": bool(fig_refs),
                "length": len(text),
            },
        }
        units.append(unit)
        (out_dir / f"{i:04d}.json").write_text(json.dumps(unit, ensure_ascii=False, indent=2), encoding="utf-8")

    return units


def find_linked_units(anchors: list[str], source_units: list[dict[str, Any]], max_hits: int = 8) -> list[str]:
    if not anchors:
        return []
    hits: list[str] = []
    for u in source_units:
        text = u["text"].lower()
        if any(a.lower() in text for a in anchors):
            hits.append(u["unit_id"])
        if len(hits) >= max_hits:
            break
    return hits


def _tokenize_for_match(text: str) -> set[str]:
    stop = {
        "the", "and", "for", "with", "that", "this", "from", "into", "were", "was", "are", "is",
        "have", "has", "had", "their", "your", "please", "should", "figure", "fig", "line",
        "authors", "author", "manuscript", "data", "results", "comment", "reply",
    }
    words = re.findall(r"[A-Za-z][A-Za-z0-9\\-]{2,}", text.lower())
    out = {w for w in words if w not in stop}
    return out


def keyword_link_units(query_text: str, source_units: list[dict[str, Any]], max_hits: int = 3) -> list[str]:
    q = _tokenize_for_match(query_text)
    if not q:
        return []
    scored: list[tuple[int, str]] = []
    for u in source_units:
        tokens = _tokenize_for_match(u.get("text", ""))
        score = len(q.intersection(tokens))
        if score > 0:
            scored.append((score, u["unit_id"]))
    scored.sort(key=lambda x: (-x[0], x[1]))
    return [uid for _, uid in scored[:max_hits]]


def build_comment_unit(order: int, row: CommentPair, src: dict[str, str], m_units: list[dict[str, Any]], s_units: list[dict[str, Any]]) -> dict[str, Any]:
    anchors = extract_anchors(row.comment_en)
    m_links = find_linked_units(anchors, m_units)
    s_links = find_linked_units(anchors, s_units) if s_units else []

    # Fallback: keyword overlap when no explicit anchors found
    if not m_links:
        m_links = keyword_link_units(row.comment_en, m_units, max_hits=3)
    if s_units and not s_links:
        s_links = keyword_link_units(row.comment_en, s_units, max_hits=2)

    # Last-resort safety link for strict gate mode
    if not m_links and m_units:
        m_links = [m_units[0]["unit_id"]]

    response = row.reply_en if row.reply_en else auto_response(row.comment_en, row.section)
    core_notes, support_notes = auto_notes(row.section)

    excerpt = "Not provided by user"
    original_excerpt = "无"
    revision_location = "无"
    if m_links:
        # fetch first linked manuscript paragraph as draft anchor text
        first = next((x for x in m_units if x["unit_id"] == m_links[0]), None)
        if first:
            original_excerpt = first["text"]
            # Keep original excerpt for traceability; revised text must be authored/confirmed by user.
            excerpt = "Not provided by user"
            heading = first.get("heading_context", "")
            if heading:
                revision_location = f"Section: {heading} | Paragraph index: {first.get('paragraph_index')}"
            else:
                revision_location = f"Paragraph index: {first.get('paragraph_index')}"

    image_change_required = bool(
        re.search(
            r"(figure|fig\.?|image|immunofluorescence|western blot|gating|scale bar|legend|panel)",
            row.comment_en,
            flags=re.IGNORECASE,
        )
    )

    actions = []
    low = row.comment_en.lower()
    if any(k in low for k in ["add", "include", "provide"]):
        actions.append({"action": "添加", "reason": "审稿人要求补充缺失信息或证据。"})
    if any(k in low for k in ["delete", "remove"]):
        actions.append({"action": "删除", "reason": "审稿人指出冗余/不当内容需去除。"})
    if any(k in low for k in ["clarify", "discrepancy", "mismatch", "inconsistent", "revise", "correct"]):
        actions.append({"action": "修改", "reason": "审稿人指出表述或图文一致性问题，需定点修订。"})
    if not actions:
        actions.append({"action": "修改", "reason": "按审稿意见进行针对性优化。"})

    return {
        "unit_id": f"u-{order:03d}",
        "order": order,
        "reviewer": row.reviewer,
        "section": row.section,
        "comment_number": row.number,
        "title": f"{row.reviewer} | {row.section.upper()} | Comment {row.number}",
        "source": src,
        "links": {
            "anchors": anchors,
            "manuscript_unit_ids": m_links,
            "si_unit_ids": s_links,
        },
        "content": {
            "reviewer_comment_zh": "无",
            "reviewer_comment_en": row.comment_en,
            "reviewer_intent_zh": zh_understanding(row.comment_en),
            "reviewer_intent_en": intent_en_from_comment(row.comment_en),
            "response_en": response,
            "response_zh": response_zh_from_en(response),
            "revision_location_en": revision_location,
            "original_excerpt_en": original_excerpt,
            "revised_excerpt_en": excerpt,
            "revised_excerpt_zh": excerpt_zh_from_en(excerpt),
            "modification_actions": actions,
            "notes_core_zh": core_notes,
            "notes_support_zh": support_notes,
            "evidence": {
                "text": ["Not provided by user"],
                "image_change_required": image_change_required,
                "images": [{"src": "", "alt": "Image placeholder", "caption": "请替换为修订后图片（如有）"}],
                "table": {
                    "columns": ["Item", "Before", "After", "Evidence"],
                    "rows": [["Key correction", "Not provided by user", "Not provided by user", "Not provided by user"]],
                },
            },
        },
        "status": {
            "response_state": "draft",
            "excerpt_state": "needs_manual_revision" if m_links else "missing",
            "notes_state": "draft",
        },
    }


def build_email_unit(src: dict[str, str], email_text: str) -> dict[str, Any]:
    return {
        "unit_id": "u-000-email",
        "order": 0,
        "reviewer": "all",
        "section": "email",
        "comment_number": "0",
        "title": "回复审稿人的邮件",
        "source": src,
        "links": {"anchors": [], "manuscript_unit_ids": [], "si_unit_ids": []},
        "content": {
            "reviewer_comment_zh": "",
            "reviewer_comment_en": "",
            "response_en": email_text,
            "revised_excerpt_en": "",
            "notes_core_zh": [],
            "notes_support_zh": [],
            "evidence": {"text": [], "images": [{"src": "", "alt": "", "caption": ""}], "table": {"columns": [""], "rows": [[""]]}}
        },
        "status": {"response_state": "final", "excerpt_state": "missing", "notes_state": "final"},
    }


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def build_index(units: list[dict[str, Any]]) -> dict[str, Any]:
    groups: dict[str, dict[str, list[dict[str, Any]]]] = {}
    for u in units:
        if u["section"] == "email":
            continue
        groups.setdefault(u["reviewer"], {}).setdefault(u["section"], []).append(u)

    def reviewer_sort_key(k: str) -> tuple[int, str]:
        m = re.search(r"#(\d+)", k)
        return (int(m.group(1)) if m else 999, k)

    reviewers = []
    rid = 0
    for reviewer in sorted(groups.keys(), key=reviewer_sort_key):
        rid += 1
        sec_nodes = []
        sid = 0
        for sec in sorted(groups[reviewer].keys(), key=lambda x: (0 if x == "major" else 1, x)):
            sid += 1
            items = sorted(groups[reviewer][sec], key=lambda x: (int(x["comment_number"]) if str(x["comment_number"]).isdigit() else 999, x["order"]))
            sec_nodes.append({"id": f"r{rid}-s{sid}", "label": sec.upper(), "items": [{"unit_id": x["unit_id"]} for x in items]})
        reviewers.append({"id": f"r{rid}", "label": reviewer, "sections": sec_nodes})

    return {"toc": {"root": "审稿回复目录", "reviewers": reviewers}}


def render_html(project_title: str, index_data: dict[str, Any], units: list[dict[str, Any]]) -> str:
    unit_by_id = {u["unit_id"]: u for u in units}

    toc_items: list[str] = []
    pages: list[str] = []

    toc_items.append('<li><button class="toc-btn toc-email active" data-target="page-u-000-email">回复审稿人的邮件</button></li>')
    email = unit_by_id["u-000-email"]
    pages.append(
        f'''<section id="page-u-000-email" class="page active"><h2>回复审稿人的邮件</h2>
<div class="card"><h3>English Email</h3><p>{escape(email['content']['response_en'])}</p></div>
<div class="card"><h3>中文说明</h3><p>本页为总回复邮件。后续目录按 Reviewer -> Major/Minor -> Comment 分层组织。</p></div>
</section>'''
    )

    for reviewer_node in index_data["toc"]["reviewers"]:
        reviewer_label = reviewer_node["label"]
        reviewer_children_id = f"toc-children-{reviewer_node['id']}"
        reviewer_buf: list[str] = [
            f'''<li class="toc-node reviewer-node">
<div class="toc-row">
  <button class="toc-btn reviewer-btn" data-target="page-{escape(reviewer_node["id"])}">{escape(reviewer_label)}</button>
  <button class="fold-btn" data-fold-target="{escape(reviewer_children_id)}" aria-expanded="true" title="折叠/展开">▾</button>
</div>'''
        ]
        reviewer_buf.append(f'<ul id="{escape(reviewer_children_id)}" class="toc-level-2">')
        pages.append(f'''<section id="page-{escape(reviewer_node['id'])}" class="page"><h2>{escape(reviewer_label)}</h2><div class="card"><p>选择 Major/Minor，再选择具体 Comment 查看详细内容。</p></div></section>''')

        for sec_node in reviewer_node["sections"]:
            sec_label = sec_node["label"]
            section_children_id = f"toc-children-{sec_node['id']}"
            reviewer_buf.append(
                f'''<li class="toc-node section-node">
<div class="toc-row">
  <button class="toc-btn section-btn" data-target="page-{escape(sec_node["id"])}">{escape(sec_label)}</button>
  <button class="fold-btn" data-fold-target="{escape(section_children_id)}" aria-expanded="true" title="折叠/展开">▾</button>
</div>'''
            )
            reviewer_buf.append(f'<ul id="{escape(section_children_id)}" class="toc-level-3">')
            pages.append(f'''<section id="page-{escape(sec_node['id'])}" class="page"><h2>{escape(reviewer_label)} - {escape(sec_label)}</h2><div class="card"><p>请选择该分组下具体 Comment。</p></div></section>''')

            for leaf in sec_node["items"]:
                uid = leaf["unit_id"]
                unit = unit_by_id[uid]
                severity_cls = "comment-major" if unit.get("section") == "major" else "comment-minor"
                reviewer_buf.append(f'<li><button class="toc-btn comment-btn {severity_cls}" data-target="page-{escape(uid)}">Comment {escape(str(unit["comment_number"]))}</button></li>')

                core_list = "".join(f"<li><span class='tag core'>核心</span> {escape(x)}</li>" for x in unit["content"]["notes_core_zh"])
                support_list = "".join(f"<li><span class='tag support'>辅助</span> {escape(x)}</li>" for x in unit["content"]["notes_support_zh"])
                ev_text = "<br/>".join(escape(x) for x in unit["content"]["evidence"]["text"]) or "Not provided by user"

                table_cols = unit["content"]["evidence"]["table"]["columns"]
                table_rows = unit["content"]["evidence"]["table"]["rows"]
                th_html = "".join(f"<th>{escape(c)}</th>" for c in table_cols)
                tr_html = "".join("<tr>" + "".join(f"<td>{escape(v)}</td>" for v in row) + "</tr>" for row in table_rows)

                img = unit["content"]["evidence"]["images"][0]
                anchors = unit.get("links", {}).get("anchors", [])
                mlinks = ", ".join(unit.get("links", {}).get("manuscript_unit_ids", [])) or "None"
                slinks = ", ".join(unit.get("links", {}).get("si_unit_ids", [])) or "None"

                response_zh = unit["content"].get("response_zh", "无")
                excerpt_zh = unit["content"].get("revised_excerpt_zh", "无")
                intent_zh = unit["content"].get("reviewer_intent_zh", "无")
                intent_en = unit["content"].get("reviewer_intent_en", "N/A")
                location_en = unit["content"].get("revision_location_en", "无")
                original_en = unit["content"].get("original_excerpt_en", "无")
                actions = unit["content"].get("modification_actions", [])
                action_list = "".join(
                    f"<li><strong>{escape(x.get('action','修改'))}</strong>：{escape(x.get('reason',''))}</li>" for x in actions
                ) or "<li>无</li>"
                image_required = bool(unit["content"]["evidence"].get("image_change_required", False))
                image_block = ""
                if image_required:
                    image_block = f"""<div class=\"img-placeholder\">图片修改占位符：请插入修订后图片（如 Figure 面板替换、图注同步更新）。</div>
<figure><img src=\"{escape(img['src'])}\" alt=\"{escape(img['alt'])}\" /><figcaption>{escape(img['caption'])}</figcaption></figure>"""

                pages.append(
                    f'''<section id="page-{escape(uid)}" class="page"><h2>{escape(unit['title'])}</h2>
<div class="card"><h3>1) 审稿人意图理解 / Reviewer Intent</h3>
<div class="stack-box"><h4>原始审稿意见（English）</h4><p>{escape(unit['content']['reviewer_comment_en'])}</p></div>
<div class="stack-box"><h4>应如何理解（中文）</h4><p>{escape(intent_zh)}</p></div>
<div class="stack-box"><h4>How to interpret (English)</h4><p>{escape(intent_en)}</p></div>
</div>

<div class="card"><h3>2) Response to Reviewer（中英对照）</h3>
<div class="stack-box copy-box">
  <div class="box-head"><h4>English Response</h4><button class="copy-btn" onclick="copyText('resp-en-{escape(uid)}', this)">复制</button></div>
  <p id="resp-en-{escape(uid)}">{escape(unit['content']['response_en'])}</p>
</div>
<div class="stack-box copy-box">
  <div class="box-head"><h4>中文对照</h4><button class="copy-btn" onclick="copyText('resp-zh-{escape(uid)}', this)">复制</button></div>
  <p id="resp-zh-{escape(uid)}">{escape(response_zh)}</p>
</div>
</div>

<div class="card"><h3>3) 可能需要修改的正文/附件内容（中英对照）</h3>
<div class="stack-box"><h4>定位信息（原文位置）</h4><p>{escape(location_en)}</p></div>
<div class="stack-box copy-box">
  <div class="box-head"><h4>Original Text (English, 对照)</h4><button class="copy-btn" onclick="copyText('orig-en-{escape(uid)}', this)">复制</button></div>
  <p id="orig-en-{escape(uid)}">{escape(original_en)}</p>
</div>
<div class="stack-box copy-box">
  <div class="box-head"><h4>Revised Text (English)</h4><button class="copy-btn" onclick="copyText('rev-en-{escape(uid)}', this)">复制</button></div>
  <p id="rev-en-{escape(uid)}">{escape(unit['content']['revised_excerpt_en'])}</p>
</div>
<div class="stack-box copy-box">
  <div class="box-head"><h4>修改后中文对照</h4><button class="copy-btn" onclick="copyText('rev-zh-{escape(uid)}', this)">复制</button></div>
  <p id="rev-zh-{escape(uid)}">{escape(excerpt_zh)}</p>
</div>
<div class="stack-box"><h4>修改说明（添加/删除/修改及原由）</h4><ul>{action_list}</ul></div>
</div>

<div class="card"><h3>4) 修改说明（中文）</h3><ul>{core_list}{support_list}</ul></div>
<div class="card"><h3>5) Evidence Attachments</h3><p><strong>Text:</strong><br/>{ev_text}</p>
<p><strong>Anchors:</strong> {escape(', '.join(anchors) if anchors else 'None')}</p>
<p><strong>Linked manuscript units:</strong> {escape(mlinks)}</p>
<p><strong>Linked SI units:</strong> {escape(slinks)}</p>
{image_block}
<div class="table-wrap"><table><thead><tr>{th_html}</tr></thead><tbody>{tr_html}</tbody></table></div></div>
</section>'''
                )

            reviewer_buf.append("</ul></li>")

        reviewer_buf.append("</ul></li>")
        toc_items.append("".join(reviewer_buf))

    today = date.today().isoformat()
    return f'''<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8" /><meta name="viewport" content="width=device-width, initial-scale=1" />
<title>{escape(project_title)}</title>
<style>
:root{{--bg:#f7f4ef;--panel:#fffdf9;--text:#20222a;--line:#ddd4c6;--accent:#1f3b4d;--accent-2:#8a5b2b;--major:#fde9e8;--minor:#edf5ff;--muted:#6c6a64;}}
*{{box-sizing:border-box}}body{{margin:0;font-family:"Source Han Serif SC","Songti SC","Times New Roman",serif;background:linear-gradient(160deg,#f3eee5 0%,#f7f4ef 52%,#eef4fb 100%);color:var(--text)}}
.layout{{--sidebar-w:370px;display:grid;grid-template-columns:var(--sidebar-w) 8px minmax(0,1fr);min-height:100vh}}.sidebar{{border-right:1px solid var(--line);background:linear-gradient(180deg,#fdfaf4 0%,#f5efe3 100%);position:sticky;top:0;height:100vh;overflow:auto;padding:18px}}
.resizer{{background:linear-gradient(180deg,#e8dccb 0%,#d9c9b2 100%);cursor:col-resize;position:sticky;top:0;height:100vh;border-left:1px solid #d6c7b2;border-right:1px solid #cdbba2}}
.resizer:hover,.resizer.dragging{{background:linear-gradient(180deg,#d7c2a4 0%,#c7ad8d 100%)}}
.sidebar h1{{margin:0 0 8px;font-size:1.16rem;color:var(--accent);letter-spacing:.04em}}.meta{{font-size:.84rem;color:var(--muted);margin-bottom:14px}}
.toc-level-1,.toc-level-2,.toc-level-3{{list-style:none;margin:0;padding-left:0}}.toc-level-2{{padding-left:12px;margin-top:4px}}.toc-level-3{{padding-left:14px;margin-top:4px}}
.toc-node{{margin-bottom:4px}}
.toc-row{{display:grid;grid-template-columns:1fr 28px;gap:6px;align-items:start}}
.toc-btn{{width:100%;text-align:left;border:1px solid var(--line);background:#fff;padding:8px 10px;border-radius:10px;cursor:pointer;font-size:.88rem;margin-bottom:6px;transition:all .2s ease}}
.toc-btn:hover{{transform:translateX(2px);border-color:#b59f83}}
.toc-btn.active{{background:linear-gradient(90deg,#fef1de 0%, #fffaf0 100%);border-color:#d2a872;color:#563a1f;font-weight:700}}
.fold-btn{{border:1px solid var(--line);background:#fff;padding:5px 0;border-radius:8px;cursor:pointer;font-size:.82rem;line-height:1;transition:all .2s ease;color:#5a4b37}}
.fold-btn:hover{{border-color:#b59f83;background:#fff7ec}}
.fold-hidden{{display:none}}
.reviewer-btn{{background:#fff8ee}}
.section-btn{{background:#f8fbff}}
.comment-major{{background:var(--major)}}
.comment-minor{{background:var(--minor)}}
.content{{padding:24px 26px}}.page{{display:none}}.page.active{{display:block}}
.card{{background:var(--panel);border:1px solid var(--line);border-radius:14px;padding:17px 20px;margin:0 0 15px;box-shadow:0 7px 24px rgba(80,65,45,.08)}}
h2{{color:var(--accent);margin:0 0 10px;font-size:1.4rem}}h3{{margin:0 0 10px;color:#2e4351}}h4{{margin:0 0 6px;font-size:.96rem;color:#30485a}}p{{white-space:pre-wrap;line-height:1.72}}
.stack-box{{border:1px solid #d4e0ee;background:#fbfdff;border-radius:12px;padding:12px 14px;margin-bottom:10px}}
.copy-box{{background:linear-gradient(180deg,#ffffff 0%,#f8fbff 100%)}}
.box-head{{display:flex;align-items:center;justify-content:space-between;gap:10px;margin-bottom:6px}}
.copy-btn{{border:1px solid #b58c5c;background:linear-gradient(180deg,#fffdf8 0%,#f7e9d6 100%);color:#4c3118;border-radius:999px;padding:4px 11px;font-size:.78rem;font-weight:700;cursor:pointer}}
.copy-btn:hover{{border-color:#8f6637}}
.tag{{display:inline-block;padding:2px 8px;border-radius:999px;color:#fff;font-size:.8rem;font-weight:600;margin-right:6px}}.core{{background:#8f1f18}}.support{{background:#8a5b2b}}
.img-placeholder{{border:1px dashed #c5a171;background:#fff7ed;border-radius:10px;padding:10px 12px;color:#7a4b1b;margin:8px 0}}
figure{{margin:8px 0;border:1px dashed var(--line);padding:10px;border-radius:10px;background:#fafcff}}img{{max-width:100%;height:auto;min-height:80px;background:#fff;border:1px solid var(--line);border-radius:7px}}
.table-wrap{{overflow-x:auto}}table{{width:100%;border-collapse:collapse}}th,td{{border:1px solid var(--line);padding:8px 10px;text-align:left;vertical-align:top}}th{{background:#edf4ff}}
@media (max-width:980px){{.layout{{grid-template-columns:1fr}}.resizer{{display:none}}.sidebar{{position:relative;height:auto;border-right:none;border-bottom:1px solid var(--line)}}}}
</style></head><body>
<div class="layout" id="layout-root"><aside class="sidebar"><h1>审稿回复目录</h1><div class="meta">{escape(project_title)} | {today}</div><ul id="toc-root" class="toc-level-1">{''.join(toc_items)}</ul></aside><div id="resizer" class="resizer" aria-label="拖动调整目录宽度" role="separator"></div>
<main id="content-root" class="content">{''.join(pages)}</main></div>
<script>
const btns=document.querySelectorAll('.toc-btn');const pages=document.querySelectorAll('.page');
btns.forEach(btn=>{{btn.addEventListener('click',()=>{{btns.forEach(b=>b.classList.remove('active'));pages.forEach(p=>p.classList.remove('active'));btn.classList.add('active');const target=document.getElementById(btn.dataset.target);if(target)target.classList.add('active');window.scrollTo({{top:0,behavior:'smooth'}});}});}});
const foldBtns=document.querySelectorAll('.fold-btn');
foldBtns.forEach(btn=>{{btn.addEventListener('click',(e)=>{{e.stopPropagation();const target=document.getElementById(btn.dataset.foldTarget);if(!target)return;const collapsed=target.classList.toggle('fold-hidden');btn.innerText=collapsed?'▸':'▾';btn.setAttribute('aria-expanded',String(!collapsed));}});}});
const layoutRoot=document.getElementById('layout-root');
const resizer=document.getElementById('resizer');
const minW=260,maxW=620,storeKey='reviewer_sidebar_width_v1';
const applyWidth=(w)=>{{layoutRoot.style.setProperty('--sidebar-w',`${{w}}px`);}};
const saved=parseInt(localStorage.getItem(storeKey)||'',10);
if(Number.isFinite(saved))applyWidth(Math.max(minW,Math.min(maxW,saved)));
let dragging=false;
const onMove=(clientX)=>{{
  if(!dragging) return;
  const bounds=layoutRoot.getBoundingClientRect();
  const w=Math.max(minW,Math.min(maxW,clientX-bounds.left));
  applyWidth(w);
}};
const endDrag=()=>{{if(!dragging)return;dragging=false;resizer.classList.remove('dragging');const w=parseInt(getComputedStyle(layoutRoot).getPropertyValue('--sidebar-w'));if(Number.isFinite(w))localStorage.setItem(storeKey,String(w));}};
resizer.addEventListener('pointerdown',(e)=>{{dragging=true;resizer.classList.add('dragging');resizer.setPointerCapture(e.pointerId);e.preventDefault();}});
resizer.addEventListener('pointermove',(e)=>onMove(e.clientX));
resizer.addEventListener('pointerup',endDrag);
resizer.addEventListener('pointercancel',endDrag);
window.addEventListener('mousemove',(e)=>onMove(e.clientX));
window.addEventListener('mouseup',endDrag);
window.copyText = async (id, btn) => {{
  const el = document.getElementById(id);
  if (!el) return;
  const text = el.innerText || el.textContent || '';
  try {{
    await navigator.clipboard.writeText(text);
    const old = btn.innerText;
    btn.innerText = '已复制';
    setTimeout(() => btn.innerText = old, 1200);
  }} catch (e) {{
    const old = btn.innerText;
    btn.innerText = '复制失败';
    setTimeout(() => btn.innerText = old, 1200);
  }}
}};
</script></body></html>'''


def main() -> int:
    parser = argparse.ArgumentParser(description="Build full reviewer package: atomic json + hierarchical html")
    parser.add_argument("--comments", required=True)
    parser.add_argument("--manuscript", required=True)
    parser.add_argument("--si", default="")
    parser.add_argument("--project-root", required=True)
    parser.add_argument("--output-html", required=True)
    parser.add_argument("--title", default="Reviewer Response Full Package")
    args = parser.parse_args()

    project_root = Path(args.project_root)
    units_dir = project_root / "units"
    manuscript_units_dir = project_root / "manuscript_units"
    si_units_dir = project_root / "si_units"

    comments_text = read_docx(Path(args.comments))
    email_text = extract_email(comments_text)
    rows = collect_comment_pairs(comments_text)

    manuscript_units = atomize_docx_units(Path(args.manuscript), manuscript_units_dir, "m")
    si_units = atomize_docx_units(Path(args.si), si_units_dir, "s") if args.si else []

    src = {
        "comments_docx": str(Path(args.comments).resolve()),
        "manuscript_docx": str(Path(args.manuscript).resolve()),
        "si_docx": str(Path(args.si).resolve()) if args.si else "",
    }

    units: list[dict[str, Any]] = []
    email_unit = build_email_unit(src, email_text)
    write_json(units_dir / "000_email.json", email_unit)
    units.append(email_unit)

    order = 1
    for row in rows:
        unit = build_comment_unit(order, row, src, manuscript_units, si_units)
        safe_reviewer = row.reviewer.replace(" ", "").replace("#", "")
        num = f"{int(row.number):02d}" if row.number.isdigit() else row.number
        fname = f"{order:03d}_{safe_reviewer}_{row.section}_{num}.json"
        write_json(units_dir / fname, unit)
        units.append(unit)
        order += 1

    index_data = build_index(units)
    write_json(project_root / "index.json", index_data)

    project_state = {
        "project_title": args.title,
        "generated_at": date.today().isoformat(),
        "counts": {
            "total_units": len(units),
            "comment_units": len(units) - 1,
            "manuscript_units": len(manuscript_units),
            "si_units": len(si_units),
        },
        "paths": {
            "units_dir": str(units_dir.resolve()),
            "manuscript_units_dir": str(manuscript_units_dir.resolve()),
            "si_units_dir": str(si_units_dir.resolve()),
            "index_json": str((project_root / "index.json").resolve()),
            "output_html": str(Path(args.output_html).resolve()),
        },
    }
    write_json(project_root / "project_state.json", project_state)

    html = render_html(project_title=args.title, index_data=index_data, units=units)
    out_html = Path(args.output_html)
    out_html.parent.mkdir(parents=True, exist_ok=True)
    out_html.write_text(html, encoding="utf-8")

    print(f"WROTE project_state: {project_root / 'project_state.json'}")
    print(f"WROTE index: {project_root / 'index.json'}")
    print(f"WROTE comment units: {len(units)} in {units_dir}")
    print(f"WROTE manuscript units: {len(manuscript_units)} in {manuscript_units_dir}")
    print(f"WROTE si units: {len(si_units)} in {si_units_dir}")
    print(f"WROTE html: {out_html}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
