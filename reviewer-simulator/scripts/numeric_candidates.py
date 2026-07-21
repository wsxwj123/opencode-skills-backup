#!/usr/bin/env python3
"""数值候选抽取（数值一致性核查·第 1 层确定性锚，4 家 vendored 逐字节一致）。

从成稿（docx/md/markdown/txt）客观抽取"带上下文的数值候选"——数字 + 单位 + 所在句
+ 指标名线索 + 分组/时间点线索 + 位置，写 <project-root>/numeric_candidates.json +
stdout 末行摘要 JSON。**只列稿子里出现了哪些指标数值，不判任何两个数是否矛盾、不判是否
同一测量、不判容差**——判断全留给第 2/3 层。

复用同目录 manuscript_index.py（8 家逐字节共享文件，绝不改，只 import）的读稿/区分正文与
参考区能力；docx 分支额外遍历 doc.tables 抽表格单元格数值（python-docx doc.paragraphs
不下钻表格，是"表格 IC50 vs 正文"头号用例的命门），仅在本脚本补，不改 manuscript_index。

护栏方向＝**宁抽勿拒**（漏抽真数值 → 真矛盾核不出，危害大于误列一个无害候选）；仅"必排"
清单（p 值/CI/图表号/引用号/年份/参考区）偏保守拒。样本量 n=/N= 纳入抽取（BRIEF 招牌用例）。

CLI:
  python numeric_candidates.py --manuscript <docx|md|markdown|txt> --project-root <root>

退出码：0=正常（含空稿/无候选）；2=用法/输入错（缺参 / 文件不存在 / 不支持类型 / 解析失败）。
契约见 .devflow/INTERFACE-numeric-consistency.md §1。
"""
import sys as _sys
try:  # Windows GBK 控制台/管道捕获下防 UnicodeEncodeError
    _sys.stdout.reconfigure(encoding="utf-8")
    _sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

import argparse
import bisect
import json
import re
import sys
from pathlib import Path
from typing import Any, Optional

# 同目录 manuscript_index.py 是 8 家逐字节共享文件，绝不改，只 import 复用其读稿层。
from manuscript_index import (
    normalize_ws,
    write_json,
    read_manuscript_paragraphs,
    reference_section_spans,
    body_row_indices,
    is_heading,
    FIG_INTEXT_RE,
    CITATION_GROUP_RE,
)

SUPPORTED_SUFFIXES = {".docx", ".md", ".markdown", ".txt"}

# 每候选存的 `sentence` 上限（I-1 性能修复的根因所在）：无断句符的超长"句"（现实=
# markdown 宽表整行 / 对抗输入）里可有成千上万个数值候选，若每个都存整句，输出体积
# = O(候选数 × 句长) = O(n²)（JSON 编码 + 落盘随之 O(n²)）。真实句远短于此上限，≤上限
# 时整句原样存（与修前逐字节一致）；仅超上限时按各数值取窗口，把输出/耗时压回线性。
SENTENCE_CAP = 2000

# 表内文引用 / 题注（本地新增，manuscript_index 只认图不认表）。
TABLE_INTEXT_RE = re.compile(r"\b(?:Table|Tab\.?)\s*(\d+)|表\s*(\d+)", re.IGNORECASE)
TABLE_CAPTION_RE = re.compile(
    r"^(?:Table|Tab\.?)\s*(\d+)\s*[.:]\s*(\S.*)$|^表\s*(\d+)\s*[.:：]\s*(\S.*)$",
    re.IGNORECASE,
)
TABLE_BARE_RE = re.compile(r"^(?:Table|Tab\.?)\s*(\d+)\s*$|^表\s*(\d+)\s*$", re.IGNORECASE)

# 计量单位表（借 structure_outline 的 _UNITS，补数值场景常见的浓度/时间单位）。
_UNITS = {
    "mg", "kg", "g", "l", "ml", "μl", "µl", "ul", "μm", "µm", "um", "mm", "cm", "nm",
    "μg", "µg", "ug", "ng", "pg", "mol", "mmol", "nmol", "μmol", "µmol", "umol",
    "kda", "da", "bp", "kb", "℃", "°c", "u", "iu", "rpm", "hz", "v", "w", "pa", "n",
}
_TIME_UNITS = {"h", "hr", "min", "s", "ms", "sec", "d"}
# 摩尔浓度 M / mM / μM / nM / pM / uM（区分大小写：mM≠mm 毫米）。
_MOLAR_RE = re.compile(r"^[munpμµ]?M$")
# 上标字符（D1，仅 docx 上标 run 重建时出现）：把上标 run 里的数字/正负号还原成 Unicode
# 上标字符，让 _SCI_RE 认出 1×10⁶、单位识别认出 mm³。md/txt 无上标概念，对其无影响。
_SUP_ALL = "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻"
_SUP_TO_ASCII = str.maketrans(_SUP_ALL, "0123456789+-")
_ASCII_TO_SUP = str.maketrans("0123456789+-", _SUP_ALL)
# 数字后紧邻的 ASCII/单位符号串（遇空格/数字/CJK 即止）；含上标字符以吃下 mm³ 的体积指数。
_UNIT_AFTER_RE = re.compile(rf"\s*([A-Za-zμµ%℃°/·{_SUP_ALL}]+)")

_YEAR_RE = re.compile(r"^(?:19|20)\d{2}$")
# 样本量线索：数字前紧邻 n= / N=（保守只认 n/N，避免误吞 p=/α= 等统计惯用）。
_SAMPLE_RE = re.compile(r"[nN]\s*=\s*$")

# 必排（几乎不可能是要比的指标）—— 这些 span 内的数字一律不进候选。
_PVAL_RE = re.compile(r"\bp\s*[<>=≤≥]\s*\d*\.?\d+", re.IGNORECASE)
_ALPHA_RE = re.compile(r"(?:α|alpha)\s*=\s*\d*\.?\d+", re.IGNORECASE)
_R2_RE = re.compile(r"\bR\s*(?:²|\^?2)\s*=\s*\d*\.?\d+", re.IGNORECASE)
_CI_RE = re.compile(r"\d+(?:\.\d+)?\s*%?\s*(?:CI\b|置信区间)", re.IGNORECASE)
_FORBIDDEN_RES = (_PVAL_RE, _ALPHA_RE, _R2_RE, _CI_RE,
                  FIG_INTEXT_RE, TABLE_INTEXT_RE, CITATION_GROUP_RE)

# 数值形态词法（按优先级从"最具体的复合形态"到"裸数字"，逐位置 match，先中者胜）。
_NUM = r"\d+(?:\.\d+)?"
_SCI_RE = re.compile(
    rf"({_NUM})\s*(?:[×xX*]\s*10\s*(?:\^\s*([+-]?\d+)|([{_SUP_ALL}]+))|[eE]([+-]?\d+))"
)
_MEANSD_RE = re.compile(rf"({_NUM})\s*±\s*({_NUM})")
# 含 ASCII `-`（D2）：仅"数字-数字"两侧纯数值才当范围；负号/日期年份区间另有护栏（见 _classify）。
_RANGE_RE = re.compile(rf"({_NUM})\s*[–—~～-]\s*({_NUM})\s*(%)?")
_RATIO_RE = re.compile(rf"({_NUM})\s*:\s*({_NUM})")
_PCT_RE = re.compile(rf"({_NUM})\s*(%)")
_THOUSANDS_RE = re.compile(r"(\d{1,3}(?:,\d{3})+(?:\.\d+)?)")
_PLAIN_RE = re.compile(rf"({_NUM})")

# 分组/时间点线索（判"是否同一测量"的关键上下文，宁多带勿漏）。
_GROUP_RE = re.compile(r"[一-鿿]{1,5}?组")
_TIMEPOINT_RE = re.compile(r"\d+\s*(?:h|hr|min|d|小时|分钟|天|周)")

# metric_clue 提取时可从数字前缀剥掉的连接词（剥完取尾部指标名 token）。
_CONNECTORS = ("约为", "达到", "等于", "为", "是", "的", "达", "约", "均", "：", ":", "=", "，", ",")
# 英文停用词（D4）：介词/冠词/连词不是指标线索，取到就跳过、继续往前找真实指标词。
_STOPWORDS = {"at", "and", "for", "in", "of", "to", "by", "with", "a", "the"}

# 粗区段关键词（location.region 启发式，供人核，非确定性）。
_REGION_KEYS = [
    (("abstract", "摘要"), "Abstract"),
    (("results", "结果"), "Results"),
    (("methods", "materials and methods", "material and methods", "材料与方法", "方法"), "Methods"),
    (("discussion", "讨论"), "Discussion"),
    (("introduction", "background", "引言", "前言", "背景"), "Body"),
]


def read_md_lines(path: Path) -> list[dict[str, Any]]:
    """md/txt 逐行成行（一非空行=一 row），不按空行粗切块。

    仿 structure_outline.read_md_lines：manuscript_index.read_md_paragraphs 按空行把多行
    粘成一 block，无空行的 Word 导出 txt 会把区段标题与正文粘连、扰乱区段/句上下文。逐行发
    row 让每段独立、区段跟踪与句切分正确。`#` 前缀标题去号并标 Heading 样式。
    """
    raw = path.read_text(encoding="utf-8", errors="ignore")
    rows: list[dict[str, Any]] = []
    para_index = 0
    for line in raw.splitlines():
        if not line.strip():
            continue
        is_md_heading = line.lstrip().startswith("#")
        cleaned = normalize_ws(re.sub(r"^#{1,6}\s*", "", line))
        if not cleaned:
            continue
        rows.append({
            "paragraph_index": para_index,
            "text": cleaned,
            "style_name": "Heading" if is_md_heading else "Normal",
        })
        para_index += 1
    return rows


def _match_region(text: str) -> Optional[str]:
    t = re.sub(r"^\d+(?:\.\d+)*[.、]?\s*", "", normalize_ws(text)).strip().lower()
    for keys, region in _REGION_KEYS:
        if t in keys:
            return region
    # 合并标题（Results and Discussion / 结果与讨论）：整行**仅由区段词 + 连接词**组成才算标题
    # （有其它实词则是正文句，不当区段，防误判）；命中的第一个区段词定 region（_REGION_KEYS
    # 顺序：results 先于 discussion，故合并标题归 results）。(D5)
    tokens = [w for w in re.split(r"[\s/&，,]+|\band\b|\bor\b|与|和|及", t) if w]
    if not tokens or len(tokens) > 5:
        return None
    key_of = {k: region for keys, region in _REGION_KEYS for k in keys}
    hit = None
    for tok in tokens:
        if tok not in key_of:
            return None
        if hit is None:
            hit = key_of[tok]
    return hit


def _split_sentences(text: str) -> list[str]:
    """段内中英混排断句：CJK 终止符后 / 拉丁 .!? 且后接空白处切。小数点(0.05)不被切碎。"""
    parts = re.split(r"(?<=[。！？；;])|(?<=[.!?])(?=\s)", text or "")
    return [p for p in (s.strip() for s in parts) if p]


def _clean_unit(tok: str) -> str:
    """把数字后紧邻 token 判为计量单位则原样返回，否则空串（不误把英文散文词当单位）。"""
    if not tok:
        return ""
    if tok == "%":
        return "%"
    base = tok.rstrip(_SUP_ALL)  # 剥尾部上标指数(mm³→mm)做单位判定，返回时保留上标(D1)
    if not base:
        return ""
    if "μ" in base or "µ" in base:  # micro 前缀单位（μM/μg/μmol...）
        return tok
    if _MOLAR_RE.match(base):  # 摩尔浓度 M/mM/nM/pM
        return tok
    if base.lower() in _TIME_UNITS:
        return tok
    parts = [p for p in base.split("/") if p]
    if parts and all(p.lower() in _UNITS for p in parts):
        return tok
    return ""


def _overlaps(span: tuple[int, int], forbidden: list[tuple[int, int]]) -> bool:
    s, e = span
    return any(s < fe and fs < e for fs, fe in forbidden)


def _build_clue_index(sent: str) -> dict[str, Any]:
    """句级预计算一次分组/时间点线索索引（I-1 性能修复：避免每候选重扫整句 O(n²)）。

    保持旧 _group_clue 的语义：命中集按"组线索(升序) ++ 时间点线索(升序)"拼接，取
    start<num_start 中的最后一个 → 即时间点里最靠近的（若有），否则组里最靠近的；
    整句无前置命中时取全局最早的一处。分开存两份升序 start 列表供 bisect O(log n) 查。
    """
    group_hits = sorted((m.start(), m.group()) for m in _GROUP_RE.finditer(sent))
    tp_hits = sorted((m.start(), m.group()) for m in _TIMEPOINT_RE.finditer(sent))
    earliest = min(group_hits + tp_hits)[1] if (group_hits or tp_hits) else ""
    return {
        "group_hits": group_hits, "group_starts": [s for s, _ in group_hits],
        "tp_hits": tp_hits, "tp_starts": [s for s, _ in tp_hits],
        "earliest": earliest,
    }


def _group_clue(idx: dict[str, Any], num_start: int) -> str:
    i = bisect.bisect_left(idx["tp_starts"], num_start)
    if i > 0:  # 最靠近的时间点线索（时间点在拼接串尾，优先）
        return idx["tp_hits"][i - 1][1]
    j = bisect.bisect_left(idx["group_starts"], num_start)
    if j > 0:  # 否则最靠近的组线索
        return idx["group_hits"][j - 1][1]
    return idx["earliest"]  # 无前置命中 → 全局最早


def _strip_connectors(s: str) -> str:
    s = s.strip()
    changed = True
    while changed:
        changed = False
        for c in _CONNECTORS:
            if s.endswith(c):
                s = s[: -len(c)].strip()
                changed = True
    return s


def _metric_clue(sent: str, num_start: int, group_clue: str, is_sample: bool) -> str:
    """指标名线索：样本量优先；否则剥掉分组/连接词后取数字前紧邻的指标名 token。不强行聚类。"""
    if is_sample:
        return "样本量"
    # 只看数字前有界窗口（I-1 性能修复：避免每候选切整段前缀 O(n²)）。指标 token 至多 8 个
    # CJK / 一段标识符，紧贴数字；窗口 96 远超之，对真实文本与旧全前缀切片结果一致。
    before = sent[max(0, num_start - 96):num_start]
    if group_clue and group_clue in before:
        before = before.replace(group_clue, "")
    before = _strip_connectors(before)
    # 取尾部指标 token；命中英文停用词(at/of/for...)则跳过它继续往前找真实指标词，抽不到为空(D4)。
    while True:
        m = re.search(r"([A-Za-z][A-Za-z0-9]*|[一-鿿]{1,8})$", before)
        if not m:
            return ""
        tok = m.group(1)
        if tok.lower() in _STOPWORDS:
            before = _strip_connectors(before[:m.start()])
            continue
        return tok


def _classify(form: str, m: "re.Match", sent: str) -> Optional[dict[str, Any]]:
    """把一个形态匹配转成候选骨架（value/value_secondary/norm/unit/form/raw + start/end）。"""
    start = m.start()
    num_end = m.end()
    unit = ""
    norm = None
    value_secondary: Optional[float] = None

    if form == "scientific":
        base = float(m.group(1))
        if m.group(2) is not None:      # ^N 尖角指数
            exp = int(m.group(2))
        elif m.group(3) is not None:    # 上标指数 ¹⁰⁶（D1，docx 上标 run 重建）
            exp = int(m.group(3).translate(_SUP_TO_ASCII))
        else:                           # eN 科学计数
            exp = int(m.group(4))
        value = base * (10 ** exp)
    elif form == "mean_sd":
        value = float(m.group(1))
        value_secondary = float(m.group(2))
    elif form == "range":
        # 日期年份区间（ASCII `-` 连两个四位年，如 2019-2021）不当数值范围抽（D2 防误伤日期）。
        # en/em dash/~（–—~～）区间不受此限——它们本就不用于写年份区间。
        if "-" in m.group(0) and _YEAR_RE.match(m.group(1)) and _YEAR_RE.match(m.group(2)):
            return None
        value = float(m.group(1))
        value_secondary = float(m.group(2))
        if m.group(3):
            unit = "%"
    elif form == "ratio":
        value = float(m.group(1))
        value_secondary = float(m.group(2))
    elif form == "percent":
        value = float(m.group(1))
        unit = "%"
        norm = value / 100.0
    else:  # plain / thousands
        value = float(m.group(1).replace(",", ""))

    raw_end = num_end
    # 尾随单位（percent/带%的range 已定单位，不再追）。
    if not unit:
        um = _UNIT_AFTER_RE.match(sent, num_end)
        if um:
            cu = _clean_unit(um.group(1))
            if cu:
                unit = cu
                raw_end = um.end()

    return {
        "form": form,
        "value": value,
        "value_secondary": value_secondary,
        "norm": norm,
        "unit": unit,
        "raw": sent[start:raw_end],
        "_start": start,
        "_end": num_end,
    }


_PATTERNS = [
    ("scientific", _SCI_RE),
    ("mean_sd", _MEANSD_RE),
    ("range", _RANGE_RE),
    ("ratio", _RATIO_RE),
    ("percent", _PCT_RE),
    ("plain", _THOUSANDS_RE),
    ("plain", _PLAIN_RE),
]


def _extract_from_sentence(sent: str) -> list[dict[str, Any]]:
    forbidden = []
    for rx in _FORBIDDEN_RES:
        for mm in rx.finditer(sent):
            forbidden.append((mm.start(), mm.end()))

    clue_idx = _build_clue_index(sent)  # 句级一次预计算，句内各候选复用（I-1）
    out: list[dict[str, Any]] = []
    i = 0
    n = len(sent)
    while i < n:
        matched = None
        for form, rx in _PATTERNS:
            m = rx.match(sent, i)
            if m:
                matched = (form, m)
                break
        if not matched:
            i += 1
            continue
        form, m = matched
        # 数字紧贴前面的 ASCII 字母 → 标识符/名称的一部分（IC50/CD4/p53），非测量值，整体跳过。
        if m.start() > 0 and sent[m.start() - 1].isascii() and sent[m.start() - 1].isalpha():
            i = m.end()
            continue
        # 必排 span（p 值/CI/图表号/引用号）内的数字跳过。
        if _overlaps((m.start(), m.end()), forbidden):
            i = m.end()
            continue

        cand = _classify(form, m, sent)
        if cand is None:
            i = m.end()
            continue

        # 样本量线索只看数字前有界窗口（I-1：避免切整段前缀）。归一化文本无长空白串，
        # "n = " 至多几字符，窗口 16 足够且与旧全前缀 `$` 锚定结果一致。
        is_sample = bool(_SAMPLE_RE.search(sent[max(0, cand["_start"] - 16):cand["_start"]]))
        # 年份（四位、无单位、非样本量）排除。
        if (cand["form"] == "plain" and not cand["unit"] and not is_sample
                and _YEAR_RE.match(m.group(1).replace(",", ""))):
            i = m.end()
            continue

        group_clue = _group_clue(clue_idx, cand["_start"])
        metric_clue = _metric_clue(sent, cand["_start"], group_clue, is_sample)
        cand["group_clue"] = group_clue
        cand["metric_clue"] = metric_clue
        out.append(cand)
        i = cand["_end"]
    return out


def _read_docx_tables(path: Path) -> list[tuple[str, Optional[str]]]:
    """遍历 docx 表格单元格，返回 [(cell_text, table_ref)]。table_ref 由文档顺序里就近题注推断，
    推不出为 None。**只在本脚本补，不改 manuscript_index。**

    确定性纪律：直接按文档序遍历底层 w:tr/w:tc XML 元素、逐 cell 拼 w:t 文本，**不碰
    python-docx 的 row.cells / cell._tc**——后者对合并单元格会重复返回同一 tc，若用
    `id(cell._tc)` 去重则依赖 lxml 瞬时代理对象的内存地址，地址被 GC 复用会让不同 cell 的
    id() 偶发相撞、把真单元格误当"已见"漏抽（与 PYTHONHASHSEED 无关的按次随机 flaky）。
    直接遍历 w:tc 每个物理单元格只出现一次（横向合并=单 tc+gridSpan，纵向合并续格 vMerge
    内容为空自然跳过），无需去重，输出对文档序恒定。
    """
    from docx import Document
    from docx.oxml.ns import qn

    w_p, w_tbl, w_tr, w_tc, w_t = (qn("w:p"), qn("w:tbl"), qn("w:tr"),
                                   qn("w:tc"), qn("w:t"))

    def _el_text(el) -> str:
        return normalize_ws("".join(t.text or "" for t in el.iter(w_t)))

    doc = Document(str(path))
    out: list[tuple[str, Optional[str]]] = []
    last_caption: Optional[str] = None
    for child in doc.element.body.iterchildren():
        if child.tag == w_p:
            text = _el_text(child)
            if not text:
                continue
            cm = TABLE_CAPTION_RE.match(text) or TABLE_BARE_RE.match(text)
            if cm:
                no = next((g for g in cm.groups() if g and g.isdigit()), None)
                last_caption = f"Table {no}" if no else None
        elif child.tag == w_tbl:
            for tr in child.findall(w_tr):  # 直接子 w:tr，不下钻嵌套表
                for tc in tr.findall(w_tc):
                    ct = _el_text(tc)
                    if ct:
                        out.append((ct, last_caption))
            last_caption = None
    return out


def _para_text_with_sup(para) -> str:
    """段落文本，但**上标 run** 的数字/正负号还原成 Unicode 上标字符（D1）。

    只重建上标（`1×10`+上标`6`→`1×10⁶`、`200mm`+上标`3`→`200mm³`）；**下标及普通 run 原样
    拼接、绝不插分隔**——化学式/指标名下标(`IC`+下标`50`)须仍拼成 `IC50` 紧贴前导字母，
    才继续被"数字贴前导字母=标识符跳过"规则挡住，不误抽成量值（下标回归锁）。
    """
    parts = []
    for run in para.runs:
        t = run.text or ""
        if t and getattr(getattr(run, "font", None), "superscript", None):
            t = t.translate(_ASCII_TO_SUP)
        parts.append(t)
    return "".join(parts)


def _read_docx_body_rows(manuscript_path: Path) -> list[dict[str, Any]]:
    """自读 docx 正文段落，行结构与 manuscript_index.read_docx_paragraphs 一致（paragraph_index
    = enumerate 序、同 style_name），差别只在上标 run 重建（D1）。不改 manuscript_index，只在本
    脚本补；非上标内容逐字节等价于原读稿层，故 ref/heading/表格等下游行为不变。"""
    from docx import Document
    doc = Document(str(manuscript_path))
    rows: list[dict[str, Any]] = []
    for i, para in enumerate(doc.paragraphs):
        text = normalize_ws(_para_text_with_sup(para))
        if not text:
            continue
        style_name = normalize_ws(getattr(getattr(para, "style", None), "name", "") or "")
        rows.append({"paragraph_index": i, "text": text, "style_name": style_name})
    return rows


def build_candidates(manuscript_path: Path) -> dict[str, Any]:
    suffix = manuscript_path.suffix.lower()
    if suffix == ".docx":
        rows = _read_docx_body_rows(manuscript_path)
    else:
        rows = read_md_lines(manuscript_path)

    ref_spans = reference_section_spans(rows)
    body_set = set(body_row_indices(rows, ref_spans))

    raw_cands: list[dict[str, Any]] = []
    current_region = "Body"
    for idx, row in enumerate(rows):
        text = row.get("text", "")
        reg = _match_region(text)
        if reg:
            current_region = reg
        if idx not in body_set or is_heading(row):
            continue
        para_index = row.get("paragraph_index")
        for sent in _split_sentences(text):
            sentence_norm = normalize_ws(sent)  # 每句一次，句内各候选复用（I-1：勿每候选重跑整句 sub）
            short = len(sentence_norm) <= SENTENCE_CAP
            for cand in _extract_from_sentence(sent):
                cand["sentence"] = (
                    sentence_norm if short
                    else normalize_ws(sent[max(0, cand["_start"] - 1000):cand["_start"] + 1000])
                )
                cand["location"] = {
                    "region": current_region, "para_index": para_index,
                    "source": "body", "table_ref": None,
                }
                raw_cands.append(cand)

    from_tables = 0
    if suffix == ".docx":
        for cell_text, table_ref in _read_docx_tables(manuscript_path):
            cell_norm = normalize_ws(cell_text)  # 每单元格一次，复用（I-1）
            if len(cell_norm) > SENTENCE_CAP:  # 超长单元格同样封顶，防 O(n²)
                cell_norm = cell_norm[:SENTENCE_CAP]
            for sent in _split_sentences(cell_text):
                for cand in _extract_from_sentence(sent):
                    cand["sentence"] = cell_norm
                    cand["location"] = {
                        "region": "Table", "para_index": None,
                        "source": "table", "table_ref": table_ref,
                    }
                    raw_cands.append(cand)
                    from_tables += 1

    candidates = []
    for i, c in enumerate(raw_cands, 1):
        candidates.append({
            "id": f"num-{i:04d}",
            "raw": c["raw"],
            "value": c["value"],
            "value_secondary": c["value_secondary"],
            "norm": c["norm"],
            "unit": c["unit"],
            "form": c["form"],
            "sentence": c["sentence"],
            "metric_clue": c["metric_clue"],
            "group_clue": c["group_clue"],
            "location": c["location"],
        })

    return {
        "candidates": candidates,
        "summary": {"candidates": len(candidates), "from_tables": from_tables},
    }


def main() -> int:
    parser = argparse.ArgumentParser(
        description="从成稿抽取带上下文的数值候选，当数值一致性核查第 1 层确定性锚。"
    )
    parser.add_argument("--manuscript", required=True, help="成稿路径（docx / md / markdown / txt）。")
    parser.add_argument("--project-root", required=True, help="输出根目录，numeric_candidates.json 写这里。")
    args = parser.parse_args()

    manuscript_path = Path(args.manuscript).expanduser().resolve()
    project_root = Path(args.project_root).expanduser().resolve()

    # 用法/输入错：显式 sys.exit(2)（不用 raise SystemExit(字符串)，那会变 exit 1）。
    if not manuscript_path.exists():
        print(f"manuscript not found: {manuscript_path}", file=sys.stderr)
        sys.exit(2)
    if manuscript_path.suffix.lower() not in SUPPORTED_SUFFIXES:
        print(f"unsupported manuscript type: {manuscript_path.suffix}", file=sys.stderr)
        sys.exit(2)

    # 解析失败（损坏/加密/空包 docx、编码错等）套 try/except 兜成 exit 2，禁止裸抛成 exit 1。
    try:
        result = build_candidates(manuscript_path)
    except Exception as exc:  # noqa: BLE001  # 契约要求：解析失败 → exit 2
        print(f"failed to parse manuscript: {exc}", file=sys.stderr)
        sys.exit(2)

    write_json(project_root / "numeric_candidates.json", result)
    print(json.dumps({"ok": True, **result["summary"]}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
