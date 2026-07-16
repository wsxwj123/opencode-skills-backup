#!/usr/bin/env python3
"""Merge polished units in order into polished_manuscript.md.

按 units_index 顺序合并 polished/<idx>.json 的 polished_text,在 section_type
变化处插入对应小节标题。可选 --docx 导出(需 python-docx;失败仅警告,md 仍产出)。
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from common import read_json, write_text

# 默认字体:西文 Times New Roman(行内 run 的中文字体由 md_runs 处理)
LATIN_FONT = "Times New Roman"
# 中文字体:显式设 w:eastAsia,否则 Word 用默认 CJK 字体回退(非宋体)。对齐 sci2doc/revise-sci 用宋体。
CJK_FONT = "宋体"

# 行内标记:**bold** / *italic* / <sup>…</sup> / <sub>…</sub>
# 注意 **bold** 必须在 *italic* 之前匹配,否则 ** 会被当成两个 *
_INLINE_RE = re.compile(
    r"\*\*(?P<bold>.+?)\*\*"
    r"|\*(?P<italic>[^*]+?)\*"
    r"|<sup>(?P<sup>.*?)</sup>"
    r"|<sub>(?P<sub>.*?)</sub>",
    re.DOTALL,
)


SECTION_TITLE = {
    "abstract": "Abstract",
    "intro": "Introduction",
    "methods": "Methods",
    "results": "Results",
    "discussion": "Discussion",
    "other": "",
}


def build_markdown(project_root: Path) -> tuple[str, list[dict]]:
    index = read_json(project_root / "units_index.json", {"units": []})
    polished_dir = project_root / "polished"
    parts: list[str] = []
    ordered_units: list[dict] = []
    last_heading = None
    for entry in index.get("units", []):
        idx = entry["idx"]
        unit = read_json(polished_dir / f"{idx}.json", None)
        if unit is None:
            continue
        heading = unit.get("heading") or entry.get("heading") or ""
        if not heading:
            heading = SECTION_TITLE.get(unit.get("section_type", "other"), "")
        if heading and heading != last_heading:
            # heading_level 还原 `#` 层级:优先 polished unit,回退 index entry,
            # 旧 unit 无此字段则回退 2(沿用历史 `## ` 行为)。clamp 到 1-6。
            level = unit.get("heading_level", entry.get("heading_level", 2)) or 2
            level = max(1, min(6, level))
            parts.append(f"{'#' * level} {heading}")
            last_heading = heading
        text = unit.get("polished_text", "").strip()
        if text:
            parts.append(text)
        ordered_units.append(unit)
    return "\n\n".join(p for p in parts if p).strip() + "\n", ordered_units


_HEADING_LINE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_LINE = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBER_LINE = re.compile(r"^\s*\d+[.)]\s+(.*)$")


def export_docx(md_text: str, docx_path: Path) -> bool:
    """手搓 md 重建导出(回退路径,不保原格式)。逐行解析:标题/列表/散文段落,
    行内标记经 md_runs.inline_md_to_runs 去净并转 run 级格式。polish 英文 SCI:
    显式锁 Normal 样式 Times New Roman + 段前段后=0,每个 run 强制 TNR + 黑。"""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return False
    from md_runs import clamp_heading_level, inline_md_to_runs

    doc = Document()
    # rebuild 路径锁 Normal 样式:Times New Roman + 零段距
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for block in md_text.split("\n\n"):
        if not block.strip():
            continue
        for line in block.split("\n"):
            if not line.strip():
                continue
            m = _HEADING_LINE.match(line)
            if m:
                level = clamp_heading_level(len(m.group(1)), 6)
                heading = doc.add_heading("", level=level)
                inline_md_to_runs(heading, m.group(2).strip(),
                                  latin=LATIN_FONT, east_asia=CJK_FONT, black=True)
                continue
            mb = _BULLET_LINE.match(line)
            if mb:
                p = doc.add_paragraph(style="List Bullet")
                inline_md_to_runs(p, mb.group(1).strip(), latin=LATIN_FONT, east_asia=CJK_FONT, black=True)
                continue
            mn = _NUMBER_LINE.match(line)
            if mn:
                p = doc.add_paragraph(style="List Number")
                inline_md_to_runs(p, mn.group(1).strip(), latin=LATIN_FONT, east_asia=CJK_FONT, black=True)
                continue
            p = doc.add_paragraph()
            inline_md_to_runs(p, line.strip(), latin=LATIN_FONT, east_asia=CJK_FONT, black=True)
    doc.save(str(docx_path))
    return True


def _capture_base_rpr(paragraph):
    """克隆段落首个非空 run 的 rPr 元素作为基础字符格式模板。

    克隆整个 rPr(而非逐属性枚举)可保留颜色、下划线、字距、eastAsia 等所有原格式;
    重建 run 时套用此模板再叠加语义标记。无 rPr 或全空段落返回 None。"""
    from copy import deepcopy
    from docx.oxml.ns import qn

    for run in paragraph.runs:
        if not run.text:
            continue
        rpr = run._element.find(qn("w:rPr"))
        if rpr is not None:
            return deepcopy(rpr)
        break
    return None


def _apply_base_rpr(run, base_rpr) -> None:
    """把克隆的 rPr 模板套到新 run 上(rPr 须为 <w:r> 的首子元素)。

    先移除 python-docx 建 run 时生成的空 rPr,再插入克隆模板,使颜色/下划线/字距/
    eastAsia 等原格式全部保留。语义标记(italic/sup/sub/bold)在此之后叠加。"""
    from copy import deepcopy
    from docx.oxml.ns import qn

    if base_rpr is None:
        return
    r = run._element
    existing = r.find(qn("w:rPr"))
    if existing is not None:
        r.remove(existing)
    r.insert(0, deepcopy(base_rpr))


def _rewrite_paragraph_runs(paragraph, marked_text: str, base_rpr) -> None:
    """清空段落 runs,按行内标记重建,每个新 run 继承段落基础字体再叠加语义格式。
    段落 pPr(对齐/样式/缩进)不动:只删 w:r 子元素,保留 w:pPr。"""
    from docx.oxml.ns import qn

    p = paragraph._element
    for run_el in list(p.findall(qn("w:r"))):
        p.remove(run_el)
    pos = 0
    spans: list[tuple[str, dict]] = []
    for m in _INLINE_RE.finditer(marked_text):
        if m.start() > pos:
            spans.append((marked_text[pos:m.start()], {}))
        if m.group("bold") is not None:
            spans.append((m.group("bold"), {"bold": True}))
        elif m.group("italic") is not None:
            spans.append((m.group("italic"), {"italic": True}))
        elif m.group("sup") is not None:
            spans.append((m.group("sup"), {"superscript": True}))
        elif m.group("sub") is not None:
            spans.append((m.group("sub"), {"subscript": True}))
        pos = m.end()
    if pos < len(marked_text):
        spans.append((marked_text[pos:], {}))
    if not spans:
        spans.append((marked_text, {}))
    for text, fmt in spans:
        if not text:
            continue
        run = paragraph.add_run(text)
        # 先套基础 rPr(保留颜色/下划线/字体等),再叠加语义标记,顺序不可颠倒。
        _apply_base_rpr(run, base_rpr)
        if fmt.get("bold"):
            run.bold = True
        if fmt.get("italic"):
            run.italic = True
        if fmt.get("superscript"):
            run.font.superscript = True
        if fmt.get("subscript"):
            run.font.subscript = True


def _paragraph_has_run_color_or_underline(paragraph) -> bool:
    """段内是否存在 run 携带 w:color(显式颜色)或 w:u(下划线,且非 none)。

    这两种 run 级格式不在 serialize_runs_to_marked_text 的标记范围内(只序列化
    italic/bold/sup/sub),polished_text 里既无对应标记、LLM 也从未见过它们,因此
    无法在 in-place 重建时还原。检测到此类段落且文字未被改动时,调用方应跳过破坏性
    重建、保留原 runs,以无损保住颜色/下划线。"""
    from docx.oxml.ns import qn

    for run in paragraph.runs:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            continue
        if rpr.find(qn("w:color")) is not None:
            return True
        u = rpr.find(qn("w:u"))
        if u is not None and u.get(qn("w:val")) not in (None, "none"):
            return True
    return False


def _paragraph_has_embedded_image(paragraph) -> bool:
    """段落是否含内嵌图片(<w:drawing>)或旧式嵌入对象(<w:object>)。
    清空此类段落的 run(in-place 重建)会删掉图并留下孤儿图片关系,故调用方
    必须对这类段落跳过 run 重建以保图。口径与 revise-sci 的同名函数一致。"""
    from docx.oxml.ns import qn

    p = paragraph._p
    return bool(
        p.findall(".//" + qn("w:drawing"))
        or p.findall(".//" + qn("w:object"))
    )


def export_inplace(src_docx: Path, project_root: Path, out_docx: Path) -> dict:
    """In-place 保格式导出:打开原始 docx,只把 prose 段落文字替换为 polished 文本,
    其余(标题/表格/图片/页眉页脚/非散文段落)完全不动。
    映射来源:units/<idx>.json 的 source_para_index + polished/<idx>.json 的 polished_text。
    映射对不齐 -> fail-closed 抛错,绝不错位写入。"""
    from docx import Document

    doc = Document(str(src_docx))
    paragraphs = doc.paragraphs
    n_paras = len(paragraphs)

    index = read_json(project_root / "units_index.json", {"units": []})
    units_dir = project_root / "units"
    polished_dir = project_root / "polished"

    plan: list[dict] = []          # 待写回:{para_index, text}
    seen_para_indices: set[int] = set()
    skipped_nonprose = 0

    for entry in index.get("units", []):
        idx = entry["idx"]
        src_unit = read_json(units_dir / f"{idx}.json", None)
        pol_unit = read_json(polished_dir / f"{idx}.json", None)
        if src_unit is None:
            raise ValueError(f"in-place: units/{idx}.json 缺失,无法定位源段落,fail-closed")
        if pol_unit is None:
            raise ValueError(f"in-place: polished/{idx}.json 缺失,fail-closed")

        para_index = src_unit.get("source_para_index")
        is_prose = pol_unit.get("prose", src_unit.get("prose", True)) is not False

        if not is_prose:
            # 非散文段落(参考文献/致谢等)原文不动,跳过写回
            skipped_nonprose += 1
            continue
        if para_index is None:
            raise ValueError(
                f"in-place: unit {idx} 无 source_para_index(可能源自 md 输入),"
                f"无法 in-place 写回,fail-closed。请改用 md 重建路径或提供原始 docx。"
            )
        if not isinstance(para_index, int) or para_index < 0 or para_index >= n_paras:
            raise ValueError(
                f"in-place: unit {idx} 的 source_para_index={para_index} 超出原 docx 段落范围 "
                f"[0,{n_paras}),fail-closed"
            )
        if para_index in seen_para_indices:
            raise ValueError(
                f"in-place: source_para_index={para_index} 被多个 unit 引用,映射冲突,fail-closed"
            )
        seen_para_indices.add(para_index)
        polished_text = pol_unit.get("polished_text", "")
        raw_text = pol_unit.get("raw_text", src_unit.get("raw_text", ""))
        plan.append({"para_index": para_index, "text": polished_text, "raw_text": raw_text})

    if not plan:
        raise ValueError("in-place: 没有可写回的 prose 段落(plan 为空),fail-closed")

    import sys

    rewritten = 0
    skipped_images: list[int] = []
    skipped_color_underline: list[int] = []
    for item in plan:
        para_index = item["para_index"]
        paragraph = paragraphs[para_index]
        # fail-safe:含内嵌图片的段落无法 in-place 重建 run —— _rewrite_paragraph_runs
        # 会清掉 <w:drawing>/<w:object> 并留下孤儿图片关系。改写文字时也无法确定图在新
        # 文字里的位置,强行重排会放错图。故跳过该段文字改写、保留原 runs(图)不动,提示人工。
        if _paragraph_has_embedded_image(paragraph):
            skipped_images.append(para_index)
            print(
                f"[export_inplace] warning: 段 {para_index} 含内嵌图片,"
                f"跳过 in-place 改写以保图,请人工处理该段文字。",
                file=sys.stderr,
            )
            continue
        # fail-safe:段内 run 级颜色/下划线无法经 marked_text 还原(不在标记范围内)。
        # 若文字未被润色改动,重建会把这些格式抹平 -> 直接跳过重建、保留原 runs 无损保色。
        # 若文字已改动,原颜色/下划线锚定的词可能已不在,无可靠映射,只能按已知局限重建丢失
        #(见 SKILL.md「docx 往返已知局限」)。
        text_unchanged = item["text"] == item["raw_text"]
        if text_unchanged and _paragraph_has_run_color_or_underline(paragraph):
            skipped_color_underline.append(para_index)
            continue
        base_rpr = _capture_base_rpr(paragraph)
        _rewrite_paragraph_runs(paragraph, item["text"], base_rpr)
        rewritten += 1

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    return {
        "rewritten_paragraphs": rewritten,
        "skipped_nonprose": skipped_nonprose,
        "paragraphs_skipped_images": skipped_images,
        "paragraphs_skipped_color_underline": skipped_color_underline,
        "total_src_paragraphs": n_paras,
    }


def _print_inplace_skip_manifest(result: dict, project_root: Path) -> None:
    """交付前把 in-place 静默跳过的段落升为逐条点名的显式清单(醒目 stderr),
    不让作者漏掉:含图段跳过润色、改动段丢run级颜色/下划线。"""
    skipped_images = result.get("paragraphs_skipped_images", []) or []
    skipped_color = result.get("paragraphs_skipped_color_underline", []) or []
    if not skipped_images and not skipped_color:
        return
    # para_index -> raw 预览:从 units/*.json 的 source_para_index 反查
    preview: dict[int, str] = {}
    units_dir = project_root / "units"
    if units_dir.exists():
        for uf in units_dir.glob("*.json"):
            u = read_json(uf, None) or {}
            pidx = u.get("source_para_index")
            if isinstance(pidx, int):
                preview[pidx] = re.sub(r"\s+", " ", u.get("raw_text", "")).strip()[:120]
    sys.stderr.write("\n" + "=" * 70 + "\n")
    sys.stderr.write("[in-place] 交付前必读:以下段落未被完整处理,请逐条人工复核\n")
    sys.stderr.write("=" * 70 + "\n")
    if skipped_images:
        sys.stderr.write(f"● 含内嵌图片、跳过文字润色({len(skipped_images)} 段,保留原文保图):\n")
        for pidx in skipped_images:
            sys.stderr.write(f"    · 段 {pidx}: {preview.get(pidx, '(无预览)')}\n")
    if skipped_color:
        sys.stderr.write(f"● 丢失 run 级颜色/下划线({len(skipped_color)} 段,文字未改故保留原格式):\n")
        for pidx in skipped_color:
            sys.stderr.write(f"    · 段 {pidx}: {preview.get(pidx, '(无预览)')}\n")
    sys.stderr.write("=" * 70 + "\n")


def main() -> int:
    parser = argparse.ArgumentParser(description="Merge polished units into manuscript")
    parser.add_argument("--project-root", required=True)
    parser.add_argument("--output-md", default="")
    parser.add_argument("--docx", default="", help="optional lossy docx output (md 重建路径)")
    parser.add_argument(
        "--in-place-src", default="",
        help="原始输入 docx 路径;提供则走保格式 in-place 导出(只换 prose 段落文字,保留原稿所有格式/表格/图片)",
    )
    args = parser.parse_args()

    project_root = Path(args.project_root)
    output_md = Path(args.output_md) if args.output_md else project_root / "polished_manuscript.md"
    md_text, units = build_markdown(project_root)
    write_text(output_md, md_text)

    docx_ok = None
    inplace_result = None
    if args.in_place_src:
        # in-place 保格式导出:--docx 指定输出路径,缺省则 polished_inplace.docx
        out_docx = Path(args.docx) if args.docx else project_root / "polished_inplace.docx"
        try:
            from docx import Document  # noqa: F401
        except ImportError:
            inplace_result = {"error": "python-docx 未安装,无法 in-place 导出"}
            docx_ok = False
        else:
            inplace_result = export_inplace(Path(args.in_place_src), project_root, out_docx)
            inplace_result["out_docx"] = str(out_docx.resolve())
            docx_ok = True
            _print_inplace_skip_manifest(inplace_result, project_root)
    elif args.docx:
        docx_ok = export_docx(md_text, Path(args.docx))

    print(json.dumps(
        {"ok": True, "output_md": str(output_md.resolve()), "units": len(units),
         "docx": (str(Path(args.docx).resolve()) if args.docx else None),
         "docx_ok": docx_ok, "in_place": inplace_result},
        ensure_ascii=False,
    ))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
