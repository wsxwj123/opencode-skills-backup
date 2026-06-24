#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
polish-sci 格式契约回归测试（开发者维护工具，非运行时流程）。

固化 in-place 保格式导出的 run 级格式行为，防止以后改 run 重建逻辑时退化。
纯 assert，无 pytest 依赖，所有输入在测试内现造（内存 docx + tempfile）。
不依赖任何外部真稿路径。

用法：
    python3 test_format_contract.py
失败抛 AssertionError（returncode != 0），全过打印 ALL OK。
"""

import json
import os
import sys
import tempfile
from pathlib import Path

# 本测试文件与被测脚本同目录，确保脚本目录在 import 路径上
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPT_DIR not in sys.path:
    sys.path.insert(0, _SCRIPT_DIR)


def _make_project(src_doc, units):
    """落地一个最小 polish 项目：源 docx + units/<idx>.json + polished/<idx>.json。

    units: list of dict，每个含 idx/source_para_index/raw_text/polished_text。
    返回 (project_root, src_docx_path)。
    """
    import common

    root = tempfile.mkdtemp()
    os.makedirs(os.path.join(root, "units"))
    os.makedirs(os.path.join(root, "polished"))
    src = os.path.join(root, "src.docx")
    src_doc.save(src)

    index = {"units": [{"idx": u["idx"]} for u in units]}
    json.dump(index, open(os.path.join(root, "units_index.json"), "w"))

    for u in units:
        pidx = u["source_para_index"]
        json.dump(
            {"idx": u["idx"], "source_para_index": pidx, "prose": True,
             "raw_text": u["raw_text"]},
            open(os.path.join(root, "units", f"{u['idx']}.json"), "w"),
        )
        json.dump(
            {"idx": u["idx"], "source_para_index": pidx, "prose": True,
             "raw_text": u["raw_text"], "polished_text": u["polished_text"]},
            open(os.path.join(root, "polished", f"{u['idx']}.json"), "w"),
        )
    return root, src


# ===========================================================================
# 契约 ⑥：in-place 导出对 run 级颜色/下划线、上下标/斜体的保真
#
# bug：被润色段含「中间 run 红色(w:color)/下划线(w:u)」时，in-place 重建后丢失，
#      因为颜色/下划线不在 marked_text 标记范围（只序列化 italic/bold/sup/sub），
#      且 _capture_base_rpr 只取首 run rPr 作整段统一模板。
# 修：文字未被改动且段内存在 run 级颜色/下划线 → 跳过破坏性重建、保留原 runs 无损。
#
# 已知局限（非退化，无可靠映射）：若该段文字被润色改动，原颜色/下划线锚定的词可能
#      已不在，marked_text 又不带这两类标记，故重建后丢失。SKILL.md 已声明此局限，
#      建议这类强调改用 markdown 标记（*斜体* / **加粗** / 上下标）。
# ===========================================================================

def test_inplace_color_underline_preserved_when_unchanged():
    from docx import Document
    from docx.shared import RGBColor
    import merge_manuscript as mm

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("plain start ")
    r = p.add_run("RED UNDERLINE mid")
    r.font.color.rgb = RGBColor(0xFF, 0, 0)
    r.font.underline = True
    p.add_run(" plain end")

    raw = "plain start RED UNDERLINE mid plain end"
    # polished == raw：文字未改动
    root, src = _make_project(
        doc, [{"idx": 0, "source_para_index": 0, "raw_text": raw, "polished_text": raw}]
    )
    out = os.path.join(root, "out.docx")
    res = mm.export_inplace(Path(src), Path(root), Path(out))

    assert res["paragraphs_skipped_color_underline"] == [0], \
        f"含颜色/下划线且文字未改的段应被跳过重建：{res!r}"

    d = Document(out)
    p3 = d.paragraphs[0]
    assert p3.text == raw, f"文字应原样保留：{p3.text!r}"
    has_red = any(
        (run.font.color is not None and run.font.color.type is not None
         and str(run.font.color.rgb) == "FF0000")
        for run in p3.runs
    )
    has_underline = any(bool(run.font.underline) for run in p3.runs)
    assert has_red, f"中间 run 红色(w:color)应保留，实际 runs：{[(r.text, r.font.color.rgb if r.font.color and r.font.color.type is not None else None) for r in p3.runs]!r}"
    assert has_underline, "中间 run 下划线(w:u)应保留"


def test_inplace_sup_sub_italic_not_degraded_when_edited():
    """硬底线：修颜色/下划线不得破坏 sup/sub/italic 重建（文字已改动的常规路径）。"""
    from docx import Document
    import merge_manuscript as mm
    import common

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("H")
    rsub = p.add_run("2")
    rsub.font.subscript = True
    p.add_run("O at 10")
    rsup = p.add_run("3")
    rsup.font.superscript = True
    ri = p.add_run(" italic")
    ri.italic = True

    raw = common.serialize_runs_to_marked_text(p)  # 带 sup/sub/italic 标记
    # 模拟润色改写但保留行内标记
    polished = "H<sub>2</sub>O measured at 10<sup>3</sup> *italic*"
    root, src = _make_project(
        doc, [{"idx": 0, "source_para_index": 0, "raw_text": raw, "polished_text": polished}]
    )
    out = os.path.join(root, "out.docx")
    res = mm.export_inplace(Path(src), Path(root), Path(out))
    assert res["rewritten_paragraphs"] == 1, f"sup/sub/italic 段应正常重建：{res!r}"

    d = Document(out)
    p3 = d.paragraphs[0]
    assert any(r.font.subscript for r in p3.runs), "下标退化"
    assert any(r.font.superscript for r in p3.runs), "上标退化"
    assert any(r.italic for r in p3.runs), "斜体退化"


def test_inplace_plain_paragraph_rewritten():
    """普通段（无颜色/下划线/sup/sub）文字改动时应正常重写。"""
    from docx import Document
    import merge_manuscript as mm

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("ordinary text to be polished")

    raw = "ordinary text to be polished"
    polished = "polished ordinary text now"
    root, src = _make_project(
        doc, [{"idx": 0, "source_para_index": 0, "raw_text": raw, "polished_text": polished}]
    )
    out = os.path.join(root, "out.docx")
    res = mm.export_inplace(Path(src), Path(root), Path(out))
    assert res["rewritten_paragraphs"] == 1

    d = Document(out)
    assert d.paragraphs[0].text == polished, \
        f"普通段应正常润色重写：{d.paragraphs[0].text!r}"


def test_inplace_embedded_image_paragraph_skipped():
    """回归：含内嵌图片段落仍被跳过（保图），不被颜色/下划线分支误吞。"""
    from docx import Document
    from docx.shared import Inches
    import merge_manuscript as mm
    import io
    import struct
    import zlib

    def _minimal_png():
        """生成一个合法的 1x1 白色 PNG 字节流。"""
        def chunk(typ, data):
            c = typ + data
            return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        sig = b"\x89PNG\r\n\x1a\n"
        ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        raw = b"\x00\xff\xff\xff"  # filter byte + 1 RGB pixel
        idat = chunk(b"IDAT", zlib.compress(raw))
        iend = chunk(b"IEND", b"")
        return sig + ihdr + idat + iend

    png = _minimal_png()

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(io.BytesIO(png), width=Inches(0.5))

    raw = ""
    polished = "should not be written into image paragraph"
    root, src = _make_project(
        doc, [{"idx": 0, "source_para_index": 0, "raw_text": raw, "polished_text": polished}]
    )
    out = os.path.join(root, "out.docx")
    res = mm.export_inplace(Path(src), Path(root), Path(out))
    assert res["paragraphs_skipped_images"] == [0], \
        f"含图段应被跳过保图：{res!r}"


# ===========================================================================
# 入口
# ===========================================================================

if __name__ == "__main__":
    test_inplace_color_underline_preserved_when_unchanged()
    print("OK 契约⑥a in-place 颜色/下划线（文字未改）无损保留")

    test_inplace_sup_sub_italic_not_degraded_when_edited()
    print("OK 契约⑥b in-place sup/sub/italic（文字已改）不退化")

    test_inplace_plain_paragraph_rewritten()
    print("OK 契约⑥c in-place 普通段正常润色重写")

    test_inplace_embedded_image_paragraph_skipped()
    print("OK 契约⑥d in-place 含图段跳过保图")

    print("ALL OK")
