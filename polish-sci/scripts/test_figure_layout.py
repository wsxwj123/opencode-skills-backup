#!/usr/bin/env python3
"""Two-layout figure-fidelity regression for polish-sci's in-place export.

polish-sci is a faithful-polish skill: the delivered docx must mirror the input
structure exactly. Its canonical docx path (`merge_manuscript.py --in-place-src`)
opens the ORIGINAL docx and only rewrites prose-paragraph runs, so figures and
captions are never disassembled — position fidelity is structural, not
anchor-reconstructed like revise-sci. This test locks that guarantee for both
figure layouts the skill must survive:

  Layout A: figure interleaved inside body prose.
  Layout B: figures in body, captions collected at the very end.

Per layout: atomize -> extract images -> synthesize polished units -> merge
in-place, then assert zero image loss, unchanged figure/caption paragraph
indices, and that prose was actually rewritten.

Run directly:  python3 test_figure_layout.py   (or via pytest)
"""
from __future__ import annotations

import json
import re
import struct
import subprocess
import sys
import tempfile
import zipfile
import zlib
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

SCRIPTS = Path(__file__).resolve().parent
ATOMIZE = SCRIPTS / "atomize_manuscript.py"
EXTRACT = SCRIPTS / "extract_docx_images.py"
MERGE = SCRIPTS / "merge_manuscript.py"


def _png(rgb: tuple[int, int, int]) -> bytes:
    def chunk(typ: bytes, data: bytes) -> bytes:
        body = typ + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + bytes(rgb))
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _run(cmd: list) -> subprocess.CompletedProcess:
    r = subprocess.run([sys.executable, *map(str, cmd)], cwd=str(SCRIPTS), capture_output=True, text=True)
    assert r.returncode == 0, f"{Path(cmd[0]).name} failed: {r.stderr or r.stdout}"
    return r


def _media_count(docx: Path) -> int:
    with zipfile.ZipFile(docx) as z:
        return len([n for n in z.namelist() if n.startswith("word/media/") and not n.endswith("/")])


def _drawing_para_indices(docx: Path) -> list[int]:
    doc = Document(str(docx))
    return [i for i, p in enumerate(doc.paragraphs) if p._p.findall(".//" + qn("w:drawing"))]


def _para_texts(docx: Path) -> list[str]:
    return [p.text for p in Document(str(docx)).paragraphs]


def _synthesize_polished(root: Path) -> None:
    """Stand in for the LLM polish step: tag prose text as [POLISHED]."""
    pol_dir = root / "polished"
    pol_dir.mkdir(exist_ok=True)
    for uf in (root / "units").glob("*.json"):
        u = json.loads(uf.read_text(encoding="utf-8"))
        u["polished_text"] = (u["raw_text"] + " [POLISHED]") if u.get("prose", True) else u["raw_text"]
        (pol_dir / uf.name).write_text(json.dumps(u, ensure_ascii=False), encoding="utf-8")


def _pipeline(root: Path, src: Path, out: Path) -> dict:
    _run([ATOMIZE, "--manuscript", src, "--project-root", root])
    _run([EXTRACT, "--manuscript", src, "--project-root", root])
    _synthesize_polished(root)
    r = _run([MERGE, "--project-root", root, "--in-place-src", src, "--docx", out])
    return json.loads(r.stdout.strip().splitlines()[-1])


def test_layout_a_figure_interleaved_in_body() -> None:
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        src, out = root / "in.docx", root / "out.docx"
        imgp = root / "a.png"
        imgp.write_bytes(_png((255, 0, 0)))

        doc = Document()
        doc.add_paragraph("A Minimal Study Title")
        doc.add_heading("Results", level=1)
        doc.add_paragraph("Body prose one before the figure appears here.")
        doc.add_picture(str(imgp), width=Inches(1))  # figure INSIDE the body
        doc.add_paragraph("Figure 1. Inline caption in results section.")
        doc.add_paragraph("Body prose two after the figure appears here.")
        doc.add_heading("Discussion", level=1)
        doc.add_paragraph("Discussion prose closing the manuscript.")
        doc.save(str(src))

        in_draw = _drawing_para_indices(src)
        report = _pipeline(root, src, out)
        out_txt = _para_texts(out)

        assert _media_count(out) == _media_count(src) == 1, "image lost or added"
        assert _drawing_para_indices(out) == in_draw, "figure paragraph moved"
        assert any("[POLISHED]" in t for t in out_txt), "prose not rewritten"
        assert any(t.startswith("Figure 1.") for t in out_txt), "inline caption lost"
        assert report["in_place"]["paragraphs_skipped_images"] == [], report


def test_layout_b_captions_collected_at_end() -> None:
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        src, out = root / "in.docx", root / "out.docx"
        i1, i2 = root / "b1.png", root / "b2.png"
        i1.write_bytes(_png((0, 255, 0)))
        i2.write_bytes(_png((0, 0, 255)))

        doc = Document()
        doc.add_paragraph("A Minimal Study Title")
        doc.add_heading("Results", level=1)
        doc.add_paragraph("Results prose referencing Figure 1 and Figure 2.")
        doc.add_picture(str(i1), width=Inches(1))
        doc.add_picture(str(i2), width=Inches(1))
        doc.add_heading("Discussion", level=1)
        doc.add_paragraph("Discussion prose closing the manuscript.")
        doc.add_heading("Figure Captions", level=1)          # captions COLLECTED AT END
        doc.add_paragraph("Figure 1. First caption collected at the end.")
        doc.add_paragraph("Figure 2. Second caption collected at the end.")
        doc.save(str(src))

        in_draw = _drawing_para_indices(src)
        report = _pipeline(root, src, out)
        out_txt = _para_texts(out)

        assert _media_count(out) == _media_count(src) == 2, "image lost or added"
        assert _drawing_para_indices(out) == in_draw, "figures moved"
        cap_idx = [i for i, t in enumerate(out_txt) if re.match(r"Figure \d+\.", t)]
        assert cap_idx == sorted(cap_idx) and len(cap_idx) == 2, f"caption order/pos broken: {cap_idx}"
        assert all(idx > max(_drawing_para_indices(out)) for idx in cap_idx), "captions no longer after figures"
        assert any("[POLISHED]" in t for t in out_txt), "prose not rewritten"
        assert report["in_place"]["paragraphs_skipped_images"] == [], report


if __name__ == "__main__":
    test_layout_a_figure_interleaved_in_body()
    test_layout_b_captions_collected_at_end()
    print("OK: zero image loss, figure+caption positions faithful, prose polished (both layouts)")
