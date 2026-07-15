#!/usr/bin/env python3
"""Regression for atomize_manuscript.py — the pipeline's first step.

Two failure modes that silently drop text if unguarded:
  ① tracked-changes fail-closed firewall: a docx carrying <w:ins>/<w:del>
     makes python-docx drop the <w:ins> text, so atomize must exit 1 with
     error=tracked_changes_present unless --allow-tracked-changes is given.
  ② section/prose classification: References/致谢/Author Contributions must be
     nonprose; numbered/named headings must be recognized (and typed) as
     headings, while numbered affiliation lines and reference entries must NOT.

Baseline = current code behavior. Runs standalone (`python3`) or via pytest.
"""
from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPT_DIR not in sys.path:
    sys.path.insert(0, _SCRIPT_DIR)

import atomize_manuscript as A

_SKIP = False
try:
    import docx  # noqa: F401
except Exception:
    _SKIP = True


# ── ② section / prose classification (pure functions, no docx needed) ─────────
def test_nonprose_sections() -> None:
    for h in ("References", "致谢", "Author Contributions", "参考文献", "Acknowledgements"):
        assert A.is_nonprose_section(h), f"{h!r} should be nonprose"


def test_prose_sections() -> None:
    for h in ("1. Introduction", "Results and Discussion", "方法", "2.1 Foo"):
        assert not A.is_nonprose_section(h), f"{h!r} should be prose"


def test_section_type_inference() -> None:
    assert A.infer_section_type("1. Introduction") == "intro"
    assert A.infer_section_type("Experimental Section") == "methods"
    assert A.infer_section_type("Results and Discussion") == "results"
    assert A.infer_section_type("3. Conclusions") == "discussion"
    assert A.infer_section_type("2.1 Foo") == "other"


def test_headings_recognized() -> None:
    for h in ("1. Introduction", "2.1 Foo", "Experimental Section",
              "Results and Discussion", "References"):
        assert A.looks_like_heading(h), f"{h!r} should be a heading"


def test_non_headings_rejected() -> None:
    # 编号机构/地址行不是标题
    assert not A.looks_like_heading("3 X University Key Laboratory of Bar")
    # 编号参考文献条目不是标题
    assert not A.looks_like_heading("1. Smith J, et al. Nature. 2020;12:3-5.")
    # 普通长句不是标题
    assert not A.looks_like_heading(
        "This paragraph is an ordinary sentence that clearly runs well past ten words."
    )


# ── ① tracked-changes fail-closed firewall (needs docx) ───────────────────────
def _make_clean_docx(path: str) -> None:
    d = docx.Document()
    d.add_heading("Introduction", level=1)
    d.add_paragraph("This study examined the effect of the treatment on cells.")
    d.save(path)


def _inject_tracked_change(src: str, dst: str, kind: str) -> None:
    """Copy src->dst then splice a <w:ins> or <w:del> paragraph into document.xml."""
    shutil.copy(src, dst)
    with zipfile.ZipFile(dst) as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}
    xml = data["word/document.xml"].decode("utf-8")
    if kind == "ins":
        frag = ('<w:p><w:ins w:id="99" w:author="a" w:date="2020-01-01T00:00:00Z">'
                '<w:r><w:t>inserted words</w:t></w:r></w:ins></w:p>')
    else:
        frag = ('<w:p><w:del w:id="98" w:author="a" w:date="2020-01-01T00:00:00Z">'
                '<w:r><w:delText>removed words</w:delText></w:r></w:del></w:p>')
    xml = xml.replace("</w:body>", frag + "</w:body>")
    data["word/document.xml"] = xml.encode("utf-8")
    with zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])


def _run_atomize(src: str, extra=None) -> tuple[int, dict]:
    root = tempfile.mkdtemp()
    cmd = [sys.executable, os.path.join(_SCRIPT_DIR, "atomize_manuscript.py"),
           "--manuscript", src, "--project-root", root] + (extra or [])
    p = subprocess.run(cmd, capture_output=True, text=True)
    try:
        out = json.loads(p.stdout.strip().splitlines()[-1])
    except Exception:
        out = {"_raw": p.stdout, "_err": p.stderr}
    return p.returncode, out


def test_count_tracked_changes_counts_ins_and_del() -> None:
    if _SKIP:
        print("SKIP: python-docx unavailable")
        return
    d = tempfile.mkdtemp()
    clean = os.path.join(d, "clean.docx")
    _make_clean_docx(clean)
    assert A.count_tracked_changes(Path(clean)) == {"ins": 0, "del": 0}
    ins = os.path.join(d, "ins.docx")
    _inject_tracked_change(clean, ins, "ins")
    tc = A.count_tracked_changes(Path(ins))
    assert tc["ins"] >= 1 and tc["del"] == 0, tc
    dele = os.path.join(d, "del.docx")
    _inject_tracked_change(clean, dele, "del")
    tc = A.count_tracked_changes(Path(dele))
    assert tc["del"] >= 1, tc


def test_tracked_changes_rejected_fail_closed() -> None:
    if _SKIP:
        print("SKIP: python-docx unavailable")
        return
    d = tempfile.mkdtemp()
    clean = os.path.join(d, "clean.docx")
    _make_clean_docx(clean)
    tracked = os.path.join(d, "tracked.docx")
    _inject_tracked_change(clean, tracked, "ins")
    code, out = _run_atomize(tracked)
    assert code == 1, f"expected exit 1, got {code}: {out}"
    assert out.get("ok") is False and out.get("error") == "tracked_changes_present", out


def test_allow_tracked_changes_overrides() -> None:
    if _SKIP:
        print("SKIP: python-docx unavailable")
        return
    d = tempfile.mkdtemp()
    clean = os.path.join(d, "clean.docx")
    _make_clean_docx(clean)
    tracked = os.path.join(d, "tracked.docx")
    _inject_tracked_change(clean, tracked, "ins")
    code, out = _run_atomize(tracked, ["--allow-tracked-changes"])
    assert code == 0 and out.get("ok") is True, out


def test_clean_docx_passes() -> None:
    if _SKIP:
        print("SKIP: python-docx unavailable")
        return
    d = tempfile.mkdtemp()
    clean = os.path.join(d, "clean.docx")
    _make_clean_docx(clean)
    code, out = _run_atomize(clean)
    assert code == 0 and out.get("ok") is True and out.get("unit_count", 0) >= 1, out


if __name__ == "__main__":
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            fn()
    print("OK: atomize tracked-changes firewall + section classification")
