#!/usr/bin/env python3
"""
试卷 Word 生成器 —— teacher-paper skill 自包含组件
输入一个 JSON 文件，输出两个 .docx：① 学生试卷 ② 参考答案及解析。
不依赖任何其它 skill，任何电脑装好 python-docx 即可运行。

用法：
    pip3 install python-docx
    python3 make_paper.py <content.json>

JSON 结构：
{
  "paper_path":  "九年级语文试卷.docx",          # 试卷输出文件名
  "answer_path": "九年级语文参考答案及解析.docx",  # 答案输出文件名
  "paper":   [ <block>, ... ],   # 学生试卷的有序内容块
  "answers": [ <block>, ... ],   # 参考答案及解析的有序内容块
  "meta":    { "page_number": true, ... }  # 可选；page_number=false 关闭页码
}

block 类型（type 字段）——试卷与答案通用：
  {"type":"title","text":"..."}                          大标题（居中,宋体三号16加粗）
  {"type":"subtitle","text":"（满分120分 时间120分钟）"}  副标题（居中,宋体小四12）
  {"type":"info"}                                         学校/班级/姓名/考号/座位号 填写行
  {"type":"sealing"}                                      密封线（印刷版横排虚线，仅试卷）
  {"type":"notice","items":["...","..."]}                注意事项编号列表
  {"type":"section","text":"一、积累运用（20分）"}        大题标题（宋体小四12加粗）
  {"type":"sub","text":"（一）积累"}                      小标题（宋体小四12加粗）
  {"type":"para","text":"...","indent":true}             正文段落,indent=首行缩进
  {"type":"material","label":"【材料一】","title":"老街","author":"",
   "paras":["...","..."],"source":"（选自《读者》2024年第6期）",
   "layout":"prose","font":"楷体"}                        阅读选文（楷体正文）
        - title/author 楷体居中；paras 楷体首行缩进2字（prose）
        - layout="verse" → 古诗词，paras 整体居中（不缩进）
        - source → 出处标注，右对齐宋体（非连/小说/古诗文/名著必标）
        - 选文正文默认楷体，可用 font 覆盖
  {"type":"table","rows":[["a","b"],["c","d"]],"header":true}  表格
  {"type":"question","num":"1","score":"（2分）","text":"..."}  题干
  {"type":"options","items":["A. ...","B. ...","C. ...","D. ..."]}  选项（每项一行）
  {"type":"blank_lines","count":3}                       答题横线 N 行
  {"type":"essay_grid","note":"...","cols":20,"rows":30}  作文方格纸(真格子,默认20×30=600格)
  {"type":"answer","num":"1","score":"（2分）","text":"..."}    答案条目
  {"type":"analysis","text":"..."}                       解析/评分说明
  {"type":"spacer"}                                      空行
  {"type":"pagebreak"}                                   分页
"""
import sys
import os
import re
import json

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[缺少依赖] 需要 python-docx。请运行：pip3 install python-docx")
    sys.exit(1)


# 字体常量（对标 2025 长沙中考真题原卷版排版）：
# 标题/大题/小题用宋体加粗，正文题干用宋体，阅读选文正文用楷体。
SONGTI = "宋体"
HEITI = "黑体"
KAITI = "楷体"

_CJK = re.compile(r"[一-鿿]")


def normalize_quotes(text):
    """规范中文引号，修正部分模型（如 deepseek）把书名/称谓写成单引号 '我的母亲'
    的问题。规则（保守，避免破坏正确的引号嵌套与英文缩写）：
    1. ASCII 直双引号 " → 交替成中文弯双引号 “ ”；
    2. 含中文、且成对（偶数个）的 ASCII 直单引号 ' → 交替成中文双引号（命题里
       单引号包中文基本都是该用双引号的误写）；
    3. 文本只有中文单引号 ‘ ’ 而无中文双引号时，单引号被误当主引号 → 升级为双引号。
    正确的双内嵌单（如 “他说‘好’”）因双引号已存在而被原样保留。"""
    if not text or not _CJK.search(text):
        return text  # 不含中文的纯英文/数字文本不动，避免误伤英文撇号
    if '"' in text:                                   # 1) ASCII 直双引号交替
        out, open_d = [], True
        for ch in text:
            if ch == '"':
                out.append("“" if open_d else "”")
                open_d = not open_d
            else:
                out.append(ch)
        text = "".join(out)
    sgl = text.count("'")                              # 2) 成对 ASCII 直单引号
    if sgl >= 2 and sgl % 2 == 0:
        out, open_d = [], True
        for ch in text:
            if ch == "'":
                out.append("“" if open_d else "”")
                open_d = not open_d
            else:
                out.append(ch)
        text = "".join(out)
    if ("‘" in text or "’" in text) and ("“" not in text and "”" not in text):
        text = text.replace("‘", "“").replace("’", "”")  # 3) 单引号误当主引号
    return text


def set_font(run, name=SONGTI, size=10.5, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    # 用官方 API 确保 rPr 存在，避免依赖赋值副作用隐式创建（否则 rPr 为 None 会崩）
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


# 上下标语法：_{...} 下标、^{...} 上标。只认带花括号形式——填空线 ____ 不含 { 不误伤，
# 单字符 Unicode 上下标（H₂O/x²）仍优先；本机制专为 Unicode 无法表达的多字符/中文下标
# （v_{max}、F_{合}）与显式上标提供真富文本渲染。
_SUBSUP_RE = re.compile(r"([_^])\{([^{}]*)\}")


def _split_subsup(text):
    """把含 _{...}/^{...} 的字符串拆成 [(片段, kind)]，kind ∈ normal/sub/sup。
    无标记则原样返回 [(text, 'normal')]，向后兼容。"""
    out, pos = [], 0
    for m in _SUBSUP_RE.finditer(text):
        if m.start() > pos:
            out.append((text[pos:m.start()], "normal"))
        out.append((m.group(2), "sub" if m.group(1) == "_" else "sup"))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], "normal"))
    if not out:
        out.append((text, "normal"))
    return out


def _emit_runs(p, line_text, name, size, bold):
    """把一行文本按 _{}/^{} 拆段加 run，下标/上标段置 vertAlign。"""
    for seg, kind in _split_subsup(line_text):
        if seg == "":
            continue
        r = p.add_run(seg)
        set_font(r, name=name, size=size, bold=bold)
        if kind == "sub":
            r.font.subscript = True
        elif kind == "sup":
            r.font.superscript = True


def add_para(doc, text="", align=None, indent_chars=0, size=10.5,
             name=SONGTI, bold=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if indent_chars:
        p.paragraph_format.first_line_indent = Pt(size * indent_chars)
    if text:
        # 先规范中文引号，再保留文本内换行：拆成多行，行间插软换行，避免 \n 被吞成空格；
        # 每行再按 _{}/^{} 拆出上下标 run。
        parts = normalize_quotes(str(text)).split("\n")
        _emit_runs(p, parts[0], name, size, bold)
        for seg in parts[1:]:
            br = p.add_run()
            set_font(br, name=name, size=size, bold=bold)
            br.add_break()
            _emit_runs(p, seg, name, size, bold)
    return p


def add_table(doc, rows, header=False):
    rows = [r for r in (rows or []) if r]  # 过滤 None/空行，避免 len(None) 崩溃
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            val = row[ci] if ci < len(row) else ""
            cells[ci].text = ""
            para = cells[ci].paragraphs[0]
            run = para.add_run(str(val))
            set_font(run, size=10.5, bold=(header and ri == 0))


def add_essay_grid(doc, cols=20, rows=30):
    """生成作文方格纸：cols×rows 的正方格表格，每格供写一字。
    对异常入参做夹紧：列 8-30、行 1-60，避免 0/负数报错或超大表卡死。"""
    try:
        cols = int(cols)
        rows = int(rows)
    except (TypeError, ValueError):
        cols, rows = 20, 30
    cols = max(8, min(cols, 30))
    rows = max(1, min(rows, 60))
    t = doc.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    t.autofit = False
    cell_w = Cm(0.75)  # 每格边长，20列约15cm，适配A4正文宽
    for row in t.rows:
        row.height = cell_w
        for cell in row.cells:
            cell.width = cell_w
            # 收紧单元格内边距，让格子接近正方形
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0


def render(doc, blocks):
    for b in blocks:
        t = b.get("type")
        if t == "title":
            # 真题：大标题 宋体三号(16)加粗居中
            add_para(doc, b["text"], align="center", size=16, name=SONGTI,
                     bold=True, space_after=6)
        elif t == "subtitle":
            add_para(doc, b["text"], align="center", size=12, space_after=8)
        elif t == "info":
            add_para(doc,
                     "学校__________ 班级__________ 姓名__________ "
                     "考号__________ 座位号______",
                     align="center", size=10.5, space_after=10)
        elif t == "sealing":
            _add_sealing_line(doc)
        elif t == "notice":
            # 真题：注意事项 宋体小四(12)加粗
            add_para(doc, "注意事项：", size=12, name=SONGTI, bold=True,
                     space_after=2)
            for i, it in enumerate(b.get("items", []), 1):
                add_para(doc, f"{i}.{it}", size=10.5, space_after=1)
        elif t == "section":
            # 真题：大题标题 宋体小四(12)加粗
            add_para(doc, b["text"], size=12, name=SONGTI, bold=True,
                     space_after=4)
        elif t == "sub":
            # 真题：小标题 宋体小四(12)加粗
            add_para(doc, b["text"], size=12, name=SONGTI, bold=True,
                     space_after=3)
        elif t == "para":
            add_para(doc, b["text"], indent_chars=2 if b.get("indent") else 0)
        elif t == "material":
            # 真题排版：选文标题/作者楷体居中；正文楷体首行缩进2字；
            # 古诗词(layout=verse)整体居中；出处右对齐宋体。
            layout = b.get("layout", "prose")     # prose 散文/记叙/说明 | verse 古诗词
            pfont = b.get("font", KAITI)          # 选文正文字体，默认楷体
            label = b.get("label", "")
            if label:
                add_para(doc, label, size=10.5, name=SONGTI, space_after=2)
            if b.get("title"):
                add_para(doc, b["title"], align="center", size=10.5, name=pfont,
                         space_after=1)
            if b.get("author"):
                add_para(doc, b["author"], align="center", size=10.5, name=pfont,
                         space_after=2)
            for para in b.get("paras", []):
                if layout == "verse":
                    add_para(doc, para, align="center", size=10.5, name=pfont)
                else:
                    add_para(doc, para, indent_chars=2, size=10.5, name=pfont)
            if b.get("source"):
                # 出处标注：右对齐宋体（如"（材料均改编自《科学之友》）"）
                add_para(doc, b["source"], align="right", size=10.5,
                         name=SONGTI, space_after=4)
        elif t == "table":
            add_table(doc, b.get("rows", []), header=b.get("header", False))
            add_para(doc, "", space_after=2)
        elif t == "question":
            num = b.get("num", "")
            score = b.get("score", "")
            head = f"{num}. " if num else ""
            text = f"{head}{b.get('text','')}{score}"
            add_para(doc, text, size=10.5, space_after=3)
        elif t == "options":
            for opt in b.get("items", []):
                add_para(doc, opt, indent_chars=2, size=10.5, space_after=1)
        elif t == "blank_lines":
            for _ in range(b.get("count", 3)):
                add_para(doc, "_" * 46, size=10.5, space_after=6)
        elif t == "essay_grid":
            note = b.get("note", "请在作文格内作答。")
            add_para(doc, note, size=10.5, space_after=4)
            # 真作文方格纸：cols 列 × rows 行的网格表格（每格写一字）
            cols = b.get("cols", 20)
            rows = b.get("rows", 30)
            add_essay_grid(doc, cols=cols, rows=rows)
        elif t == "answer":
            num = b.get("num", "")
            score = b.get("score", "")
            head = f"{num}. " if num else ""
            add_para(doc, f"{head}{score}{b.get('text','')}", size=10.5,
                     space_after=3)
        elif t == "analysis":
            add_para(doc, b["text"], size=10.5, space_after=4)
        elif t == "figure":
            # figure 块的 alt/src/caption/width_cm 直接写在块顶层（见文档与 SKILL.md）；
            # 兼容历史上可能的 spec 嵌套写法。
            _render_figure(doc, b.get("spec", b))
        elif t == "spacer":
            add_para(doc, "", space_after=4)
        elif t == "pagebreak":
            doc.add_page_break()


# 配图输出目录（main 按输出路径设定；理科配图渲染器 make_figure.py 后续接入）
FIG_DIR = None


def _render_figure(doc, spec):
    """插入配图，三级优先：① spec.src 指向已有图片文件 → 直接插入；
    ② 否则调 make_figure.render_figure 按 spec.kind 自动渲染
    （function/geometry/number_line/bar/line/pie/scatter/vector/climate/pyramid/svg，
    需本机有 matplotlib，svg 另需 PyMuPDF）→ 成功则插入；
    ③ 都不行 → 降级 "［图：alt］" 文字占位，绝不阻断出卷。
    v3.24.0 R1：spec 必须是 dict；若上游误传字符串（如把 "<svg>" 直接当 spec），
    转成占位而非在 spec.get(...) 处崩溃。"""
    if not isinstance(spec, dict):
        # 上游写法错误：figure 块的 spec 不是对象（常见把 svg 字符串直接当 spec）
        add_para(doc, "［图：spec 格式错误，请按 figure 块规范填写］",
                 align="center", size=9, space_after=4)
        return
    src = spec.get("src")
    width = spec.get("width_cm", 6.0)
    png = None
    if src and os.path.exists(src):
        png = src
    else:
        try:
            from make_figure import render_figure  # 下一轮才有；现在通常 ImportError
            out_dir = FIG_DIR or os.getcwd()
            os.makedirs(out_dir, exist_ok=True)
            png = render_figure(spec, out_dir)
        except Exception:
            png = None
    if png and os.path.exists(png):
        try:
            doc.add_picture(png, width=Cm(width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if spec.get("caption"):
                add_para(doc, spec["caption"], align="center", size=9,
                         space_after=4)
            return
        except Exception:
            pass
    # 降级：占位 + alt 文字（无障碍/可手动补图）
    alt = spec.get("alt", "此处应有配图，请手动补充")
    add_para(doc, f"［图：{alt}］", align="center", size=9, space_after=4)


def _add_sealing_line(doc):
    """密封线（横排印刷版）：填写区下方一条带提示的虚线，提示密封线内不要答题。
    真考卷的竖排左边距密封线在 Word 里实现脆弱，这里用可靠的横排印刷版替代。"""
    add_para(doc,
             "…………………………密封…………线…………内…………不…………要…………"
             "答…………题…………………………",
             align="center", size=9, name=SONGTI, space_after=2)


def _add_page_footer(doc):
    """给每个 section 的页脚加居中页码「第 X 页　共 Y 页」（用 Word 域，打开自动更新）。"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in list(p.runs):                       # 清空已有内容
            r._element.getparent().remove(r._element)

        def _field(instr):
            run = p.add_run()
            set_font(run, size=9)
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            instr_el = OxmlElement("w:instrText")
            instr_el.set(qn("xml:space"), "preserve")
            instr_el.text = instr
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            run._r.append(begin)
            run._r.append(instr_el)
            run._r.append(end)

        r1 = p.add_run("第 ")
        set_font(r1, size=9)
        _field("PAGE")
        r2 = p.add_run(" 页　共 ")
        set_font(r2, size=9)
        _field("NUMPAGES")
        r3 = p.add_run(" 页")
        set_font(r3, size=9)


def new_doc():
    doc = Document()
    sec = doc.sections[0]
    # 页边距对标 2025 长沙中考真题原卷版（上1.61/下2.54/左右1.91 cm）
    sec.top_margin = Cm(1.61)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(1.91)
    sec.right_margin = Cm(1.91)
    # 默认正文字体
    style = doc.styles["Normal"]
    style.font.name = SONGTI
    style.font.size = Pt(10.5)
    # get_or_add_rPr 避免 rPr 为 None 时崩溃（不同 python-docx 版本行为不一）
    style.element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), SONGTI)
    return doc


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    try:
        with open(sys.argv[1], encoding="utf-8") as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        print(f"[错误] content.json 不是合法 JSON：{e}")
        sys.exit(1)
    if not isinstance(data, dict):
        print("[错误] content.json 顶层必须是对象 {paper_path, answer_path, paper, answers}")
        sys.exit(1)

    def _ensure_dir(path):
        d = os.path.dirname(os.path.abspath(path))
        if d:
            os.makedirs(d, exist_ok=True)

    # 配图缓存目录：输出目录下的 figures/
    global FIG_DIR
    FIG_DIR = os.path.join(
        os.path.dirname(os.path.abspath(data.get("paper_path", "."))), "figures")

    meta = data.get("meta", {}) if isinstance(data.get("meta"), dict) else {}
    want_page_num = meta.get("page_number", True)  # 默认加页码

    # Bug Batch6-L1：拒绝把 listening_transcript material 渲染到试卷文档
    # 听力稿应放 answer 字段供老师朗读，印在试卷上等于把听力题退化为阅读题
    paper_blocks = data.get("paper", [])
    _bad = [b for b in paper_blocks if isinstance(b, dict)
            and b.get("type") == "material"
            and b.get("layout") == "listening_transcript"]
    if _bad:
        print("[错误] paper 字段包含听力录音稿 material 块（layout=listening_transcript）：",
              file=sys.stderr)
        for b in _bad:
            print(f"   - title={b.get('title') or b.get('label') or '(无标题)'}",
                  file=sys.stderr)
        print("\n  录音稿必须写到原子文件的 answer 字段，由老师按答案文档朗读。", file=sys.stderr)
        print("  把听力稿印在试卷上会让听力题退化为阅读题，整张卷废掉。", file=sys.stderr)
        print("  修复：把听力 material 块从 paper:[...] 移到 answer:[...]，重跑 build。",
              file=sys.stderr)
        sys.exit(2)
    paper = new_doc()
    render(paper, paper_blocks)
    if want_page_num:
        _add_page_footer(paper)
    paper_path = data.get("paper_path", "试卷.docx")
    _ensure_dir(paper_path)
    paper.save(paper_path)

    ans = new_doc()
    render(ans, data.get("answers", []))
    if want_page_num:
        _add_page_footer(ans)
    ans_path = data.get("answer_path", "参考答案及解析.docx")
    _ensure_dir(ans_path)
    ans.save(ans_path)

    print(f"[完成] 试卷已生成：{paper_path}")
    print(f"[完成] 参考答案及解析已生成：{ans_path}")


if __name__ == "__main__":
    main()
