import json
import subprocess
import tempfile
import unittest
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE


ROOT = str(Path(__file__).resolve().parent.parent)
COUNT_SCRIPT = f"{ROOT}/scripts/count_words.py"
MD_SCRIPT = f"{ROOT}/scripts/markdown_to_docx.py"
ATOMIC_MD_SCRIPT = f"{ROOT}/scripts/atomic_md_workflow.py"
MERGE_SCRIPT = f"{ROOT}/scripts/merge_chapters.py"
MERGE_COMPAT_SCRIPT = f"{ROOT}/scripts/merge_documents.py"
QUALITY_SCRIPT = f"{ROOT}/scripts/check_quality.py"
STATE_SCRIPT = f"{ROOT}/scripts/state_manager.py"


def run_py(script, args, cwd=None):
    return subprocess.run(
        ["python3", script] + args,
        cwd=cwd,
        capture_output=True,
        text=True,
    )


def extract_json_from_stdout(stdout_text):
    start = stdout_text.find("{")
    end = stdout_text.rfind("}")
    if start != -1 and end != -1 and end > start:
        candidate = stdout_text[start : end + 1]
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            pass
    lines = [line for line in stdout_text.splitlines() if line.strip()]
    for i in range(len(lines)):
        for j in range(len(lines), i, -1):
            candidate = "\n".join(lines[i:j])
            try:
                return json.loads(candidate)
            except json.JSONDecodeError:
                continue
    raise ValueError("No JSON payload found in stdout")


class Sci2DocScriptsTests(unittest.TestCase):
    def test_atomic_md_validate_detects_numbering_gap(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            cdir = root / "atomic_md" / "第2章"
            cdir.mkdir(parents=True, exist_ok=True)
            (cdir / "2.1_背景.md").write_text("## 2.1 背景\n### 2.1.1 引言\n内容", encoding="utf-8")
            (cdir / "2.3_方法.md").write_text("## 2.3 方法\n### 2.3.1 引言\n内容", encoding="utf-8")

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                ["--project-root", str(root), "validate", "--chapter", "2"],
            )
            self.assertNotEqual(proc.returncode, 0)
            payload = json.loads(proc.stdout)
            self.assertFalse(payload.get("ok", True))
            self.assertGreater(len(payload.get("errors", [])), 0)

    def test_atomic_md_merge_outputs_chapter_md(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            cdir = root / "atomic_md" / "第2章"
            cdir.mkdir(parents=True, exist_ok=True)
            (cdir / "2.1_背景.md").write_text("## 2.1 背景\n### 2.1.1 引言\nA", encoding="utf-8")
            (cdir / "2.2_方法.md").write_text("## 2.2 方法\n### 2.2.1 引言\nB", encoding="utf-8")

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                ["--project-root", str(root), "merge", "--chapter", "2"],
            )
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertTrue(payload.get("ok"))
            out = Path(payload["output_md"])
            self.assertTrue(out.exists())
            merged_text = out.read_text(encoding="utf-8")
            self.assertIn("## 2.1 背景", merged_text)
            self.assertIn("## 2.2 方法", merged_text)

    def test_atomic_md_validate_research_structure_detects_missing_sections(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            cdir = root / "atomic_md" / "第2章"
            cdir.mkdir(parents=True, exist_ok=True)
            (cdir / "2.1_引言.md").write_text("## 2.1 引言\n内容", encoding="utf-8")
            (cdir / "2.2_材料与方法.md").write_text("## 2.2 材料与方法\n内容", encoding="utf-8")

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                [
                    "--project-root",
                    str(root),
                    "validate",
                    "--chapter",
                    "2",
                    "--enforce-research-structure",
                ],
            )
            self.assertNotEqual(proc.returncode, 0)
            payload = json.loads(proc.stdout)
            self.assertFalse(payload.get("ok", True))
            self.assertIn("研究章结构缺项", json.dumps(payload.get("errors", []), ensure_ascii=False))

    def test_atomic_md_validate_research_structure_passes_when_complete(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            cdir = root / "atomic_md" / "第2章"
            cdir.mkdir(parents=True, exist_ok=True)
            (cdir / "2.1_引言.md").write_text("## 2.1 引言\n内容", encoding="utf-8")
            (cdir / "2.2_材料与方法.md").write_text("## 2.2 材料与方法\n内容", encoding="utf-8")
            (cdir / "2.3_结果讨论.md").write_text("## 2.3 结果讨论\n内容", encoding="utf-8")
            (cdir / "2.4_实验结论.md").write_text("## 2.4 实验结论\n内容", encoding="utf-8")
            (cdir / "2.5_小结.md").write_text("## 2.5 小结\n内容", encoding="utf-8")

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                [
                    "--project-root",
                    str(root),
                    "validate",
                    "--chapter",
                    "2",
                    "--enforce-research-structure",
                ],
            )
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertTrue(payload.get("ok"))

    def test_atomic_md_merge_full_outputs_whole_md(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            md_dir = root / "02_分章节文档_md"
            md_dir.mkdir(parents=True, exist_ok=True)
            (md_dir / "第1章_绪论_合并.md").write_text("# 第一章\nA", encoding="utf-8")
            (md_dir / "第2章_研究_合并.md").write_text("# 第二章\nB", encoding="utf-8")

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                ["--project-root", str(root), "merge-full"],
            )
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertTrue(payload.get("ok"))
            out = Path(payload["output_md"])
            self.assertTrue(out.exists())
            merged_text = out.read_text(encoding="utf-8")
            self.assertIn("第一章", merged_text)
            self.assertIn("第二章", merged_text)

    def test_atomic_md_self_check_returns_word_and_quality(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            docx_path = root / "第2章_测试.docx"
            d = Document()
            d.add_heading("第二章 方法", level=1)
            d.add_paragraph("这是章节正文内容。")
            d.save(str(docx_path))

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                ["--project-root", str(root), "self-check", "--docx", str(docx_path)],
            )
            self.assertIn(proc.returncode, (0, 1), msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertIn("word_count", payload)
            self.assertIn("quality_check", payload)

    def test_atomic_md_validate_experiment_map_detects_missing_mapping(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            cdir = root / "atomic_md" / "第2章"
            cdir.mkdir(parents=True, exist_ok=True)
            (cdir / "2.1_材料与方法.md").write_text(
                "## 2.1 材料与方法\n[实验] EXP-2-1\n方法内容",
                encoding="utf-8",
            )
            (cdir / "2.2_结果与讨论.md").write_text(
                "## 2.2 结果与讨论\n结果内容（缺少对应实验和图表标记）",
                encoding="utf-8",
            )

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                ["--project-root", str(root), "validate-experiment-map", "--chapter", "2"],
            )
            self.assertNotEqual(proc.returncode, 0)
            payload = json.loads(proc.stdout)
            self.assertFalse(payload.get("ok", True))
            self.assertGreater(len(payload.get("errors", [])), 0)

    def test_atomic_md_validate_experiment_map_passes_when_complete(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            cdir = root / "atomic_md" / "第2章"
            cdir.mkdir(parents=True, exist_ok=True)
            (cdir / "2.1_材料与方法.md").write_text(
                "## 2.1 材料与方法\n[实验] EXP-2-1\n方法内容",
                encoding="utf-8",
            )
            (cdir / "2.2_结果与讨论.md").write_text(
                "## 2.2 结果与讨论\n[对应实验] EXP-2-1\n[图] 图2-1\n结果讨论内容",
                encoding="utf-8",
            )

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                ["--project-root", str(root), "validate-experiment-map", "--chapter", "2"],
            )
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertTrue(payload.get("ok"))

    def test_atomic_md_validate_experiment_map_counts_figure_only_in_results_discussion(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            cdir = root / "atomic_md" / "第2章"
            cdir.mkdir(parents=True, exist_ok=True)
            (cdir / "2.1_材料与方法.md").write_text(
                "## 2.1 材料与方法\n[实验] EXP-2-1\n[图] 图2-1\n方法内容",
                encoding="utf-8",
            )
            (cdir / "2.2_结果与讨论.md").write_text(
                "## 2.2 结果与讨论\n[对应实验] EXP-2-1\n结果讨论内容（故意无图表标记）",
                encoding="utf-8",
            )

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                ["--project-root", str(root), "validate-experiment-map", "--chapter", "2"],
            )
            self.assertNotEqual(proc.returncode, 0)
            payload = json.loads(proc.stdout)
            self.assertFalse(payload.get("ok", True))
            self.assertIn("缺少 [图] 或 [表]", json.dumps(payload.get("errors", []), ensure_ascii=False))

    def test_atomic_md_validate_experiment_map_detects_result_discussion_title_variant(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            cdir = root / "atomic_md" / "第2章"
            cdir.mkdir(parents=True, exist_ok=True)
            (cdir / "2.1_材料与方法.md").write_text(
                "## 2.1 材料与方法\n[实验] EXP-2-1\n方法内容",
                encoding="utf-8",
            )
            (cdir / "2.2_实验A_结果讨论.md").write_text(
                "## 2.2 实验A 结果讨论\n结果内容（缺少[对应实验]）",
                encoding="utf-8",
            )

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                ["--project-root", str(root), "validate-experiment-map", "--chapter", "2"],
            )
            self.assertNotEqual(proc.returncode, 0)
            payload = json.loads(proc.stdout)
            self.assertFalse(payload.get("ok", True))
            self.assertIn("缺少 [对应实验] 标记", json.dumps(payload.get("errors", []), ensure_ascii=False))

    def test_atomic_md_section_snapshot_creates_snapshot(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            init_proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "init", "--title", "原子快照测试"],
            )
            self.assertEqual(init_proc.returncode, 0, msg=init_proc.stdout + init_proc.stderr)

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                [
                    "--project-root",
                    str(root),
                    "section-snapshot",
                    "--chapter",
                    "2",
                    "--section",
                    "2.1",
                ],
            )
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertTrue(payload.get("ok"))
            self.assertTrue(Path(payload["snapshot_dir"]).exists())

    def test_state_init_command_creates_core_files(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td) / "thesis_project"
            proc = run_py(
                STATE_SCRIPT,
                [
                    "--project-root",
                    str(root),
                    "init",
                    "--title",
                    "测试题目",
                    "--author",
                    "张三",
                ],
            )
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = extract_json_from_stdout(proc.stdout)
            self.assertTrue(payload.get("ok"))
            for rel in [
                "project_state.json",
                "thesis_profile.json",
                "context_memory.md",
                "history_log.json",
                "chapter_index.json",
                "literature_index.json",
                "figures_index.json",
            ]:
                self.assertTrue((root / rel).exists(), msg=f"Missing {rel}")

    def test_count_words_handles_abstract_then_body(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "test.md"
            p.write_text(
                "# 中文摘要\n\n这里是摘要内容\n\n# 第一章 绪论\n\n这里是正文内容\n",
                encoding="utf-8",
            )

            proc = run_py(COUNT_SCRIPT, [str(p), "--output", "json"])
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertGreater(payload["body_text"]["chinese_chars"], 0)
            self.assertEqual(payload["schema_version"], "3.0")
            self.assertEqual(payload["targets"]["body_target"], 80000)

    def test_count_words_recognizes_review_heading_with_chapter_prefix(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "review.md"
            p.write_text(
                "# 第二章 文献综述\n\n这是综述内容\n\n# 第三章 研究方法\n\n这是正文内容\n",
                encoding="utf-8",
            )

            proc = run_py(COUNT_SCRIPT, [str(p), "--output", "json"])
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertGreater(payload["review"]["chinese_chars"], 0)
            self.assertGreater(payload["body_text"]["chinese_chars"], 0)

    def test_count_words_recognizes_chinese_heading_style_name(self):
        """md 模式下中文标题也能被识别为综述章节。"""
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "review_cn_style.md"
            p.write_text(
                "# 第二章 文献综述\n\n这是综述内容\n\n# 第三章 研究方法\n\n这是正文内容\n",
                encoding="utf-8",
            )

            proc = run_py(COUNT_SCRIPT, [str(p), "--output", "json"])
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertGreater(payload["review"]["chinese_chars"], 0)
            self.assertGreater(payload["body_text"]["chinese_chars"], 0)

    def test_count_words_auto_finds_profile_from_docx_parent_chain(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            (root / "thesis_profile.json").write_text(
                json.dumps({"targets": {"body_target_chars": 86000}}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            out_dir = root / "atomic_md"
            out_dir.mkdir(parents=True, exist_ok=True)
            p = out_dir / "全文.md"
            p.write_text("# 第一章 绪论\n\n这里是正文\n", encoding="utf-8")

            proc = run_py(COUNT_SCRIPT, [str(p), "--output", "json"])
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertEqual(payload.get("targets", {}).get("body_target"), 86000)

    def test_markdown_headings_use_heading_styles(self):
        with tempfile.TemporaryDirectory() as td:
            md = Path(td) / "sample.md"
            out = Path(td) / "sample.docx"
            md.write_text(
                "# 第一章 绪论\n## 1.1 背景\n### 1.1.1 小节\n正文内容\n",
                encoding="utf-8",
            )
            proc = run_py(MD_SCRIPT, [str(md), "-o", str(out)])
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)

            doc = Document(str(out))
            self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")
            self.assertEqual(doc.paragraphs[1].style.name, "Heading 2")
            self.assertEqual(doc.paragraphs[2].style.name, "Heading 3")

    def test_merge_chapters_accepts_front_matter_args(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            input_dir = root / "chapters"
            input_dir.mkdir(parents=True, exist_ok=True)
            output = root / "merged.docx"

            for name, text in [
                ("cover.docx", "封面"),
                ("zh_abstract.docx", "中文摘要"),
                ("en_abstract.docx", "英文摘要"),
                ("第1章_绪论.docx", "第一章内容"),
                ("第2章_方法.docx", "第二章内容"),
            ]:
                d = Document()
                d.add_paragraph(text)
                d.save(str(input_dir / name))

            proc = run_py(
                MERGE_SCRIPT,
                [
                    "--input-dir",
                    str(input_dir),
                    "--output",
                    str(output),
                    "--cover",
                    "cover.docx",
                    "--abstract",
                    "zh_abstract.docx",
                    "--abstract-en",
                    "en_abstract.docx",
                ],
            )
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            self.assertTrue(output.exists())

    def test_merge_documents_compat_wrapper_uses_current_layout(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            chapter_dir = root / "02_分章节文档"
            chapter_dir.mkdir(parents=True, exist_ok=True)
            output = root / "03_合并文档" / "完整博士论文.docx"

            for name, text in [
                ("第1章_绪论.docx", "第一章内容"),
                ("第2章_方法.docx", "第二章内容"),
            ]:
                d = Document()
                d.add_paragraph(text)
                d.save(str(chapter_dir / name))

            proc = run_py(
                MERGE_COMPAT_SCRIPT,
                [
                    "--project-root",
                    str(root),
                ],
            )
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            self.assertTrue(output.exists())

    def test_quality_check_no_false_spacing_warning_for_20pt(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "fmt.docx"
            d = Document()
            para = d.add_paragraph("正文测试段落")
            para.style = "Normal"
            para.paragraph_format.line_spacing = Pt(20)
            para.paragraph_format.first_line_indent = Cm(0.74)
            d.save(str(p))

            proc = run_py(QUALITY_SCRIPT, [str(p), "--output", "json"])
            self.assertNotEqual(proc.stdout.strip(), "")
            data = extract_json_from_stdout(proc.stdout)
            msg_text = json.dumps(data.get("issues", []), ensure_ascii=False)
            self.assertNotIn("行距不符合要求", msg_text)

    def test_check_quality_json_output_is_pure_json(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "q.docx"
            d = Document()
            d.add_heading("第一章 绪论", level=1)
            d.add_paragraph("这是正文段落。")
            d.save(str(p))

            proc = run_py(QUALITY_SCRIPT, [str(p), "--output", "json"])
            self.assertIn(proc.returncode, (0, 1), msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertIn("success", payload)
            self.assertIn("statistics", payload)
            self.assertEqual(payload.get("targets", {}).get("body_target_chars"), 80000)

    def test_check_quality_recognizes_review_heading_with_chapter_prefix(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "q_review.docx"
            d = Document()
            d.add_heading("第二章 文献综述", level=1)
            d.add_paragraph("这是综述段落内容")
            d.add_heading("第三章 研究方法", level=1)
            d.add_paragraph("这是正文段落内容")
            d.save(str(p))

            proc = run_py(QUALITY_SCRIPT, [str(p), "--output", "json"])
            self.assertIn(proc.returncode, (0, 1), msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertGreater(payload["statistics"].get("review_words", 0), 0)

    def test_check_quality_recognizes_chinese_heading_style_name(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "q_review_cn_style.docx"
            d = Document()
            h1 = d.styles.add_style("标题 1", WD_STYLE_TYPE.PARAGRAPH)
            p1 = d.add_paragraph("第二章 文献综述")
            p1.style = h1
            d.add_paragraph("这是综述段落内容")
            p2 = d.add_paragraph("第三章 研究方法")
            p2.style = h1
            d.add_paragraph("这是正文段落内容")
            d.save(str(p))

            proc = run_py(QUALITY_SCRIPT, [str(p), "--output", "json"])
            self.assertIn(proc.returncode, (0, 1), msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertGreater(payload["statistics"].get("review_words", 0), 0)

    def test_check_quality_auto_finds_profile_from_docx_parent_chain(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            (root / "thesis_profile.json").write_text(
                json.dumps({"targets": {"body_target_chars": 87000}}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            out_dir = root / "03_合并文档"
            out_dir.mkdir(parents=True, exist_ok=True)
            p = out_dir / "完整博士论文.docx"
            d = Document()
            d.add_heading("第一章 绪论", level=1)
            d.add_paragraph("这是正文段落内容")
            d.save(str(p))

            proc = run_py(QUALITY_SCRIPT, [str(p), "--output", "json"])
            self.assertIn(proc.returncode, (0, 1), msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertEqual(payload.get("targets", {}).get("body_target_chars"), 87000)

    def test_check_quality_references_min_reads_from_profile(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            (root / "thesis_profile.json").write_text(
                json.dumps({"targets": {"references_min_count": 120}}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            out_dir = root / "03_合并文档"
            out_dir.mkdir(parents=True, exist_ok=True)
            p = out_dir / "完整博士论文.docx"
            d = Document()
            d.add_heading("参考文献", level=1)
            d.add_paragraph("[1] ref one")
            d.save(str(p))

            proc = run_py(QUALITY_SCRIPT, [str(p), "--output", "json"])
            self.assertIn(proc.returncode, (0, 1), msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertEqual(payload.get("targets", {}).get("references_min_count"), 120)
            issues_text = json.dumps(payload.get("issues", []), ensure_ascii=False)
            self.assertIn("/ 120 篇", issues_text)

    def test_check_quality_reference_position_detects_chapter_after_references(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "q_ref_pos.docx"
            d = Document()
            d.add_heading("第一章 绪论", level=1)
            d.add_paragraph("正文")
            d.add_heading("参考文献", level=1)
            d.add_paragraph("[1] ref one")
            d.add_heading("第二章 方法", level=1)
            d.add_paragraph("不应在参考文献后出现")
            d.save(str(p))

            proc = run_py(QUALITY_SCRIPT, [str(p), "--output", "json"])
            self.assertIn(proc.returncode, (0, 1), msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            issues_text = json.dumps(payload.get("issues", []), ensure_ascii=False)
            self.assertIn("参考文献后出现章节标题", issues_text)

    def test_check_quality_reference_position_detects_multiple_reference_sections(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "q_ref_multi.docx"
            d = Document()
            d.add_heading("第一章 绪论", level=1)
            d.add_paragraph("正文")
            d.add_heading("参考文献", level=1)
            d.add_paragraph("[1] ref one")
            d.add_heading("参考文献", level=1)
            d.add_paragraph("[2] ref two")
            d.save(str(p))

            proc = run_py(QUALITY_SCRIPT, [str(p), "--output", "json"])
            self.assertIn(proc.returncode, (0, 1), msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            issues_text = json.dumps(payload.get("issues", []), ensure_ascii=False)
            self.assertIn("多个参考文献标题", issues_text)

    def test_check_quality_full_structure_gate_detects_insufficient_chapters(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "q_structure.docx"
            d = Document()
            d.add_heading("第一章 绪论", level=1)
            d.add_paragraph("正文")
            d.add_heading("第二章 研究结果", level=1)
            d.add_paragraph("正文")
            d.add_heading("第三章 结论", level=1)
            d.add_paragraph("正文")
            d.save(str(p))

            proc = run_py(
                QUALITY_SCRIPT,
                [str(p), "--output", "json", "--enforce-full-structure", "--min-chapters", "5"],
            )
            self.assertIn(proc.returncode, (0, 1), msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertTrue(payload.get("targets", {}).get("enforce_full_structure"))
            self.assertEqual(payload.get("targets", {}).get("min_chapters"), 5)
            issues_text = json.dumps(payload.get("issues", []), ensure_ascii=False)
            self.assertIn("章节数不足", issues_text)

    def test_check_quality_auto_enforces_full_structure_for_full_doc_name(self):
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "完整博士论文.docx"
            d = Document()
            d.add_heading("第一章 绪论", level=1)
            d.add_paragraph("正文")
            d.add_heading("第二章 方法", level=1)
            d.add_paragraph("正文")
            d.add_heading("第三章 结论", level=1)
            d.add_paragraph("正文")
            d.save(str(p))

            proc = run_py(QUALITY_SCRIPT, [str(p), "--output", "json"])
            self.assertIn(proc.returncode, (0, 1), msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertTrue(payload.get("targets", {}).get("enforce_full_structure"))
            issues_text = json.dumps(payload.get("issues", []), ensure_ascii=False)
            self.assertIn("章节数不足", issues_text)

    def test_count_words_help_works(self):
        proc = run_py(COUNT_SCRIPT, ["--help"])
        self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
        self.assertIn("usage:", proc.stdout.lower())

    def test_count_words_missing_file_json_output_is_machine_readable(self):
        missing = "/tmp/__sci2doc_missing_count_words__.docx"
        proc = run_py(COUNT_SCRIPT, [missing, "--output", "json"])
        self.assertNotEqual(proc.returncode, 0)
        payload = json.loads(proc.stdout)
        self.assertFalse(payload.get("success", True))
        self.assertEqual(payload.get("error"), "file_not_found")

    def test_check_quality_missing_file_json_output_is_machine_readable(self):
        missing = "/tmp/__sci2doc_missing_quality__.docx"
        proc = run_py(QUALITY_SCRIPT, [missing, "--output", "json"])
        self.assertNotEqual(proc.returncode, 0)
        payload = json.loads(proc.stdout)
        self.assertFalse(payload.get("success", True))
        self.assertEqual(payload.get("error"), "file_not_found")

    def test_state_gate_requires_write_cycle_origin(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            (root / "project_state.json").write_text(
                json.dumps(
                    {
                        "project_info": {"save_path": str(root)},
                        "progress": {"status": "writing"},
                        "stats": {},
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
            (root / "thesis_profile.json").write_text(
                json.dumps({"targets": {"body_target_chars": 80000}}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )

            p1 = run_py(STATE_SCRIPT, ["--project-root", str(root), "preflight", "--chapter", "1"])
            self.assertEqual(p1.returncode, 0)

            p2 = run_py(STATE_SCRIPT, ["--project-root", str(root), "load", "--chapter", "1"])
            self.assertEqual(p2.returncode, 0)

            p3 = run_py(STATE_SCRIPT, ["--project-root", str(root), "gate-check", "--chapter", "1", "--phase", "prewrite"])
            self.assertNotEqual(p3.returncode, 0)

            p4 = run_py(STATE_SCRIPT, ["--project-root", str(root), "write-cycle", "--chapter", "1", "--preflight-lenient"])
            self.assertEqual(p4.returncode, 0, msg=p4.stdout + p4.stderr)

    def test_state_write_cycle_json_summary_emits_single_payload(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            init_proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "init", "--title", "汇总输出测试"],
            )
            self.assertEqual(init_proc.returncode, 0, msg=init_proc.stdout + init_proc.stderr)

            proc = run_py(
                STATE_SCRIPT,
                [
                    "--project-root",
                    str(root),
                    "write-cycle",
                    "--chapter",
                    "1",
                    "--preflight-lenient",
                    "--json-summary",
                ],
            )
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertTrue(payload.get("ok"))
            self.assertEqual(payload.get("command"), "write-cycle")
            self.assertEqual(payload.get("chapter"), "1")
            steps = payload.get("steps", [])
            self.assertGreaterEqual(len(steps), 3)
            names = {s.get("step") for s in steps if isinstance(s, dict)}
            self.assertIn("preflight", names)
            self.assertIn("load", names)
            self.assertIn("gate-prewrite", names)

    def test_state_gate_uninitialized_has_actionable_hint(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "gate-check", "--chapter", "1", "--phase", "prewrite"],
            )
            self.assertNotEqual(proc.returncode, 0)
            payload = extract_json_from_stdout(proc.stdout)
            self.assertIn("init", payload.get("reason", ""))

    def test_state_preflight_reports_dependency_checks(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            (root / "project_state.json").write_text(
                json.dumps({"project_info": {"save_path": str(root)}}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            (root / "thesis_profile.json").write_text(
                json.dumps({"targets": {"body_target_chars": 80000}}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            proc = run_py(STATE_SCRIPT, ["--project-root", str(root), "preflight", "--chapter", "1"])
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = extract_json_from_stdout(proc.stdout)
            self.assertIn("checks", payload)
            names = {item.get("name") for item in payload["checks"] if isinstance(item, dict) and item.get("name")}
            self.assertIn("python-docx", names)
            self.assertIn("docxcompose", names)
            self.assertFalse(payload.get("dependency_require_high_fidelity", True))

    def test_state_preflight_strict_fails_when_optional_files_missing(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            (root / "project_state.json").write_text(
                json.dumps({"project_info": {"save_path": str(root)}}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            (root / "thesis_profile.json").write_text(
                json.dumps({"targets": {"body_target_chars": 80000}}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "preflight", "--chapter", "1", "--strict"],
            )
            self.assertNotEqual(proc.returncode, 0)
            payload = extract_json_from_stdout(proc.stdout)
            self.assertFalse(payload.get("ok", True))
            missing_items = [
                item
                for item in payload.get("checks", [])
                if isinstance(item, dict) and item.get("error") == "missing"
            ]
            self.assertGreater(len(missing_items), 0)

    def test_state_preflight_high_fidelity_marks_docxcompose_required(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            (root / "project_state.json").write_text(
                json.dumps({"project_info": {"save_path": str(root)}}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            (root / "thesis_profile.json").write_text(
                json.dumps({"targets": {"body_target_chars": 80000}}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            proc = run_py(
                STATE_SCRIPT,
                [
                    "--project-root",
                    str(root),
                    "preflight",
                    "--chapter",
                    "1",
                    "--require-high-fidelity",
                ],
            )
            payload = extract_json_from_stdout(proc.stdout)
            self.assertTrue(payload.get("dependency_require_high_fidelity", False))
            docxcompose_item = next(
                (
                    item
                    for item in payload.get("checks", [])
                    if isinstance(item, dict) and item.get("name") == "docxcompose"
                ),
                None,
            )
            self.assertIsNotNone(docxcompose_item)
            self.assertTrue(docxcompose_item.get("required", False))
            if docxcompose_item.get("available", False):
                self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            else:
                self.assertNotEqual(proc.returncode, 0)
                self.assertFalse(payload.get("ok", True))

    def test_state_word_count_syncs_into_project_state(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            atomic_dir = root / "atomic_md"
            atomic_dir.mkdir(parents=True, exist_ok=True)
            md_path = atomic_dir / "1.1_绪论.md"
            md_path.write_text("# 第一章 绪论\n\n这里是正文内容\n", encoding="utf-8")

            proc = run_py(STATE_SCRIPT, ["--project-root", str(root), "word-count"])
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = extract_json_from_stdout(proc.stdout)
            self.assertTrue(payload.get("synced_to_project_state"))

            updated_state = json.loads((root / "project_state.json").read_text(encoding="utf-8"))
            self.assertGreater(updated_state["stats"].get("total_body_words", 0), 0)

    def test_state_profile_command_updates_targets(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            init_proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "init", "--title", "配置测试"],
            )
            self.assertEqual(init_proc.returncode, 0, msg=init_proc.stdout + init_proc.stderr)

            proc = run_py(
                STATE_SCRIPT,
                [
                    "--project-root",
                    str(root),
                    "profile",
                    "--body-target",
                    "86000",
                    "--abstract-min",
                    "1600",
                    "--abstract-max",
                    "2400",
                    "--chapter-target",
                    "2:18000",
                ],
            )
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = extract_json_from_stdout(proc.stdout)
            self.assertTrue(payload.get("updated"))

            profile = json.loads((root / "thesis_profile.json").read_text(encoding="utf-8"))
            self.assertEqual(profile["targets"]["body_target_chars"], 86000)
            self.assertEqual(profile["targets"]["abstract_min_chars"], 1600)
            self.assertEqual(profile["targets"]["abstract_max_chars"], 2400)
            self.assertEqual(profile["chapter_targets"]["2"], 18000)

    def test_state_profile_command_rejects_body_target_below_80000(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            init_proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "init", "--title", "配置校验测试"],
            )
            self.assertEqual(init_proc.returncode, 0, msg=init_proc.stdout + init_proc.stderr)

            proc = run_py(
                STATE_SCRIPT,
                [
                    "--project-root",
                    str(root),
                    "profile",
                    "--body-target",
                    "79000",
                ],
            )
            self.assertNotEqual(proc.returncode, 0)
            payload = extract_json_from_stdout(proc.stdout)
            self.assertEqual(payload.get("error"), "state_file_error")
            self.assertIn("must be >= 80000", payload.get("detail", ""))

    def test_atomic_md_self_check_uses_profile_targets(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            init_proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "init", "--title", "自检配置测试"],
            )
            self.assertEqual(init_proc.returncode, 0, msg=init_proc.stdout + init_proc.stderr)
            profile_path = root / "thesis_profile.json"
            profile = json.loads(profile_path.read_text(encoding="utf-8"))
            profile.setdefault("targets", {})["body_target_chars"] = 86000
            profile["targets"]["references_min_count"] = 120
            profile_path.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")

            md_path = root / "第2章_测试.md"
            md_path.write_text("# 第二章 方法\n\n这是章节正文内容。\n", encoding="utf-8")

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                ["--project-root", str(root), "self-check", "--target", str(md_path)],
            )
            self.assertIn(proc.returncode, (0, 1), msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertEqual(
                payload["word_count"].get("targets", {}).get("body_target"),
                86000,
            )
            self.assertTrue(payload["checks"].get("quality_skipped"))
            self.assertEqual(payload.get("effective_body_target"), 86000)

    def test_atomic_md_self_check_uses_chapter_target_and_can_pass(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            init_proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "init", "--title", "章节目标测试"],
            )
            self.assertEqual(init_proc.returncode, 0, msg=init_proc.stdout + init_proc.stderr)

            profile_path = root / "thesis_profile.json"
            profile = json.loads(profile_path.read_text(encoding="utf-8"))
            profile.setdefault("targets", {})["body_target_chars"] = 80000
            profile.setdefault("chapter_targets", {})["2"] = 10
            profile_path.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")

            md_path = root / "第2章_测试.md"
            md_path.write_text("# 第二章 方法\n\n这是章节正文内容，超过十个字。\n", encoding="utf-8")

            proc = run_py(
                ATOMIC_MD_SCRIPT,
                ["--project-root", str(root), "self-check", "--target", str(md_path)],
            )
            self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)
            payload = json.loads(proc.stdout)
            self.assertTrue(payload.get("ok"))
            self.assertEqual(payload.get("effective_body_target"), 10)
            self.assertTrue(payload.get("checks", {}).get("word_passed"))
            self.assertTrue(payload["checks"].get("quality_skipped"))

    def test_state_reports_structured_error_for_invalid_json(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            root.mkdir(parents=True, exist_ok=True)
            (root / "project_state.json").write_text("{invalid", encoding="utf-8")

            proc = run_py(STATE_SCRIPT, ["--project-root", str(root), "stats"])
            self.assertNotEqual(proc.returncode, 0)
            payload = extract_json_from_stdout(proc.stdout)
            self.assertEqual(payload.get("error"), "state_file_error")
            self.assertEqual(payload.get("reason"), "invalid_json")
            self.assertIn("rollback --target snapshot", payload.get("hint", ""))

    def test_state_word_count_concurrent_updates_keep_history_intact(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            init_proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "init", "--title", "并发测试"],
            )
            self.assertEqual(init_proc.returncode, 0, msg=init_proc.stdout + init_proc.stderr)

            atomic_dir = root / "atomic_md"
            atomic_dir.mkdir(parents=True, exist_ok=True)
            md_path = atomic_dir / "1.1_绪论.md"
            md_path.write_text("# 第一章 绪论\n\n这里是并发字数统计正文内容\n", encoding="utf-8")

            calls = 6
            with ThreadPoolExecutor(max_workers=calls) as pool:
                futures = [
                    pool.submit(
                        run_py,
                        STATE_SCRIPT,
                        ["--project-root", str(root), "word-count"],
                    )
                    for _ in range(calls)
                ]
                results = [f.result() for f in futures]

            for proc in results:
                self.assertEqual(proc.returncode, 0, msg=proc.stdout + proc.stderr)

            history = json.loads((root / "history_log.json").read_text(encoding="utf-8"))
            events = [item for item in history if isinstance(item, dict) and item.get("event") == "word_count"]
            self.assertEqual(len(events), calls)

    def test_state_snapshot_paths_are_unique_when_called_rapidly(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            init_proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "init", "--title", "快照测试"],
            )
            self.assertEqual(init_proc.returncode, 0, msg=init_proc.stdout + init_proc.stderr)

            p1 = run_py(STATE_SCRIPT, ["--project-root", str(root), "snapshot"])
            p2 = run_py(STATE_SCRIPT, ["--project-root", str(root), "snapshot"])
            self.assertEqual(p1.returncode, 0, msg=p1.stdout + p1.stderr)
            self.assertEqual(p2.returncode, 0, msg=p2.stdout + p2.stderr)

            s1 = extract_json_from_stdout(p1.stdout).get("snapshot_dir")
            s2 = extract_json_from_stdout(p2.stdout).get("snapshot_dir")
            self.assertTrue(s1 and s2)
            self.assertNotEqual(s1, s2)
            self.assertTrue(Path(s1).exists())
            self.assertTrue(Path(s2).exists())

    def test_state_rollback_non_strict_keeps_extra_files(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            init_proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "init", "--title", "回滚测试"],
            )
            self.assertEqual(init_proc.returncode, 0, msg=init_proc.stdout + init_proc.stderr)

            snapshot_proc = run_py(STATE_SCRIPT, ["--project-root", str(root), "snapshot"])
            self.assertEqual(snapshot_proc.returncode, 0, msg=snapshot_proc.stdout + snapshot_proc.stderr)
            snapshot_dir = extract_json_from_stdout(snapshot_proc.stdout).get("snapshot_dir")
            self.assertTrue(snapshot_dir)

            extra = root / "02_分章节文档" / "extra.tmp"
            extra.parent.mkdir(parents=True, exist_ok=True)
            extra.write_text("extra", encoding="utf-8")
            self.assertTrue(extra.exists())

            rollback_proc = run_py(
                STATE_SCRIPT,
                [
                    "--project-root",
                    str(root),
                    "rollback",
                    "--target",
                    "snapshot",
                    "--snapshot-dir",
                    str(snapshot_dir),
                ],
            )
            self.assertEqual(rollback_proc.returncode, 0, msg=rollback_proc.stdout + rollback_proc.stderr)
            self.assertTrue(extra.exists())

    def test_state_rollback_strict_mirror_removes_extra_files(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            init_proc = run_py(
                STATE_SCRIPT,
                ["--project-root", str(root), "init", "--title", "严格回滚测试"],
            )
            self.assertEqual(init_proc.returncode, 0, msg=init_proc.stdout + init_proc.stderr)

            snapshot_proc = run_py(STATE_SCRIPT, ["--project-root", str(root), "snapshot"])
            self.assertEqual(snapshot_proc.returncode, 0, msg=snapshot_proc.stdout + snapshot_proc.stderr)
            snapshot_dir = extract_json_from_stdout(snapshot_proc.stdout).get("snapshot_dir")
            self.assertTrue(snapshot_dir)

            extra = root / "02_分章节文档" / "extra.tmp"
            extra.parent.mkdir(parents=True, exist_ok=True)
            extra.write_text("extra", encoding="utf-8")
            self.assertTrue(extra.exists())

            rollback_proc = run_py(
                STATE_SCRIPT,
                [
                    "--project-root",
                    str(root),
                    "rollback",
                    "--target",
                    "snapshot",
                    "--snapshot-dir",
                    str(snapshot_dir),
                    "--strict-mirror",
                ],
            )
            self.assertEqual(rollback_proc.returncode, 0, msg=rollback_proc.stdout + rollback_proc.stderr)
            payload = extract_json_from_stdout(rollback_proc.stdout)
            self.assertTrue(payload.get("strict_mirror"))
            self.assertFalse(extra.exists())


class TestMethodsTableCheck(unittest.TestCase):
    """测试 check_methods_sections_have_tables 函数"""

    def _make_section_file(self, tmp, filename, content):
        """创建一个模拟的 section file 对象"""
        import sys, os
        script_dir = os.path.join(ROOT, "scripts")
        if script_dir not in sys.path:
            sys.path.insert(0, script_dir)
        from atomic_md_workflow import SectionFile
        p = Path(tmp) / filename
        p.write_text(content, encoding="utf-8")
        # 解析编号
        parts = filename.split("_", 1)
        number_text = parts[0]
        title = parts[1].replace(".md", "") if len(parts) > 1 else ""
        number = tuple(int(x) for x in number_text.split("."))
        return SectionFile(path=p, number=number, number_text=number_text, title=title)

    def _get_checker(self):
        import sys, os
        script_dir = os.path.join(ROOT, "scripts")
        if script_dir not in sys.path:
            sys.path.insert(0, script_dir)
        from atomic_md_workflow import check_methods_sections_have_tables
        return check_methods_sections_have_tables

    def test_reagent_section_with_table_no_warning(self):
        check = self._get_checker()
        with tempfile.TemporaryDirectory() as tmp:
            sf = self._make_section_file(tmp, "2.2_实验试剂与耗材.md",
                "## 2.2 实验试剂与耗材\n\n表 2-1：主要试剂\n\n"
                "| 试剂名称 | 规格 | 厂家 |\n|---|---|---|\n| FBS | 500mL | Gibco |\n")
            warnings = check([sf])
            self.assertEqual(warnings, [])

    def test_reagent_section_without_table_warns(self):
        check = self._get_checker()
        with tempfile.TemporaryDirectory() as tmp:
            sf = self._make_section_file(tmp, "2.2_实验试剂与耗材.md",
                "## 2.2 实验试剂与耗材\n\n本实验使用了FBS和DMEM培养基。\n")
            warnings = check([sf])
            self.assertEqual(len(warnings), 1)
            self.assertIn("试剂", warnings[0])
            self.assertIn("管道表格", warnings[0])

    def test_instrument_section_without_table_warns(self):
        check = self._get_checker()
        with tempfile.TemporaryDirectory() as tmp:
            sf = self._make_section_file(tmp, "2.3_实验仪器与设备.md",
                "## 2.3 实验仪器与设备\n\n使用了流式细胞仪。\n")
            warnings = check([sf])
            self.assertEqual(len(warnings), 1)
            self.assertIn("仪器", warnings[0])

    def test_grouping_section_without_table_warns(self):
        check = self._get_checker()
        with tempfile.TemporaryDirectory() as tmp:
            sf = self._make_section_file(tmp, "2.4_实验分组设计.md",
                "## 2.4 实验分组设计\n\n分为对照组和实验组。\n")
            warnings = check([sf])
            self.assertEqual(len(warnings), 1)
            self.assertIn("分组", warnings[0])

    def test_unrelated_section_no_warning(self):
        check = self._get_checker()
        with tempfile.TemporaryDirectory() as tmp:
            sf = self._make_section_file(tmp, "2.1_引言.md",
                "## 2.1 引言\n\n本章介绍研究背景。\n")
            warnings = check([sf])
            self.assertEqual(warnings, [])

    def test_mixed_sections(self):
        check = self._get_checker()
        with tempfile.TemporaryDirectory() as tmp:
            sf1 = self._make_section_file(tmp, "2.2_实验试剂与耗材.md",
                "## 2.2 实验试剂与耗材\n\n| 名称 | 规格 |\n|---|---|\n| FBS | 500mL |\n")
            sf2 = self._make_section_file(tmp, "2.3_主要仪器.md",
                "## 2.3 主要仪器\n\n使用了PCR仪。\n")
            sf3 = self._make_section_file(tmp, "2.1_引言.md",
                "## 2.1 引言\n\n背景介绍。\n")
            warnings = check([sf1, sf2, sf3])
            self.assertEqual(len(warnings), 1)
            self.assertIn("2.3_主要仪器.md", warnings[0])

    def test_validate_includes_table_warnings(self):
        """validate 命令输出应包含 table_warnings 字段"""
        with tempfile.TemporaryDirectory() as tmp:
            proj = Path(tmp)
            ch_dir = proj / "atomic_md" / "第2章"
            ch_dir.mkdir(parents=True)
            (ch_dir / "2.1_引言.md").write_text("## 2.1 引言\n\n背景。\n", encoding="utf-8")
            (ch_dir / "2.2_实验试剂.md").write_text("## 2.2 实验试剂\n\n无表格。\n", encoding="utf-8")
            proc = run_py(ATOMIC_MD_SCRIPT, [
                "--project-root", str(proj), "validate", "--chapter", "2"
            ])
            payload = extract_json_from_stdout(proc.stdout)
            self.assertIn("table_warnings", payload)
            self.assertTrue(len(payload["table_warnings"]) >= 1)


if __name__ == "__main__":
    unittest.main()
