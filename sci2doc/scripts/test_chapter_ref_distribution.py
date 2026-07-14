#!/usr/bin/env python3
"""按章软文献门 + 总量硬门冒烟测试（Part1）。

覆盖：
  1) 研究/实验章 [n] 引用 < research 地板(博=12) → level:warning（不阻断）。
  2) 绪论章引用 < intro_review 地板(博=30) → warning。
  3) 研究章引用 >= 地板 → 不告警。
  4) 结论章不设文献地板（即便引用为 0 也不告警）。
  5) 总量 < references_min_count(80) 仍为 level:error（硬门未被软门取代）。
  6) thesis_profile 硕地板 < 博地板。

standalone: `python3 test_chapter_ref_distribution.py`。
"""
from __future__ import annotations

import os
import sys

from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from check_quality import (  # noqa: E402
    check_chapter_reference_distribution,
    check_reference_count,
)
from thesis_profile import _derive_per_chapter_ref_floor  # noqa: E402


def _h1(doc, text):
    doc.add_paragraph(text, style="Heading 1")


def _body(doc, text):
    doc.add_paragraph(text, style="Normal")


def _cites(n):
    return "研究表明" + "".join(f"[{i}]" for i in range(1, n + 1)) + "。"


DOCTOR_FLOOR = {"intro_review": 30, "research": 12}


def test_research_chapter_below_floor_warns():
    doc = Document()
    _h1(doc, "第2章 材料的实验研究")
    _body(doc, _cites(3))  # 3 < 12
    issues = check_chapter_reference_distribution(doc, DOCTOR_FLOOR)
    assert len(issues) == 1, issues
    assert issues[0]["level"] == "warning", issues[0]
    assert "研究/实验章" in issues[0]["message"], issues[0]


def test_intro_chapter_below_floor_warns():
    doc = Document()
    _h1(doc, "第1章 绪论")
    _body(doc, _cites(5))  # 5 < 30
    issues = check_chapter_reference_distribution(doc, DOCTOR_FLOOR)
    assert len(issues) == 1 and "绪论/文献综述章" in issues[0]["message"], issues


def test_research_chapter_meets_floor_no_warn():
    doc = Document()
    _h1(doc, "第3章 器件性能研究")
    _body(doc, _cites(12))  # 12 >= 12
    issues = check_chapter_reference_distribution(doc, DOCTOR_FLOOR)
    assert issues == [], issues


def test_conclusion_chapter_no_floor():
    doc = Document()
    _h1(doc, "第6章 总结与展望")
    _body(doc, "本文工作到此结束，无引用。")
    issues = check_chapter_reference_distribution(doc, DOCTOR_FLOOR)
    assert issues == [], issues


def test_total_reference_shortfall_still_error():
    doc = Document()
    _h1(doc, "参考文献")
    for i in range(1, 6):  # 仅 5 条 << 80
        _body(doc, f"[{i}] 某作者. 某文献. 某期刊, 2020.")
    issues, count = check_reference_count(doc, min_reference_count=80)
    assert count == 5, count
    assert any(it["level"] == "error" for it in issues), issues


def test_master_floor_below_doctor():
    master = _derive_per_chapter_ref_floor(
        {"format_profile": {"degree_type": "硕士学位论文"}, "targets": {}}
    )["targets"]["per_chapter_ref_floor"]
    doctor = _derive_per_chapter_ref_floor(
        {"format_profile": {"degree_type": "博士学位论文"}, "targets": {}}
    )["targets"]["per_chapter_ref_floor"]
    assert master["intro_review"] < doctor["intro_review"], (master, doctor)
    assert master["research"] < doctor["research"], (master, doctor)


if __name__ == "__main__":
    test_research_chapter_below_floor_warns()
    test_intro_chapter_below_floor_warns()
    test_research_chapter_meets_floor_no_warn()
    test_conclusion_chapter_no_floor()
    test_total_reference_shortfall_still_error()
    test_master_floor_below_doctor()
    print("OK: chapter-ref soft gate (research/intro<floor warn, meets no-warn, "
          "conclusion no-floor) + total<80 still error + master<doctor floor")
