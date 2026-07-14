#!/usr/bin/env python3
"""F2 上下标 run-aware 检测 + F5 [图]/[表] 前缀剥离 回归测试。

F2：check_char_level 的 D2 上下标检测改以"非上下标 run 的拼接"为输入——
  已被 markdown_to_docx 渲染成 subscript/superscript run 的字符不再误报 subsup_bare，
  真裸写(全 plain run)照报。
F5：parse_markdown_line 剥行首 [图]/[表] 前缀 + figure 分支剥方括号——
  前缀式/括号式题注都能落进居中题注分支，无 [ ] / [图] 残留。

standalone: `python3 test_subsup_run_aware.py`
"""
from __future__ import annotations

import os
import sys
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from markdown_to_docx import markdown_to_docx  # noqa: E402
from check_quality import check_char_level  # noqa: E402


def _subsup_codes(doc):
    return [i for i in check_char_level(doc) if i.get("code") == "subsup_bare"]


def test_f2_rendered_not_flagged():
    """① md 含 CO<sub>2</sub>/10<sup>6</sup> → 渲染后零 subsup_bare。"""
    md = "# 第一章 测试\n\n实验中 CO<sub>2</sub> 浓度达到 10<sup>6</sup> 个单位。\n"
    with tempfile.TemporaryDirectory() as d:
        out = os.path.join(d, "a.docx")
        markdown_to_docx(md, out)
        doc = Document(out)
        hits = _subsup_codes(doc)
        assert not hits, f"渲染正确的上下标不应报 subsup_bare，实报：{[h['message'] for h in hits]}"


def test_f2_bare_flagged():
    """② python-docx 直接造裸写 CO2 段落 → 照报 subsup_bare。"""
    doc = Document()
    doc.add_paragraph("实验中 CO2 浓度升高明显。")
    hits = _subsup_codes(doc)
    assert hits, "全 plain run 的裸写 CO2 必须报 subsup_bare"
    assert any("CO2" in h["message"] for h in hits), hits


def test_f2_mixed_only_bare():
    """③ 混合文档：一段正确渲染 + 一段裸写 → 只报裸写那段。"""
    doc = Document()
    md_out = None
    with tempfile.TemporaryDirectory() as d:
        out = os.path.join(d, "b.docx")
        markdown_to_docx("# 章\n\n正确写法 CO<sub>2</sub> 在此。\n", out)
        rendered = Document(out)
    # 把渲染好的段落搬进一个新文档，再加一段裸写
    target = Document()
    for p in rendered.paragraphs:
        np = target.add_paragraph()
        np.style = p.style
        for r in p.runs:
            nr = np.add_run(r.text)
            nr.font.subscript = r.font.subscript
            nr.font.superscript = r.font.superscript
    target.add_paragraph("裸写 CO2 应被抓。")
    hits = _subsup_codes(target)
    assert len(hits) == 1, f"只应报裸写那 1 段，实报 {len(hits)}：{[h['message'] for h in hits]}"
    assert "CO2" in hits[0]["message"]


def _caption_paras(doc, needle):
    return [p for p in doc.paragraphs if needle in p.text]


def test_f5_prefix_and_bracket_forms():
    """F5：[图] 前缀式 与 [图 X-X：] 括号式 → 题注居中、无 [ ] / [图] 残留。"""
    md = "# 章\n\n[图] 图2-1：测试图\n\n[图 2-2：直接式]\n"
    with tempfile.TemporaryDirectory() as d:
        out = os.path.join(d, "c.docx")
        markdown_to_docx(md, out)
        doc = Document(out)

        p1 = _caption_paras(doc, "测试图")
        assert p1, "前缀式题注段未找到"
        assert p1[0].alignment == WD_ALIGN_PARAGRAPH.CENTER, "前缀式题注应居中"

        p2 = _caption_paras(doc, "直接式")
        assert p2, "括号式题注段未找到"
        assert p2[0].alignment == WD_ALIGN_PARAGRAPH.CENTER, "括号式题注应居中"

        for p in p1 + p2:
            assert "[" not in p.text and "]" not in p.text, f"题注残留方括号：{p.text!r}"
            assert "[图]" not in p.text, f"题注残留 [图] 前缀：{p.text!r}"


if __name__ == "__main__":
    test_f2_rendered_not_flagged()
    test_f2_bare_flagged()
    test_f2_mixed_only_bare()
    test_f5_prefix_and_bracket_forms()
    print("ALL PASS: F2(3) + F5(2)")
