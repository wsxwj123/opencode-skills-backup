#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合并章节文档并生成完整论文

功能：
1. 按顺序合并多个 docx 章节文件
2. 保持原有样式和格式
3. 自动插入分页符
4. 生成目录
5. 添加页眉页脚

作者：Sci2Doc Team
日期：2024-03-15
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os
import re

try:
    from docxcompose.composer import Composer
except ImportError:
    Composer = None


def add_page_break(doc):
    """在文档中添加分页符"""
    doc.add_page_break()


def extract_chapter_number(filename):
    """从文件名中提取章节号"""
    match = re.search(r'第(\d+)章', filename)
    if match:
        return int(match.group(1))
    return 999  # 未识别的放到最后


# ---------------------------------------------------------------------------
# 全文合并排序：前置部分 → 正文章节 → 后置部分
# ---------------------------------------------------------------------------

_FRONT_MATTER_ORDER = [
    ("封面", 1),
    ("题名", 2),
    ("独创性", 3),
    ("授权", 4),
    ("中文摘要", 5),
    ("摘要", 6),
    ("英文摘要", 7),
    ("abstract", 8),
    ("目录", 9),
    ("缩略", 10),
    ("符号", 11),
]

_BACK_MATTER_ORDER = [
    ("参考文献", 1),
    ("致谢", 2),
    ("攻读", 3),
    ("成果", 4),
    ("附录", 5),
]


def _docx_merge_sort_key(filepath):
    """为 docx 全文合并生成排序键: (大类, 子序号, 文件名)。

    大类:
      0 = 前置部分 (封面/摘要/目录等)
      1 = 正文章节 (第X章)
      2 = 后置部分 (参考文献/致谢等)
      3 = 未识别
    """
    name = os.path.basename(filepath).lower().replace(" ", "")

    # 正文章节
    m = re.search(r"第(\d+)章", name)
    if m:
        return (1, int(m.group(1)), name)

    # 前置部分
    for keyword, order in _FRONT_MATTER_ORDER:
        if keyword in name:
            return (0, order, name)

    # 后置部分
    for keyword, order in _BACK_MATTER_ORDER:
        if keyword in name:
            return (2, order, name)

    # 未识别 → 放在正文章节之后、后置部分之前
    return (1, 10**6, name)


try:
    from shared_utils import heading_level
except ImportError:  # pragma: no cover
    _script_dir = os.path.dirname(os.path.abspath(__file__))
    if _script_dir not in sys.path:
        sys.path.insert(0, _script_dir)
    from shared_utils import heading_level


def merge_docx_files(file_list, output_path, require_high_fidelity=False):
    """
    合并多个 docx 文件
    
    Args:
        file_list: docx 文件路径列表
        output_path: 输出文件路径
    
    Returns:
        dict: 合并结果
    """
    if not file_list:
        return {
            'success': False,
            'error': '文件列表为空'
        }
    
    # 按章节号排序（支持前置/后置部分智能排序）
    file_list_sorted = sorted(file_list, key=_docx_merge_sort_key)
    
    # 优先使用高保真合并（docxcompose）
    if require_high_fidelity and Composer is None:
        return {
            "success": False,
            "error": "high-fidelity merge requires docxcompose; install with 'pip3 install docxcompose'",
        }

    if Composer is not None:
        try:
            master_doc = Document(file_list_sorted[0])
            composer = Composer(master_doc)
            for file_path in file_list_sorted[1:]:
                composer.append(Document(file_path))
                print(f"✅ 已合并：{os.path.basename(file_path)}")

            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)

            composer.save(output_path)
            return {
                "success": True,
                "output_path": output_path,
                "merged_files": len(file_list_sorted),
                "file_list": [os.path.basename(f) for f in file_list_sorted],
                "merge_engine": "docxcompose",
            }
        except Exception as e:
            if require_high_fidelity:
                return {
                    "success": False,
                    "error": f"high-fidelity merge failed: {str(e)}",
                }
            print(f"⚠️  docxcompose 合并失败，将回退简化合并：{str(e)}")

    # 创建主文档（回退：简化复制）
    try:
        master_doc = Document(file_list_sorted[0])
    except Exception as e:
        return {
            'success': False,
            'error': f'无法打开第一个文件：{str(e)}'
        }
    
    # 添加后续文件
    for file_path in file_list_sorted[1:]:
        try:
            # 添加分页符
            add_page_break(master_doc)
            
            # 读取章节文档
            chapter_doc = Document(file_path)
            
            # 复制段落
            for para in chapter_doc.paragraphs:
                # 创建新段落
                new_para = master_doc.add_paragraph()
                try:
                    new_para.style = para.style.name
                except Exception:
                    pass
                new_para.alignment = para.alignment
                
                # 复制段落格式
                new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
                new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                new_para.paragraph_format.space_before = para.paragraph_format.space_before
                new_para.paragraph_format.space_after = para.paragraph_format.space_after
                new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
                
                # 复制文本和格式
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                    if run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
            
            # 复制表格
            for table in chapter_doc.tables:
                # 简化：仅复制文本内容
                # 完整实现需要复制表格样式、单元格合并等
                rows = len(table.rows)
                cols = len(table.columns)
                new_table = master_doc.add_table(rows=rows, cols=cols)
                
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.rows[i].cells[j].text = cell.text
            
            print(f"✅ 已合并：{os.path.basename(file_path)}")
            
        except Exception as e:
            print(f"⚠️  合并失败：{os.path.basename(file_path)} - {str(e)}")
            continue
    
    # 保存合并后的文档
    try:
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        master_doc.save(output_path)
        return {
            'success': True,
            'output_path': output_path,
            'merged_files': len(file_list_sorted),
            'file_list': [os.path.basename(f) for f in file_list_sorted],
            'merge_engine': 'fallback'
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'保存文件失败：{str(e)}'
        }


def generate_toc(doc):
    """
    生成目录（简化版）
    注意：完整的 Word 目录需要 VBA 或 COM 自动化

    格式：
    - 标题"目  录"：三号黑体加粗居中
    - 章标题：小四号黑体
    - 节标题：小四号宋体
    - 1.5 倍行距

    Args:
        doc: Document 对象
    """
    toc_entries = []

    for para in doc.paragraphs:
        level = heading_level(getattr(para.style, "name", ""))
        if level is not None:
            toc_entries.append({
                'text': para.text,
                'level': level
            })

    # 在文档开头插入目录标题
    toc_title = doc.paragraphs[0].insert_paragraph_before()
    run_title = toc_title.add_run('目  录')
    run_title.bold = True
    run_title.font.name = 'SimHei'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run_title.font.size = Pt(16)  # 三号
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_after = Pt(12)
    from docx.enum.text import WD_LINE_SPACING
    toc_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for entry in toc_entries:
        toc_line = doc.paragraphs[1].insert_paragraph_before()
        indent = (entry['level'] - 1) * 0.74
        toc_line.paragraph_format.left_indent = Cm(indent)
        toc_line.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        run = toc_line.add_run(entry['text'])
        run.font.size = Pt(12)  # 小四号
        if entry['level'] == 1:
            run.font.name = 'SimHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        else:
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.name = 'Times New Roman'  # 西文字体

    print("✅ 目录已生成（页码需要在 Word 中手动更新字段）")


def _set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=10.5, bold=None):
    """设置 run 的中西文字体、字号、加粗。"""
    run.font.name = latin
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _add_page_field(run):
    """向 run 中插入 PAGE 域代码。"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _set_page_number_format(section, fmt='decimal', start_at=None):
    """设置节的页码格式。fmt: 'lowerRoman' | 'decimal'"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start_at is not None:
        pgNumType.set(qn('w:start'), str(start_at))
    elif qn('w:start') in pgNumType.attrib:
        del pgNumType.attrib[qn('w:start')]


def _insert_section_break_before(paragraph):
    """在段落前插入分节符（下一页）。

    实现方式：在前一个段落的 pPr 中添加 <w:sectPr><w:type val="nextPage"/></w:sectPr>。
    如果该段落是文档第一个段落，则不插入。
    """
    p_elem = paragraph._element
    prev = p_elem.getprevious()
    if prev is None:
        return
    # prev 可能是 <w:p> 或 <w:tbl>，只处理 <w:p>
    if prev.tag != qn('w:p'):
        return
    pPr = prev.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        prev.insert(0, pPr)
    sectPr = OxmlElement('w:sectPr')
    sectType = OxmlElement('w:type')
    sectType.set(qn('w:val'), 'nextPage')
    sectPr.append(sectType)
    pPr.append(sectPr)


def _classify_para_zone(para_text):
    """判断段落属于前置/正文/后置。返回 'front'/'body'/'back'/None。"""
    text = para_text.strip()
    if re.search(r'第\d+章', text):
        return 'body'
    front_kw = ['封面', '题名', '独创性', '授权', '摘要', 'abstract', '目录', '缩略', '符号']
    for kw in front_kw:
        if kw in text.lower():
            return 'front'
    back_kw = ['参考文献', '致谢', '攻读', '成果', '附录']
    for kw in back_kw:
        if kw in text:
            return 'back'
    return None


def add_header_footer(doc, thesis_title, university_name="中南大学"):
    """
    添加页眉页脚（中南大学 2022 规范），按章节分节。

    流程：
    1. 扫描所有 Heading 1 段落，在其前方插入分节符
    2. 为每个 section 设置独立页眉（左：大学名+博士学位论文，右：章名）
    3. 前置部分页码用 lowerRoman，正文起用 decimal（从 1 开始）

    Args:
        doc: Document 对象
        thesis_title: 论文标题（前置部分页眉右侧备用）
        university_name: 大学名称（默认：中南大学）
    """
    from docx.enum.text import WD_TAB_ALIGNMENT
    from shared_utils import heading_level

    # ---- 第一步：收集 H1 段落及其章名 ----
    h1_paragraphs = []
    for para in doc.paragraphs:
        lvl = heading_level(getattr(para.style, 'name', ''))
        if lvl == 1:
            h1_paragraphs.append(para)

    # ---- 第二步：在每个 H1 前插入分节符（跳过第一个 H1） ----
    for para in h1_paragraphs[1:]:
        _insert_section_break_before(para)

    # ---- 第三步：建立 section → 章名 / 区域 映射 ----
    # 重新遍历，因为插入分节符后 sections 数量变了
    section_info = []  # [(chapter_name, zone)]
    current_name = thesis_title or ''
    current_zone = 'front'
    h1_idx = 0

    for para in doc.paragraphs:
        lvl = heading_level(getattr(para.style, 'name', ''))
        if lvl == 1:
            zone = _classify_para_zone(para.text)
            if zone:
                current_zone = zone
            current_name = para.text.strip()
            if h1_idx > 0:
                section_info.append((current_name, current_zone))
            else:
                section_info.append((current_name, current_zone))
            h1_idx += 1

    # 如果没有 H1，整个文档算一个 section
    if not section_info:
        section_info.append((thesis_title or '', 'front'))

    # 补齐：sections 可能比 section_info 多（文档原有 section）
    while len(section_info) < len(doc.sections):
        section_info.append((thesis_title or '', 'front'))

    # ---- 第四步：逐 section 设置页眉页脚和页码格式 ----
    left_text = f'{university_name}博士学位论文'
    body_started = False

    for i, section in enumerate(doc.sections):
        if i < len(section_info):
            chapter_name, zone = section_info[i]
        else:
            chapter_name, zone = thesis_title or '', 'body'

        # ---- 页眉 ----
        section.header_distance = Cm(1.5)
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0]
        header_para.clear()

        text_width = section.page_width - section.left_margin - section.right_margin
        header_para.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)

        right_text = chapter_name if zone == 'body' else (thesis_title or '')
        run_h = header_para.add_run(f'{left_text}\t{right_text}')
        _set_run_font(run_h, latin='Times New Roman', east_asia='SimSun', size_pt=10.5)

        # ---- 页脚 ----
        section.footer_distance = Cm(1.75)
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0]
        footer_para.clear()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_f = footer_para.add_run()
        _add_page_field(run_f)
        _set_run_font(run_f, latin='Times New Roman', east_asia='SimSun', size_pt=9)

        # ---- 页码格式 ----
        if zone == 'body' and not body_started:
            _set_page_number_format(section, fmt='decimal', start_at=1)
            body_started = True
        elif zone == 'front':
            if i == 0:
                _set_page_number_format(section, fmt='lowerRoman', start_at=1)
            else:
                _set_page_number_format(section, fmt='lowerRoman')
        else:
            # body 后续章节 / back matter：续前页码
            _set_page_number_format(section, fmt='decimal')

    print(f"✅ 页眉页脚已添加（{len(doc.sections)} 个分节，正文从第 1 页起用阿拉伯数字）")


def resolve_merge_order(input_dir, cover=None, abstract=None, abstract_en=None):
    """
    组装最终合并文件顺序：
    1) 可选前置文件（cover/abstract/abstract_en）
    2) input_dir 下章节文件（按第X章排序）
    """
    ordered = []
    seen = set()
    for p in [cover, abstract, abstract_en]:
        if not p:
            continue
        full = p if os.path.isabs(p) else os.path.join(input_dir, p)
        if os.path.exists(full):
            if full not in seen:
                ordered.append(full)
                seen.add(full)
        else:
            print(f"⚠️  前置文件不存在，已跳过：{full}")

    chapters = []
    for filename in os.listdir(input_dir):
        if filename.endswith(".docx") and not filename.startswith("~"):
            chapters.append(os.path.join(input_dir, filename))
    chapters = sorted(chapters, key=_docx_merge_sort_key)
    for p in chapters:
        if p not in seen:
            ordered.append(p)
            seen.add(p)
    return ordered


def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='合并 Word 章节文档')
    parser.add_argument('--input-dir', required=True, help='章节文件所在目录')
    parser.add_argument('--output', required=True, help='输出文件路径')
    parser.add_argument('--title', default='论文标题', help='论文标题（用于页眉）')
    parser.add_argument('--add-toc', action='store_true', help='生成目录')
    parser.add_argument('--add-header', action='store_true', help='添加页眉页脚')
    parser.add_argument('--cover', help='封面文件（可选，支持绝对路径或相对 input-dir）')
    parser.add_argument('--abstract', help='中文摘要文件（可选，支持绝对路径或相对 input-dir）')
    parser.add_argument('--abstract-en', help='英文摘要文件（可选，支持绝对路径或相对 input-dir）')
    parser.add_argument('--require-high-fidelity', action='store_true', help='要求使用 docxcompose 高保真合并；不可回退')
    
    args = parser.parse_args()
    
    # 获取所有 docx 文件
    if not os.path.exists(args.input_dir):
        print(f"❌ 目录不存在：{args.input_dir}")
        sys.exit(1)
    
    file_list = resolve_merge_order(
        input_dir=args.input_dir,
        cover=args.cover,
        abstract=args.abstract,
        abstract_en=args.abstract_en,
    )
    
    if not file_list:
        print(f"❌ 目录中没有找到 docx 文件：{args.input_dir}")
        sys.exit(1)
    
    print(f"📁 找到 {len(file_list)} 个章节文件")
    print(f"🔄 开始合并...")
    
    # 执行合并
    result = merge_docx_files(
        file_list,
        args.output,
        require_high_fidelity=args.require_high_fidelity,
    )
    
    if not result['success']:
        print(f"❌ 合并失败：{result['error']}")
        sys.exit(1)
    
    print(f"\n✅ 合并完成！")
    print(f"📄 输出文件：{result['output_path']}")
    print(f"📊 合并章节数：{result['merged_files']}")
    print(f"📝 章节列表：")
    for filename in result['file_list']:
        print(f"   - {filename}")
    
    # 可选：生成目录
    if args.add_toc:
        print(f"\n📑 正在生成目录...")
        doc = Document(args.output)
        generate_toc(doc)
        doc.save(args.output)
    
    # 可选：添加页眉页脚
    if args.add_header:
        print(f"\n📋 正在添加页眉页脚...")
        doc = Document(args.output)
        add_header_footer(doc, args.title)
        doc.save(args.output)
    
    print(f"\n🎉 所有操作完成！")


if __name__ == '__main__':
    main()
