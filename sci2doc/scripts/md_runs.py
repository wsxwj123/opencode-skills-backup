#!/usr/bin/env python3
"""共享 Markdown 行内解析 → python-docx run(学术技能通用)。

三家手搓 md→docx 转换器(revise-sci / polish-sci / sci2doc)此前各自维护一份
只认 `**`/`*`/sup/sub 的行内正则,导致下划线强调、列表、code、链接的字面标记残留
在 Word 里。本模块把「行内标记 → 带格式的 run」这一层统一,一处补全、三家复用,
避免「改了 revise 忘了 polish」。**只共享行内解析这一层**,各家外层结构(track
changes / 三线表 / in-place)不动。

核心:
- inline_md_to_runs(paragraph, text, ...):把 text 拆成若干 run 加进段落,行内标记
  转成 run 级格式并**去掉字面标记**;支持 `**b**`/`__b__`/`*i*`/`_i_`/`` `code` ``/
  `[t](u)`/`<sup>`/`<sub>`。
- set_run_font(run, latin, east_asia, black, ...):设 ascii/hAnsi/cs 拉丁字体 + eastAsia
  中文字体 + 可选强制黑色。
- strip_inline_markers(text):要纯字符串时(不建 run)去标记。
- clamp_heading_level(n, max_level):把 h5/h6 收敛到已有最深层级。

保护规则(防误吞):
- 统计显著性标记 `*p<0.05` / `**P<0.01` 等不当粗体/斜体处理。
- 词内下划线(如变量 `x_i`、`p_value`)不当强调;`_` / `__` 仅在词边界成对时才算强调。
- 正文引用 `[12]` 不会被当链接(链接需 `](url)` 结构,`[12]` 后无 `(` )。
"""
from __future__ import annotations

import re

from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# 显著性标记保护:1-4 个星号 + 可选空格 + p/P + 比较符
_SIG_PROTECT = re.compile(r"(\*{1,4})(\s*)([pP][<>≤≥=])")

# 行内片段。顺序要紧:`**` 先于 `*`、`__` 先于 `_`,链接先于其它括号。
# 下划线强调用 (?<![\w]) / (?![\w]) 词边界锚定,避开 x_i / p_value。
_INLINE = re.compile(
    r"\*\*(?P<b1>.+?)\*\*"
    r"|(?<![A-Za-z0-9])__(?P<b2>.+?)__(?![A-Za-z0-9])"
    r"|\*(?P<i1>[^*]+?)\*"
    r"|(?<![A-Za-z0-9])_(?P<i2>[^_]+?)_(?![A-Za-z0-9])"
    r"|`(?P<code>[^`]+?)`"
    r"|\[(?P<ltext>[^\]]+?)\]\((?P<lurl>(?:https?://|www\.|mailto:|ftp://|doi:|/|#)[^)\s]*?)\)"
    r"|<sup>(?P<sup>.*?)</sup>"
    r"|<sub>(?P<sub>.*?)</sub>",
    re.IGNORECASE,
)


def set_run_font(run, latin="Times New Roman", east_asia=None,
                 size_pt=None, bold=None, black=False):
    """给 run 设字体。latin 走 ascii/hAnsi/cs;east_asia 走 w:eastAsia(中文);
    black=True 强制黑色。size_pt/bold 可选。"""
    f = run.font
    f.name = latin
    if size_pt is not None:
        f.size = Pt(size_pt)
    if bold is not None:
        f.bold = bold
    if black:
        f.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), latin)
    if east_asia:
        rfonts.set(qn("w:eastAsia"), east_asia)


def _iter_segments(text):
    """生成 (segment_text, kind) 序列;kind ∈ plain/bold/italic/code/link/sup/sub。
    保护显著性标记与词内下划线后再切分。"""
    ph = []

    def _protect(m):
        ph.append(m.group(0))
        return f"§SIG{len(ph) - 1}§"

    protected = _SIG_PROTECT.sub(_protect, text)

    def _restore(s):
        for i, orig in enumerate(ph):
            s = s.replace(f"§SIG{i}§", orig)
        return s

    pos = 0
    for m in _INLINE.finditer(protected):
        if m.start() > pos:
            yield _restore(protected[pos:m.start()]), "plain"
        if m.group("b1") is not None:
            yield _restore(m.group("b1")), "bold"
        elif m.group("b2") is not None:
            yield _restore(m.group("b2")), "bold"
        elif m.group("i1") is not None:
            yield _restore(m.group("i1")), "italic"
        elif m.group("i2") is not None:
            yield _restore(m.group("i2")), "italic"
        elif m.group("code") is not None:
            yield _restore(m.group("code")), "code"
        elif m.group("ltext") is not None:
            yield _restore(m.group("ltext")), "link"
        elif m.group("sup") is not None:
            yield _restore(m.group("sup")), "sup"
        elif m.group("sub") is not None:
            yield _restore(m.group("sub")), "sub"
        pos = m.end()
    if pos < len(protected):
        yield _restore(protected[pos:]), "plain"


def inline_md_to_runs(paragraph, text, *, latin="Times New Roman",
                      east_asia=None, size_pt=None, black=False):
    """把 text 的行内 md 拆成 run 加进 paragraph(去掉字面标记)。返回加入的 run 列表。"""
    runs = []
    for seg, kind in _iter_segments(text):
        if seg == "":
            continue
        r = paragraph.add_run(seg)
        set_run_font(r, latin=latin, east_asia=east_asia, size_pt=size_pt, black=black)
        if kind == "bold":
            r.bold = True
        elif kind == "italic":
            r.italic = True
        elif kind == "code":
            r.font.name = "Consolas"  # code 用等宽,仍保留文本
        elif kind == "sup":
            r.font.superscript = True
        elif kind == "sub":
            r.font.subscript = True
        # link / plain: 纯文本(链接只留文字、丢 url)
        runs.append(r)
    return runs


def strip_inline_markers(text):
    """去掉行内 md 标记只留纯文本(需要 str 而非 run 的场合)。"""
    return "".join(seg for seg, _ in _iter_segments(text))


def clamp_heading_level(n, max_level=4):
    """h5/h6 等超深层级收敛到已有最深层级;下限 1。"""
    try:
        n = int(n)
    except (TypeError, ValueError):
        return max_level
    return max(1, min(n, max_level))


def demo():
    """assert 自测:标记去净 + 格式正确 + 保护规则生效。"""
    from docx import Document

    doc = Document()
    p = doc.add_paragraph()
    inline_md_to_runs(
        p,
        "背景 **bold** 与 *italic* 和 __b2__ 及 _i2_ 还有 `code` 与 [text](http://x) "
        "上标<sup>2</sup> 下标<sub>i</sub> 变量 x_i 显著 *p<0.05 引用 [12] "
        "参考 Ref [3](Smith 2020) 结束",
        latin="Times New Roman", black=True,
    )
    full = "".join(r.text for r in p.runs)
    # 1. 字面标记去净(*/#/`/方括号链接语法)
    assert "**" not in full and "__" not in full, full
    assert "`" not in full, full
    assert "](http" not in full, full
    assert "<sup>" not in full and "<sub>" not in full, full
    # 2. 保护:x_i 与 *p<0.05 与 [12] 原样保留
    assert "x_i" in full, "词内下划线被误吞: " + full
    assert "*p<0.05" in full, "显著性标记被误吞: " + full
    assert "[12]" in full, "引用被误当链接: " + full
    # 编号引用紧跟非 url 圆括号(如 [3](Smith 2020))不当链接、内容不丢
    assert "[3]" in full and "(Smith 2020)" in full, "编号引用+圆括号被误吞: " + full
    # 3. 格式:bold/italic 各至少一个 run
    assert any(r.bold for r in p.runs), "无粗体 run"
    assert any(r.italic for r in p.runs), "无斜体 run"
    assert any(r.font.superscript for r in p.runs), "无上标 run"
    # 4. 链接只留文字
    assert "text" in full and "http://x" not in full, full
    # 5. 字体+黑色
    assert all(r.font.name in ("Times New Roman", "Consolas") for r in p.runs)
    assert any(r.font.color and r.font.color.rgb == RGBColor(0, 0, 0) for r in p.runs)
    # strip 版
    assert strip_inline_markers("**a** _b_ `c`") == "a b c"
    assert clamp_heading_level(6, 4) == 4 and clamp_heading_level(0) == 1
    print("md_runs demo: ALL PASS")


if __name__ == "__main__":
    demo()
