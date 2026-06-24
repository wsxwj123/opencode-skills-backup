#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
sci2doc 格式契约回归测试（开发者维护工具，非运行时流程）。

固化本轮已修的 4 个 bug，防止以后改正则/字段读取时退化。
纯 assert，无 pytest 依赖，所有输入在测试内现造（字符串 / 内存 docx / tempfile）。
不依赖任何外部真稿路径。

用法：
    python3 test_format_contract.py
失败抛 AssertionError（returncode != 0），全过打印 OK。
"""

import json
import os
import sys
import tempfile

# 本测试文件与被测脚本同目录，确保脚本目录在 import 路径上
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPT_DIR not in sys.path:
    sys.path.insert(0, _SCRIPT_DIR)


# ===========================================================================
# 契约 1：reference_renderer raw_vancouver 回退
# bug：SCI 参考库条目常只有完整 raw_vancouver 字符串而无结构化 authors/journal，
#      旧逻辑直接渲染会输出"佚名 / 空刊名"。修复后应回退原样输出 raw_vancouver。
# ===========================================================================

def test_reference_renderer_raw_vancouver_fallback():
    from reference_renderer import render_entry

    # (a) 只有 raw_vancouver、无结构化字段 → 应回退原文，不出现"佚名"
    raw_text = ("Smith J, Doe A, Roe B, et al. A landmark study on tumor immunology. "
                "Nature. 2021;590(7844):123-130. DOI: 10.1038/s41586-021-00001-2.")
    entry_raw = {"id": "ref_1", "raw_vancouver": raw_text}
    rendered = render_entry(entry_raw)
    assert "A landmark study on tumor immunology" in rendered, \
        f"raw_vancouver 片段应原样保留，实际：{rendered!r}"
    assert "佚名" not in rendered, \
        f"无结构化字段时不应出现'佚名'，实际：{rendered!r}"
    assert "Nature" in rendered, \
        f"raw_vancouver 刊名片段应保留，实际：{rendered!r}"

    # 带编号也应保留原文片段
    rendered_idx = render_entry(entry_raw, index=7)
    assert rendered_idx.startswith("[7] "), f"编号前缀缺失：{rendered_idx!r}"
    assert "佚名" not in rendered_idx
    assert "A landmark study on tumor immunology" in rendered_idx

    # (b) 完整结构化字段 → 走正常 GB/T 7714 渲染（不退化为 raw_vancouver 原文）
    entry_struct = {
        "id": "ref_2",
        "type": "journal",
        "authors": ["Zhang S", "Li W", "Wang H"],
        "title": "Structured rendering must still work",
        "journal": "Journal of Test",
        "year": "2022",
        "volume": "10",
        "issue": "3",
        "pages": "200-210",
        "doi": "10.1000/test.2022",
        # 故意附带一个不同的 raw_vancouver，确认正常路径不会误用它
        "raw_vancouver": "RAW_SHOULD_NOT_APPEAR_IN_OUTPUT",
    }
    rendered_struct = render_entry(entry_struct)
    assert "RAW_SHOULD_NOT_APPEAR_IN_OUTPUT" not in rendered_struct, \
        f"有结构化字段时不应退化为 raw_vancouver，实际：{rendered_struct!r}"
    assert "Structured rendering must still work" in rendered_struct
    assert "[J]" in rendered_struct, f"期刊类应带 [J] 标识，实际：{rendered_struct!r}"
    assert "Journal of Test" in rendered_struct
    assert "佚名" not in rendered_struct


# ===========================================================================
# 契约 2：markdown_to_docx 输出 A4 页面
# bug：python-docx 默认 Letter（21.59×27.94cm），check_quality 硬要求 A4。
#      修复后应显式设 page_width≈21.0cm、page_height≈29.7cm。
# ===========================================================================

def test_markdown_to_docx_a4_page_size():
    from markdown_to_docx import markdown_to_docx
    from docx import Document
    from docx.shared import Cm

    md = "# 第1章 绪论\n\n这是一段用于生成最小 docx 的正文内容。\n"

    tmp_path = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        ok = markdown_to_docx(md, tmp_path)
        assert ok is True, "markdown_to_docx 应返回 True"

        doc = Document(tmp_path)
        section = doc.sections[0]
        # EMU 容差比较：1cm ≈ 360000 EMU，允许 ±0.05cm（=18000 EMU）误差
        tol = Cm(0.05)
        expected_w = Cm(21.0)
        expected_h = Cm(29.7)
        assert abs(section.page_width - expected_w) <= tol, \
            f"page_width 应≈A4 宽 21.0cm，实际 EMU={section.page_width}"
        assert abs(section.page_height - expected_h) <= tol, \
            f"page_height 应≈A4 高 29.7cm，实际 EMU={section.page_height}"
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


# ===========================================================================
# 契约 3：abbreviation_registry 连字符全称
# bug：旧正则全称部分不接受连字符，"Triple-Negative Breast Cancer" 抓不到。
#      修复后含连字符的全称仍能抓到 abbr；无连字符的全称不退化。
# ===========================================================================

def test_abbreviation_hyphenated_full_name():
    from abbreviation_registry import extract_abbreviations

    # (a) 含连字符全称
    found_hyphen = extract_abbreviations('三阴性乳腺癌（Triple-Negative Breast Cancer, TNBC）')
    abbrs = {item["abbr"] for item in found_hyphen}
    assert "TNBC" in abbrs, f"含连字符全称应抓到 TNBC，实际：{found_hyphen!r}"
    tnbc = next(item for item in found_hyphen if item["abbr"] == "TNBC")
    assert "Triple-Negative" in tnbc["full_en"], \
        f"全称应保留连字符片段，实际：{tnbc!r}"

    # (b) 无连字符全称仍正常（不退化）
    found_plain = extract_abbreviations('磁共振成像（Magnetic Resonance Imaging, MRI）')
    abbrs_plain = {item["abbr"] for item in found_plain}
    assert "MRI" in abbrs_plain, f"无连字符全称应正常抓到 MRI，实际：{found_plain!r}"
    mri = next(item for item in found_plain if item["abbr"] == "MRI")
    assert "Magnetic Resonance Imaging" in mri["full_en"], \
        f"全称应完整保留，实际：{mri!r}"

    # (c) 小写英文全称（期刊惯例 focused ultrasound，非首字母大写）应抓到 abbr 且 full_en 完整
    found_lower = extract_abbreviations('聚焦超声（focused ultrasound，FUS）')
    fus = next((item for item in found_lower if item["abbr"] == "FUS"), None)
    assert fus is not None, f"小写英文全称应抓到 FUS，实际：{found_lower!r}"
    assert fus["full_en"] == "focused ultrasound", \
        f"小写全称应完整保留，实际：{fus!r}"
    assert fus["full_cn"] == "聚焦超声", f"中文全称应正确，实际：{fus!r}"

    # (d) 希腊字母（γ/β）不得在缩写或全称处截断，且缩写不留悬空连字符
    found_greek = extract_abbreviations('转化生长因子β（transforming growth factor-β，TGF-β）')
    tgf = next((item for item in found_greek if item["abbr"] == "TGF-β"), None)
    assert tgf is not None, \
        f"希腊字母缩写应完整抓到 TGF-β（不是 TGF-），实际：{found_greek!r}"
    assert not tgf["abbr"].endswith("-"), \
        f"缩写不得以悬空连字符结尾，实际：{tgf!r}"
    assert "transforming growth factor-β" in tgf["full_en"], \
        f"含希腊字母的英文全称应完整不截断，实际：{tgf!r}"


# ===========================================================================
# 契约 4：check_quality 表号 + 图号去重
# bug：同一表/图号被题注 + 正文多次引用时，旧逻辑按出现次数排序导致 expected
#      错位，产生假"不连续"警告。修复后按 (chapter, number) 去重，真缺号仍报错。
# ===========================================================================

def _build_docx_with_paragraphs(lines):
    """在内存里造一个最小 docx Document，每行一段。"""
    from docx import Document
    doc = Document()
    for ln in lines:
        doc.add_paragraph(ln)
    return doc


def _table_issues(issues):
    return [i for i in issues if "表编号不连续" in i.get("message", "")]


def _figure_issues(issues):
    return [i for i in issues if "图编号不连续" in i.get("message", "")]


def test_check_quality_table_figure_dedup():
    from check_quality import check_figure_numbering

    # ---- 表号：4 个连续表号 2-1~2-4，且每个被正文多次引用 → 0 假阳性 ----
    table_ok_lines = [
        "表 2-1：第一张表的题注",
        "如表 2-1 所示，结果良好。",        # 重复引用 2-1
        "表 2-2：第二张表的题注",
        "见表 2-2 与表 2-1 的对比。",        # 重复引用 2-2 和 2-1
        "表 2-3：第三张表的题注",
        "表 2-3 进一步说明问题。",            # 重复引用 2-3
        "表 2-4：第四张表的题注",
        "综合表 2-1 至表 2-4 可知。",        # 同段同时引用 2-1 和 2-4
    ]
    issues_ok = check_figure_numbering(_build_docx_with_paragraphs(table_ok_lines))
    assert _table_issues(issues_ok) == [], \
        f"连续表号被多次引用不应误报，实际表号告警：{_table_issues(issues_ok)!r}"

    # ---- 表号：真缺号 2-1, 2-3（缺 2-2）→ 仍正确报错 ----
    table_gap_lines = [
        "表 2-1：第一张表",
        "表 2-3：第三张表",   # 缺 2-2
    ]
    issues_gap = check_figure_numbering(_build_docx_with_paragraphs(table_gap_lines))
    assert len(_table_issues(issues_gap)) >= 1, \
        f"真缺号（缺表 2-2）应报错，实际：{issues_gap!r}"

    # ---- 图号：4 个连续图号 2-1~2-4，且每个被正文多次引用 → 0 假阳性 ----
    figure_ok_lines = [
        "图 2-1：第一张图的题注",
        "如图 2-1 所示。",                    # 重复引用 2-1
        "图 2-2：第二张图的题注",
        "图 2-2 与图 2-1 对比。",             # 重复引用 2-2 和 2-1
        "图 2-3：第三张图的题注",
        "见图 2-3。",                         # 重复引用 2-3
        "图 2-4：第四张图的题注",
        "图 2-1 至图 2-4 共同说明。",         # 同段引用 2-1 和 2-4
    ]
    issues_fig_ok = check_figure_numbering(_build_docx_with_paragraphs(figure_ok_lines))
    assert _figure_issues(issues_fig_ok) == [], \
        f"连续图号被多次引用不应误报，实际图号告警：{_figure_issues(issues_fig_ok)!r}"

    # ---- 图号：真缺号 2-1, 2-3（缺 2-2）→ 仍正确报错 ----
    figure_gap_lines = [
        "图 2-1：第一张图",
        "图 2-3：第三张图",   # 缺 2-2
    ]
    issues_fig_gap = check_figure_numbering(_build_docx_with_paragraphs(figure_gap_lines))
    assert len(_figure_issues(issues_fig_gap)) >= 1, \
        f"真缺号（缺图 2-2）应报错，实际：{issues_fig_gap!r}"


# ===========================================================================
# 契约 5：citation_guard 引用门禁（J4 完整 / J5 自引 / J7 时效 / A4 双向）
# J4/A4 fail-closed（exit_code=2 阻断），J5/J7 advisory（WARN，不改 exit_code）。
# 调 citation_guard_core 三新函数 + check_bidirectional，仿 gsw --gates。
# 双向断言：缺字段/僵尸 → 阻断；齐全/全引 → 放行；raw_only 不误杀。
# ===========================================================================

def test_citation_gates_j4_a4_j5_j7():
    from citation_guard import run_integrity_gates

    # ---- J4：title 缺失 → incomplete（fail-closed，exit_code=2）----
    entries_bad = [
        {"citation_number": 1, "type": "journal", "title": "完整条目",
         "authors": ["张三"], "journal": "测试学报", "year": 2024,
         "volume": "10", "pages": "1-9", "doi": "10.1/x"},
        {"citation_number": 2, "authors": ["李四"]},  # 无 title 无 handle → incomplete
    ]
    rep = run_integrity_gates(entries_bad, drafts_dir=None,
                              manuscript_authors=[], current_year=2026)
    assert rep["exit_code"] == 2, f"J4 缺字段应 fail-closed：{rep}"
    assert len(rep["j4_completeness"]["incomplete"]) >= 1
    assert rep["j4_completeness"]["strength"] == "fail-closed"

    # ---- J4 raw_only 不误杀：有 title + raw_vancouver、无结构化字段 → raw_only（放行）----
    entries_raw = [
        {"citation_number": 1, "title": "Some real paper",
         "raw_vancouver": "Wang H, et al. Some real paper. Nature. 2023;1:1-2.",
         "doi": "10.1/y"},
    ]
    rep_raw = run_integrity_gates(entries_raw, drafts_dir=None,
                                  manuscript_authors=[], current_year=2026)
    assert rep_raw["exit_code"] == 0, f"raw_only 不应误杀：{rep_raw}"
    assert rep_raw["j4_completeness"]["raw_only_count"] == 1
    assert rep_raw["j4_completeness"]["incomplete"] == []

    # ---- J5 自引：作者命中过半 → warn，但 exit_code 仍 0（advisory）----
    entries_self = [
        {"citation_number": 1, "title": "A", "authors": ["张三"], "doi": "10.1/a",
         "journal": "J", "year": 2024, "volume": "1", "pages": "1-2"},
        {"citation_number": 2, "title": "B", "authors": ["王五"], "doi": "10.1/b",
         "journal": "J", "year": 2024, "volume": "1", "pages": "3-4"},
    ]
    rep_self = run_integrity_gates(entries_self, drafts_dir=None,
                                   manuscript_authors=["张三"], current_year=2026)
    assert rep_self["j5_self_citation"]["status"] == "warn", \
        f"自引过半应 warn：{rep_self['j5_self_citation']}"
    assert rep_self["j5_self_citation"]["strength"] == "warn"
    assert rep_self["exit_code"] == 0, "J5 是 advisory，不应改 exit_code"

    # ---- J5 无作者 → skip（不报错）----
    rep_noauth = run_integrity_gates(entries_self, drafts_dir=None,
                                     manuscript_authors=[], current_year=2026)
    assert rep_noauth["j5_self_citation"]["status"] == "skipped"

    # ---- J7 时效：全部老文献 → warn，exit_code 仍 0（advisory）----
    entries_old = [
        {"citation_number": 1, "title": "old1", "year": 2005, "doi": "10.1/c",
         "journal": "J", "authors": ["A"], "volume": "1", "pages": "1-2"},
        {"citation_number": 2, "title": "old2", "year": 2006, "doi": "10.1/d",
         "journal": "J", "authors": ["B"], "volume": "1", "pages": "3-4"},
    ]
    rep_old = run_integrity_gates(entries_old, drafts_dir=None,
                                  manuscript_authors=[], current_year=2026)
    assert rep_old["j7_recency"]["status"] == "warn", \
        f"全老文献应 warn：{rep_old['j7_recency']}"
    assert rep_old["exit_code"] == 0, "J7 是 advisory，不应改 exit_code"

    # ---- A4 双向：扫描 drafts 比对 [n] ----
    tmp_dir = tempfile.mkdtemp()
    try:
        from pathlib import Path
        md_dir = Path(tmp_dir) / "manuscripts"
        md_dir.mkdir()
        good_entries = [
            {"citation_number": n, "title": f"t{n}", "doi": f"10.1/{n}",
             "journal": "J", "authors": ["A"], "year": 2024,
             "volume": "1", "pages": "1-2"}
            for n in (1, 2, 3)
        ]
        # (a) 正文全引 1,2,3 → A4 ok，exit_code 0
        (md_dir / "ch1.md").write_text("# 绪论\n结论[1]，方法[2,3]。\n", encoding="utf-8")
        rep_ok = run_integrity_gates(good_entries, drafts_dir=md_dir,
                                     manuscript_authors=[], current_year=2026)
        assert rep_ok["a4_bidirectional"]["status"] == "ok", \
            f"全引应 A4 ok：{rep_ok['a4_bidirectional']}"
        assert rep_ok["exit_code"] == 0

        # (b) 正文漏引 3（僵尸）+ 引用不存在的 9（孤儿）→ A4 fail，exit_code 2
        (md_dir / "ch1.md").write_text("# 绪论\n结论[1]，另见[9]。\n", encoding="utf-8")
        rep_fail = run_integrity_gates(good_entries, drafts_dir=md_dir,
                                       manuscript_authors=[], current_year=2026)
        a4 = rep_fail["a4_bidirectional"]
        assert a4["status"] == "fail", f"漏引/孤儿应 A4 fail：{a4}"
        assert 9 in a4["orphans"], f"9 应为孤儿：{a4}"
        assert set(a4["zombies"]) >= {2, 3}, f"2,3 应为僵尸：{a4}"
        assert rep_fail["exit_code"] == 2, "A4 fail 应 fail-closed"
        assert a4["strength"] == "fail-closed"
    finally:
        import shutil
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ===========================================================================
# 契约 6：check_quality 字符级（D1 半角标点 / D2 上下标裸写 / F1 中文错别字）
# 全 WARN（level=warning）。双向：违规命中 + 合法不误报 + 已标注剥离生效。
# ===========================================================================

def _build_docx(lines):
    from docx import Document
    doc = Document()
    for ln in lines:
        doc.add_paragraph(ln)
    return doc


def test_check_quality_char_level():
    from check_quality import check_char_level

    doc = _build_docx([
        "本文研究方法,采用对照实验。",       # D1 半角逗号（汉字,汉字）
        "测得水分子H2O含量与CO2浓度。",      # D2 裸写 H2O、CO2
        "实验登陆系统后按装软件。",           # F1 登陆→登录、按装→安装
        "已正确标注的 H~2~O 不应再报。",      # D2 已标注 → 剥离不报
        "正常陈述句没有任何问题。",           # 合法 → 不误报
    ])
    issues = check_char_level(doc)
    codes = {}
    for i in issues:
        codes[i["code"]] = codes.get(i["code"], 0) + 1

    # D1：半角逗号命中
    assert codes.get("halfwidth_punct_in_cn", 0) >= 1, f"D1 应命中半角逗号：{codes}"
    # D2：恰为 2（H2O+CO2）；已标注 H~2~O 被剥离不计 → 证明剥离生效
    assert codes.get("subsup_bare", 0) == 2, f"D2 应恰为2(剥离生效)：{codes}"
    # F1：登陆 + 按装
    assert codes.get("chinese_typo", 0) >= 2, f"F1 应命中错别字：{codes}"
    # 全部 WARN，不阻断
    assert all(i["level"] == "warning" for i in issues), "字符级必须全为 warning"

    # 合法对照：全角标点 + 标注规范 + 无错字 → 0 命中
    clean = _build_docx([
        "本文研究方法，采用对照实验。",
        "测得水分子 H~2~O 含量与 CO~2~ 浓度。",
        "实验登录系统后安装软件。",
    ])
    clean_issues = check_char_level(clean)
    clean_codes = {i["code"] for i in clean_issues}
    assert "halfwidth_punct_in_cn" not in clean_codes, f"合法全角不应报：{clean_issues}"
    assert "subsup_bare" not in clean_codes, f"已标注不应报：{clean_issues}"
    assert "chinese_typo" not in clean_codes, f"无错字不应报：{clean_issues}"


# ===========================================================================
# 契约 7：A1 公式编号连续（check_figure_numbering 内新增）
# 写法对齐图/表：按 (chapter, number) 去重后比对 expected=i+1。
# 双向：连续(多次引用)不误报 + 真缺号报错；不破坏图/表去重。
# ===========================================================================

def _a5_issues(issues):
    return [i for i in issues if i.get("category") == "交叉引用"]


def test_check_quality_a5_crossref_validity():
    """A5 章节交叉引用有效性：双向。
    引用存在的目标 → 不报；引用不存在的目标（断链）→ WARN 报。
    """
    from check_quality import check_markdown_quality

    def _run_md(text):
        with tempfile.NamedTemporaryFile(
                "w", suffix=".md", delete=False, encoding="utf-8") as f:
            f.write(text)
            path = f.name
        try:
            issues, _ = check_markdown_quality(path)
        finally:
            os.unlink(path)
        return issues

    # ---- 干净稿：所有交叉引用目标都存在 → A5 零报 ----
    clean = (
        "# 第1章 绪论\n本研究方法见第2章。\n\n"
        "# 第2章 方法\n\n## 2.1 流程\n图 2-1：实验流程图\n\n"
        "如图 2-1 所示。详见第 2.1 节。\n\n"
        "表 2-1：参数配置表\n\n见表 2-1，参数已列出。\n\n"
        "# 附录 A 原始数据\n原始数据见附录 A。\n"
    )
    a5_clean = _a5_issues(_run_md(clean))
    assert a5_clean == [], f"有效交叉引用不应报断链：{[i['message'] for i in a5_clean]}"

    # ---- 脏稿：图/表/节/附录均断链 → 4 类各至少 1 报，且全为 warning ----
    dirty = (
        "# 第1章 绪论\n\n"
        "# 第2章 实验\n见图 9-9 不存在。见表 8-8 不存在。\n"
        "见第 99 节不存在。见附录 Z 不存在。\n"
    )
    a5_dirty = _a5_issues(_run_md(dirty))
    msgs = " ".join(i["message"] for i in a5_dirty)
    assert "图 9-9" in msgs, f"断链图 9-9 应报：{msgs}"
    assert "表 8-8" in msgs, f"断链表 8-8 应报：{msgs}"
    assert "节 99" in msgs, f"断链第 99 节应报：{msgs}"
    assert "附录 Z" in msgs, f"断链附录 Z 应报：{msgs}"
    assert all(i["level"] == "warning" for i in a5_dirty), \
        f"A5 断链必须全为 warning（保守不阻断）：{a5_dirty}"


def _formula_issues(issues):
    return [i for i in issues if i.get("category") == "公式编号"]


def test_check_quality_formula_numbering():
    from check_quality import check_figure_numbering

    # ---- 连续公式 2-1~2-3，每个被定义+正文多次引用 → 0 假阳性 ----
    ok_lines = [
        "如公式(2-1)所示，能量守恒。",
        "代入公式(2-1)得结果。",          # 重复引用 2-1
        "公式(2-2)描述动量变化。",
        "由式(2-2)与式(2-1)联立。",       # 重复引用 2-2、2-1（含"式(N-M)"变体）
        "公式(2-3)为最终形式。",
    ]
    issues_ok = check_figure_numbering(_build_docx(ok_lines))
    assert _formula_issues(issues_ok) == [], \
        f"连续公式被多次引用不应误报：{_formula_issues(issues_ok)!r}"

    # ---- 真缺号 2-1, 2-3（缺 2-2）→ 报错 ----
    gap_lines = ["公式(2-1)定义。", "公式(2-3)缺了 2-2。"]
    issues_gap = check_figure_numbering(_build_docx(gap_lines))
    assert len(_formula_issues(issues_gap)) >= 1, \
        f"真缺号（缺公式 2-2）应报错：{issues_gap!r}"

    # ---- 不破坏图/表去重（回归）：连续表号多次引用仍 0 假阳性 ----
    tab_lines = ["表 2-1：题注", "如表 2-1 所示。", "表 2-2：题注", "表 2-3：题注"]
    tab_issues = check_figure_numbering(_build_docx(tab_lines))
    assert [i for i in tab_issues if "表编号不连续" in i.get("message", "")] == [], \
        "新增公式检查不应破坏表号去重"


# ===========================================================================
# 契约 9：figure_registry 锁文件 encoding（Windows 中文不崩）
# bug：save_figure_map 的 open(lock_file, "a+") 未指定 encoding，Windows 默认
#      GBK，含中文路径/内容场景在锁文件读写处可能 UnicodeDecodeError 崩溃。
# 修：open(..., "a+", encoding="utf-8")。本测试做含中文内容的 save→load 往返，
#      确保锁文件路径走通且不崩。
# ===========================================================================

def test_figure_registry_chinese_save_load_roundtrip():
    import figure_registry as fr

    root = tempfile.mkdtemp()
    figure_map = {
        "图1-1": {"cn_id": "图1-1", "caption": "中文图注：实验流程示意图", "src": "图片/流程图.png"},
        "图2-3": {"cn_id": "图2-3", "caption": "另一张图的中文说明", "src": "结果/对比.png"},
    }
    # 写入（触发锁文件 open）
    fr.save_figure_map(root, figure_map)
    # 锁文件确实被创建过
    lock = fr._figure_map_path(root) + ".lock"
    assert os.path.exists(lock), "锁文件应被创建"
    # 读回应完整保留中文
    loaded = fr.load_figure_map(root)
    assert loaded == figure_map, f"中文 figure_map 往返应无损：{loaded!r}"
    assert "中文图注：实验流程示意图" in loaded["图1-1"]["caption"]


# ===========================================================================
# 入口
# ===========================================================================

if __name__ == "__main__":
    test_reference_renderer_raw_vancouver_fallback()
    print("OK 契约1 reference_renderer raw_vancouver 回退")

    test_markdown_to_docx_a4_page_size()
    print("OK 契约2 markdown_to_docx A4 页面")

    test_abbreviation_hyphenated_full_name()
    print("OK 契约3 abbreviation_registry 连字符全称")

    test_check_quality_table_figure_dedup()
    print("OK 契约4 check_quality 表号+图号去重")

    test_citation_gates_j4_a4_j5_j7()
    print("OK 契约5 citation_guard 门禁 J4/J5/J7/A4")

    test_check_quality_char_level()
    print("OK 契约6 check_quality 字符级 D1/D2/F1")

    test_check_quality_formula_numbering()
    print("OK 契约7 check_quality A1 公式编号连续")

    test_check_quality_a5_crossref_validity()
    print("OK 契约8 check_quality A5 章节交叉引用有效性")

    test_figure_registry_chinese_save_load_roundtrip()
    print("OK 契约9 figure_registry 锁文件 encoding 中文往返不崩")

    print("ALL OK")
