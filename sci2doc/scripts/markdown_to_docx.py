#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 转 Word 工具（应用中南大学样式）

功能：
1. 解析 Markdown 文本
2. 转换为 Word 文档
3. 自动应用中南大学博士论文样式
4. 处理图表占位符

作者：Sci2Doc Team
日期：2024-03-15
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import re
import sys
import os

try:
    from abbreviation_registry import load_registry, get_all as get_all_abbreviations
except Exception:  # pragma: no cover
    _script_dir = os.path.dirname(os.path.abspath(__file__))
    if _script_dir not in sys.path:
        sys.path.insert(0, _script_dir)
    try:
        from abbreviation_registry import load_registry, get_all as get_all_abbreviations
    except ImportError:
        load_registry = None
        get_all_abbreviations = None


# ---------------------------------------------------------------------------
# Markdown 管道表格解析
# ---------------------------------------------------------------------------

_PIPE_TABLE_RE = re.compile(r'^\|(.+)\|$')
_SEPARATOR_RE = re.compile(r'^[\|\s\-:]+$')


def _parse_pipe_row(line):
    """解析管道表格行，返回单元格文本列表。支持 \\| 转义。"""
    stripped = line.strip()
    m = _PIPE_TABLE_RE.match(stripped)
    if not m:
        return None
    inner = m.group(1)
    # 先将转义管道替换为占位符，分割后再还原
    placeholder = '\x00PIPE\x00'
    inner = inner.replace('\\|', placeholder)
    cells = [c.strip().replace(placeholder, '|') for c in inner.split('|')]
    return cells


def _is_separator_row(line):
    """判断是否为分隔行（如 |---|---|）"""
    return bool(_SEPARATOR_RE.match(line.strip()))


def parse_markdown_line(line):
    """
    解析 Markdown 行，识别类型
    
    Returns:
        tuple: (type, content, level)
    """
    line = line.rstrip()
    
    # 空行
    if not line.strip():
        return ('empty', '', 0)
    
    # 一级标题 # Title
    if line.startswith('# ') and not line.startswith('## '):
        return ('heading1', line[2:].strip(), 1)
    
    # 二级标题 ## Title
    elif line.startswith('## ') and not line.startswith('### '):
        return ('heading2', line[3:].strip(), 2)
    
    # 三级标题 ### Title
    elif line.startswith('### '):
        return ('heading3', line[4:].strip(), 3)
    
    # 图片占位符 [图 1-1：标题] 或裸格式 图 1-1：标题
    elif re.match(r'\[图\s*\d+-\d+[：:].+\]', line):
        return ('figure', line.strip(), 0)
    elif re.match(r'^图\s*\d+-\d+[：:]', line.strip()):
        return ('figure', line.strip(), 0)
    
    # 表格标题 [表 1-1：标题] 或裸格式 表 1-1：标题
    elif re.match(r'\[表\s*\d+-\d+[：:].+\]', line):
        return ('table', line.strip(), 0)
    elif re.match(r'^表\s*\d+-\d+[：:]', line.strip()):
        return ('table', line.strip(), 0)
    
    # 正文段落
    else:
        return ('paragraph', line.strip(), 0)


def set_run_font(run, latin, east_asia, size_pt, bold=None):
    """
    同时设置拉丁字体和东亚字体，避免 Word 回退到意外字体。
    """
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), east_asia)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def apply_csu_heading1_style(paragraph):
    """应用一级标题样式"""
    paragraph.style = "Heading 1"
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    
    for run in paragraph.runs:
        set_run_font(run, latin='Times New Roman', east_asia='SimHei', size_pt=16, bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)


def apply_csu_heading2_style(paragraph):
    """应用二级标题样式"""
    paragraph.style = "Heading 2"
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    
    for run in paragraph.runs:
        set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=14, bold=False)


def apply_csu_heading3_style(paragraph):
    """应用三级标题样式"""
    paragraph.style = "Heading 3"
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    
    for run in paragraph.runs:
        set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=12, bold=False)


def apply_csu_normal_style(paragraph):
    """应用正文样式"""
    paragraph.style = "Normal"
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)  # 2字符
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    
    for run in paragraph.runs:
        set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=12, bold=False)


def apply_csu_caption_style(paragraph):
    """应用图表题注样式"""
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    
    for run in paragraph.runs:
        set_run_font(run, latin='Times New Roman', east_asia='KaiTi', size_pt=10.5, bold=False)


# ---------------------------------------------------------------------------
# 页眉 / 页脚 / 页码
# ---------------------------------------------------------------------------

def _add_page_number_field(run):
    """在 run 中插入 PAGE 域代码（动态页码）"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def set_page_number_format(section, fmt='decimal', start_at=None):
    """
    设置节的页码格式。
    fmt: 'lowerRoman' | 'decimal'
    start_at: 起始页码（int），None 表示续前节
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start_at is not None:
        pgNumType.set(qn('w:start'), str(start_at))
    elif qn('w:start') in pgNumType.attrib:
        del pgNumType.attrib[qn('w:start')]


def setup_header(section, left_text, right_text):
    """
    设置页眉：左侧 left_text，右侧 right_text，宋体五号(10.5pt)。
    距顶端 1.5cm。
    """
    section.header_distance = Cm(1.5)
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = ""
    # 右对齐制表位 = 页面文本区宽度
    text_width = section.page_width - section.left_margin - section.right_margin
    p.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(f"{left_text}\t{right_text}")
    set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=10.5, bold=False)


def setup_footer(section, page_num_fmt='decimal', start_at=None):
    """
    设置页脚：居中页码，TNR 小五号(9pt)。
    距底端 1.75cm。
    page_num_fmt: 'lowerRoman' (前置部分) | 'decimal' (正文)
    start_at: 起始页码
    """
    section.footer_distance = Cm(1.75)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    _add_page_number_field(run)
    set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=9, bold=False)
    # 设置页码格式
    set_page_number_format(section, fmt=page_num_fmt, start_at=start_at)


# ---------------------------------------------------------------------------
# 前置部分专用样式：摘要 / 英文摘要 / 目录
# ---------------------------------------------------------------------------

def add_abstract_section(doc, abstract_body, keywords=None):
    """
    添加中文摘要页。
    - 标题"摘  要"：三号黑体加粗，居中
    - "摘要："标识：四号黑体加粗，顶格
    - 正文：四号宋体，1.5倍行距
    - 关键词：四号黑体加粗"关键词："+ 四号宋体内容，全角分号分隔
    """
    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(0)
    title_run = title_para.add_run('摘  要')
    set_run_font(title_run, latin='Times New Roman', east_asia='SimHei', size_pt=16, bold=True)

    # "摘要："标识行
    label_para = doc.add_paragraph()
    label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label_para.paragraph_format.first_line_indent = Cm(0)
    label_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    label_run = label_para.add_run('摘要：')
    set_run_font(label_run, latin='Times New Roman', east_asia='SimHei', size_pt=14, bold=True)

    # 正文
    if abstract_body:
        for para_text in abstract_body.split('\n'):
            para_text = para_text.strip()
            if not para_text:
                continue
            body_para = doc.add_paragraph()
            body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_para.paragraph_format.first_line_indent = Cm(0.74)
            body_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            body_run = body_para.add_run(para_text)
            set_run_font(body_run, latin='Times New Roman', east_asia='SimSun', size_pt=14, bold=False)

    # 关键词
    if keywords:
        kw_para = doc.add_paragraph()
        kw_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        kw_para.paragraph_format.first_line_indent = Cm(0)
        kw_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        kw_label = kw_para.add_run('关键词：')
        set_run_font(kw_label, latin='Times New Roman', east_asia='SimHei', size_pt=14, bold=True)
        kw_content = kw_para.add_run(keywords)
        set_run_font(kw_content, latin='Times New Roman', east_asia='SimSun', size_pt=14, bold=False)


def add_english_abstract_section(doc, abstract_body, keywords=None):
    """
    添加英文摘要页。
    - 标题"ABSTRACT"：三号 TNR 加粗，居中
    - "Abstract："标识：四号 TNR 加粗，顶格
    - 正文：四号 TNR，1.5倍行距
    - Keywords：四号 TNR 加粗"Keywords："+ 四号 TNR 内容，半角分号分隔
    """
    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(0)
    title_run = title_para.add_run('ABSTRACT')
    set_run_font(title_run, latin='Times New Roman', east_asia='Times New Roman', size_pt=16, bold=True)

    # "Abstract："标识行
    label_para = doc.add_paragraph()
    label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label_para.paragraph_format.first_line_indent = Cm(0)
    label_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    label_run = label_para.add_run('Abstract：')
    set_run_font(label_run, latin='Times New Roman', east_asia='Times New Roman', size_pt=14, bold=True)

    # 正文
    if abstract_body:
        for para_text in abstract_body.split('\n'):
            para_text = para_text.strip()
            if not para_text:
                continue
            body_para = doc.add_paragraph()
            body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_para.paragraph_format.first_line_indent = Cm(0.74)
            body_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            body_run = body_para.add_run(para_text)
            set_run_font(body_run, latin='Times New Roman', east_asia='Times New Roman', size_pt=14, bold=False)

    # Keywords
    if keywords:
        kw_para = doc.add_paragraph()
        kw_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        kw_para.paragraph_format.first_line_indent = Cm(0)
        kw_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        kw_label = kw_para.add_run('Keywords：')
        set_run_font(kw_label, latin='Times New Roman', east_asia='Times New Roman', size_pt=14, bold=True)
        kw_content = kw_para.add_run(keywords)
        set_run_font(kw_content, latin='Times New Roman', east_asia='Times New Roman', size_pt=14, bold=False)


def add_toc_section(doc, toc_entries=None):
    """
    添加目录页。
    - 标题"目  录"：三号黑体加粗，居中（中间空两格）
    - 章标题：小四号黑体
    - 节标题：小四号宋体
    如果 toc_entries 为 None，插入 TOC 域代码（Word 打开后按 F9 更新）。
    toc_entries 格式：[(level, title, page_str), ...]  level=1 章, level=2 节, level=3 小节
    """
    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(12)
    title_run = title_para.add_run('目  录')
    set_run_font(title_run, latin='Times New Roman', east_asia='SimHei', size_pt=16, bold=True)

    if toc_entries is None:
        # 插入 TOC 域代码，用户在 Word 中按 F9 更新
        toc_para = doc.add_paragraph()
        run = toc_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
    else:
        for level, title, page_str in toc_entries:
            entry_para = doc.add_paragraph()
            entry_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            if level == 1:
                # 章标题：小四号黑体，无缩进
                entry_para.paragraph_format.first_line_indent = Cm(0)
                entry_run = entry_para.add_run(f'{title}{"." * 20}{page_str}')
                set_run_font(entry_run, latin='Times New Roman', east_asia='SimHei', size_pt=12, bold=False)
            else:
                # 节/小节标题：小四号宋体，缩进
                indent = Cm(0.74 * (level - 1))
                entry_para.paragraph_format.first_line_indent = indent
                entry_run = entry_para.add_run(f'{title}{"." * 20}{page_str}')
                set_run_font(entry_run, latin='Times New Roman', east_asia='SimSun', size_pt=12, bold=False)


# ---------------------------------------------------------------------------
# 正文 Markdown 粗体/斜体处理
# ---------------------------------------------------------------------------

# 匹配 **bold** 或 __bold__（双星号/双下划线粗体）
_BOLD_RE = re.compile(r'\*\*(.+?)\*\*|__(.+?)__')
# 匹配显著性标记保护：星号(1-4个) + 可选空格 + P/p + 比较符
_SIGNIFICANCE_PROTECT_RE = re.compile(r'(\*{1,4})(\s*)([pP][<>≤≥=])')


def strip_bold_markers(text):
    """去除正文中的 **粗体** 标记，但保留统计学显著性标记 *p<0.05 等。

    逻辑：
    1. 保护显著性标记（如 **P<0.01），替换为占位符。
    2. 执行去除粗体操作。
    3. 还原占位符。
    这防止了 "A(**P<0.01)B(**P<0.01)" 被误判为粗体包裹。
    """
    placeholders = []

    def protect(m):
        # 将整个匹配串（如 "**P<"）替换为占位符
        token = f"§SIG{len(placeholders)}§"
        placeholders.append(m.group(0))
        return token

    # 1. 保护
    temp_text = _SIGNIFICANCE_PROTECT_RE.sub(protect, text)

    # 2. 去粗体
    stripped_text = _BOLD_RE.sub(lambda m: m.group(1) or m.group(2), temp_text)

    # 3. 还原
    for i, original in enumerate(placeholders):
        stripped_text = stripped_text.replace(f"§SIG{i}§", original)

    return stripped_text


# ---------------------------------------------------------------------------
# 三线表工具
# ---------------------------------------------------------------------------


def _set_cell_border(cell, **kwargs):
    """
    设置单元格边框。

    kwargs 示例: top={"sz": 12, "val": "single", "color": "000000"}
    sz 单位为 1/8 pt，所以 12 = 1.5pt, 4 = 0.5pt
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        for attr_name, attr_val in attrs.items():
            element.set(qn(f'w:{attr_name}'), str(attr_val))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _no_border():
    return {"sz": "0", "val": "none", "color": "auto"}


def _thick_border():
    return {"sz": "12", "val": "single", "color": "000000"}  # 1.5pt


def _thin_border():
    return {"sz": "6", "val": "single", "color": "000000"}   # 0.75pt


def is_table_row(line):
    """检测是否为 Markdown 表格行: | col1 | col2 |"""
    stripped = line.strip()
    return stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 3


def is_table_separator(line):
    """检测是否为 Markdown 表格分隔行: |---|---|"""
    stripped = line.strip()
    if not (stripped.startswith('|') and stripped.endswith('|')):
        return False
    inner = stripped[1:-1]
    cells = inner.split('|')
    return all(re.match(r'^[\s\-:]+$', c) for c in cells)


def parse_table_rows(lines):
    """
    从连续的 Markdown 表格行中提取表头和数据行。

    Returns:
        tuple: (headers: list[str], rows: list[list[str]])
    """
    headers = []
    rows = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        cells = [c.strip() for c in stripped.strip('|').split('|')]
        if i == 0:
            headers = cells
        elif is_table_separator(line):
            continue
        else:
            rows.append(cells)
    return headers, rows


def create_three_line_table(doc, headers, rows, caption=None):
    """
    创建三线表格式的 Word 表格。

    - 顶线: 1.5pt
    - 表头下线: 0.5pt
    - 底线: 1.5pt
    - 无竖线

    Args:
        doc: Document 对象
        headers: 表头列表
        rows: 数据行列表
        caption: 表格标题（可选，置于表格上方）
    """
    # 添加标题（如果有）
    if caption:
        cap_para = doc.add_paragraph(caption)
        cap_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_before = Pt(12)  # 表题注：段前1行
        cap_para.paragraph_format.space_after = Pt(0)   # 表题注：段后0行
        cap_para.paragraph_format.first_line_indent = Cm(0)
        for run in cap_para.runs:
            set_run_font(run, latin='Times New Roman', east_asia='KaiTi', size_pt=10.5, bold=False)

    num_cols = len(headers)
    num_rows = 1 + len(rows)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 填充表头
    for j, header_text in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=10.5, bold=True)

    # 填充数据行
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i + 1].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_run_font(run, latin='Times New Roman', east_asia='SimSun', size_pt=10.5, bold=False)

    # 应用三线表边框
    no = _no_border()
    thick = _thick_border()
    thin = _thin_border()

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if i == 0:
                # 表头行：顶线粗，底线细，无竖线
                _set_cell_border(cell, top=thick, bottom=thin, left=no, right=no)
            elif i == num_rows - 1:
                # 最后一行：底线粗，无竖线
                _set_cell_border(cell, top=no, bottom=thick, left=no, right=no)
            else:
                # 中间行：无边框
                _set_cell_border(cell, top=no, bottom=no, left=no, right=no)

    return table


# ---------------------------------------------------------------------------
# 缩略语对照表页
# ---------------------------------------------------------------------------


def create_abbreviation_table_page(doc, project_root):
    """
    在文档中插入缩略语对照表页（三线表格式）。

    Args:
        doc: Document 对象
        project_root: 项目根目录（用于读取注册表）

    Returns:
        bool: 是否成功插入（无缩略语时返回 False）
    """
    if get_all_abbreviations is None:
        return False

    try:
        items = get_all_abbreviations(project_root)
    except Exception:
        return False

    if not items:
        return False

    # 标题
    heading = doc.add_heading("主要缩略语对照表", level=1)
    apply_csu_heading1_style(heading)

    # 构建表格数据
    headers = ["缩略语", "英文全称", "中文全称"]
    rows = []
    for abbr, info in items:
        full_en = info.get("full_en", "") or ""
        full_cn = info.get("full_cn", "") or ""
        rows.append([abbr, full_en, full_cn])

    create_three_line_table(doc, headers, rows)

    # 分页符
    doc.add_page_break()

    return True


def markdown_to_docx(md_content, output_path, chapter_num=None, project_root=None,
                     include_abbreviation_table=False,
                     header_right_text=None, page_num_fmt='decimal',
                     page_num_start=None):
    """
    将 Markdown 内容转换为 Word 文档
    
    Args:
        md_content: Markdown 文本内容
        output_path: 输出文件路径
        chapter_num: 章节号（可选）
        project_root: 项目根目录（可选，用于缩略语表）
        include_abbreviation_table: 是否在文档开头插入缩略语对照表
        header_right_text: 页眉右侧文字（如"第1章 绪论"），None 则自动从首个 H1 提取
        page_num_fmt: 页码格式 'decimal'(阿拉伯) | 'lowerRoman'(罗马)
        page_num_start: 起始页码（int），None 续前节
    
    Returns:
        bool: 转换是否成功
    """
    try:
        # 创建文档
        doc = Document()
        
        # 设置页面
        section = doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

        # 自动提取页眉右侧文字（从首个 H1）
        if header_right_text is None:
            for ln in md_content.split('\n'):
                ln_s = ln.strip()
                if ln_s.startswith('# ') and not ln_s.startswith('## '):
                    header_right_text = ln_s[2:].strip()
                    break
            if header_right_text is None:
                header_right_text = ""

        # 设置页眉
        setup_header(section, '中南大学博士学位论文', header_right_text)
        # 设置页脚（页码）
        setup_footer(section, page_num_fmt=page_num_fmt, start_at=page_num_start)

        # 插入缩略语对照表（如果需要）
        if include_abbreviation_table and project_root:
            create_abbreviation_table_page(doc, project_root)
        
        # 逐行解析，支持表格累积
        lines = md_content.split('\n')
        table_buffer = []       # 累积连续的表格行
        table_caption = None    # 表格标题（表 X-X）

        def flush_table():
            """将累积的表格行渲染为三线表"""
            nonlocal table_buffer, table_caption
            if not table_buffer:
                return
            headers, rows = parse_table_rows(table_buffer)
            if headers:
                create_three_line_table(doc, headers, rows, caption=table_caption)
            table_buffer = []
            table_caption = None

        for line in lines:
            # 检测是否为表格行
            if is_table_row(line) or (table_buffer and is_table_separator(line)):
                table_buffer.append(line)
                continue

            # 非表格行：先刷新之前累积的表格
            if table_buffer:
                flush_table()

            # 如果之前设置了 table_caption 但没有紧跟管道表格，
            # 将其作为普通题注段落渲染，避免泄漏到后续表格
            if table_caption is not None:
                line_type_cur, content_cur, _ = parse_markdown_line(line)
                if line_type_cur != 'empty':
                    # 非空行且不是表格行 → caption 后面没有紧跟管道表格
                    cap_para = doc.add_paragraph(table_caption)
                    apply_csu_caption_style(cap_para)
                    table_caption = None
                    # 继续处理当前行（不 skip）

            line_type, content, level = parse_markdown_line(line)
            
            if line_type == 'empty':
                continue
            
            elif line_type == 'heading1':
                para = doc.add_heading(content, level=1)
                apply_csu_heading1_style(para)
            
            elif line_type == 'heading2':
                para = doc.add_heading(content, level=2)
                apply_csu_heading2_style(para)
            
            elif line_type == 'heading3':
                para = doc.add_heading(content, level=3)
                apply_csu_heading3_style(para)
            
            elif line_type == 'figure':
                para = doc.add_paragraph(content)
                apply_csu_caption_style(para)
            
            elif line_type == 'table':
                # 表格占位符 [表 X-X：标题] — 可能是后续 Markdown 表格的标题
                table_caption = content.strip('[]')
                # 不立即渲染，等待后续表格行
            
            elif line_type == 'paragraph':
                if content:
                    cleaned = strip_bold_markers(content)
                    para = doc.add_paragraph(cleaned)
                    apply_csu_normal_style(para)

        # 文件末尾：刷新残留的表格
        if table_buffer:
            flush_table()
        
        # 保存文档
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        doc.save(output_path)
        print(f"✅ 转换成功：{output_path}")
        return True
    
    except Exception as e:
        print(f"❌ 转换失败：{str(e)}")
        return False


def convert_markdown_file(md_file_path, output_path=None, project_root=None,
                          include_abbreviation_table=False,
                          header_right_text=None, page_num_fmt='decimal',
                          page_num_start=None):
    """
    转换 Markdown 文件
    
    Args:
        md_file_path: Markdown 文件路径
        output_path: 输出文件路径（可选）
        project_root: 项目根目录（可选，用于缩略语表）
        include_abbreviation_table: 是否插入缩略语对照表
        header_right_text: 页眉右侧文字，None 自动提取
        page_num_fmt: 页码格式 'decimal' | 'lowerRoman'
        page_num_start: 起始页码
    """
    if not os.path.exists(md_file_path):
        print(f"❌ 文件不存在：{md_file_path}")
        return False
    
    # 读取 Markdown 文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 确定输出路径
    if output_path is None:
        output_path = md_file_path.replace('.md', '.docx')
    
    # 提取章节号（如果文件名包含）
    chapter_num = None
    match = re.search(r'第(\d+)章', os.path.basename(md_file_path))
    if match:
        chapter_num = int(match.group(1))
    
    # 执行转换
    return markdown_to_docx(
        md_content, output_path, chapter_num=chapter_num,
        project_root=project_root,
        include_abbreviation_table=include_abbreviation_table,
        header_right_text=header_right_text,
        page_num_fmt=page_num_fmt,
        page_num_start=page_num_start,
    )


def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Markdown 转 Word（中南大学样式）')
    parser.add_argument('input', help='输入 Markdown 文件路径')
    parser.add_argument('-o', '--output', help='输出 Word 文件路径（可选）')
    parser.add_argument('-c', '--chapter', type=int, help='章节号（可选）')
    parser.add_argument('--project-root', help='项目根目录（用于缩略语表）')
    parser.add_argument('--abbreviation-table', action='store_true',
                        help='在文档开头插入缩略语对照表')
    parser.add_argument('--header-right', help='页眉右侧文字（如"第1章 绪论"），默认自动提取')
    parser.add_argument('--page-num-fmt', choices=['decimal', 'lowerRoman'],
                        default='decimal', help='页码格式（默认阿拉伯数字）')
    parser.add_argument('--page-num-start', type=int, help='起始页码')
    
    args = parser.parse_args()
    
    # 执行转换
    success = convert_markdown_file(
        args.input, args.output,
        project_root=args.project_root,
        include_abbreviation_table=args.abbreviation_table,
        header_right_text=args.header_right,
        page_num_fmt=args.page_num_fmt,
        page_num_start=args.page_num_start,
    )
    
    if success:
        print("\n📋 样式说明：")
        print("  - 一级标题：三号黑体加粗，居中")
        print("  - 二级标题：四号宋体，顶格")
        print("  - 三级标题：小四号宋体，顶格")
        print("  - 正文：小四号宋体/Times New Roman，首行缩进")
        print("  - 图表题注：五号楷体，居中")
        sys.exit(0)
    else:
        sys.exit(1)


if __name__ == '__main__':
    main()
