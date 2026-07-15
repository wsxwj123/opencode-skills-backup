#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
merge_chapters.py 回归测试（开发者维护工具，非运行时流程）。

覆盖：
- 全文合并排序键 _docx_merge_sort_key / extract_chapter_number / resolve_merge_order
  （前置封面/摘要 + 第1/2/10章 顺序，第10章不得排到第2章前）
- fallback（无 docxcompose）单章异常时的行为——安全不变量：
  正常章节内容不得因某一章损坏而全部丢失。

纯 assert，无框架依赖。缺 python-docx 时打印 SKIP 并返回 0。
用法：python3 test_merge_chapters.py
"""

import os
import sys
import tempfile
import shutil

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPT_DIR not in sys.path:
    sys.path.insert(0, _SCRIPT_DIR)

try:
    from docx import Document
except Exception:
    print("SKIP test_merge_chapters：缺少 python-docx")
    sys.exit(0)

import merge_chapters as mc


def _make_chapter_docx(path, marker_text):
    doc = Document()
    doc.add_paragraph(marker_text)
    doc.save(path)


# ---------------------------------------------------------------------------
# 排序键
# ---------------------------------------------------------------------------

def test_extract_chapter_number():
    assert mc.extract_chapter_number("第1章_绪论.docx") == 1
    assert mc.extract_chapter_number("第10章_总结.docx") == 10
    assert mc.extract_chapter_number("封面.docx") == 999


def test_docx_merge_sort_key_zones():
    # 大类：前置=0，正文=1，后置=2
    assert mc._docx_merge_sort_key("封面.docx")[0] == 0
    assert mc._docx_merge_sort_key("中文摘要.docx")[0] == 0
    assert mc._docx_merge_sort_key("第1章.docx")[0] == 1
    assert mc._docx_merge_sort_key("参考文献.docx")[0] == 2
    # 正文按章号排序，10 不得排到 2 前
    k2 = mc._docx_merge_sort_key("第2章.docx")
    k10 = mc._docx_merge_sort_key("第10章.docx")
    assert k2 < k10, f"第2章应排在第10章之前：{k2} vs {k10}"
    assert k2[1] == 2 and k10[1] == 10


def test_docx_merge_sort_full_order():
    files = ["第10章.docx", "第2章.docx", "封面.docx", "第1章.docx", "中文摘要.docx", "致谢.docx"]
    ordered = sorted(files, key=mc._docx_merge_sort_key)
    assert ordered == ["封面.docx", "中文摘要.docx", "第1章.docx", "第2章.docx", "第10章.docx", "致谢.docx"], ordered


def test_resolve_merge_order_front_then_chapters():
    tmp = tempfile.mkdtemp()
    try:
        for name in ["第10章.docx", "第2章.docx", "第1章.docx", "封面.docx", "中文摘要.docx"]:
            open(os.path.join(tmp, name), "w").close()
        ordered = mc.resolve_merge_order(input_dir=tmp)
        bases = [os.path.basename(p) for p in ordered]
        # 前置在前，正文按章号（10 在 2 后）
        assert bases.index("封面.docx") < bases.index("第1章.docx")
        assert bases.index("中文摘要.docx") < bases.index("第1章.docx")
        assert bases.index("第2章.docx") < bases.index("第10章.docx"), bases
        # 显式前置参数去重：不重复出现
        ordered2 = mc.resolve_merge_order(input_dir=tmp, cover="封面.docx")
        bases2 = [os.path.basename(p) for p in ordered2]
        assert bases2.count("封面.docx") == 1, bases2
        assert bases2[0] == "封面.docx"
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


# ---------------------------------------------------------------------------
# fallback（无 docxcompose）单章损坏
# ---------------------------------------------------------------------------

def _merged_paragraph_texts(output_path):
    doc = Document(output_path)
    return "\n".join(p.text for p in doc.paragraphs)


def test_fallback_all_chapters_ok_counts_match():
    """fallback 路径 3 章全部有效：merged_files 应等于实际写入的 3 章。"""
    tmp = tempfile.mkdtemp()
    saved_composer = mc.Composer
    try:
        mc.Composer = None  # 强制走 fallback 简化合并
        paths = []
        for i in (1, 2, 3):
            p = os.path.join(tmp, f"第{i}章.docx")
            _make_chapter_docx(p, f"MARK_CHAPTER_{i}")
            paths.append(p)
        out = os.path.join(tmp, "merged.docx")
        result = mc.merge_docx_files(paths, out)
        assert result["success"] is True
        assert result.get("merge_engine") == "fallback"
        assert result["merged_files"] == 3
        text = _merged_paragraph_texts(out)
        for i in (1, 2, 3):
            assert f"MARK_CHAPTER_{i}" in text, f"章{i}内容应在合并结果中"
    finally:
        mc.Composer = saved_composer
        shutil.rmtree(tmp, ignore_errors=True)


def test_fallback_corrupt_chapter_does_not_drop_good_chapters():
    """安全不变量：某一章损坏时，其余有效章内容不得整体丢失（不崩溃）。

    注意：当前实现对损坏章仅 print+continue，且 merged_files 仍返回传入总数、
    success 仍为 True——此处不把该(疑似 bug)行为锁进断言当正确，只断言
    「有效章不丢」这一安全底线，count/success 的过报问题在测试报告中单列。
    """
    tmp = tempfile.mkdtemp()
    saved_composer = mc.Composer
    try:
        mc.Composer = None
        p1 = os.path.join(tmp, "第1章.docx")
        _make_chapter_docx(p1, "MARK_CHAPTER_1")
        p2 = os.path.join(tmp, "第2章.docx")  # 损坏：非法 docx
        with open(p2, "w", encoding="utf-8") as f:
            f.write("this is not a valid docx zip")
        p3 = os.path.join(tmp, "第3章.docx")
        _make_chapter_docx(p3, "MARK_CHAPTER_3")

        out = os.path.join(tmp, "merged.docx")
        result = mc.merge_docx_files([p1, p2, p3], out)
        # 不崩溃、产出文件存在
        assert os.path.exists(out)
        text = _merged_paragraph_texts(out)
        # 有效章内容保留（第2章损坏被跳过）
        assert "MARK_CHAPTER_1" in text
        assert "MARK_CHAPTER_3" in text
        # 特征化当前行为（非断言其正确）：仅记录，不作硬性期望
        actually_written = sum(1 for i in (1, 3) if f"MARK_CHAPTER_{i}" in text)
        assert actually_written == 2
    finally:
        mc.Composer = saved_composer
        shutil.rmtree(tmp, ignore_errors=True)


def test_empty_file_list():
    r = mc.merge_docx_files([], "/tmp/x.docx")
    assert r["success"] is False


if __name__ == "__main__":
    test_extract_chapter_number()
    print("OK extract_chapter_number")
    test_docx_merge_sort_key_zones()
    print("OK _docx_merge_sort_key 分区/章号")
    test_docx_merge_sort_full_order()
    print("OK 全序：封面→摘要→1→2→10→致谢")
    test_resolve_merge_order_front_then_chapters()
    print("OK resolve_merge_order 前置+章号+去重")
    test_fallback_all_chapters_ok_counts_match()
    print("OK fallback 三章全有效 count 匹配")
    test_fallback_corrupt_chapter_does_not_drop_good_chapters()
    print("OK fallback 单章损坏不丢有效章（安全不变量）")
    test_empty_file_list()
    print("OK 空文件列表 success=False")
    print("ALL OK")
