#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
check_quality.py 只读 QC 门禁回归测试（开发者维护工具，非运行时流程）。

原则：合规稿→零 issue；违规稿→命中对应 level/category（重点防“该报没报”）。
覆盖：check_full_thesis_structure / check_section_order / check_figure_map_consistency /
      check_cross_chapter_coherence / check_word_format_compliance / check_table_format

纯 assert，无框架依赖。缺 python-docx 时打印 SKIP。
用法：python3 test_check_quality_gates.py
"""

import json
import os
import sys
import tempfile
import shutil

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPT_DIR not in sys.path:
    sys.path.insert(0, _SCRIPT_DIR)

try:
    from docx import Document
    from docx.shared import Cm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    print("SKIP test_check_quality_gates：缺少 python-docx")
    sys.exit(0)

import check_quality as cq


def _levels(issues):
    return [i.get("level") for i in issues]


def _messages(issues):
    return " | ".join(i.get("message", "") for i in issues)


# ---------------------------------------------------------------------------
# check_full_thesis_structure
# ---------------------------------------------------------------------------

def _thesis_doc(chapter_titles):
    doc = Document()
    for t in chapter_titles:
        doc.add_heading(t, level=1)
        doc.add_paragraph("正文内容占位。")
    return doc


def test_full_structure_compliant_zero_error():
    doc = _thesis_doc([
        "第1章 绪论", "第2章 研究一", "第3章 研究二", "第4章 研究三", "第5章 结论与展望",
    ])
    issues = cq.check_full_thesis_structure(doc, min_chapters=5)
    assert issues == [], _messages(issues)


def test_full_structure_too_few_chapters_errors():
    doc = _thesis_doc(["第1章 绪论", "第2章 结论"])
    issues = cq.check_full_thesis_structure(doc, min_chapters=5)
    assert any(i["level"] == "error" for i in issues)
    assert "章节数不足" in _messages(issues)


def test_full_structure_first_not_intro_errors():
    doc = _thesis_doc([
        "第1章 材料与方法", "第2章 研究一", "第3章 研究二", "第4章 研究三", "第5章 结论",
    ])
    issues = cq.check_full_thesis_structure(doc, min_chapters=5)
    assert "第一章标题不符合绪论要求" in _messages(issues)


def test_full_structure_last_not_conclusion_errors():
    doc = _thesis_doc([
        "第1章 绪论", "第2章 研究一", "第3章 研究二", "第4章 研究三", "第5章 更多实验",
    ])
    issues = cq.check_full_thesis_structure(doc, min_chapters=5)
    assert "最后一章标题不符合独立总结章要求" in _messages(issues)


# ---------------------------------------------------------------------------
# check_section_order
# ---------------------------------------------------------------------------

def test_section_order_compliant():
    doc = Document()
    for t in ["封面", "中文摘要", "Abstract", "目录", "第1章 绪论", "第2章 结论", "参考文献", "致谢"]:
        doc.add_heading(t, level=1)
    issues = cq.check_section_order(doc)
    assert issues == [], _messages(issues)


def test_section_order_violation_detected():
    # 参考文献 出现在 正文 之前 → 乱序
    doc = Document()
    for t in ["封面", "参考文献", "第1章 绪论", "第2章 结论"]:
        doc.add_heading(t, level=1)
    issues = cq.check_section_order(doc)
    assert any(i["level"] == "error" for i in issues), _messages(issues)
    assert "顺序" in _messages(issues)


# ---------------------------------------------------------------------------
# check_figure_map_consistency
# ---------------------------------------------------------------------------

def test_figure_map_consistency_flags_unregistered():
    tmp = tempfile.mkdtemp()
    try:
        # figure_map 只注册 图2-1
        with open(os.path.join(tmp, "figure_map.json"), "w", encoding="utf-8") as f:
            json.dump({"图2-1": {"chapter": 2, "seq": 1, "source_figure": "Figure 1A"}}, f, ensure_ascii=False)
        doc = Document()
        doc.add_paragraph("如图2-1和图2-9所示。")  # 图2-9 未注册
        issues = cq.check_figure_map_consistency(doc, project_root=tmp)
        msg = _messages(issues)
        assert "图2-9" in msg and "未在 figure_map" in msg, msg
        assert not any("图2-1" in i["message"] and "未在 figure_map" in i["message"] for i in issues)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_figure_map_consistency_all_referenced_no_warning():
    tmp = tempfile.mkdtemp()
    try:
        with open(os.path.join(tmp, "figure_map.json"), "w", encoding="utf-8") as f:
            json.dump({"图2-1": {"chapter": 2, "seq": 1, "source_figure": "Figure 1A"}}, f, ensure_ascii=False)
        doc = Document()
        doc.add_paragraph("如图2-1所示。")
        issues = cq.check_figure_map_consistency(doc, project_root=tmp)
        # 无 warning（无未注册图）；也无 unreferenced（图2-1 已引用）
        assert not any(i["level"] == "warning" for i in issues), _messages(issues)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


# ---------------------------------------------------------------------------
# check_cross_chapter_coherence
# ---------------------------------------------------------------------------

def _outline_project(root):
    outline = [
        {"chapter": 1, "title": "绪论"},
        {"chapter": 2, "title": "锚点甲", "core_argument": "论点甲"},
        {"chapter": 3, "title": "结论"},
    ]
    with open(os.path.join(root, "project_state.json"), "w", encoding="utf-8") as f:
        json.dump({"outline": outline}, f, ensure_ascii=False)


def test_cross_chapter_coherence_flags_missing_anchor():
    tmp = tempfile.mkdtemp()
    try:
        _outline_project(tmp)
        doc = Document()
        doc.add_heading("第1章 绪论", level=1)
        doc.add_paragraph("本文开展多项研究。")  # 未预告“锚点甲”
        doc.add_heading("第2章 锚点甲", level=1)
        doc.add_paragraph("本章围绕锚点甲展开。")
        doc.add_heading("第3章 结论", level=1)
        doc.add_paragraph("总体而言完成了工作。")  # 未跨章综合“锚点甲”
        issues = cq.check_cross_chapter_coherence(doc, project_root=tmp)
        msg = _messages(issues)
        assert "绪论未预告" in msg, msg
        assert "结论未跨章综合" in msg, msg
        assert all(i["level"] in ("warning", "info") for i in issues)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_cross_chapter_coherence_ok_when_anchors_present():
    tmp = tempfile.mkdtemp()
    try:
        _outline_project(tmp)
        doc = Document()
        doc.add_heading("第1章 绪论", level=1)
        doc.add_paragraph("本文预告了锚点甲的贡献。")
        doc.add_heading("第2章 锚点甲", level=1)
        doc.add_paragraph("承接上一章，本章围绕锚点甲展开。")
        doc.add_heading("第3章 结论", level=1)
        doc.add_paragraph("结论综合了锚点甲的核心发现。")
        issues = cq.check_cross_chapter_coherence(doc, project_root=tmp)
        assert issues == [], _messages(issues)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


# ---------------------------------------------------------------------------
# check_word_format_compliance
# ---------------------------------------------------------------------------

def test_word_format_page_layout_violation():
    # 默认 python-docx 文档为 Letter 尺寸，宽度 ≠ 21.0cm → 页面布局 error
    doc = Document()
    doc.add_paragraph("正文")
    issues = cq.check_word_format_compliance(doc)
    assert any(i["category"] == "页面布局" for i in issues), _messages(issues)


def test_word_format_a4_compliant_no_layout_issue():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    doc.add_paragraph("正文段落，无显式字号。")
    issues = cq.check_word_format_compliance(doc)
    assert not any(i["category"] == "页面布局" for i in issues), _messages(issues)


# ---------------------------------------------------------------------------
# check_table_format
# ---------------------------------------------------------------------------

def _set_cell_border(cell, edge, val="single", sz=12):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    el = tcBorders.find(qn(f"w:{edge}"))
    if el is None:
        el = OxmlElement(f"w:{edge}")
        tcBorders.append(el)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))


def test_table_format_no_borders_flagged():
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    issues = cq.check_table_format(doc)
    msg = _messages(issues)
    assert any(i["level"] == "error" for i in issues), msg
    assert "顶部边框" in msg or "缺少" in msg, msg


def test_table_format_three_line_compliant():
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    # 顶线 1.5pt(sz=12)，末行底线 1.5pt，表头分隔线 0.75pt(sz=6)，无竖线
    for cell in table.rows[0].cells:
        _set_cell_border(cell, "top", sz=12)
        _set_cell_border(cell, "bottom", sz=6)
    for cell in table.rows[-1].cells:
        _set_cell_border(cell, "bottom", sz=12)
    issues = cq.check_table_format(doc)
    # 不得有三线表 error（顶/底/表头分隔缺失）
    errs = [i for i in issues if i["level"] == "error" and i["category"] == "三线表"]
    assert errs == [], _messages(errs)


if __name__ == "__main__":
    test_full_structure_compliant_zero_error()
    print("OK full_structure 合规零错")
    test_full_structure_too_few_chapters_errors()
    print("OK full_structure 章节不足报错")
    test_full_structure_first_not_intro_errors()
    print("OK full_structure 首章非绪论报错")
    test_full_structure_last_not_conclusion_errors()
    print("OK full_structure 末章非总结报错")
    test_section_order_compliant()
    print("OK section_order 合规")
    test_section_order_violation_detected()
    print("OK section_order 乱序检出")
    test_figure_map_consistency_flags_unregistered()
    print("OK figure_map 未注册图检出")
    test_figure_map_consistency_all_referenced_no_warning()
    print("OK figure_map 全引用无告警")
    test_cross_chapter_coherence_flags_missing_anchor()
    print("OK cross_chapter 缺预告/综合检出")
    test_cross_chapter_coherence_ok_when_anchors_present()
    print("OK cross_chapter 锚点齐全零问题")
    test_word_format_page_layout_violation()
    print("OK word_format 非A4页面布局报错")
    test_word_format_a4_compliant_no_layout_issue()
    print("OK word_format A4 合规无布局问题")
    test_table_format_no_borders_flagged()
    print("OK table_format 无边框检出")
    test_table_format_three_line_compliant()
    print("OK table_format 三线表合规")
    print("ALL OK")
