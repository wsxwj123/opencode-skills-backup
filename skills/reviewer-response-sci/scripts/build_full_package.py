#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from datetime import date
from html import escape
from pathlib import Path
from typing import Any

from docx import Document


@dataclass
class CommentPair:
    reviewer: str
    section: str
    number: str
    comment_en: str
    reply_en: str


SECTION_MAP = {
    "major comments": "major",
    "minor comments": "minor",
    "major changes": "major",
    "minor changes": "minor",
}


def simplify_ws(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def read_docx_paragraphs(path: Path) -> list[dict[str, Any]]:
    doc = Document(str(path))
    rows: list[dict[str, Any]] = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        style_name = ""
        try:
            style_name = simplify_ws(getattr(getattr(p, "style", None), "name", "") or "")
        except Exception:
            style_name = ""
        rows.append({"paragraph_index": i, "text": simplify_ws(t), "style_name": style_name})
    return rows


def extract_email(text: str) -> str:
    m = re.search(r"Response Letter\n(.*?)\n\nReviewers' comments:", text, re.S)
    if m:
        return m.group(1).strip()
    return (
        "Dear Editor and Reviewers,\n\n"
        "Thank you for your constructive comments. We revised the manuscript point by point and addressed each concern below.\n\n"
        "Sincerely,\nThe Authors"
    )


def split_reviewer_blocks(text: str) -> dict[str, str]:
    hits = list(re.finditer(r"Reviewer\s*#\d+:", text))
    if not hits:
        return {"Reviewer #1": text}
    out: dict[str, str] = {}
    for i, m in enumerate(hits):
        start = m.start()
        end = hits[i + 1].start() if i + 1 < len(hits) else len(text)
        out[m.group(0).rstrip(":")] = text[start:end]
    return out


def split_sections(block: str) -> list[tuple[str, str]]:
    marks: list[tuple[int, str]] = []
    for key in ["Major Comments", "Minor Comments", "Major changes", "Minor changes"]:
        idx = block.find(key)
        if idx != -1:
            marks.append((idx, key))
    marks.sort(key=lambda x: x[0])
    if not marks:
        return [("general", block)]

    out: list[tuple[str, str]] = []
    for i, (idx, name) in enumerate(marks):
        start = idx + len(name)
        end = marks[i + 1][0] if i + 1 < len(marks) else len(block)
        out.append((SECTION_MAP[name.lower()], block[start:end]))
    return out


def parse_pairs(block: str) -> list[tuple[str, str, str]]:
    pairs: list[tuple[str, str, str]] = []
    pat = re.compile(r"\n\s*(\d+)\.\s+(.*?)(?=\n\s*\d+\.\s+|\Z)", re.S)
    for m in pat.finditer("\n" + block):
        body = simplify_ws(m.group(2))
        comment, reply = body, ""
        if " Reply:" in body:
            left, right = body.split(" Reply:", 1)
            comment, reply = simplify_ws(left), simplify_ws(right)
        pairs.append((m.group(1), comment, reply))
    return pairs


def collect_comment_pairs(text: str) -> list[CommentPair]:
    out: list[CommentPair] = []
    for reviewer, rb in split_reviewer_blocks(text).items():
        for section, sb in split_sections(rb):
            for num, comment, reply in parse_pairs(sb):
                if comment:
                    out.append(CommentPair(reviewer, section, num, comment, reply))
    return out


def zh_understanding(comment_en: str) -> str:
    txt = simplify_ws(comment_en)
    low = txt.lower()

    concerns: list[str] = []
    actions: list[str] = []

    if any(k in low for k in ["discrepancy", "mismatch", "inconsistent", "contradict"]):
        concerns.append("数据结果与文字表述（或图注）存在一致性风险")
        actions.append("逐项对齐原始数据、图表和正文结论")
    if any(k in low for k in ["figure", "fig.", "fig ", "image", "panel", "label", "legend"]):
        concerns.append("图像标注、分组标签或图文对应关系需要复核")
        actions.append("核对图号、面板标签与对应描述")
    if any(k in low for k in ["please provide", "please include", "add", "supplementary", "western blot"]):
        concerns.append("关键证据或方法细节说明不足")
        actions.append("补充必要的实验信息、证据或说明文本")
    if any(k in low for k in ["clarify", "explain", "interpretation"]):
        concerns.append("论证链条表达不够清晰")
        actions.append("增强因果逻辑并明确结论边界")
    if any(k in low for k in ["grammar", "proofread", "language editing"]):
        concerns.append("语言与学术表达规范性不足")
        actions.append("进行专业语言润色和术语统一")
    if any(k in low for k in ["format", "superscript", "notation"]):
        concerns.append("术语或格式书写不符合期刊规范")
        actions.append("按期刊规范统一格式（如上标、符号、缩写）")

    if not concerns:
        concerns.append("该意见要求作者就特定科学点提供更充分、可核查的修订")
    if not actions:
        actions.append("在对应段落给出定点修改并说明修改依据")

    anchors = extract_anchors(txt)
    anchor_txt = f"；重点定位：{', '.join(anchors[:3])}" if anchors else ""
    return f"审稿人核心关切：{'；'.join(concerns)}。建议处理：{'；'.join(actions)}{anchor_txt}。"


def _regex_replace_ci(text: str, pattern: str, repl: str) -> str:
    return re.sub(pattern, repl, text, flags=re.IGNORECASE)


def _figure_ref_to_zh(text: str) -> str:
    s = text
    # Pair references first to avoid partial replacements like "图s".
    s = _regex_replace_ci(
        s,
        r"\b(?:Figures?|Figs?\.?)[\s]*([S]?\d+[A-Za-z]?)\s*(?:and|&)\s*([S]?\d+[A-Za-z]?)\b",
        r"图\1和图\2",
    )
    s = _regex_replace_ci(s, r"\b(?:Figure|Fig\.?)\s*([S]?\d+[A-Za-z]?)\b", r"图\1")
    s = _regex_replace_ci(s, r"\b(Fig)(?=[S]?\d+[A-Za-z]?)", "图")
    return s


def _collapse_spaces_and_punct(text: str) -> str:
    s = text
    s = re.sub(r"\s+", " ", s).strip()
    # Punctuation normalization around Chinese text.
    s = s.replace(", ", "，")
    s = s.replace("; ", "；")
    s = s.replace(": ", "：")
    s = s.replace("( ", "(").replace(" )", ")")
    s = s.replace(" ，", "，").replace(" 。", "。")
    s = s.replace("。。", "。")
    s = s.replace("，，", "，")
    s = s.rstrip(". ")
    return s


def _translation_quality_tune(text: str) -> str:
    s = text
    # Light cleanup only: avoid introducing template-like Chinese function words.
    cleanup_patterns = [
        (r"\bthe\b", ""),
        (r"\ba\b", ""),
        (r"\ban\b", ""),
        (r"\bis\b", ""),
        (r"\bare\b", ""),
        (r"\bwas\b", ""),
        (r"\bwere\b", ""),
        (r"\bbe\b", ""),
        (r"\bthat\b", ""),
        (r"\bwhich\b", ""),
        (r"\bwhile\b", "尽管"),
        (r"\bhowever\b", "然而"),
        (r"\bmoreover\b", "此外"),
        (r"\btherefore\b", "因此"),
        (r"\bcould\b", ""),
        (r"\bwould\b", ""),
        (r"\bcan\b", ""),
        (r"\bnot\b", "不"),
        (r"\bbetween\b", "在…之间"),
        (r"\bfrom\b", "来自"),
        (r"\bto\b", ""),
        (r"\bby\b", "通过"),
    ]
    for pat, repl in cleanup_patterns:
        s = _regex_replace_ci(s, pat, repl)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _rule_based_zh_translation(comment_en: str) -> str:
    txt = simplify_ws(comment_en)
    s = txt.replace("，", ", ")
    s = _figure_ref_to_zh(s)

    phrase_patterns = [
        (r"\bthere is a serious discrepancy\b", "存在严重差异"),
        (r"\bthere is a clear mismatch\b", "存在明显不匹配"),
        (r"\bthere is\b", "存在"),
        (r"\bclearly indicated that\b", "清楚表明"),
        (r"\blower than\b", "低于"),
        (r"\bdirectly opposite to\b", "与…直接相反"),
        (r"\bstated conclusion\b", "所述结论"),
        (r"\bis observed in both\b", "在两者中均可观察到"),
        (r"\bsuggesting it is not an isolated error\b", "这表明这并非孤立错误"),
        (r"\bif the data is correct as plotted\b", "如果图示数据正确"),
        (r"\bit would imply that\b", "这将意味着"),
        (r"\brather than\b", "而不是"),
        (r"\bfundamentally contradict\b", "从根本上与…矛盾"),
        (r"\bcore narrative of the paper\b", "论文核心叙述"),
        (r"\bregarding\b", "关于"),
        (r"\burgently\b", "紧急地"),
        (r"\bdemonstrated that\b", "表明"),
        (r"\bshowed that\b", "显示"),
        (r"\bsignificantly attenuated\b", "显著减弱了"),
        (r"\bnuclear translocation\b", "核转位"),
        (r"\bprotein knockdown in\b", "在…中的蛋白敲低："),
        (r"\bplease\b", "请"),
        (r"\bplease provide\b", "请提供"),
        (r"\bplease include\b", "请补充"),
        (r"\bplease add\b", "请添加"),
        (r"\bplease clarify\b", "请澄清"),
        (r"\bcould the authors explain whether\b", "请作者解释是否"),
        (r"\bcould authors explain whether\b", "请作者解释是否"),
        (r"\bit is imperative that\b", "必须"),
        (r"\bit is strongly recommended to\b", "强烈建议"),
        (r"\bit is strongly recommended that\b", "强烈建议"),
        (r"\bit is recommended to\b", "建议"),
        (r"\bmust\b", "必须"),
        (r"\bshould\b", "应当"),
        (r"\bneed to\b", "需要"),
        (r"\bused for the experiments in\b", "用于以下实验："),
        (r"\bused for experiments in\b", "用于以下实验："),
        (r"\bdemonstrating the efficiency of\b", "以证明……效率："),
        (r"\bto further confirm\b", "以进一步确认"),
    ]
    for pat, repl in phrase_patterns:
        s = _regex_replace_ci(s, pat, repl)

    term_patterns = [
        (r"\bthe authors\b", "作者"),
        (r"\bauthors\b", "作者"),
        (r"\bmanuscript\b", "稿件"),
        (r"\bsupplementary information\b", "补充信息"),
        (r"\bsupplementary\b", "补充材料"),
        (r"\bcell viability\b", "细胞活力"),
        (r"\bprotective effect\b", "保护作用"),
        (r"\bwestern blot analysis\b", "Western blot分析"),
        (r"\bwestern blot\b", "Western blot"),
        (r"\bprotein knockdown\b", "蛋白敲低"),
        (r"\befficiency\b", "效率"),
        (r"\bmismatch\b", "不匹配"),
        (r"\binconsistent\b", "不一致"),
        (r"\bclarify\b", "澄清"),
        (r"\bcorrect(?:ed|ion)?\b", "更正"),
        (r"\bverify\b", "验证"),
        (r"\bprovide(?:d)?\b", "提供"),
        (r"\binclude(?:d)?\b", "补充"),
        (r"\badd(?:ed)?\b", "补充"),
        (r"\bproofread(?:ing)?\b", "校对"),
        (r"\blanguage editing\b", "语言润色"),
        (r"\bgrammar\b", "语法"),
        (r"\bresults\b", "结果"),
        (r"\bdata\b", "数据"),
        (r"\btext\b", "正文"),
        (r"\blabel(?:s|ing)?\b", "标签"),
        (r"\blegend(?:s)?\b", "图例"),
        (r"\bpanel(?:s)?\b", "面板"),
        (r"\bflow cytometry\b", "流式细胞术"),
        (r"\bgating strategy\b", "门控策略"),
        (r"\bscale bars\b", "比例尺"),
        (r"\bcellular uptake\b", "细胞摄取"),
        (r"\bconfocal microscopy\b", "共聚焦显微镜"),
        (r"\bserum stability\b", "血清稳定性"),
        (r"\bparticle concentration\b", "颗粒浓度"),
        (r"\bparticle size\b", "粒径"),
        (r"\bmorphology\b", "形态"),
        (r"\bmajor organs\b", "主要器官"),
        (r"\bliver\b", "肝脏"),
        (r"\bkidney\b", "肾脏"),
        (r"\bspleen\b", "脾脏"),
        (r"\blung\b", "肺脏"),
        (r"\bheart\b", "心脏"),
        (r"\bcorrection(?:s)?\b", "更正"),
        (r"\bcited\b", "引用"),
        (r"\breferenced\b", "引用"),
        (r"\bdiscussed\b", "讨论"),
    ]
    for pat, repl in term_patterns:
        s = _regex_replace_ci(s, pat, repl)

    small_words = [
        (r"\band\b", "和"),
        (r"\bor\b", "或"),
        (r"\bin\b", "在"),
        (r"\bfor\b", "用于"),
        (r"\bof\b", "的"),
        (r"\bwith\b", "伴随"),
    ]
    for pat, repl in small_words:
        s = _regex_replace_ci(s, pat, repl)

    s = _translation_quality_tune(s)
    s = _collapse_spaces_and_punct(s)
    if not s.endswith("。"):
        s += "。"
    return "中文翻译（直译）：" + s


def zh_translation(comment_en: str) -> str:
    # Final Chinese translation must be generated by the model, not script.
    return "【待AI直译：请在生成HTML前由模型完成中文直译】"


def auto_response(comment_en: str, section: str) -> str:
    low = comment_en.lower()
    if section == "minor":
        return (
            "Thank you for this helpful suggestion. "
            "We revised the relevant text or figure and rechecked consistency across the manuscript."
        )
    if any(k in low for k in ["discrepancy", "mismatch", "inconsistent", "contradict", "clarify"]):
        return (
            "Thank you for pointing out this critical issue. "
            "We rechecked the source data and corrected the text-figure alignment. "
            "The interpretation now matches the evidence."
        )
    if any(k in low for k in ["please provide", "please include", "recommended", "add"]):
        return (
            "Thank you for this constructive recommendation. "
            "We added the requested clarification and updated the corresponding section in the revised manuscript."
        )
    return (
        "Thank you for this valuable comment. "
        "We revised the manuscript accordingly and clarified the related scientific point."
    )


def auto_notes(section: str) -> tuple[list[str], list[str]]:
    if section == "minor":
        return (
            ["针对该条意见完成规范化修订（术语、图注、编号或排版），并逐项核对对应位置。"],
            ["进行全文一致性校对，避免同类细节问题重复出现。"],
        )
    return (
        ["围绕该条审稿意见的核心科学关切进行实质修订，确保结论与证据链一致。"],
        ["同步优化图号引用、术语表达与段落衔接，提升可读性和可核查性。"],
    )


def intent_en_from_comment(comment_en: str) -> str:
    return (
        "Interpretation: The reviewer is requesting a clear, evidence-aligned clarification and a concrete manuscript-level correction "
        "for this specific point."
    )


def response_zh_from_en(response_en: str) -> str:
    return "【待AI翻译：请在生成HTML前由模型完成中文回应】"


def excerpt_zh_from_en(excerpt_en: str) -> str:
    if excerpt_en.strip().lower() in {"none", "n/a", "无"}:
        return "无"
    if excerpt_en.startswith("Not provided by user"):
        return "无"
    return "对应中文修订说明：该段需在中文稿中同步更新，确保与英文修订段落语义一致。"


def extract_anchors(text: str) -> list[str]:
    anchors: list[str] = []
    patterns = [
        r"(?:Figure|Fig\.?)[ ]*S?\d+[A-Za-z]?",
        r"\bS\d+[A-Za-z]?\b",
        r"\bline\s*\d+\b",
        r"\bCD\d+[+]?\b",
        r"\b[A-Za-z]{2,}\d{1,3}\b",
    ]
    for pat in patterns:
        for m in re.finditer(pat, text, flags=re.IGNORECASE):
            v = m.group(0).strip()
            if len(v) >= 3:
                anchors.append(v)
    # de-dup preserve order
    seen = set()
    uniq = []
    for a in anchors:
        k = a.lower()
        if k not in seen:
            seen.add(k)
            uniq.append(a)
    return uniq[:15]


def is_figure_caption(text: str) -> bool:
    t = simplify_ws(text)
    return bool(re.match(r"^(Figure|Fig\.?)\s*S?\d+[A-Za-z]?[:.\s]", t, flags=re.IGNORECASE))


def parse_figure_ids(text: str) -> list[str]:
    ids = [m.group(0) for m in re.finditer(r"(?:Figure|Fig\.?)\s*S?\d+[A-Za-z]?", text, flags=re.IGNORECASE)]
    out: list[str] = []
    seen: set[str] = set()
    for x in ids:
        k = simplify_ws(x).lower()
        if k not in seen:
            seen.add(k)
            out.append(simplify_ws(x))
    return out


BACK_MATTER_PATTERNS = [
    r"references",
    r"acknowledg(?:e)?ments?",
    r"author contributions?",
    r"funding",
    r"conflicts? of interest",
    r"declaration of competing interest",
    r"data availability",
    r"ethics statement",
    r"参考文献",
    r"致谢",
    r"作者贡献",
    r"资金支持",
    r"利益冲突",
    r"数据可用性",
    r"伦理声明",
]


def _back_matter_regex() -> str:
    # Allow optional numeric prefix like "6. AUTHOR CONTRIBUTIONS".
    english = (
        r"references|acknowledg(?:e)?ments?|author contributions?|funding|"
        r"conflicts? of interest|declaration of competing interest|"
        r"data availability|ethics statement"
    )
    chinese = r"参考文献|致谢|作者贡献|资金支持|利益冲突|数据可用性|伦理声明"
    return (
        r"^(?:\d+\.\s*)?(?:(?:" + english + r")\b|(?:" + chinese + r"))"
    )


def is_back_matter_heading_text(text: str) -> bool:
    t = simplify_ws(text)
    if not t:
        return False
    return bool(re.match(_back_matter_regex(), t, flags=re.IGNORECASE))


def split_inline_back_matter_heading(text: str) -> tuple[str, str]:
    """
    Split 'References ...' style inline paragraph into heading + remainder.
    Returns (heading, remainder). If no inline split needed, remainder is empty.
    """
    t = simplify_ws(text)
    m = re.match(
        r"^(?:\d+\.\s*)?(references|acknowledg(?:e)?ments?|author contributions?|funding|conflicts? of interest|"
        r"declaration of competing interest|data availability|ethics statement|"
        r"参考文献|致谢|作者贡献|资金支持|利益冲突|数据可用性|伦理声明)"
        r"(?:\s*[:：.\-]\s*|\s+)(.+)$",
        t,
        flags=re.IGNORECASE,
    )
    if not m:
        return (t, "")
    return (simplify_ws(m.group(1)), simplify_ws(m.group(2)))


def split_row_by_inline_back_matter_headings(row: dict[str, Any]) -> list[dict[str, Any]]:
    """
    Split one paragraph row when back-matter headings appear inline in the middle
    of a long paragraph (common in exported DOCX content).
    """
    text = simplify_ws(row.get("text", ""))
    if not text:
        return []
    pat = re.compile(
        r"(?:^|(?<=\s))(?:\d+\.\s*)?"
        r"(references|acknowledg(?:e)?ments?|author contributions?|"
        r"conflicts? of interest|declaration of competing interest|"
        r"data availability|ethics statement|"
        r"参考文献|致谢|作者贡献|利益冲突|数据可用性|伦理声明)"
        r"(?=\s|[:：.\-]|$)",
        flags=re.IGNORECASE,
    )
    starts = [m.start() for m in pat.finditer(text)]
    if not starts:
        return [row]
    # Only split when heading marker appears after non-empty prefix text.
    starts = sorted(set(i for i in starts if i > 0))
    if not starts:
        return [row]

    points = [0] + starts + [len(text)]
    out: list[dict[str, Any]] = []
    for i in range(len(points) - 1):
        seg = simplify_ws(text[points[i]:points[i + 1]])
        if not seg:
            continue
        out.append(
            {
                "paragraph_index": row["paragraph_index"],
                "text": seg,
                "style_name": row.get("style_name", ""),
            }
        )
    return out or [row]


def is_section_heading_text(text: str, style_name: str = "") -> bool:
    t = simplify_ws(text)
    if not t:
        return False
    if is_figure_caption(t):
        return False
    if is_back_matter_heading_text(t):
        return True
    style_low = simplify_ws(style_name).lower()
    if style_low and ("heading" in style_low or "标题" in style_low):
        return True
    if len(t) > 180:
        return False
    if re.match(r"^\d+(?:\.\d+)*\s+", t):
        return True
    if re.match(
        r"^(abstract|keywords?|introduction|materials and methods|methods|results(?: and discussion)?|discussion|conclusion|references)\\b",
        t,
        flags=re.IGNORECASE,
    ):
        return True
    if re.match(r"^[A-Z][A-Za-z0-9\s,;:()\-/.]{1,120}$", t) and len(t.split()) <= 14 and not t.endswith('.'):
        return True
    return False


def split_sentences(text: str) -> list[str]:
    txt = simplify_ws(text)
    if not txt:
        return []
    # Keep punctuation-boundary sentence splitting lightweight and deterministic.
    parts = re.split(r"(?<=[。！？；.!?;])\s+", txt)
    return [p.strip() for p in parts if p.strip()]


def choose_sentence(comment_en: str, paragraph_text: str) -> tuple[int, str]:
    sents = split_sentences(paragraph_text)
    if not sents:
        return (0, paragraph_text)
    q = _tokenize_for_match(comment_en)
    best_idx = 0
    best_score = -1
    for i, s in enumerate(sents):
        score = len(q.intersection(_tokenize_for_match(s)))
        if score > best_score:
            best_score = score
            best_idx = i
    return (best_idx, sents[best_idx])


def unit_json_relpath(unit_id: str, kind: str) -> str:
    # unit_id like m-0007 / s-0003 -> manuscript_units/0007.json
    suffix = unit_id.split("-", 1)[1] if "-" in unit_id else unit_id
    folder = "manuscript_units" if kind == "m" else "si_units"
    return f"{folder}/{suffix}.json"


def atomize_docx_units(docx_path: Path, out_dir: Path, prefix: str) -> list[dict[str, Any]]:
    raw_rows = read_docx_paragraphs(docx_path)
    rows: list[dict[str, Any]] = []
    for r in raw_rows:
        rows.extend(split_row_by_inline_back_matter_headings(r))
    units: list[dict[str, Any]] = []
    out_dir.mkdir(parents=True, exist_ok=True)

    sections: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None

    def _new_section(heading: str, heading_para_idx: int, is_preamble: bool = False) -> dict[str, Any]:
        return {
            "heading": heading,
            "heading_paragraph_index": heading_para_idx,
            "paragraphs": [],
            "captions": [],
            "is_preamble": is_preamble,
        }

    def _flush_section(sec: dict[str, Any] | None) -> None:
        if not sec:
            return
        if not sec.get("paragraphs") and not sec.get("captions") and not sec.get("heading"):
            return
        sections.append(sec)

    for row in rows:
        text = row["text"]
        # Force back-matter sections (References/Acknowledgements/Author contributions, etc.)
        # to be isolated even if they are not styled as headings in DOCX.
        if is_back_matter_heading_text(text):
            heading, remainder = split_inline_back_matter_heading(text)
            _flush_section(current)
            current = _new_section(heading, row["paragraph_index"], is_preamble=False)
            if remainder:
                current["paragraphs"].append(
                    {
                        "paragraph_index": row["paragraph_index"],
                        "text": remainder,
                        "style_name": row.get("style_name", ""),
                    }
                )
            continue

        if is_section_heading_text(text, row.get("style_name", "")):
            _flush_section(current)
            current = _new_section(text, row["paragraph_index"], is_preamble=False)
            continue

        if current is None:
            current = _new_section("Front matter", row["paragraph_index"], is_preamble=True)

        if is_figure_caption(text):
            current["captions"].append(row)
        else:
            current["paragraphs"].append(row)

    _flush_section(current)

    unit_no = 0
    max_paragraphs_per_chunk = 12

    for sec_i, sec in enumerate(sections, start=1):
        heading = simplify_ws(sec.get("heading", ""))
        body_rows = sec.get("paragraphs", [])
        cap_rows = sec.get("captions", [])
        all_rows_sorted = sorted(body_rows + cap_rows, key=lambda x: x["paragraph_index"])

        # Split oversized sections into readable atomic chunks while keeping section context.
        chunks: list[list[dict[str, Any]]] = []
        cur: list[dict[str, Any]] = []
        for row in all_rows_sorted:
            is_cap = is_figure_caption(row["text"])
            if cur and len(cur) >= max_paragraphs_per_chunk and (is_cap or not is_figure_caption(cur[-1]["text"])):
                chunks.append(cur)
                cur = []
            cur.append(row)
        if cur:
            chunks.append(cur)
        if not chunks:
            chunks = [[]]

        chunk_total = len(chunks)
        for chunk_idx, chunk_rows in enumerate(chunks, start=1):
            unit_no += 1
            unit_id = f"{prefix}-{unit_no:04d}"

            heading_context = heading if heading and heading != "Front matter" else "Front matter"
            if chunk_total > 1:
                heading_context = f"{heading_context} [part {chunk_idx}/{chunk_total}]"

            combined_parts: list[str] = []
            if heading and heading != "Front matter":
                combined_parts.append(heading)
            combined_parts.extend(r["text"] for r in chunk_rows)
            combined_text = "\n".join([x for x in combined_parts if x]).strip()
            if not combined_text:
                continue

            para_indices = [r["paragraph_index"] for r in chunk_rows] or [sec.get("heading_paragraph_index", 0)]
            para_start = min(para_indices)
            para_end = max(para_indices)

            fig_refs: list[str] = []
            for txt in combined_parts:
                fig_refs.extend([m.group(0) for m in re.finditer(r"(?:Figure|Fig\\.?)[ ]*S?\\d+[A-Za-z]?|\\bS\\d+[A-Za-z]?\\b", txt, flags=re.IGNORECASE)])

            seen_refs: set[str] = set()
            anchors: list[str] = []
            for ref in fig_refs:
                k = simplify_ws(ref).lower()
                if k not in seen_refs:
                    seen_refs.add(k)
                    anchors.append(simplify_ws(ref))

            figure_captions = [r["text"] for r in chunk_rows if is_figure_caption(r["text"])]
            figure_ids: list[str] = []
            seen_ids: set[str] = set()
            for cap in figure_captions:
                for fid in parse_figure_ids(cap):
                    k = fid.lower()
                    if k not in seen_ids:
                        seen_ids.add(k)
                        figure_ids.append(fid)

            sents = split_sentences(combined_text)
            unit_type = "preamble_block" if sec.get("is_preamble") else "section_block"

            unit = {
                "unit_id": unit_id,
                "order": unit_no,
                "section_index": sec_i,
                "section_chunk_index": chunk_idx,
                "section_chunk_total": chunk_total,
                "paragraph_index": para_start,
                "paragraph_start_index": para_start,
                "paragraph_end_index": para_end,
                "text": combined_text,
                "sentences": [{"sentence_index": j, "text": st} for j, st in enumerate(sents)],
                "unit_type": unit_type,
                "section_unit_id": f"{prefix}-sec-{sec_i:04d}",
                "heading_context": heading_context,
                "section_title": heading if heading else heading_context,
                "title_en": heading_context,
                "source_paragraph_indices": para_indices,
                "anchors": anchors,
                "figure_ids": figure_ids,
                "figure_captions": figure_captions,
                "tags": {
                    "has_figure_ref": bool(anchors),
                    "length": len(combined_text),
                    "sentence_count": len(sents),
                    "paragraph_count": len(chunk_rows),
                    "is_heading": False,
                    "is_figure_caption": False,
                    "has_figure_caption": bool(figure_captions),
                },
            }
            units.append(unit)
            (out_dir / f"{unit_no:04d}.json").write_text(json.dumps(unit, ensure_ascii=False, indent=2), encoding="utf-8")

    return units


def find_linked_units(anchors: list[str], source_units: list[dict[str, Any]], max_hits: int = 8) -> list[str]:
    if not anchors:
        return []
    hits: list[str] = []
    for u in source_units:
        text = u["text"].lower()
        if any(a.lower() in text for a in anchors):
            hits.append(u["unit_id"])
        if len(hits) >= max_hits:
            break
    return hits


def _tokenize_for_match(text: str) -> set[str]:
    stop = {
        "the", "and", "for", "with", "that", "this", "from", "into", "were", "was", "are", "is",
        "have", "has", "had", "their", "your", "please", "should", "figure", "fig", "line",
        "authors", "author", "manuscript", "data", "results", "comment", "reply",
    }
    words = re.findall(r"[A-Za-z][A-Za-z0-9\\-]{2,}", text.lower())
    out = {w for w in words if w not in stop}
    return out


def keyword_link_units(query_text: str, source_units: list[dict[str, Any]], max_hits: int = 3) -> list[str]:
    q = _tokenize_for_match(query_text)
    if not q:
        return []
    scored: list[tuple[int, str]] = []
    for u in source_units:
        tokens = _tokenize_for_match(u.get("text", ""))
        score = len(q.intersection(tokens))
        if score > 0:
            scored.append((score, u["unit_id"]))
    scored.sort(key=lambda x: (-x[0], x[1]))
    return [uid for _, uid in scored[:max_hits]]


def build_comment_unit(order: int, row: CommentPair, src: dict[str, str], m_units: list[dict[str, Any]], s_units: list[dict[str, Any]]) -> dict[str, Any]:
    anchors = extract_anchors(row.comment_en)
    m_links = find_linked_units(anchors, m_units)
    s_links = find_linked_units(anchors, s_units) if s_units else []

    # Fallback: keyword overlap when no explicit anchors found
    if not m_links:
        m_links = keyword_link_units(row.comment_en, m_units, max_hits=3)
    if s_units and not s_links:
        s_links = keyword_link_units(row.comment_en, s_units, max_hits=2)

    # Last-resort safety link for strict gate mode
    if not m_links and m_units:
        m_links = [m_units[0]["unit_id"]]

    response_en = row.reply_en if row.reply_en else "[AI_FILL_REQUIRED] Response to reviewer in English."

    excerpt = "Not provided by user"
    original_excerpt = "无"
    revision_location = "无"
    atomic_location = {
        "manuscript_unit_id": "",
        "manuscript_unit_json": "",
        "manuscript_unit_type": "",
        "manuscript_paragraph_index": None,
        "manuscript_heading_context": "",
        "manuscript_sentence_index": None,
        "manuscript_sentence_text": "",
        "manuscript_figure_caption_unit_id": "",
        "manuscript_figure_caption_json": "",
        "manuscript_figure_caption_text": "",
        "si_unit_id": "",
        "si_unit_json": "",
        "si_unit_type": "",
        "si_paragraph_index": None,
        "si_heading_context": "",
        "si_figure_caption_unit_id": "",
        "si_figure_caption_json": "",
        "si_figure_caption_text": "",
    }
    if m_links:
        # fetch first linked manuscript paragraph as draft anchor text
        first = next((x for x in m_units if x["unit_id"] == m_links[0]), None)
        if first:
            original_excerpt = first["text"]
            # Keep original excerpt for traceability; revised text must be authored/confirmed by user.
            excerpt = "Not provided by user"
            heading = first.get("heading_context", "")
            if heading:
                revision_location = f"Section: {heading} | Paragraph index: {first.get('paragraph_index')}"
            else:
                revision_location = f"Paragraph index: {first.get('paragraph_index')}"
            sent_idx, sent_txt = choose_sentence(row.comment_en, first["text"])
            atomic_location.update(
                {
                    "manuscript_unit_id": first["unit_id"],
                    "manuscript_unit_json": unit_json_relpath(first["unit_id"], "m"),
                    "manuscript_unit_type": first.get("unit_type", ""),
                    "manuscript_paragraph_index": first.get("paragraph_index"),
                    "manuscript_heading_context": heading,
                    "manuscript_sentence_index": sent_idx,
                    "manuscript_sentence_text": sent_txt,
                }
            )
    if m_links:
        m_caption = next((x for x in m_units if x["unit_id"] in m_links and x.get("unit_type") == "figure_caption"), None)
        if not m_caption:
            m_caption = next((x for x in m_units if x["unit_id"] in m_links and x.get("figure_captions")), None)
        if m_caption:
            caption_text = m_caption.get("text", "")
            if not caption_text and m_caption.get("figure_captions"):
                caption_text = m_caption["figure_captions"][0]
            atomic_location.update(
                {
                    "manuscript_figure_caption_unit_id": m_caption["unit_id"],
                    "manuscript_figure_caption_json": unit_json_relpath(m_caption["unit_id"], "m"),
                    "manuscript_figure_caption_text": caption_text,
                }
            )
    if s_links:
        s_first = next((x for x in s_units if x["unit_id"] == s_links[0]), None)
        if s_first:
            atomic_location.update(
                {
                    "si_unit_id": s_first["unit_id"],
                    "si_unit_json": unit_json_relpath(s_first["unit_id"], "s"),
                    "si_unit_type": s_first.get("unit_type", ""),
                    "si_paragraph_index": s_first.get("paragraph_index"),
                    "si_heading_context": s_first.get("heading_context", ""),
                }
            )
    if s_links:
        s_caption = next((x for x in s_units if x["unit_id"] in s_links and x.get("unit_type") == "figure_caption"), None)
        if not s_caption:
            s_caption = next((x for x in s_units if x["unit_id"] in s_links and x.get("figure_captions")), None)
        if s_caption:
            caption_text = s_caption.get("text", "")
            if not caption_text and s_caption.get("figure_captions"):
                caption_text = s_caption["figure_captions"][0]
            atomic_location.update(
                {
                    "si_figure_caption_unit_id": s_caption["unit_id"],
                    "si_figure_caption_json": unit_json_relpath(s_caption["unit_id"], "s"),
                    "si_figure_caption_text": caption_text,
                }
            )

    image_change_required = bool(
        re.search(
            r"(figure|fig\.?|image|immunofluorescence|western blot|gating|scale bar|legend|panel)",
            row.comment_en,
            flags=re.IGNORECASE,
        )
    )

    actions = [{"action": "修改", "reason": "【待AI填写：添加/删除/修改及原因】"}]

    return {
        "unit_id": f"u-{order:03d}",
        "order": order,
        "reviewer": row.reviewer,
        "section": row.section,
        "comment_number": row.number,
        "title": f"{row.reviewer} | {row.section.upper()} | Comment {row.number}",
        "source": src,
        "links": {
            "anchors": anchors,
            "manuscript_unit_ids": m_links,
            "si_unit_ids": s_links,
        },
        "content": {
            "reviewer_comment_zh": "【待AI直译：请按原文直译，不要意译】",
            "reviewer_comment_en": row.comment_en,
            "reviewer_intent_zh": "【待AI填写：审稿意见中文理解】",
            "reviewer_intent_en": intent_en_from_comment(row.comment_en),
            "response_en": response_en,
            "response_zh": "【待AI翻译：Response to Reviewer中文对应】",
            "revision_location_en": revision_location,
            "atomic_location": atomic_location,
            "original_excerpt_en": original_excerpt,
            "revised_excerpt_en": "[AI_FILL_REQUIRED] Revised manuscript/SI text in English.",
            "revised_excerpt_zh": "【待AI翻译：修订后英文文本中文对应】",
            "modification_actions": actions,
            "notes_core_zh": ["【待AI填写：🔴核心修改说明】"],
            "notes_support_zh": ["【待AI填写：🟡辅助修改说明】"],
            "evidence": {
                "text": ["Not provided by user"],
                "image_change_required": image_change_required,
                "images": [{"src": "", "alt": "Image placeholder", "caption": "请替换为修订后图片（如有）"}],
                "table": {
                    "columns": ["Item", "Before", "After", "Evidence"],
                    "rows": [["Key correction", "Not provided by user", "Not provided by user", "Not provided by user"]],
                },
            },
        },
        "status": {
            "response_state": "draft",
            "excerpt_state": "needs_manual_revision" if m_links else "missing",
            "notes_state": "draft",
        },
    }


def build_email_unit(src: dict[str, str], email_text: str) -> dict[str, Any]:
    return {
        "unit_id": "u-000-email",
        "order": 0,
        "reviewer": "all",
        "section": "email",
        "comment_number": "0",
        "title": "回复审稿人的邮件",
        "source": src,
        "links": {"anchors": [], "manuscript_unit_ids": [], "si_unit_ids": []},
        "content": {
            "reviewer_comment_zh": "",
            "reviewer_comment_en": "",
            "response_en": email_text,
            "atomic_location": {},
            "revised_excerpt_en": "",
            "notes_core_zh": [],
            "notes_support_zh": [],
            "evidence": {"text": [], "images": [{"src": "", "alt": "", "caption": ""}], "table": {"columns": [""], "rows": [[""]]}}
        },
        "status": {"response_state": "final", "excerpt_state": "missing", "notes_state": "final"},
    }


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def build_index(units: list[dict[str, Any]]) -> dict[str, Any]:
    groups: dict[str, dict[str, list[dict[str, Any]]]] = {}
    for u in units:
        if u["section"] == "email":
            continue
        groups.setdefault(u["reviewer"], {}).setdefault(u["section"], []).append(u)

    def reviewer_sort_key(k: str) -> tuple[int, str]:
        m = re.search(r"#(\d+)", k)
        return (int(m.group(1)) if m else 999, k)

    reviewers = []
    rid = 0
    for reviewer in sorted(groups.keys(), key=reviewer_sort_key):
        rid += 1
        sec_nodes = []
        sid = 0
        for sec in sorted(groups[reviewer].keys(), key=lambda x: (0 if x == "major" else 1, x)):
            sid += 1
            items = sorted(groups[reviewer][sec], key=lambda x: (int(x["comment_number"]) if str(x["comment_number"]).isdigit() else 999, x["order"]))
            sec_nodes.append({"id": f"r{rid}-s{sid}", "label": sec.upper(), "items": [{"unit_id": x["unit_id"]} for x in items]})
        reviewers.append({"id": f"r{rid}", "label": reviewer, "sections": sec_nodes})

    return {"toc": {"root": "审稿回复目录", "reviewers": reviewers}}


def render_html(project_title: str, index_data: dict[str, Any], units: list[dict[str, Any]]) -> str:
    unit_by_id = {u["unit_id"]: u for u in units}

    def _pick_focus_paragraph(original_text: str, anchor_sentence: str) -> str:
        txt = (original_text or "").strip()
        if not txt:
            return "无"
        if txt == "无":
            return txt

        # Prefer paragraph-level context instead of dumping whole section.
        paras = [simplify_ws(p) for p in re.split(r"\n{2,}|\r\n\r\n", txt) if simplify_ws(p)]
        if not paras:
            paras = [simplify_ws(p) for p in re.split(r"\n+", txt) if simplify_ws(p)]
        if not paras:
            paras = [simplify_ws(txt)]

        anchor = simplify_ws(anchor_sentence or "")
        if not anchor or anchor == "N/A":
            return paras[0]

        # 1) direct containment (prefix) first
        prefix = anchor[:48]
        if prefix:
            for p in paras:
                if prefix in p:
                    return p

        # 2) token overlap fallback
        q = _tokenize_for_match(anchor)
        best = paras[0]
        best_score = -1
        for p in paras:
            s = len(q.intersection(_tokenize_for_match(p)))
            if s > best_score:
                best_score = s
                best = p
        return best

    toc_items: list[str] = []
    pages: list[str] = []

    toc_items.append('<li><button class="toc-btn toc-email active" data-target="page-u-000-email">回复审稿人的邮件</button></li>')
    email = unit_by_id["u-000-email"]
    pages.append(
        f'''<section id="page-u-000-email" class="page active"><h2>回复审稿人的邮件</h2>
<div class="card"><h3>English Email</h3><p>{escape(email['content']['response_en'])}</p></div>
<div class="card"><h3>中文说明</h3><p>本页为总回复邮件。后续目录按 Reviewer -> Major/Minor -> Comment 分层组织。</p></div>
</section>'''
    )

    for reviewer_node in index_data["toc"]["reviewers"]:
        reviewer_label = reviewer_node["label"]
        reviewer_children_id = f"toc-children-{reviewer_node['id']}"
        reviewer_buf: list[str] = [
            f'''<li class="toc-node reviewer-node">
<div class="toc-row">
  <button class="toc-btn reviewer-btn" data-target="page-{escape(reviewer_node["id"])}">{escape(reviewer_label)}</button>
  <button class="fold-btn" data-fold-target="{escape(reviewer_children_id)}" aria-expanded="true" title="折叠/展开">▾</button>
</div>'''
        ]
        reviewer_buf.append(f'<ul id="{escape(reviewer_children_id)}" class="toc-level-2">')
        pages.append(f'''<section id="page-{escape(reviewer_node['id'])}" class="page"><h2>{escape(reviewer_label)}</h2><div class="card"><p>选择 Major/Minor，再选择具体 Comment 查看详细内容。</p></div></section>''')

        for sec_node in reviewer_node["sections"]:
            sec_label = sec_node["label"]
            section_children_id = f"toc-children-{sec_node['id']}"
            reviewer_buf.append(
                f'''<li class="toc-node section-node">
<div class="toc-row">
  <button class="toc-btn section-btn" data-target="page-{escape(sec_node["id"])}">{escape(sec_label)}</button>
  <button class="fold-btn" data-fold-target="{escape(section_children_id)}" aria-expanded="true" title="折叠/展开">▾</button>
</div>'''
            )
            reviewer_buf.append(f'<ul id="{escape(section_children_id)}" class="toc-level-3">')
            pages.append(f'''<section id="page-{escape(sec_node['id'])}" class="page"><h2>{escape(reviewer_label)} - {escape(sec_label)}</h2><div class="card"><p>请选择该分组下具体 Comment。</p></div></section>''')

            for leaf in sec_node["items"]:
                uid = leaf["unit_id"]
                unit = unit_by_id[uid]
                severity_cls = "comment-major" if unit.get("section") == "major" else "comment-minor"
                reviewer_buf.append(f'<li><button class="toc-btn comment-btn {severity_cls}" data-target="page-{escape(uid)}">Comment {escape(str(unit["comment_number"]))}</button></li>')

                core_list = "".join(f"<li><span class='tag core'>核心</span> {escape(x)}</li>" for x in unit["content"]["notes_core_zh"])
                support_list = "".join(f"<li><span class='tag support'>辅助</span> {escape(x)}</li>" for x in unit["content"]["notes_support_zh"])
                ev_text = "<br/>".join(escape(x) for x in unit["content"]["evidence"]["text"]) or "Not provided by user"

                table_cols = unit["content"]["evidence"]["table"]["columns"]
                table_rows = unit["content"]["evidence"]["table"]["rows"]
                th_html = "".join(f"<th>{escape(c)}</th>" for c in table_cols)
                tr_html = "".join("<tr>" + "".join(f"<td>{escape(v)}</td>" for v in row) + "</tr>" for row in table_rows)

                img = unit["content"]["evidence"]["images"][0]
                anchors = unit.get("links", {}).get("anchors", [])
                mlinks = ", ".join(unit.get("links", {}).get("manuscript_unit_ids", [])) or "None"
                slinks = ", ".join(unit.get("links", {}).get("si_unit_ids", [])) or "None"

                response_zh = unit["content"].get("response_zh", "无")
                excerpt_zh = unit["content"].get("revised_excerpt_zh", "无")
                intent_zh = unit["content"].get("reviewer_intent_zh", "无")
                comment_zh = unit["content"].get("reviewer_comment_zh", "无")
                location_en = unit["content"].get("revision_location_en", "无")
                atomic_loc = unit["content"].get("atomic_location", {})
                original_en = unit["content"].get("original_excerpt_en", "无")
                quick_section = atomic_loc.get("manuscript_heading_context") or atomic_loc.get("si_heading_context") or "N/A"
                quick_para = (
                    str(atomic_loc.get("manuscript_paragraph_index"))
                    if atomic_loc.get("manuscript_paragraph_index") is not None
                    else (
                        f"SI:{atomic_loc.get('si_paragraph_index')}"
                        if atomic_loc.get("si_paragraph_index") is not None
                        else "N/A"
                    )
                )
                quick_sentence = atomic_loc.get("manuscript_sentence_text") or "N/A"
                quick_sentence = " ".join(str(quick_sentence).split())
                if len(quick_sentence) > 220:
                    quick_sentence = quick_sentence[:219] + "…"
                actions = unit["content"].get("modification_actions", [])
                focus_original_en = _pick_focus_paragraph(
                    original_text=original_en,
                    anchor_sentence=quick_sentence,
                )
                action_list = "".join(
                    f"<li><strong>{escape(x.get('action','修改'))}</strong>：{escape(x.get('reason',''))}</li>" for x in actions
                ) or "<li>无</li>"
                image_required = bool(unit["content"]["evidence"].get("image_change_required", False))
                image_block = ""
                if image_required:
                    img_src = str(img.get("src", "") or "").strip()
                    if img_src:
                        image_block = f"""<div class=\"img-placeholder\">图片修改占位符：请插入修订后图片（如 Figure 面板替换、图注同步更新）。</div>
<figure><img src=\"{escape(img_src)}\" alt=\"{escape(str(img.get('alt','image')))}\" /><figcaption>{escape(str(img.get('caption','')))}</figcaption></figure>"""
                    else:
                        image_block = """<div class=\"img-placeholder\">图片修改占位符：请插入修订后图片（如 Figure 面板替换、图注同步更新）。</div>"""

                pages.append(
                    f'''<section id="page-{escape(uid)}" class="page"><h2>{escape(unit['title'])}</h2>
<div class="card"><h3>1) 审稿意见与中文理解</h3>
<div class="stack-box"><h4>原始审稿意见（English）</h4><p>{escape(unit['content']['reviewer_comment_en'])}</p></div>
<div class="stack-box"><h4>审稿意见中文翻译（直译）</h4><p>{escape(comment_zh)}</p></div>
<div class="stack-box"><h4>审稿意见中文理解</h4><p>{escape(intent_zh)}</p></div>
</div>

<div class="card"><h3>2) Response to Reviewer（中英对照）</h3>
<div class="stack-box copy-box">
  <div class="box-head"><h4>中文回应</h4><button class="copy-btn" onclick="copyText('resp-zh-{escape(uid)}', this)">复制</button></div>
  <p id="resp-zh-{escape(uid)}">{escape(response_zh)}</p>
</div>
<div class="stack-box copy-box">
  <div class="box-head"><h4>English Translation</h4><button class="copy-btn" onclick="copyText('resp-en-{escape(uid)}', this)">复制</button></div>
  <p id="resp-en-{escape(uid)}">{escape(unit['content']['response_en'])}</p>
</div>
</div>

<div class="card"><h3>3) 可能需要修改的正文/附件内容（中英对照）</h3>
<div class="stack-box"><h4>快速定位（请先看这里）</h4>
<p><strong>Step 1 - 小节：</strong>{escape(str(quick_section))}</p>
<p><strong>Step 2 - 段落索引：</strong>{escape(str(quick_para))}</p>
<p><strong>Step 3 - Word检索锚句：</strong>{escape(str(quick_sentence))}</p>
<p><strong>Step 4 - 修改动作：</strong>{escape(' / '.join([x.get('action','修改') for x in actions]) if actions else '修改')}</p>
</div>
<details class="stack-box"><summary><strong>原子化定位（Atomic Location）</strong>（调试信息）</summary>
<p><strong>manuscript_unit_id:</strong> {escape(str(atomic_loc.get('manuscript_unit_id') or 'None'))}</p>
<p><strong>manuscript_unit_json:</strong> {escape(str(atomic_loc.get('manuscript_unit_json') or 'None'))}</p>
<p><strong>manuscript_unit_type:</strong> {escape(str(atomic_loc.get('manuscript_unit_type') or 'None'))}</p>
<p><strong>manuscript_paragraph_index:</strong> {escape(str(atomic_loc.get('manuscript_paragraph_index')) if atomic_loc.get('manuscript_paragraph_index') is not None else 'None')}</p>
<p><strong>manuscript_sentence_index:</strong> {escape(str(atomic_loc.get('manuscript_sentence_index')) if atomic_loc.get('manuscript_sentence_index') is not None else 'None')}</p>
<p><strong>manuscript_sentence_text:</strong> {escape(str(atomic_loc.get('manuscript_sentence_text') or 'None'))}</p>
<p><strong>manuscript_figure_caption_unit_id:</strong> {escape(str(atomic_loc.get('manuscript_figure_caption_unit_id') or 'None'))}</p>
<p><strong>manuscript_figure_caption_json:</strong> {escape(str(atomic_loc.get('manuscript_figure_caption_json') or 'None'))}</p>
<p><strong>manuscript_figure_caption_text:</strong> {escape(str(atomic_loc.get('manuscript_figure_caption_text') or 'None'))}</p>
<p><strong>si_unit_id:</strong> {escape(str(atomic_loc.get('si_unit_id') or 'None'))}</p>
<p><strong>si_unit_json:</strong> {escape(str(atomic_loc.get('si_unit_json') or 'None'))}</p>
<p><strong>si_unit_type:</strong> {escape(str(atomic_loc.get('si_unit_type') or 'None'))}</p>
<p><strong>si_figure_caption_unit_id:</strong> {escape(str(atomic_loc.get('si_figure_caption_unit_id') or 'None'))}</p>
<p><strong>si_figure_caption_json:</strong> {escape(str(atomic_loc.get('si_figure_caption_json') or 'None'))}</p>
<p><strong>si_figure_caption_text:</strong> {escape(str(atomic_loc.get('si_figure_caption_text') or 'None'))}</p>
</details>
<details class="stack-box copy-box">
  <summary><strong>Original Text (English, 对照)</strong>（已按锚句定位，仅展示相关段落）</summary>
  <div class="box-head" style="margin-top:8px"><h4>定位段落</h4><button class="copy-btn" onclick="copyText('orig-en-{escape(uid)}', this)">复制</button></div>
  <p id="orig-en-{escape(uid)}">{escape(focus_original_en)}</p>
</details>
<div class="stack-box copy-box">
  <div class="box-head"><h4>对应段落修订文本（English）</h4><button class="copy-btn" onclick="copyText('rev-en-{escape(uid)}', this)">复制</button></div>
  <p id="rev-en-{escape(uid)}">{escape(unit['content']['revised_excerpt_en'])}</p>
</div>
<div class="stack-box copy-box">
  <div class="box-head"><h4>修改后中文对照</h4><button class="copy-btn" onclick="copyText('rev-zh-{escape(uid)}', this)">复制</button></div>
  <p id="rev-zh-{escape(uid)}">{escape(excerpt_zh)}</p>
</div>
</div>

<div class="card"><h3>4) 修改说明（中文）</h3>
<div class="stack-box"><h4>细节修改（添加/删除/修改及原由）</h4><ul>{action_list}</ul></div>
<div class="stack-box"><h4>总结（核心/辅助）</h4><ul>{core_list}{support_list}</ul></div>
</div>
<div class="card"><h3>5) Evidence Attachments</h3><p><strong>Text:</strong><br/>{ev_text}</p>
<p><strong>Anchors:</strong> {escape(', '.join(anchors) if anchors else 'None')}</p>
<p><strong>Linked manuscript units:</strong> {escape(mlinks)}</p>
<p><strong>Linked SI units:</strong> {escape(slinks)}</p>
{image_block}
<div class="table-wrap"><table><thead><tr>{th_html}</tr></thead><tbody>{tr_html}</tbody></table></div></div>
</section>'''
                )

            reviewer_buf.append("</ul></li>")

        reviewer_buf.append("</ul></li>")
        toc_items.append("".join(reviewer_buf))

    today = date.today().isoformat()
    return f'''<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8" /><meta name="viewport" content="width=device-width, initial-scale=1" />
<title>{escape(project_title)}</title>
<style>
:root{{--bg:#f7f4ef;--panel:#fffdf9;--text:#20222a;--line:#ddd4c6;--accent:#1f3b4d;--accent-2:#8a5b2b;--major:#fde9e8;--minor:#edf5ff;--muted:#6c6a64;}}
*{{box-sizing:border-box}}body{{margin:0;font-family:"Source Han Serif SC","Songti SC","Times New Roman",serif;background:linear-gradient(160deg,#f3eee5 0%,#f7f4ef 52%,#eef4fb 100%);color:var(--text)}}
.layout{{--sidebar-w:370px;display:grid;grid-template-columns:var(--sidebar-w) 8px minmax(0,1fr);min-height:100vh}}.sidebar{{border-right:1px solid var(--line);background:linear-gradient(180deg,#fdfaf4 0%,#f5efe3 100%);position:sticky;top:0;height:100vh;overflow:auto;padding:18px}}
.resizer{{background:linear-gradient(180deg,#e8dccb 0%,#d9c9b2 100%);cursor:col-resize;position:sticky;top:0;height:100vh;border-left:1px solid #d6c7b2;border-right:1px solid #cdbba2}}
.resizer:hover,.resizer.dragging{{background:linear-gradient(180deg,#d7c2a4 0%,#c7ad8d 100%)}}
.sidebar h1{{margin:0 0 8px;font-size:1.16rem;color:var(--accent);letter-spacing:.04em}}.meta{{font-size:.84rem;color:var(--muted);margin-bottom:14px}}
.toc-level-1,.toc-level-2,.toc-level-3{{list-style:none;margin:0;padding-left:0}}.toc-level-2{{padding-left:12px;margin-top:4px}}.toc-level-3{{padding-left:14px;margin-top:4px}}
.toc-node{{margin-bottom:4px}}
.toc-row{{display:grid;grid-template-columns:1fr 28px;gap:6px;align-items:start}}
.toc-btn{{width:100%;text-align:left;border:1px solid var(--line);background:#fff;padding:8px 10px;border-radius:10px;cursor:pointer;font-size:.88rem;margin-bottom:6px;transition:all .2s ease}}
.toc-btn:hover{{transform:translateX(2px);border-color:#b59f83}}
.toc-btn.active{{background:linear-gradient(90deg,#fef1de 0%, #fffaf0 100%);border-color:#d2a872;color:#563a1f;font-weight:700}}
.fold-btn{{border:1px solid var(--line);background:#fff;padding:5px 0;border-radius:8px;cursor:pointer;font-size:.82rem;line-height:1;transition:all .2s ease;color:#5a4b37}}
.fold-btn:hover{{border-color:#b59f83;background:#fff7ec}}
.fold-hidden{{display:none}}
.reviewer-btn{{background:#fff8ee}}
.section-btn{{background:#f8fbff}}
.comment-major{{background:var(--major)}}
.comment-minor{{background:var(--minor)}}
.content{{padding:24px 26px}}.page{{display:none}}.page.active{{display:block}}
.card{{background:var(--panel);border:1px solid var(--line);border-radius:14px;padding:17px 20px;margin:0 0 15px;box-shadow:0 7px 24px rgba(80,65,45,.08)}}
h2{{color:var(--accent);margin:0 0 10px;font-size:1.4rem}}h3{{margin:0 0 10px;color:#2e4351}}h4{{margin:0 0 6px;font-size:.96rem;color:#30485a}}p{{white-space:pre-wrap;line-height:1.72}}
.stack-box{{border:1px solid #d4e0ee;background:#fbfdff;border-radius:12px;padding:12px 14px;margin-bottom:10px}}
.copy-box{{background:linear-gradient(180deg,#ffffff 0%,#f8fbff 100%)}}
.box-head{{display:flex;align-items:center;justify-content:space-between;gap:10px;margin-bottom:6px}}
.copy-btn{{border:1px solid #b58c5c;background:linear-gradient(180deg,#fffdf8 0%,#f7e9d6 100%);color:#4c3118;border-radius:999px;padding:4px 11px;font-size:.78rem;font-weight:700;cursor:pointer}}
.copy-btn:hover{{border-color:#8f6637}}
.mark-add{{font-weight:700;text-decoration:underline;text-decoration-thickness:2px;text-underline-offset:2px;background:#fff3c4;padding:0 2px;border-radius:3px}}
.tag{{display:inline-block;padding:2px 8px;border-radius:999px;color:#fff;font-size:.8rem;font-weight:600;margin-right:6px}}.core{{background:#8f1f18}}.support{{background:#8a5b2b}}
.img-placeholder{{border:1px dashed #c5a171;background:#fff7ed;border-radius:10px;padding:10px 12px;color:#7a4b1b;margin:8px 0}}
figure{{margin:8px 0;border:1px dashed var(--line);padding:10px;border-radius:10px;background:#fafcff}}img{{max-width:100%;height:auto;min-height:80px;background:#fff;border:1px solid var(--line);border-radius:7px}}
.table-wrap{{overflow-x:auto}}table{{width:100%;border-collapse:collapse}}th,td{{border:1px solid var(--line);padding:8px 10px;text-align:left;vertical-align:top}}th{{background:#edf4ff}}
@media (max-width:980px){{.layout{{grid-template-columns:1fr}}.resizer{{display:none}}.sidebar{{position:relative;height:auto;border-right:none;border-bottom:1px solid var(--line)}}}}
</style></head><body>
<div class="layout" id="layout-root"><aside class="sidebar"><h1>审稿回复目录</h1><div class="meta">{escape(project_title)} | {today}</div><ul id="toc-root" class="toc-level-1">{''.join(toc_items)}</ul></aside><div id="resizer" class="resizer" aria-label="拖动调整目录宽度" role="separator"></div>
<main id="content-root" class="content">{''.join(pages)}</main></div>
<script>
const btns=document.querySelectorAll('.toc-btn');const pages=document.querySelectorAll('.page');
btns.forEach(btn=>{{btn.addEventListener('click',()=>{{btns.forEach(b=>b.classList.remove('active'));pages.forEach(p=>p.classList.remove('active'));btn.classList.add('active');const target=document.getElementById(btn.dataset.target);if(target)target.classList.add('active');window.scrollTo({{top:0,behavior:'smooth'}});}});}});
const foldBtns=document.querySelectorAll('.fold-btn');
foldBtns.forEach(btn=>{{btn.addEventListener('click',(e)=>{{e.stopPropagation();const target=document.getElementById(btn.dataset.foldTarget);if(!target)return;const collapsed=target.classList.toggle('fold-hidden');btn.innerText=collapsed?'▸':'▾';btn.setAttribute('aria-expanded',String(!collapsed));}});}});
const layoutRoot=document.getElementById('layout-root');
const resizer=document.getElementById('resizer');
const minW=260,maxW=620,storeKey='reviewer_sidebar_width_v1';
const applyWidth=(w)=>{{layoutRoot.style.setProperty('--sidebar-w',`${{w}}px`);}};
const saved=parseInt(localStorage.getItem(storeKey)||'',10);
if(Number.isFinite(saved))applyWidth(Math.max(minW,Math.min(maxW,saved)));
let dragging=false;
const onMove=(clientX)=>{{
  if(!dragging) return;
  const bounds=layoutRoot.getBoundingClientRect();
  const w=Math.max(minW,Math.min(maxW,clientX-bounds.left));
  applyWidth(w);
}};
const endDrag=()=>{{if(!dragging)return;dragging=false;resizer.classList.remove('dragging');const w=parseInt(getComputedStyle(layoutRoot).getPropertyValue('--sidebar-w'));if(Number.isFinite(w))localStorage.setItem(storeKey,String(w));}};
resizer.addEventListener('pointerdown',(e)=>{{dragging=true;resizer.classList.add('dragging');resizer.setPointerCapture(e.pointerId);e.preventDefault();}});
resizer.addEventListener('pointermove',(e)=>onMove(e.clientX));
resizer.addEventListener('pointerup',endDrag);
resizer.addEventListener('pointercancel',endDrag);
window.addEventListener('mousemove',(e)=>onMove(e.clientX));
window.addEventListener('mouseup',endDrag);
window.copyText = async (id, btn) => {{
  const el = document.getElementById(id);
  if (!el) return;
  const text = el.innerText || el.textContent || '';
  try {{
    await navigator.clipboard.writeText(text);
    const old = btn.innerText;
    btn.innerText = '已复制';
    setTimeout(() => btn.innerText = old, 1200);
  }} catch (e) {{
    const old = btn.innerText;
    btn.innerText = '复制失败';
    setTimeout(() => btn.innerText = old, 1200);
  }}
}};
</script></body></html>'''


def main() -> int:
    parser = argparse.ArgumentParser(description="Build full reviewer package: atomic json + hierarchical html")
    parser.add_argument("--comments", required=True)
    parser.add_argument("--manuscript", required=True)
    parser.add_argument("--si", default="")
    parser.add_argument("--project-root", required=True)
    parser.add_argument("--output-html", required=True)
    parser.add_argument("--title", default="Reviewer Response Full Package")
    args = parser.parse_args()

    project_root = Path(args.project_root)
    units_dir = project_root / "units"
    manuscript_units_dir = project_root / "manuscript_units"
    si_units_dir = project_root / "si_units"

    comments_text = read_docx(Path(args.comments))
    email_text = extract_email(comments_text)
    rows = collect_comment_pairs(comments_text)

    manuscript_units = atomize_docx_units(Path(args.manuscript), manuscript_units_dir, "m")
    si_units = atomize_docx_units(Path(args.si), si_units_dir, "s") if args.si else []

    src = {
        "comments_docx": str(Path(args.comments).resolve()),
        "manuscript_docx": str(Path(args.manuscript).resolve()),
        "si_docx": str(Path(args.si).resolve()) if args.si else "",
    }

    units: list[dict[str, Any]] = []
    email_unit = build_email_unit(src, email_text)
    write_json(units_dir / "000_email.json", email_unit)
    units.append(email_unit)

    order = 1
    for row in rows:
        unit = build_comment_unit(order, row, src, manuscript_units, si_units)
        safe_reviewer = row.reviewer.replace(" ", "").replace("#", "")
        num = f"{int(row.number):02d}" if row.number.isdigit() else row.number
        fname = f"{order:03d}_{safe_reviewer}_{row.section}_{num}.json"
        write_json(units_dir / fname, unit)
        units.append(unit)
        order += 1

    index_data = build_index(units)
    write_json(project_root / "index.json", index_data)

    project_state = {
        "project_title": args.title,
        "generated_at": date.today().isoformat(),
        "counts": {
            "total_units": len(units),
            "comment_units": len(units) - 1,
            "manuscript_units": len(manuscript_units),
            "si_units": len(si_units),
        },
        "paths": {
            "units_dir": str(units_dir.resolve()),
            "manuscript_units_dir": str(manuscript_units_dir.resolve()),
            "si_units_dir": str(si_units_dir.resolve()),
            "index_json": str((project_root / "index.json").resolve()),
            "output_html": str(Path(args.output_html).resolve()),
        },
    }
    write_json(project_root / "project_state.json", project_state)

    html = render_html(project_title=args.title, index_data=index_data, units=units)
    out_html = Path(args.output_html)
    out_html.parent.mkdir(parents=True, exist_ok=True)
    out_html.write_text(html, encoding="utf-8")

    print(f"WROTE project_state: {project_root / 'project_state.json'}")
    print(f"WROTE index: {project_root / 'index.json'}")
    print(f"WROTE comment units: {len(units)} in {units_dir}")
    print(f"WROTE manuscript units: {len(manuscript_units)} in {manuscript_units_dir}")
    print(f"WROTE si units: {len(si_units)} in {si_units_dir}")
    print(f"WROTE html: {out_html}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
