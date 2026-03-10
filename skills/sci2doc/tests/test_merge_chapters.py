#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""merge_chapters.py 单元测试"""

import json
import sys, os, unittest, tempfile, shutil
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

_SCRIPTS = str(Path(__file__).resolve().parent.parent / "scripts")
if _SCRIPTS not in sys.path:
    sys.path.insert(0, _SCRIPTS)

from merge_chapters import (
    extract_chapter_number,
    merge_docx_files,
    generate_toc,
    add_header_footer,
    resolve_merge_order,
)


def _create_chapter_docx(path, title, body="正文内容。"):
    """Helper: create a minimal docx with a heading and body paragraph."""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(body)
    doc.save(str(path))


def _get_pg_num_type(section):
    pg_num = section._sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        return None
    return {
        "fmt": pg_num.get(qn("w:fmt")),
        "start": pg_num.get(qn("w:start")),
    }


class TestExtractChapterNumber(unittest.TestCase):
    def test_normal(self):
        self.assertEqual(extract_chapter_number("第2章_材料方法.docx"), 2)

    def test_double_digit(self):
        self.assertEqual(extract_chapter_number("第12章_总结.docx"), 12)

    def test_no_match(self):
        self.assertEqual(extract_chapter_number("abstract.docx"), 999)

    def test_chinese_only(self):
        self.assertEqual(extract_chapter_number("第1章.docx"), 1)


class TestMergeDocxFiles(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir)

    def test_empty_list(self):
        result = merge_docx_files([], os.path.join(self.tmpdir, "out.docx"))
        self.assertFalse(result["success"])
        self.assertIn("空", result["error"])

    def test_single_file(self):
        f1 = os.path.join(self.tmpdir, "第1章_绪论.docx")
        _create_chapter_docx(f1, "第1章 绪论")
        out = os.path.join(self.tmpdir, "merged.docx")
        result = merge_docx_files([f1], out)
        self.assertTrue(result["success"])
        self.assertEqual(result["merged_files"], 1)
        self.assertTrue(os.path.exists(out))

    def test_multiple_files_sorted(self):
        f3 = os.path.join(self.tmpdir, "第3章_结果.docx")
        f1 = os.path.join(self.tmpdir, "第1章_绪论.docx")
        f2 = os.path.join(self.tmpdir, "第2章_方法.docx")
        _create_chapter_docx(f1, "第1章 绪论")
        _create_chapter_docx(f2, "第2章 方法")
        _create_chapter_docx(f3, "第3章 结果")
        out = os.path.join(self.tmpdir, "merged.docx")
        # Pass in unsorted order
        result = merge_docx_files([f3, f1, f2], out)
        self.assertTrue(result["success"])
        self.assertEqual(result["merged_files"], 3)
        # Verify sorted order
        self.assertEqual(result["file_list"][0], "第1章_绪论.docx")
        self.assertEqual(result["file_list"][1], "第2章_方法.docx")
        self.assertEqual(result["file_list"][2], "第3章_结果.docx")

    def test_output_dir_created(self):
        f1 = os.path.join(self.tmpdir, "第1章.docx")
        _create_chapter_docx(f1, "第1章")
        out = os.path.join(self.tmpdir, "subdir", "nested", "out.docx")
        result = merge_docx_files([f1], out)
        self.assertTrue(result["success"])
        self.assertTrue(os.path.exists(out))


class TestGenerateToc(unittest.TestCase):
    def test_toc_inserted(self):
        doc = Document()
        doc.add_heading("第1章 绪论", level=1)
        doc.add_paragraph("正文")
        doc.add_heading("1.1 背景", level=2)
        doc.add_paragraph("更多正文")
        generate_toc(doc)
        # First paragraph should be the TOC title
        first_text = doc.paragraphs[0].text
        self.assertIn("目", first_text)
        self.assertIn("录", first_text)

    def test_toc_contains_headings(self):
        doc = Document()
        doc.add_heading("第1章 绪论", level=1)
        doc.add_heading("1.1 背景", level=2)
        generate_toc(doc)
        # Collect all paragraph texts
        texts = [p.text for p in doc.paragraphs]
        combined = "\n".join(texts)
        self.assertIn("第1章 绪论", combined)
        self.assertIn("1.1 背景", combined)


class TestAddHeaderFooter(unittest.TestCase):
    def test_header_content(self):
        doc = Document()
        doc.add_paragraph("正文")
        add_header_footer(doc, "基于深度学习的图像分割研究")
        header_text = doc.sections[0].header.paragraphs[0].text
        self.assertIn("中南大学", header_text)
        self.assertIn("基于深度学习的图像分割研究", header_text)

    def test_custom_university(self):
        doc = Document()
        doc.add_paragraph("正文")
        add_header_footer(doc, "论文标题", university_name="清华大学")
        header_text = doc.sections[0].header.paragraphs[0].text
        self.assertIn("清华大学", header_text)

    def test_footer_has_page_field(self):
        doc = Document()
        doc.add_paragraph("正文")
        add_header_footer(doc, "论文标题")
        footer = doc.sections[0].footer
        footer_xml = footer._element.xml
        self.assertIn("PAGE", footer_xml)

    def test_footer_centered(self):
        doc = Document()
        doc.add_paragraph("正文")
        add_header_footer(doc, "论文标题")
        footer_para = doc.sections[0].footer.paragraphs[0]
        self.assertEqual(footer_para.alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_page_numbering_uses_profile_config(self):
        with tempfile.TemporaryDirectory() as td:
            profile_path = Path(td) / "thesis_profile.json"
            profile_path.write_text(
                json.dumps(
                    {
                        "format_profile": {
                            "page_numbering": {
                                "front_matter": {"format": "upperRoman", "start": 3},
                                "body": {"format": "decimal", "start": 5},
                                "back_matter": {"format": "lowerLetter", "start": 2},
                            }
                        }
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
            doc = Document()
            doc.add_heading("摘  要", level=1)
            doc.add_paragraph("前置内容")
            doc.add_heading("第1章 绪论", level=1)
            doc.add_paragraph("正文")
            doc.add_heading("参考文献", level=1)
            doc.add_paragraph("[1] ref")

            add_header_footer(doc, "论文标题", project_root=td)

            self.assertEqual(_get_pg_num_type(doc.sections[0]), {"fmt": "upperRoman", "start": "3"})
            self.assertEqual(_get_pg_num_type(doc.sections[1]), {"fmt": "decimal", "start": "5"})
            self.assertEqual(_get_pg_num_type(doc.sections[2]), {"fmt": "lowerLetter", "start": "2"})


class TestResolveMergeOrder(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.tmpdir)

    def test_chapters_sorted(self):
        for name in ["第3章.docx", "第1章.docx", "第2章.docx"]:
            _create_chapter_docx(os.path.join(self.tmpdir, name), name)
        result = resolve_merge_order(self.tmpdir)
        basenames = [os.path.basename(p) for p in result]
        self.assertEqual(basenames, ["第1章.docx", "第2章.docx", "第3章.docx"])

    def test_front_matter_first(self):
        cover = os.path.join(self.tmpdir, "cover.docx")
        _create_chapter_docx(cover, "封面")
        _create_chapter_docx(os.path.join(self.tmpdir, "第1章.docx"), "第1章")
        result = resolve_merge_order(self.tmpdir, cover=cover)
        self.assertEqual(os.path.basename(result[0]), "cover.docx")
        self.assertEqual(os.path.basename(result[1]), "第1章.docx")

    def test_missing_front_matter_skipped(self):
        _create_chapter_docx(os.path.join(self.tmpdir, "第1章.docx"), "第1章")
        result = resolve_merge_order(self.tmpdir, cover="/nonexistent/cover.docx")
        self.assertEqual(len(result), 1)

    def test_no_duplicates(self):
        f = os.path.join(self.tmpdir, "第1章.docx")
        _create_chapter_docx(f, "第1章")
        # Pass same file as cover and it's also in the dir
        result = resolve_merge_order(self.tmpdir, cover=f)
        paths = [os.path.abspath(p) for p in result]
        self.assertEqual(len(paths), len(set(paths)))

    def test_skips_temp_files(self):
        _create_chapter_docx(os.path.join(self.tmpdir, "第1章.docx"), "第1章")
        _create_chapter_docx(os.path.join(self.tmpdir, "~$第2章.docx"), "temp")
        result = resolve_merge_order(self.tmpdir)
        basenames = [os.path.basename(p) for p in result]
        self.assertNotIn("~$第2章.docx", basenames)


if __name__ == "__main__":
    unittest.main()
