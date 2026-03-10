#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""markdown_to_docx.py 单元测试"""

import sys, os, unittest
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ensure scripts/ importable
_SCRIPTS = str(Path(__file__).resolve().parent.parent / "scripts")
if _SCRIPTS not in sys.path:
    sys.path.insert(0, _SCRIPTS)

from markdown_to_docx import (
    _parse_pipe_row,
    _is_separator_row,
    parse_markdown_line,
    apply_three_line_table_borders,
    set_run_font,
    _set_cell_border,
)


class TestParsePipeRow(unittest.TestCase):
    def test_normal_row(self):
        self.assertEqual(_parse_pipe_row("| A | B | C |"), ["A", "B", "C"])

    def test_whitespace_trimmed(self):
        self.assertEqual(_parse_pipe_row("|  hello  |  world  |"), ["hello", "world"])

    def test_escaped_pipe(self):
        result = _parse_pipe_row("| a\\|b | c |")
        self.assertEqual(result, ["a|b", "c"])

    def test_not_a_pipe_row(self):
        self.assertIsNone(_parse_pipe_row("just normal text"))

    def test_empty_cells(self):
        self.assertEqual(_parse_pipe_row("|  |  |"), ["", ""])

    def test_single_cell(self):
        self.assertEqual(_parse_pipe_row("| only |"), ["only"])


class TestIsSeparatorRow(unittest.TestCase):
    def test_basic_separator(self):
        self.assertTrue(_is_separator_row("|---|---|"))

    def test_aligned_separator(self):
        self.assertTrue(_is_separator_row("| :---: | ---: |"))

    def test_not_separator(self):
        self.assertFalse(_is_separator_row("| A | B |"))

    def test_empty_string(self):
        self.assertFalse(_is_separator_row(""))


class TestParseMarkdownLine(unittest.TestCase):
    def test_empty_line(self):
        self.assertEqual(parse_markdown_line(""), ("empty", "", 0))

    def test_whitespace_only(self):
        self.assertEqual(parse_markdown_line("   "), ("empty", "", 0))

    def test_heading1(self):
        self.assertEqual(parse_markdown_line("# 第一章 绪论"), ("heading1", "第一章 绪论", 1))

    def test_heading2(self):
        self.assertEqual(parse_markdown_line("## 1.1 背景"), ("heading2", "1.1 背景", 2))

    def test_heading3(self):
        self.assertEqual(parse_markdown_line("### 1.1.1 详细"), ("heading3", "1.1.1 详细", 3))

    def test_heading2_not_heading3(self):
        """## should not be parsed as heading3"""
        t, _, lvl = parse_markdown_line("## Title")
        self.assertEqual(t, "heading2")
        self.assertEqual(lvl, 2)

    def test_figure_placeholder(self):
        line = "[图 2-1：实验结果]"
        self.assertEqual(parse_markdown_line(line), ("figure", line, 0))

    def test_table_placeholder(self):
        line = "[表 3-2：参数对比]"
        self.assertEqual(parse_markdown_line(line), ("table", line, 0))

    def test_normal_paragraph(self):
        t, content, lvl = parse_markdown_line("这是一段正文。")
        self.assertEqual(t, "paragraph")
        self.assertEqual(lvl, 0)

    def test_figure_with_colon_variant(self):
        """Test colon : variant (not ：)"""
        line = "[图 1-1:标题]"
        self.assertEqual(parse_markdown_line(line)[0], "figure")


class TestApplyThreeLineTableBorders(unittest.TestCase):
    def _make_table(self, rows=3, cols=2):
        doc = Document()
        return doc.add_table(rows=rows, cols=cols)

    def _get_border(self, cell, edge):
        tc = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            return None
        borders = tcPr.find(qn('w:tcBorders'))
        if borders is None:
            return None
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            return None
        return {
            'sz': el.get(qn('w:sz')),
            'val': el.get(qn('w:val')),
            'color': el.get(qn('w:color')),
        }

    def test_top_row_has_thick_top(self):
        table = self._make_table()
        apply_three_line_table_borders(table)
        b = self._get_border(table.rows[0].cells[0], 'top')
        self.assertEqual(b['sz'], '12')
        self.assertEqual(b['val'], 'single')

    def test_header_bottom_thin(self):
        table = self._make_table()
        apply_three_line_table_borders(table, header_rows=1)
        b = self._get_border(table.rows[0].cells[0], 'bottom')
        self.assertEqual(b['sz'], '4')
        self.assertEqual(b['val'], 'single')

    def test_last_row_thick_bottom(self):
        table = self._make_table(rows=4)
        apply_three_line_table_borders(table)
        b = self._get_border(table.rows[3].cells[0], 'bottom')
        self.assertEqual(b['sz'], '12')

    def test_middle_row_no_horizontal(self):
        table = self._make_table(rows=4)
        apply_three_line_table_borders(table)
        # row 1 (not header, not last) should have top=none, bottom=none
        top = self._get_border(table.rows[1].cells[0], 'top')
        bottom = self._get_border(table.rows[1].cells[0], 'bottom')
        self.assertEqual(top['val'], 'none')
        self.assertEqual(bottom['val'], 'none')

    def test_no_vertical_lines(self):
        table = self._make_table()
        apply_three_line_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                left = self._get_border(cell, 'left')
                right = self._get_border(cell, 'right')
                self.assertEqual(left['val'], 'none')
                self.assertEqual(right['val'], 'none')

    def test_two_header_rows(self):
        table = self._make_table(rows=4)
        apply_three_line_table_borders(table, header_rows=2)
        # row 1 bottom should be thin (header separator)
        b = self._get_border(table.rows[1].cells[0], 'bottom')
        self.assertEqual(b['sz'], '4')
        self.assertEqual(b['val'], 'single')
        # row 0 bottom should be none (not the header separator)
        b0 = self._get_border(table.rows[0].cells[0], 'bottom')
        self.assertEqual(b0['val'], 'none')


class TestSetRunFont(unittest.TestCase):
    def _make_run(self, text="测试test"):
        doc = Document()
        para = doc.add_paragraph()
        return para.add_run(text)

    def test_latin_font(self):
        run = self._make_run()
        set_run_font(run, 'Times New Roman', 'SimSun', 12)
        self.assertEqual(run.font.name, 'Times New Roman')

    def test_east_asia_font(self):
        run = self._make_run()
        set_run_font(run, 'Times New Roman', 'SimHei', 14)
        rfonts = run._element.rPr.rFonts
        self.assertEqual(rfonts.get(qn('w:eastAsia')), 'SimHei')

    def test_font_size(self):
        run = self._make_run()
        set_run_font(run, 'Arial', 'SimSun', 10.5)
        self.assertEqual(run.font.size, Pt(10.5))

    def test_bold_true(self):
        run = self._make_run()
        set_run_font(run, 'Arial', 'SimSun', 12, bold=True)
        self.assertTrue(run.font.bold)

    def test_bold_false(self):
        run = self._make_run()
        set_run_font(run, 'Arial', 'SimSun', 12, bold=False)
        self.assertFalse(run.font.bold)

    def test_bold_none_unchanged(self):
        run = self._make_run()
        set_run_font(run, 'Arial', 'SimSun', 12, bold=None)
        # bold should remain unset (None)
        self.assertIsNone(run.font.bold)


if __name__ == "__main__":
    unittest.main()
