#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 转 Word 工具（默认应用中南大学样式，支持自定义格式配置）

功能：
1. 解析 Markdown 文本
2. 转换为 Word 文档
3. 自动应用论文格式配置（默认中南大学博士论文样式）
4. 处理图表占位符

作者：Sci2Doc Team
日期：2024-03-15
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import json
import re
import sys
import os

try:
    from thesis_profile import build_format_render_context, load_profile, normalize_style_profile
    from shared_utils import infer_project_root_for_profile
    from abbreviation_registry import load_registry, get_all as get_all_abbreviations
except Exception:  # pragma: no cover
    _script_dir = os.path.dirname(os.path.abspath(__file__))
    if _script_dir not in sys.path:
        sys.path.insert(0, _script_dir)
    try:
        from thesis_profile import build_format_render_context, load_profile, normalize_style_profile
        from shared_utils import infer_project_root_for_profile
        from abbreviation_registry import load_registry, get_all as get_all_abbreviations
    except ImportError:
        build_format_render_context = None
        load_profile = None
        normalize_style_profile = None
        infer_project_root_for_profile = None
        load_registry = None
        get_all_abbreviations = None


# ---------------------------------------------------------------------------
# Markdown 管道表格解析
# ---------------------------------------------------------------------------

_PIPE_TABLE_RE = re.compile(r'^\|(.+)\|$')
_SEPARATOR_RE = re.compile(r'^[\|\s\-:]+$')


def _resolve_project_root(candidate_project_root, source_path):
    if candidate_project_root:
        return os.path.abspath(candidate_project_root)
    if infer_project_root_for_profile is None:
        return os.path.abspath(os.path.dirname(source_path))
    return infer_project_root_for_profile(source_path)


def _guard_format_profile(project_root):
    if load_profile is None:
        return True, None
    profile, _ = load_profile(project_root)
    format_profile = profile.get("format_profile", {}) if isinstance(profile, dict) else {}
    if str(format_profile.get("status", "")).strip() == "pending_template":
        payload = {
            "success": False,
            "error": "pending_template",
            "project_root": project_root,
            "format_profile": format_profile,
            "message": "当前项目处于 pending_template，禁止执行 Word 导出。请先补齐自定义院校格式模板要求。",
        }
        return False, payload
    return True, None


def _load_render_context(project_root):
    fallback = {
        "page_margins_cm": {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 3.17},
        "header_distance_cm": 1.5,
        "footer_distance_cm": 1.75,
        "header_left_text": "中南大学博士学位论文",
        "style_profile": normalize_style_profile(None) if normalize_style_profile is not None else {},
    }
    if load_profile is None or build_format_render_context is None:
        return fallback
    try:
        profile, _ = load_profile(project_root)
    except Exception:
        return fallback
    format_profile = profile.get("format_profile", {}) if isinstance(profile, dict) else {}
    return build_format_render_context(format_profile)


def _resolve_style_profile(render_context=None):
    if isinstance(render_context, dict) and isinstance(render_context.get("style_profile"), dict):
        return render_context.get("style_profile")
    if normalize_style_profile is not None:
        return normalize_style_profile(None)
    return {}


def _parse_pipe_row(line):
    """解析管道表格行，返回单元格文本列表。支持 \\| 转义。"""
    stripped = line.strip()
    m = _PIPE_TABLE_RE.match(stripped)
    if not m:
        return None
    inner = m.group(1)
    # 先将转义管道替换为占位符，分割后再还原
    placeholder = '\x00PIPE\x00'
    inner = inner.replace('\\|', placeholder)
    cells = [c.strip().replace(placeholder, '|') for c in inner.split('|')]
    return cells


def _is_separator_row(line):
    """判断是否为分隔行（如 |---|---|）"""
    return bool(_SEPARATOR_RE.match(line.strip()))


def parse_markdown_line(line):
    """
    解析 Markdown 行，识别类型
    
    Returns:
        tuple: (type, content, level)
    """
    line = line.rstrip()
    
    # 空行
    if not line.strip():
        return ('empty', '', 0)
    if line.strip().startswith('<!--') and line.strip().endswith('-->'):
        return ('empty', '', 0)
    
    # 一级标题 # Title
    if line.startswith('# ') and not line.startswith('## '):
        return ('heading1', line[2:].strip(), 1)
    
    # 二级标题 ## Title
    elif line.startswith('## ') and not line.startswith('### '):
        return ('heading2', line[3:].strip(), 2)
    
    # 三级标题 ### Title
    elif line.startswith('### '):
        return ('heading3', line[4:].strip(), 3)
    
    # 图片占位符 [图 1-1：标题] 或裸格式 图 1-1：标题
    elif re.match(r'\[图\s*\d+-\d+[：:].+\]', line):
        return ('figure', line.strip(), 0)
    elif re.match(r'^图\s*\d+-\d+[：:]', line.strip()):
        return ('figure', line.strip(), 0)
    
    # 表格标题 [表 1-1：标题] 或裸格式 表 1-1：标题
    elif re.match(r'\[表\s*\d+-\d+[：:].+\]', line):
        return ('table', line.strip(), 0)
    elif re.match(r'^表\s*\d+-\d+[：:]', line.strip()):
        return ('table', line.strip(), 0)
    
    # 正文段落
    else:
        return ('paragraph', line.strip(), 0)


def set_run_font(run, latin, east_asia, size_pt, bold=None):
    """
    同时设置拉丁字体和东亚字体，避免 Word 回退到意外字体。
    """
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), east_asia)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_LINE_SPACING_RULE_MAP = {
    "exact": WD_LINE_SPACING.EXACTLY,
    "single": WD_LINE_SPACING.SINGLE,
    "one_point_five": WD_LINE_SPACING.ONE_POINT_FIVE,
}


def _apply_line_spacing(paragraph, spec):
    rule_name = spec.get("line_spacing_rule")
    rule = _LINE_SPACING_RULE_MAP.get(rule_name)
    if rule is not None:
        paragraph.paragraph_format.line_spacing_rule = rule
    if rule_name == "exact" and spec.get("line_spacing_pt") is not None:
        paragraph.paragraph_format.line_spacing = Pt(spec["line_spacing_pt"])


def _apply_paragraph_style_from_spec(paragraph, spec, word_style=None, set_text_black=False):
    if word_style:
        paragraph.style = word_style
    alignment = _ALIGNMENT_MAP.get(spec.get("alignment"))
    if alignment is not None:
        paragraph.paragraph_format.alignment = alignment
    _apply_line_spacing(paragraph, spec)
    if spec.get("space_before_pt") is not None:
        paragraph.paragraph_format.space_before = Pt(spec["space_before_pt"])
    if spec.get("space_after_pt") is not None:
        paragraph.paragraph_format.space_after = Pt(spec["space_after_pt"])
    if spec.get("first_line_indent_cm") is not None:
        paragraph.paragraph_format.first_line_indent = Cm(spec["first_line_indent_cm"])

    for run in paragraph.runs:
        set_run_font(
            run,
            latin=spec.get("font_latin", "Times New Roman"),
            east_asia=spec.get("font_east_asia", "SimSun"),
            size_pt=spec.get("font_size_pt", 12),
            bold=spec.get("bold"),
        )
        if set_text_black:
            run.font.color.rgb = RGBColor(0, 0, 0)


def apply_csu_heading1_style(paragraph, render_context=None):
    """应用一级标题样式"""
    spec = _resolve_style_profile(render_context).get("heading1", {})
    _apply_paragraph_style_from_spec(paragraph, spec, word_style="Heading 1", set_text_black=True)


def apply_csu_heading2_style(paragraph, render_context=None):
    """应用二级标题样式"""
    spec = _resolve_style_profile(render_context).get("heading2", {})
    _apply_paragraph_style_from_spec(paragraph, spec, word_style="Heading 2")


def apply_csu_heading3_style(paragraph, render_context=None):
    """应用三级标题样式"""
    spec = _resolve_style_profile(render_context).get("heading3", {})
    _apply_paragraph_style_from_spec(paragraph, spec, word_style="Heading 3")


def apply_csu_normal_style(paragraph, render_context=None):
    """应用正文样式"""
    spec = _resolve_style_profile(render_context).get("body", {})
    _apply_paragraph_style_from_spec(paragraph, spec, word_style="Normal")


def apply_csu_caption_style(paragraph, render_context=None, caption_key="figure_caption"):
    """应用图表题注样式"""
    spec = _resolve_style_profile(render_context).get(caption_key, {})
    _apply_paragraph_style_from_spec(paragraph, spec)


# ---------------------------------------------------------------------------
# 页眉 / 页脚 / 页码
# ---------------------------------------------------------------------------

def _add_page_number_field(run):
    """在 run 中插入 PAGE 域代码（动态页码）"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def set_page_number_format(section, fmt='decimal', start_at=None):
    """
    设置节的页码格式。
    fmt: 'lowerRoman' | 'decimal'
    start_at: 起始页码（int），None 表示续前节
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start_at is not None:
        pgNumType.set(qn('w:start'), str(start_at))
    elif qn('w:start') in pgNumType.attrib:
        del pgNumType.attrib[qn('w:start')]


def setup_header(section, left_text, right_text, distance_cm=1.5, style_spec=None):
    """
    设置页眉：左侧 left_text，右侧 right_text，宋体五号(10.5pt)。
    距顶端 1.5cm。
    """
    section.header_distance = Cm(distance_cm)
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = ""
    # 右对齐制表位 = 页面文本区宽度
    text_width = section.page_width - section.left_margin - section.right_margin
    p.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(f"{left_text}\t{right_text}")
    spec = style_spec or {
        "font_latin": "Times New Roman",
        "font_east_asia": "SimSun",
        "font_size_pt": 10.5,
        "bold": False,
    }
    set_run_font(
        run,
        latin=spec.get("font_latin", "Times New Roman"),
        east_asia=spec.get("font_east_asia", "SimSun"),
        size_pt=spec.get("font_size_pt", 10.5),
        bold=spec.get("bold", False),
    )


def setup_footer(section, page_num_fmt='decimal', start_at=None, distance_cm=1.75, style_spec=None):
    """
    设置页脚：居中页码，TNR 小五号(9pt)。
    距底端 1.75cm。
    page_num_fmt: 'lowerRoman' (前置部分) | 'decimal' (正文)
    start_at: 起始页码
    """
    section.footer_distance = Cm(distance_cm)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    _add_page_number_field(run)
    spec = style_spec or {
        "font_latin": "Times New Roman",
        "font_east_asia": "SimSun",
        "font_size_pt": 9,
        "bold": False,
    }
    set_run_font(
        run,
        latin=spec.get("font_latin", "Times New Roman"),
        east_asia=spec.get("font_east_asia", "SimSun"),
        size_pt=spec.get("font_size_pt", 9),
        bold=spec.get("bold", False),
    )
    # 设置页码格式
    set_page_number_format(section, fmt=page_num_fmt, start_at=start_at)


# ---------------------------------------------------------------------------
# 前置部分专用样式：摘要 / 英文摘要 / 目录
# ---------------------------------------------------------------------------

def add_abstract_section(doc, abstract_body, keywords=None, render_context=None):
    """
    添加中文摘要页。
    - 标题"摘  要"：三号黑体加粗，居中
    - "摘要："标识：四号黑体加粗，顶格
    - 正文：四号宋体，1.5倍行距
    - 关键词：四号黑体加粗"关键词："+ 四号宋体内容，全角分号分隔
    """
    # 标题
    fm_spec = _resolve_style_profile(render_context).get("front_matter", {}).get("zh_abstract", {})
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(fm_spec.get("title_text", "摘  要"))
    _apply_paragraph_style_from_spec(title_para, fm_spec.get("title", {}))

    # "摘要："标识行
    label_para = doc.add_paragraph()
    label_run = label_para.add_run(fm_spec.get("label_text", "摘要："))
    _apply_paragraph_style_from_spec(label_para, fm_spec.get("label", {}))

    # 正文
    if abstract_body:
        for para_text in abstract_body.split('\n'):
            para_text = para_text.strip()
            if not para_text:
                continue
            body_para = doc.add_paragraph()
            body_run = body_para.add_run(para_text)
            _apply_paragraph_style_from_spec(body_para, fm_spec.get("body", {}))

    # 关键词
    if keywords:
        kw_para = doc.add_paragraph()
        kw_label = kw_para.add_run(fm_spec.get("keywords_label_text", "关键词："))
        _apply_paragraph_style_from_spec(kw_para, fm_spec.get("keywords_label", {}))
        kw_content = kw_para.add_run(keywords)
        set_run_font(
            kw_content,
            latin=fm_spec.get("keywords_body", {}).get("font_latin", "Times New Roman"),
            east_asia=fm_spec.get("keywords_body", {}).get("font_east_asia", "SimSun"),
            size_pt=fm_spec.get("keywords_body", {}).get("font_size_pt", 14),
            bold=fm_spec.get("keywords_body", {}).get("bold", False),
        )


def add_english_abstract_section(doc, abstract_body, keywords=None, render_context=None):
    """
    添加英文摘要页。
    - 标题"ABSTRACT"：三号 TNR 加粗，居中
    - "Abstract："标识：四号 TNR 加粗，顶格
    - 正文：四号 TNR，1.5倍行距
    - Keywords：四号 TNR 加粗"Keywords："+ 四号 TNR 内容，半角分号分隔
    """
    # 标题
    fm_spec = _resolve_style_profile(render_context).get("front_matter", {}).get("en_abstract", {})
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(fm_spec.get("title_text", "ABSTRACT"))
    _apply_paragraph_style_from_spec(title_para, fm_spec.get("title", {}))

    # "Abstract："标识行
    label_para = doc.add_paragraph()
    label_run = label_para.add_run(fm_spec.get("label_text", "Abstract："))
    _apply_paragraph_style_from_spec(label_para, fm_spec.get("label", {}))

    # 正文
    if abstract_body:
        for para_text in abstract_body.split('\n'):
            para_text = para_text.strip()
            if not para_text:
                continue
            body_para = doc.add_paragraph()
            body_run = body_para.add_run(para_text)
            _apply_paragraph_style_from_spec(body_para, fm_spec.get("body", {}))

    # Keywords
    if keywords:
        kw_para = doc.add_paragraph()
        kw_label = kw_para.add_run(fm_spec.get("keywords_label_text", "Keywords："))
        _apply_paragraph_style_from_spec(kw_para, fm_spec.get("keywords_label", {}))
        kw_content = kw_para.add_run(keywords)
        set_run_font(
            kw_content,
            latin=fm_spec.get("keywords_body", {}).get("font_latin", "Times New Roman"),
            east_asia=fm_spec.get("keywords_body", {}).get("font_east_asia", "Times New Roman"),
            size_pt=fm_spec.get("keywords_body", {}).get("font_size_pt", 14),
            bold=fm_spec.get("keywords_body", {}).get("bold", False),
        )


def add_toc_section(doc, toc_entries=None, render_context=None):
    """
    添加目录页。
    - 标题"目  录"：三号黑体加粗，居中（中间空两格）
    - 章标题：小四号黑体
    - 节标题：小四号宋体
    如果 toc_entries 为 None，插入 TOC 域代码（Word 打开后按 F9 更新）。
    toc_entries 格式：[(level, title, page_str), ...]  level=1 章, level=2 节, level=3 小节
    """
    # 标题
    toc_spec = _resolve_style_profile(render_context).get("front_matter", {}).get("toc", {})
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(toc_spec.get("title_text", "目  录"))
    _apply_paragraph_style_from_spec(title_para, toc_spec.get("title", {}))

    if toc_entries is None:
        # 插入 TOC 域代码，用户在 Word 中按 F9 更新
        toc_para = doc.add_paragraph()
        run = toc_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
    else:
        for level, title, page_str in toc_entries:
            entry_para = doc.add_paragraph()
            if level == 1:
                entry_run = entry_para.add_run(f'{title}{"." * 20}{page_str}')
                _apply_paragraph_style_from_spec(entry_para, toc_spec.get("level1", {}))
            else:
                entry_run = entry_para.add_run(f'{title}{"." * 20}{page_str}')
                level_key = "level2" if level == 2 else "level3"
                _apply_paragraph_style_from_spec(entry_para, toc_spec.get(level_key, {}))


# ---------------------------------------------------------------------------
# 正文 Markdown 粗体/斜体处理
# ---------------------------------------------------------------------------

# 匹配 **bold** 或 __bold__（双星号/双下划线粗体）
_BOLD_RE = re.compile(r'\*\*(.+?)\*\*|__(.+?)__')
# 匹配显著性标记保护：星号(1-4个) + 可选空格 + P/p + 比较符
_SIGNIFICANCE_PROTECT_RE = re.compile(r'(\*{1,4})(\s*)([pP][<>≤≥=])')


def strip_bold_markers(text):
    """去除正文中的 **粗体** 标记，但保留统计学显著性标记 *p<0.05 等。

    逻辑：
    1. 保护显著性标记（如 **P<0.01），替换为占位符。
    2. 执行去除粗体操作。
    3. 还原占位符。
    这防止了 "A(**P<0.01)B(**P<0.01)" 被误判为粗体包裹。
    """
    placeholders = []

    def protect(m):
        # 将整个匹配串（如 "**P<"）替换为占位符
        token = f"§SIG{len(placeholders)}§"
        placeholders.append(m.group(0))
        return token

    # 1. 保护
    temp_text = _SIGNIFICANCE_PROTECT_RE.sub(protect, text)

    # 2. 去粗体
    stripped_text = _BOLD_RE.sub(lambda m: m.group(1) or m.group(2), temp_text)

    # 3. 还原
    for i, original in enumerate(placeholders):
        stripped_text = stripped_text.replace(f"§SIG{i}§", original)

    return stripped_text


# ---------------------------------------------------------------------------
# 三线表工具
# ---------------------------------------------------------------------------


def _set_cell_border(cell, **kwargs):
    """
    设置单元格边框。

    kwargs 示例: top={"sz": 12, "val": "single", "color": "000000"}
    sz 单位为 1/8 pt，所以 12 = 1.5pt, 4 = 0.5pt
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        for attr_name, attr_val in attrs.items():
            element.set(qn(f'w:{attr_name}'), str(attr_val))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _no_border():
    return {"sz": "0", "val": "none", "color": "auto"}


def _thick_border():
    return {"sz": "12", "val": "single", "color": "000000"}  # 1.5pt


def _thin_border():
    return {"sz": "6", "val": "single", "color": "000000"}   # 0.75pt


def apply_three_line_table_borders(table, header_rows=1):
    """
    对已有 Word 表格应用三线表边框。

    规则：
    - 顶线：1.5pt 粗线（sz=12）
    - 表头分隔线：0.5pt 细线（sz=4），位于最后一个表头行底部
    - 底线：1.5pt 粗线（sz=12）
    - 无竖线

    Args:
        table: python-docx Table
        header_rows: 表头行数，>=1
    """
    if table is None or not getattr(table, "rows", None):
        return

    total_rows = len(table.rows)
    if total_rows == 0:
        return

    header_rows = max(1, int(header_rows or 1))
    header_sep_idx = min(header_rows - 1, total_rows - 1)
    last_idx = total_rows - 1

    no = _no_border()
    thick = _thick_border()
    thin_05pt = {"sz": "4", "val": "single", "color": "000000"}

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            top = no
            bottom = no
            if i == 0:
                top = thick
            if i == header_sep_idx:
                bottom = thin_05pt
            if i == last_idx:
                bottom = thick
            _set_cell_border(cell, top=top, bottom=bottom, left=no, right=no)


def is_table_row(line):
    """检测是否为 Markdown 表格行: | col1 | col2 |"""
    stripped = line.strip()
    return stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 3


def is_table_separator(line):
    """检测是否为 Markdown 表格分隔行: |---|---|"""
    stripped = line.strip()
    if not (stripped.startswith('|') and stripped.endswith('|')):
        return False
    inner = stripped[1:-1]
    cells = inner.split('|')
    return all(re.match(r'^[\s\-:]+$', c) for c in cells)


def parse_table_rows(lines):
    """
    从连续的 Markdown 表格行中提取表头和数据行。

    Returns:
        tuple: (headers: list[str], rows: list[list[str]])
    """
    headers = []
    rows = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        cells = [c.strip() for c in stripped.strip('|').split('|')]
        if i == 0:
            headers = cells
        elif is_table_separator(line):
            continue
        else:
            rows.append(cells)
    return headers, rows


def create_three_line_table(doc, headers, rows, caption=None, render_context=None):
    """
    创建三线表格式的 Word 表格。

    - 顶线: 1.5pt
    - 表头下线: 0.5pt
    - 底线: 1.5pt
    - 无竖线

    Args:
        doc: Document 对象
        headers: 表头列表
        rows: 数据行列表
        caption: 表格标题（可选，置于表格上方）
    """
    # 添加标题（如果有）
    if caption:
        cap_para = doc.add_paragraph(caption)
        apply_csu_caption_style(cap_para, render_context=render_context, caption_key="table_caption")

    num_cols = len(headers)
    num_rows = 1 + len(rows)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 填充表头
    for j, header_text in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            cell_spec = _resolve_style_profile(render_context).get("table_cell", {})
            paragraph.paragraph_format.alignment = _ALIGNMENT_MAP.get(
                cell_spec.get("alignment", "center"),
                WD_ALIGN_PARAGRAPH.CENTER,
            )
            for run in paragraph.runs:
                set_run_font(
                    run,
                    latin=cell_spec.get("font_latin", "Times New Roman"),
                    east_asia=cell_spec.get("font_east_asia", "SimSun"),
                    size_pt=cell_spec.get("font_size_pt", 10.5),
                    bold=cell_spec.get("header_bold", True),
                )

    # 填充数据行
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i + 1].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    cell_spec = _resolve_style_profile(render_context).get("table_cell", {})
                    paragraph.paragraph_format.alignment = _ALIGNMENT_MAP.get(
                        cell_spec.get("alignment", "center"),
                        WD_ALIGN_PARAGRAPH.CENTER,
                    )
                    for run in paragraph.runs:
                        set_run_font(
                            run,
                            latin=cell_spec.get("font_latin", "Times New Roman"),
                            east_asia=cell_spec.get("font_east_asia", "SimSun"),
                            size_pt=cell_spec.get("font_size_pt", 10.5),
                            bold=cell_spec.get("bold", False),
                        )

    # 应用三线表边框
    no = _no_border()
    thick = _thick_border()
    thin = _thin_border()

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if i == 0:
                # 表头行：顶线粗，底线细，无竖线
                _set_cell_border(cell, top=thick, bottom=thin, left=no, right=no)
            elif i == num_rows - 1:
                # 最后一行：底线粗，无竖线
                _set_cell_border(cell, top=no, bottom=thick, left=no, right=no)
            else:
                # 中间行：无边框
                _set_cell_border(cell, top=no, bottom=no, left=no, right=no)

    return table


# ---------------------------------------------------------------------------
# 缩略语对照表页
# ---------------------------------------------------------------------------


def create_abbreviation_table_page(doc, project_root, render_context=None):
    """
    在文档中插入缩略语对照表页（三线表格式）。

    Args:
        doc: Document 对象
        project_root: 项目根目录（用于读取注册表）

    Returns:
        bool: 是否成功插入（无缩略语时返回 False）
    """
    if get_all_abbreviations is None:
        return False

    try:
        items = get_all_abbreviations(project_root)
    except Exception:
        return False

    if not items:
        return False

    # 标题
    abbr_spec = _resolve_style_profile(render_context).get("front_matter", {}).get("abbreviation_table", {})
    heading = doc.add_heading(abbr_spec.get("title_text", "主要缩略语对照表"), level=1)
    _apply_paragraph_style_from_spec(heading, abbr_spec.get("title", {}), word_style="Heading 1", set_text_black=True)

    # 构建表格数据
    headers = ["缩略语", "英文全称", "中文全称"]
    rows = []
    for abbr, info in items:
        full_en = info.get("full_en", "") or ""
        full_cn = info.get("full_cn", "") or ""
        rows.append([abbr, full_en, full_cn])

    create_three_line_table(doc, headers, rows, render_context=render_context)

    # 分页符
    doc.add_page_break()

    return True


def markdown_to_docx(md_content, output_path, chapter_num=None, project_root=None,
                     include_abbreviation_table=False,
                     header_right_text=None, page_num_fmt='decimal',
                     page_num_start=None):
    """
    将 Markdown 内容转换为 Word 文档
    
    Args:
        md_content: Markdown 文本内容
        output_path: 输出文件路径
        chapter_num: 章节号（可选）
        project_root: 项目根目录（可选，用于缩略语表）
        include_abbreviation_table: 是否在文档开头插入缩略语对照表
        header_right_text: 页眉右侧文字（如"第1章 绪论"），None 则自动从首个 H1 提取
        page_num_fmt: 页码格式 'decimal'(阿拉伯) | 'lowerRoman'(罗马)
        page_num_start: 起始页码（int），None 续前节
    
    Returns:
        bool: 转换是否成功
    """
    try:
        render_context = _load_render_context(project_root)

        # 创建文档
        doc = Document()
        
        # 设置页面
        section = doc.sections[0]
        page_margins = render_context.get("page_margins_cm", {})
        section.top_margin = Cm(page_margins.get("top", 2.54))
        section.bottom_margin = Cm(page_margins.get("bottom", 2.54))
        section.left_margin = Cm(page_margins.get("left", 3.17))
        section.right_margin = Cm(page_margins.get("right", 3.17))

        # 自动提取页眉右侧文字（从首个 H1）
        if header_right_text is None:
            for ln in md_content.split('\n'):
                ln_s = ln.strip()
                if ln_s.startswith('# ') and not ln_s.startswith('## '):
                    header_right_text = ln_s[2:].strip()
                    break
            if header_right_text is None:
                header_right_text = ""

        # 设置页眉
        setup_header(
            section,
            render_context.get("header_left_text", "中南大学博士学位论文"),
            header_right_text,
            distance_cm=render_context.get("header_distance_cm", 1.5),
            style_spec=_resolve_style_profile(render_context).get("header", {}),
        )
        # 设置页脚（页码）
        setup_footer(
            section,
            page_num_fmt=page_num_fmt,
            start_at=page_num_start,
            distance_cm=render_context.get("footer_distance_cm", 1.75),
            style_spec=_resolve_style_profile(render_context).get("footer", {}),
        )

        # 插入缩略语对照表（如果需要）
        if include_abbreviation_table and project_root:
            create_abbreviation_table_page(doc, project_root, render_context=render_context)
        
        # 逐行解析，支持表格累积
        lines = md_content.split('\n')
        table_buffer = []       # 累积连续的表格行
        table_caption = None    # 表格标题（表 X-X）

        def flush_table():
            """将累积的表格行渲染为三线表"""
            nonlocal table_buffer, table_caption
            if not table_buffer:
                return
            headers, rows = parse_table_rows(table_buffer)
            if headers:
                create_three_line_table(doc, headers, rows, caption=table_caption, render_context=render_context)
            table_buffer = []
            table_caption = None

        for line in lines:
            # 检测是否为表格行
            if is_table_row(line) or (table_buffer and is_table_separator(line)):
                table_buffer.append(line)
                continue

            # 非表格行：先刷新之前累积的表格
            if table_buffer:
                flush_table()

            # 如果之前设置了 table_caption 但没有紧跟管道表格，
            # 将其作为普通题注段落渲染，避免泄漏到后续表格
            if table_caption is not None:
                line_type_cur, content_cur, _ = parse_markdown_line(line)
                if line_type_cur != 'empty':
                    # 非空行且不是表格行 → caption 后面没有紧跟管道表格
                    cap_para = doc.add_paragraph(table_caption)
                    apply_csu_caption_style(cap_para, render_context=render_context, caption_key="table_caption")
                    table_caption = None
                    # 继续处理当前行（不 skip）

            line_type, content, level = parse_markdown_line(line)
            
            if line_type == 'empty':
                continue
            
            elif line_type == 'heading1':
                para = doc.add_heading(content, level=1)
                apply_csu_heading1_style(para, render_context=render_context)
            
            elif line_type == 'heading2':
                para = doc.add_heading(content, level=2)
                apply_csu_heading2_style(para, render_context=render_context)
            
            elif line_type == 'heading3':
                para = doc.add_heading(content, level=3)
                apply_csu_heading3_style(para, render_context=render_context)
            
            elif line_type == 'figure':
                para = doc.add_paragraph(content)
                apply_csu_caption_style(para, render_context=render_context, caption_key="figure_caption")
            
            elif line_type == 'table':
                # 表格占位符 [表 X-X：标题] — 可能是后续 Markdown 表格的标题
                table_caption = content.strip('[]')
                # 不立即渲染，等待后续表格行
            
            elif line_type == 'paragraph':
                if content:
                    cleaned = strip_bold_markers(content)
                    para = doc.add_paragraph(cleaned)
                    apply_csu_normal_style(para, render_context=render_context)

        # 文件末尾：刷新残留的表格
        if table_buffer:
            flush_table()
        
        # 保存文档
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        doc.save(output_path)
        print(f"✅ 转换成功：{output_path}")
        return True
    
    except Exception as e:
        print(f"❌ 转换失败：{str(e)}")
        return False


def convert_markdown_file(md_file_path, output_path=None, project_root=None,
                          include_abbreviation_table=False,
                          header_right_text=None, page_num_fmt='decimal',
                          page_num_start=None):
    """
    转换 Markdown 文件
    
    Args:
        md_file_path: Markdown 文件路径
        output_path: 输出文件路径（可选）
        project_root: 项目根目录（可选，用于缩略语表）
        include_abbreviation_table: 是否插入缩略语对照表
        header_right_text: 页眉右侧文字，None 自动提取
        page_num_fmt: 页码格式 'decimal' | 'lowerRoman'
        page_num_start: 起始页码
    """
    if not os.path.exists(md_file_path):
        print(f"❌ 文件不存在：{md_file_path}")
        return False

    resolved_project_root = _resolve_project_root(project_root, md_file_path)
    ok_to_convert, gate_payload = _guard_format_profile(resolved_project_root)
    if not ok_to_convert:
        print(f"❌ {gate_payload['message']}")
        print(f"pending_template: {json.dumps(gate_payload, ensure_ascii=False)}")
        return False
    
    # 读取 Markdown 文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 确定输出路径
    if output_path is None:
        output_path = md_file_path.replace('.md', '.docx')
    
    # 提取章节号（如果文件名包含）
    chapter_num = None
    match = re.search(r'第(\d+)章', os.path.basename(md_file_path))
    if match:
        chapter_num = int(match.group(1))
    
    # 执行转换
    return markdown_to_docx(
        md_content, output_path, chapter_num=chapter_num,
        project_root=resolved_project_root,
        include_abbreviation_table=include_abbreviation_table,
        header_right_text=header_right_text,
        page_num_fmt=page_num_fmt,
        page_num_start=page_num_start,
    )


def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Markdown 转 Word（默认 CSU，可读取自定义格式配置）')
    parser.add_argument('input', help='输入 Markdown 文件路径')
    parser.add_argument('-o', '--output', help='输出 Word 文件路径（可选）')
    parser.add_argument('-c', '--chapter', type=int, help='章节号（可选）')
    parser.add_argument('--project-root', help='项目根目录（用于缩略语表）')
    parser.add_argument('--abbreviation-table', action='store_true',
                        help='在文档开头插入缩略语对照表')
    parser.add_argument('--header-right', help='页眉右侧文字（如"第1章 绪论"），默认自动提取')
    parser.add_argument('--page-num-fmt', choices=['decimal', 'lowerRoman'],
                        default='decimal', help='页码格式（默认阿拉伯数字）')
    parser.add_argument('--page-num-start', type=int, help='起始页码')
    
    args = parser.parse_args()
    
    # 执行转换
    success = convert_markdown_file(
        args.input, args.output,
        project_root=args.project_root,
        include_abbreviation_table=args.abbreviation_table,
        header_right_text=args.header_right,
        page_num_fmt=args.page_num_fmt,
        page_num_start=args.page_num_start,
    )
    
    if success:
        print("\n📋 样式说明：")
        print("  - 一级标题：三号黑体加粗，居中")
        print("  - 二级标题：四号宋体，顶格")
        print("  - 三级标题：小四号宋体，顶格")
        print("  - 正文：小四号宋体/Times New Roman，首行缩进")
        print("  - 图表题注：五号楷体，居中")
        sys.exit(0)
    else:
        sys.exit(1)


if __name__ == '__main__':
    main()
