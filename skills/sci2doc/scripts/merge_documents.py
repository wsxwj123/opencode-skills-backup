#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
博士论文文档合并脚本
功能：将分章节的docx文件按顺序合并为完整论文
依赖：pip install python-docx
版本：v2.0
作者：AI Assistant
更新：2026-01-27
"""

import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.xmlchemy import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("❌ 错误：未安装 python-docx 库")
    print("请执行：pip3 install python-docx")
    sys.exit(1)


def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()


def merge_documents(file_list, output_file):
    """
    合并多个Word文档
    
    Args:
        file_list: 文件路径列表（按合并顺序）
        output_file: 输出文件路径
    """
    print(f"\n开始合并 {len(file_list)} 个文档...\n")
    
    # 创建主文档（使用第一个文件作为基础）
    merged_doc = Document(file_list[0])
    print(f"✅ 已加载基础文档：{os.path.basename(file_list[0])}")
    
    # 遍历其余文档并追加内容
    for file_path in file_list[1:]:
        if not os.path.exists(file_path):
            print(f"⚠️ 文件不存在，跳过：{os.path.basename(file_path)}")
            continue
        
        try:
            # 读取待合并文档
            sub_doc = Document(file_path)
            
            # 添加分页符
            add_page_break(merged_doc)
            
            # 复制所有段落
            for para in sub_doc.paragraphs:
                new_para = merged_doc.add_paragraph(para.text, style=para.style)
                
                # 复制段落格式
                if para.paragraph_format:
                    new_para.paragraph_format.alignment = para.paragraph_format.alignment
                    new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                    new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
                    new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
                    new_para.paragraph_format.space_before = para.paragraph_format.space_before
                    new_para.paragraph_format.space_after = para.paragraph_format.space_after
                    new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            
            # 复制所有表格
            for table in sub_doc.tables:
                new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.rows[i].cells[j].text = cell.text
            
            print(f"✅ 已合并：{os.path.basename(file_path)}")
            
        except Exception as e:
            print(f"❌ 合并失败 {os.path.basename(file_path)}：{e}")
            continue
    
    # 保存合并后的文档
    try:
        merged_doc.save(output_file)
        print(f"\n🎉 合并完成！")
        print(f"   输出文件：{output_file}")
        print(f"   📊 总计合并 {len(file_list)} 个文档")
        
        # 统计文件大小
        file_size = os.path.getsize(output_file) / 1024 / 1024  # MB
        print(f"   📦 文件大小：{file_size:.2f} MB")
        
        return True
        
    except Exception as e:
        print(f"\n❌ 保存失败：{e}")
        return False


def main():
    """主函数：自动检测并合并所有章节文件"""
    
    # 项目根目录（脚本所在目录的上级目录）
    script_dir = Path(__file__).parent.absolute()
    project_root = script_dir.parent
    
    print("=" * 70)
    print("博士论文文档合并工具")
    print("=" * 70)
    print(f"📁 项目根目录：{project_root}\n")
    
    # 定义合并顺序（按论文结构）
    file_order = [
        # 前置部分
        "01_前置部分/封面.docx",
        "01_前置部分/扉页.docx",
        "01_前置部分/原创性声明.docx",
        "01_前置部分/中文摘要.docx",
        "01_前置部分/英文摘要.docx",
        "01_前置部分/目录.docx",
        "01_前置部分/英文缩略词说明.docx",
        "01_前置部分/符号说明.docx",  # 可选
        
        # 正文章节（会自动检测）
        # ...
        
        # 后置部分
        "03_后置部分/参考文献.docx",
        "03_后置部分/综述.docx",
        "03_后置部分/附录.docx",  # 可选
        "03_后置部分/攻读学位期间成果.docx",
        "03_后置部分/致谢.docx",
    ]
    
    # 构建完整路径列表
    full_paths = []
    
    # 添加前置部分（跳过不存在的可选文件）
    print("📋 检查前置部分...")
    for rel_path in file_order[:8]:
        full_path = project_root / rel_path
        if full_path.exists():
            full_paths.append(str(full_path))
            print(f"   ✅ {os.path.basename(rel_path)}")
        elif "符号说明" not in rel_path:
            print(f"   ⚠️ 缺失：{os.path.basename(rel_path)}")
    
    # 自动检测正文章节（按章节编号排序）
    print("\n📋 检查正文章节...")
    chapter_dir = project_root / "02_正文章节"
    if chapter_dir.exists():
        chapter_files = sorted([f for f in os.listdir(chapter_dir) 
                               if f.endswith('.docx') and f.startswith('第')])
        
        if chapter_files:
            for chapter_file in chapter_files:
                full_paths.append(str(chapter_dir / chapter_file))
                print(f"   📄 {chapter_file}")
        else:
            print("   ⚠️ 未找到章节文件")
    else:
        print(f"   ❌ 正文章节目录不存在：{chapter_dir}")
        sys.exit(1)
    
    # 添加后置部分
    print("\n📋 检查后置部分...")
    for rel_path in file_order[8:]:
        full_path = project_root / rel_path
        if full_path.exists():
            full_paths.append(str(full_path))
            print(f"   ✅ {os.path.basename(rel_path)}")
        elif "附录" not in rel_path:
            print(f"   ⚠️ 缺失：{os.path.basename(rel_path)}")
    
    # 检查是否有足够文件
    if len(full_paths) < 2:
        print("\n❌ 至少需要2个文件才能进行合并")
        sys.exit(1)
    
    # 输出文件路径
    output_file = str(project_root / "完整版_最新.docx")
    
    # 执行合并
    print("\n" + "=" * 70)
    success = merge_documents(full_paths, output_file)
    print("=" * 70)
    
    if success:
        print("\n✅ 合并成功！")
        print(f"\n📄 完整版文档：{output_file}")
        print("\n后续步骤：")
        print("1. 用Word打开完整版文档检查格式")
        print("2. 如有占位标识，替换为实际图片")
        print("3. 更新目录（右键目录→更新域）")
        print("4. 检查页眉页脚和页码")
        print("5. 保存最终版本")
    else:
        print("\n❌ 合并失败，请检查错误信息")
        sys.exit(1)


if __name__ == "__main__":
    main()
