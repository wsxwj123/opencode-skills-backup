#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Render managed front matter markdown/docx files from local templates.
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm

try:
    from thesis_profile import build_format_render_context, load_profile
except Exception:  # pragma: no cover
    _script_dir = os.path.dirname(os.path.abspath(__file__))
    if _script_dir not in sys.path:
        sys.path.insert(0, _script_dir)
    from thesis_profile import build_format_render_context, load_profile

try:
    from markdown_to_docx import (
        add_abstract_section,
        add_english_abstract_section,
        add_toc_section,
        convert_markdown_file,
        create_abbreviation_table_page,
        setup_footer,
        setup_header,
    )
except Exception:  # pragma: no cover
    _script_dir = os.path.dirname(os.path.abspath(__file__))
    if _script_dir not in sys.path:
        sys.path.insert(0, _script_dir)
    from markdown_to_docx import (
        add_abstract_section,
        add_english_abstract_section,
        add_toc_section,
        convert_markdown_file,
        create_abbreviation_table_page,
        setup_footer,
        setup_header,
    )

try:
    from abbreviation_registry import generate_abbreviation_table_markdown
except Exception:  # pragma: no cover
    _script_dir = os.path.dirname(os.path.abspath(__file__))
    if _script_dir not in sys.path:
        sys.path.insert(0, _script_dir)
    from abbreviation_registry import generate_abbreviation_table_markdown


MANAGED_MARKER = "<!-- managed-by-sci2doc-front-matter -->"
TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates"

FRONT_MATTER_SPECS = [
    {
        "kind": "template",
        "template": "cover_page.md",
        "md_output": "atomic_md/封面.md",
        "docx_output": "02_分章节文档/封面.docx",
    },
    {
        "kind": "template",
        "template": "title_page.md",
        "md_output": "atomic_md/题名页.md",
        "docx_output": "02_分章节文档/题名页.docx",
    },
    {
        "kind": "template",
        "template": "declaration.md",
        "md_output": "atomic_md/独创性声明与授权书.md",
        "docx_output": "02_分章节文档/独创性声明与授权书.docx",
    },
    {
        "kind": "zh_abstract",
        "md_output": "atomic_md/中文摘要.md",
        "docx_output": "02_分章节文档/中文摘要.docx",
    },
    {
        "kind": "en_abstract",
        "md_output": "atomic_md/英文摘要.md",
        "docx_output": "02_分章节文档/英文摘要.docx",
    },
    {
        "kind": "toc",
        "md_output": "atomic_md/目录.md",
        "docx_output": "02_分章节文档/目录.docx",
    },
    {
        "kind": "abbreviation_table",
        "md_output": "atomic_md/缩略语表.md",
        "docx_output": "02_分章节文档/缩略语表.docx",
    },
]


def _read_project_info(project_root):
    state_path = Path(project_root) / "project_state.json"
    if not state_path.exists():
        return {}
    try:
        return json.loads(state_path.read_text(encoding="utf-8")).get("project_info", {}) or {}
    except Exception:
        return {}


def _join_keywords(value, separator):
    if isinstance(value, list):
        items = [str(item).strip() for item in value if str(item).strip()]
        return separator.join(items)
    if value is None:
        return ""
    return str(value).strip()


def _build_placeholder_map(project_root):
    profile, _ = load_profile(project_root)
    format_context = build_format_render_context(profile.get("format_profile", {}))
    project_info = _read_project_info(project_root)
    now = datetime.now()

    def _text(key, default=""):
        value = project_info.get(key, default)
        return str(value or default)

    return {
        "CLASSIFICATION": _text("classification"),
        "UDC": _text("udc"),
        "SCHOOL_CODE": str(format_context.get("school_code", "10533")),
        "STUDENT_ID": _text("student_id"),
        "UNIVERSITY_NAME": str(format_context.get("university_name", "")),
        "DEGREE_TYPE": str(format_context.get("degree_type", "")),
        "TITLE": _text("title"),
        "TITLE_EN": _text("title_en"),
        "AUTHOR": _text("author"),
        "SUPERVISOR": _text("supervisor"),
        "CO_SUPERVISOR": _text("co_supervisor"),
        "MAJOR": _text("major"),
        "RESEARCH_DIRECTION": _text("research_direction"),
        "DEPARTMENT": _text("department"),
        "GRADUATE_SCHOOL_NAME": str(format_context.get("graduate_school_name", "")),
        "DECLARATION_SCHOOL_NAME": str(format_context.get("declaration_authorization_school_name", "")),
        "YEAR": str(now.year),
        "MONTH": f"{now.month:02d}",
    }


def _render_template_text(template_name, placeholders):
    template_path = TEMPLATE_DIR / template_name
    text = template_path.read_text(encoding="utf-8")
    for key, value in placeholders.items():
        text = text.replace(f"[{key}]", value)
    return f"{MANAGED_MARKER}\n\n{text.strip()}\n"


def _render_zh_abstract_text(project_root, placeholders):
    project_info = _read_project_info(project_root)
    body = str(project_info.get("abstract_zh", "") or "").strip() or "待补充中文摘要正文。"
    keywords = _join_keywords(project_info.get("keywords_zh"), "；") or "待补充关键词1；待补充关键词2"
    classification = str(project_info.get("classification", "") or "").strip()
    lines = [
        "# 摘  要",
        "",
        placeholders.get("TITLE", ""),
        "",
        "摘要：",
        body,
        "",
        f"关键词：{keywords}",
    ]
    if classification:
        lines.append(f"分类号：{classification}")
    return f"{MANAGED_MARKER}\n\n" + "\n".join(lines).strip() + "\n"


def _render_en_abstract_text(project_root, placeholders):
    project_info = _read_project_info(project_root)
    body = str(project_info.get("abstract_en", "") or "").strip() or "Pending English abstract body."
    keywords = _join_keywords(project_info.get("keywords_en"), "; ") or "keyword1; keyword2"
    classification = str(project_info.get("classification", "") or "").strip()
    lines = [
        "# ABSTRACT",
        "",
        placeholders.get("TITLE_EN", "") or placeholders.get("TITLE", ""),
        "",
        "Abstract：",
        body,
        "",
        f"Keywords：{keywords}",
    ]
    if classification:
        lines.append(f"Classification：{classification}")
    return f"{MANAGED_MARKER}\n\n" + "\n".join(lines).strip() + "\n"


def _render_toc_text():
    lines = [
        "# 目  录",
        "",
        "目录页由 Word 字段自动生成，打开文档后更新域即可显示页码。",
    ]
    return f"{MANAGED_MARKER}\n\n" + "\n".join(lines).strip() + "\n"


def _render_abbreviation_table_text(project_root):
    body = generate_abbreviation_table_markdown(project_root).strip()
    if not body:
        body = "# 主要缩略语对照表\n\n暂无已注册缩略语。"
    return f"{MANAGED_MARKER}\n\n{body}\n"


def _render_managed_markdown(spec, project_root, placeholders):
    kind = spec.get("kind")
    if kind == "template":
        return _render_template_text(spec["template"], placeholders)
    if kind == "zh_abstract":
        return _render_zh_abstract_text(project_root, placeholders)
    if kind == "en_abstract":
        return _render_en_abstract_text(project_root, placeholders)
    if kind == "toc":
        return _render_toc_text()
    if kind == "abbreviation_table":
        return _render_abbreviation_table_text(project_root)
    raise ValueError(f"unsupported front matter kind: {kind}")


def _apply_doc_layout(doc, render_context, header_right_text=""):
    section = doc.sections[0]
    page_margins = render_context.get("page_margins_cm", {})
    page_numbering = render_context.get("page_numbering", {})
    front_numbering = page_numbering.get("front_matter", {}) if isinstance(page_numbering, dict) else {}
    section.top_margin = Cm(page_margins.get("top", 2.54))
    section.bottom_margin = Cm(page_margins.get("bottom", 2.54))
    section.left_margin = Cm(page_margins.get("left", 3.17))
    section.right_margin = Cm(page_margins.get("right", 3.17))
    setup_header(
        section,
        render_context.get("header_left_text", "中南大学博士学位论文"),
        header_right_text,
        distance_cm=render_context.get("header_distance_cm", 1.5),
        style_spec=render_context.get("style_profile", {}).get("header", {}),
    )
    setup_footer(
        section,
        page_num_fmt=front_numbering.get("format", "lowerRoman"),
        start_at=front_numbering.get("start", 1),
        distance_cm=render_context.get("footer_distance_cm", 1.75),
        style_spec=render_context.get("style_profile", {}).get("footer", {}),
    )


def _render_special_docx(spec, project_root, docx_path, placeholders):
    profile, _ = load_profile(project_root)
    render_context = build_format_render_context(profile.get("format_profile", {}))
    project_info = _read_project_info(project_root)
    doc = Document()
    _apply_doc_layout(doc, render_context, header_right_text="")
    kind = spec.get("kind")
    if kind == "zh_abstract":
        add_abstract_section(
            doc,
            str(project_info.get("abstract_zh", "") or "").strip() or "待补充中文摘要正文。",
            keywords=_join_keywords(project_info.get("keywords_zh"), "；") or "待补充关键词1；待补充关键词2",
            render_context=render_context,
        )
    elif kind == "en_abstract":
        add_english_abstract_section(
            doc,
            str(project_info.get("abstract_en", "") or "").strip() or "Pending English abstract body.",
            keywords=_join_keywords(project_info.get("keywords_en"), "; ") or "keyword1; keyword2",
            render_context=render_context,
        )
    elif kind == "toc":
        add_toc_section(doc, toc_entries=None, render_context=render_context)
    elif kind == "abbreviation_table":
        inserted = create_abbreviation_table_page(doc, project_root, render_context=render_context)
        if not inserted:
            fm_spec = render_context.get("style_profile", {}).get("front_matter", {}).get("abbreviation_table", {})
            heading = doc.add_heading(fm_spec.get("title_text", "主要缩略语对照表"), level=1)
            heading.text = ""
            heading.add_run(fm_spec.get("title_text", "主要缩略语对照表"))
            body = doc.add_paragraph(fm_spec.get("empty_text", "暂无已注册缩略语"))
            from markdown_to_docx import apply_csu_heading1_style, apply_csu_normal_style

            apply_csu_heading1_style(heading, render_context=render_context)
            apply_csu_normal_style(body, render_context=render_context)
    else:
        raise ValueError(f"unsupported docx render kind: {kind}")
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    return "updated"


def _is_managed_file(path):
    if not path.exists():
        return False
    try:
        content = path.read_text(encoding="utf-8")
    except Exception:
        return False
    return content.startswith(MANAGED_MARKER)


def _write_managed_markdown(path, content, overwrite=False):
    if path.exists() and not overwrite and not _is_managed_file(path):
        return "skipped_unmanaged"
    path.parent.mkdir(parents=True, exist_ok=True)
    existed = path.exists()
    path.write_text(content, encoding="utf-8")
    return "updated" if existed else "created"


def render_front_matter(project_root, overwrite=False, to_docx=True):
    project_root = os.path.abspath(project_root)
    profile, _ = load_profile(project_root)
    format_profile = profile.get("format_profile", {}) if isinstance(profile, dict) else {}
    allow_docx_generation = bool(format_profile.get("allow_docx_generation", True))
    placeholders = _build_placeholder_map(project_root)
    results = []

    for spec in FRONT_MATTER_SPECS:
        md_path = Path(project_root) / spec["md_output"]
        docx_path = Path(project_root) / spec["docx_output"]
        rendered = _render_managed_markdown(spec, project_root, placeholders)
        md_status = _write_managed_markdown(md_path, rendered, overwrite=overwrite)

        docx_status = "skipped"
        if to_docx:
            should_render_docx = (
                md_status in {"created", "updated"}
                or not docx_path.exists()
                or overwrite
                or _is_managed_file(md_path)
            )
            if not allow_docx_generation:
                docx_status = "blocked_pending_template"
            elif should_render_docx and md_status != "skipped_unmanaged":
                if spec.get("kind") == "template":
                    docx_path.parent.mkdir(parents=True, exist_ok=True)
                    ok = convert_markdown_file(str(md_path), str(docx_path), project_root=project_root)
                    docx_status = "updated" if ok else "failed"
                else:
                    docx_status = _render_special_docx(spec, project_root, docx_path, placeholders)

        results.append(
            {
                "template": spec.get("template"),
                "kind": spec.get("kind"),
                "md_output": str(md_path),
                "docx_output": str(docx_path),
                "md_status": md_status,
                "docx_status": docx_status,
            }
        )

    return {
        "ok": True,
        "project_root": project_root,
        "managed_marker": MANAGED_MARKER,
        "results": results,
    }


def main():
    parser = argparse.ArgumentParser(description="Render managed front matter files for Sci2Doc")
    parser.add_argument("--project-root", default=".", help="Sci2Doc project root")
    parser.add_argument("--overwrite", action="store_true", help="Overwrite even unmanaged markdown files")
    parser.add_argument("--no-docx", action="store_true", help="Render markdown only")
    args = parser.parse_args()

    payload = render_front_matter(
        project_root=args.project_root,
        overwrite=args.overwrite,
        to_docx=not args.no_docx,
    )
    print(json.dumps(payload, ensure_ascii=False))


if __name__ == "__main__":
    main()
