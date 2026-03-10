#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from common import read_json


PROFILE_PRESETS: dict[str, dict[str, Any]] = {
    "journal-manuscript": {
        "body_font": "Times New Roman",
        "body_size": 11,
        "title_font": "Arial",
        "title_size": 16,
        "heading_font": "Arial",
        "heading1_size": 13,
        "heading2_size": 12,
        "heading3_size": 11.5,
        "heading4_size": 10.5,
        "label_font": "Arial",
        "body_line_spacing": 1.15,
        "body_space_after": 6,
        "header_text_align": WD_ALIGN_PARAGRAPH.CENTER,
        "ref_size": 9.5,
        "abstract_size": 10.5,
        "table_text_size": 9.5,
        "table_header_fill": "D9E2F3",
        "table_header_font": "Arial",
    },
    "nature-review": {
        "body_font": "Garamond",
        "body_size": 11,
        "title_font": "Arial",
        "title_size": 17,
        "heading_font": "Arial",
        "heading1_size": 13.5,
        "heading2_size": 12.5,
        "heading3_size": 11.5,
        "heading4_size": 10.5,
        "label_font": "Arial",
        "body_line_spacing": 1.12,
        "body_space_after": 5,
        "header_text_align": WD_ALIGN_PARAGRAPH.CENTER,
        "ref_size": 9.5,
        "abstract_size": 10.5,
        "table_text_size": 9.5,
        "table_header_fill": "EAF2F8",
        "table_header_font": "Arial",
    },
    "cell-press": {
        "body_font": "Times New Roman",
        "body_size": 11,
        "title_font": "Georgia",
        "title_size": 16.5,
        "heading_font": "Georgia",
        "heading1_size": 13,
        "heading2_size": 12,
        "heading3_size": 11,
        "heading4_size": 10.5,
        "label_font": "Georgia",
        "body_line_spacing": 1.15,
        "body_space_after": 6,
        "header_text_align": WD_ALIGN_PARAGRAPH.CENTER,
        "ref_size": 9.5,
        "abstract_size": 10.5,
        "table_text_size": 9.5,
        "table_header_fill": "F2E5D7",
        "table_header_font": "Georgia",
    },
    "lancet-review": {
        "body_font": "Times New Roman",
        "body_size": 10.5,
        "title_font": "Arial",
        "title_size": 16,
        "heading_font": "Arial",
        "heading1_size": 12.5,
        "heading2_size": 11.5,
        "heading3_size": 11,
        "heading4_size": 10.5,
        "label_font": "Arial",
        "body_line_spacing": 1.15,
        "body_space_after": 5,
        "header_text_align": WD_ALIGN_PARAGRAPH.CENTER,
        "ref_size": 9,
        "abstract_size": 10,
        "table_text_size": 9,
        "table_header_fill": "F3F6F9",
        "table_header_font": "Arial",
    },
    "review-response": {
        "body_font": "Times New Roman",
        "body_size": 11,
        "title_font": "Arial",
        "title_size": 15.5,
        "heading_font": "Arial",
        "heading1_size": 13.5,
        "heading2_size": 12,
        "heading3_size": 11.5,
        "heading4_size": 10.5,
        "label_font": "Arial",
        "body_line_spacing": 1.15,
        "body_space_after": 5,
        "header_text_align": WD_ALIGN_PARAGRAPH.CENTER,
        "ref_size": 9.5,
        "abstract_size": 10.5,
        "table_text_size": 9.5,
        "table_header_fill": "DDEBF7",
        "table_header_font": "Arial",
    },
}


def doc_for_template(template_path: str) -> Document:
    if template_path:
        return Document(template_path)
    return Document()


def add_runs_with_bold(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def profile_settings(profile_name: str) -> dict[str, Any]:
    return PROFILE_PRESETS.get(profile_name, PROFILE_PRESETS["journal-manuscript"])


def set_style_font(style, name: str, size: float, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold


def ensure_custom_styles(doc: Document, profile_name: str, doc_kind: str) -> None:
    profile = profile_settings(profile_name)
    style_names = {style.name for style in doc.styles}
    if "ReviseSciTitle" not in style_names:
        style = doc.styles.add_style("ReviseSciTitle", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Title"]
    if "ReviseSciAbstractLabel" not in style_names:
        style = doc.styles.add_style("ReviseSciAbstractLabel", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
    if "ReviseSciAbstractBody" not in style_names:
        style = doc.styles.add_style("ReviseSciAbstractBody", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
    if "ReviseSciBody" not in style_names:
        style = doc.styles.add_style("ReviseSciBody", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
    if "ReviseSciRefEntry" not in style_names:
        style = doc.styles.add_style("ReviseSciRefEntry", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
    if "ReviseSciLabel" not in style_names:
        style = doc.styles.add_style("ReviseSciLabel", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
    if "ReviseSciReviewerHeading" not in style_names:
        style = doc.styles.add_style("ReviseSciReviewerHeading", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Heading 1"]
    if "ReviseSciCommentHeading" not in style_names:
        style = doc.styles.add_style("ReviseSciCommentHeading", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Heading 3"]
    normal_style = doc.styles["Normal"]
    set_style_font(normal_style, profile["body_font"], profile["body_size"])

    set_style_font(doc.styles["ReviseSciTitle"], profile["title_font"], profile["title_size"], bold=True)
    set_style_font(doc.styles["ReviseSciAbstractLabel"], profile["label_font"], profile["body_size"], bold=True)
    set_style_font(doc.styles["ReviseSciAbstractBody"], profile["body_font"], profile["abstract_size"])
    set_style_font(doc.styles["ReviseSciBody"], profile["body_font"], profile["body_size"])
    set_style_font(doc.styles["ReviseSciRefEntry"], profile["body_font"], profile["ref_size"])
    set_style_font(doc.styles["ReviseSciLabel"], profile["label_font"], profile["body_size"], bold=True)
    set_style_font(doc.styles["ReviseSciReviewerHeading"], profile["heading_font"], profile["heading1_size"], bold=True)
    set_style_font(doc.styles["ReviseSciCommentHeading"], profile["heading_font"], profile["heading3_size"], bold=True)

    title_style = doc.styles["ReviseSciTitle"]
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_before = Pt(6)
    title_style.paragraph_format.space_after = Pt(12)

    abstract_label_style = doc.styles["ReviseSciAbstractLabel"]
    abstract_label_style.paragraph_format.space_before = Pt(8)
    abstract_label_style.paragraph_format.space_after = Pt(0)

    abstract_body_style = doc.styles["ReviseSciAbstractBody"]
    abstract_body_style.paragraph_format.line_spacing = profile["body_line_spacing"]
    abstract_body_style.paragraph_format.space_after = Pt(6)
    abstract_body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    normal_style = doc.styles["Normal"]
    body_style = doc.styles["ReviseSciBody"]
    body_style.paragraph_format.line_spacing = profile["body_line_spacing"]
    body_style.paragraph_format.space_after = Pt(profile["body_space_after"])
    body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ref_style = doc.styles["ReviseSciRefEntry"]
    ref_style.paragraph_format.line_spacing = 1.0
    ref_style.paragraph_format.space_after = Pt(2)
    ref_style.paragraph_format.left_indent = Inches(0.25)
    ref_style.paragraph_format.first_line_indent = Inches(-0.25)
    ref_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    label_style = doc.styles["ReviseSciLabel"]
    label_style.paragraph_format.space_before = Pt(6)
    label_style.paragraph_format.space_after = Pt(0)
    reviewer_style = doc.styles["ReviseSciReviewerHeading"]
    reviewer_style.paragraph_format.space_before = Pt(12)
    reviewer_style.paragraph_format.space_after = Pt(6)
    reviewer_style.paragraph_format.keep_with_next = True
    comment_style = doc.styles["ReviseSciCommentHeading"]
    comment_style.paragraph_format.space_before = Pt(10)
    comment_style.paragraph_format.space_after = Pt(4)
    comment_style.paragraph_format.keep_with_next = True
    for heading_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = doc.styles[heading_name]
        set_style_font(style, profile["heading_font"], profile[f"{heading_name.lower().replace(' ', '')}_size"], bold=True)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(6)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(10)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(4)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(3)
    doc.styles["Heading 4"].paragraph_format.space_before = Pt(6)
    doc.styles["Heading 4"].paragraph_format.space_after = Pt(2)
    if doc_kind == "response":
        reviewer_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        comment_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_markdown_paragraph(doc: Document, line: str, in_references: bool = False, in_abstract: bool = False) -> None:
    label_match = re.fullmatch(r"\*\*(.+?)\*\*", line.strip())
    if label_match:
        paragraph = doc.add_paragraph(style="ReviseSciLabel")
        paragraph.add_run(label_match.group(1)).bold = True
        return
    bullet_match = re.match(r"^- (.+)$", line)
    if bullet_match:
        paragraph = doc.add_paragraph(style="List Bullet")
        add_runs_with_bold(paragraph, bullet_match.group(1))
        return
    number_match = re.match(r"^\d+\. (.+)$", line)
    if number_match:
        paragraph = doc.add_paragraph(style="ReviseSciRefEntry" if in_references else "List Number")
        add_runs_with_bold(paragraph, line if in_references else number_match.group(1))
        return
    paragraph = doc.add_paragraph(style="ReviseSciRefEntry" if in_references else ("ReviseSciAbstractBody" if in_abstract else "ReviseSciBody"))
    add_runs_with_bold(paragraph, line)


def parse_markdown_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def is_markdown_table_separator(line: str) -> bool:
    cells = parse_markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_markdown_table(doc: Document, rows: list[list[str]], profile_name: str) -> None:
    if not rows:
        return
    profile = profile_settings(profile_name)
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(column_count):
            text = row[col_idx] if col_idx < len(row) else ""
            cell_paragraph = table.cell(row_idx, col_idx).paragraphs[0]
            add_runs_with_bold(cell_paragraph, text)
            for run in cell_paragraph.runs:
                run.font.name = profile["body_font"]
                run.font.size = Pt(profile["table_text_size"])
            if row_idx == 0:
                for run in cell_paragraph.runs:
                    run.bold = True
                    run.font.name = profile["table_header_font"]
                cell = table.cell(row_idx, col_idx)
                tc_pr = cell._tc.get_or_add_tcPr()
                shade = OxmlElement("w:shd")
                shade.set(qn("w:fill"), profile["table_header_fill"])
                tc_pr.append(shade)
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def append_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def configure_document(doc: Document, header_text: str = "", profile_name: str = "journal-manuscript", doc_kind: str = "manuscript") -> None:
    profile = profile_settings(profile_name)
    ensure_custom_styles(doc, profile_name, doc_kind)
    normal_style = doc.styles["Normal"]
    normal_style.font.name = profile["body_font"]
    normal_style.font.size = Pt(profile["body_size"])
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)
        if header_text:
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = header_text
            header_para.alignment = profile["header_text_align"]
            if header_para.runs:
                header_para.runs[0].bold = True
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        append_field(footer_para, "PAGE", "1")


def add_toc_block(doc: Document) -> None:
    toc_title = doc.add_paragraph()
    toc_title.add_run("Table of Contents").bold = True
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_para = doc.add_paragraph()
    append_field(toc_para, 'TOC \\o "1-4" \\h \\z \\u', "Right-click to update field.")
    doc.add_paragraph("")


def markdown_to_docx(
    md_path: Path,
    docx_path: Path,
    template_path: str = "",
    page_break_before_comment: bool = False,
    header_text: str = "",
    include_toc: bool = False,
    profile_name: str = "journal-manuscript",
    doc_kind: str = "manuscript",
) -> None:
    doc = doc_for_template(template_path)
    configure_document(doc, header_text, profile_name=profile_name, doc_kind=doc_kind)
    lines = md_path.read_text(encoding="utf-8").splitlines()
    seen_comment = False
    toc_inserted = False
    seen_title = False
    in_references = False
    in_abstract = False
    idx = 0
    while idx < len(lines):
        raw_line = lines[idx]
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            idx += 1
            continue
        if "|" in line and idx + 1 < len(lines) and is_markdown_table_separator(lines[idx + 1]):
            table_rows = [parse_markdown_table_row(line)]
            idx += 2
            while idx < len(lines):
                next_line = lines[idx].rstrip()
                if not next_line or "|" not in next_line:
                    break
                table_rows.append(parse_markdown_table_row(next_line))
                idx += 1
            add_markdown_table(doc, table_rows, profile_name)
            continue
        if line.startswith("### Comment") and page_break_before_comment:
            if seen_comment:
                doc.add_page_break()
            seen_comment = True
        if line.startswith("#### "):
            in_abstract = False
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            in_abstract = False
            paragraph = doc.add_paragraph(style="ReviseSciCommentHeading")
            paragraph.add_run(line[4:]).bold = True
        elif line.startswith("## "):
            heading_text = line[3:]
            in_references = heading_text.lower() in {"references", "reference", "参考文献"}
            in_abstract = heading_text.lower() in {"abstract", "摘要"}
            if in_abstract:
                paragraph = doc.add_paragraph(style="ReviseSciAbstractLabel")
                paragraph.add_run(heading_text).bold = True
            else:
                doc.add_heading(heading_text, level=2)
        elif line.startswith("# "):
            heading_text = line[2:]
            in_references = False
            in_abstract = False
            if doc_kind == "manuscript" and not seen_title:
                paragraph = doc.add_paragraph(style="ReviseSciTitle")
                paragraph.add_run(heading_text).bold = True
                seen_title = True
            elif page_break_before_comment and (heading_text.startswith("Reviewer #") or heading_text.startswith("Editor")):
                paragraph = doc.add_paragraph(style="ReviseSciReviewerHeading")
                paragraph.add_run(heading_text).bold = True
            else:
                doc.add_heading(heading_text, level=1)
            if include_toc and not toc_inserted:
                add_toc_block(doc)
                toc_inserted = True
        else:
            add_markdown_paragraph(doc, line, in_references=in_references, in_abstract=in_abstract)
        idx += 1
    doc.save(str(docx_path))


def main() -> int:
    parser = argparse.ArgumentParser(description="Export markdown artifacts to docx")
    parser.add_argument("--project-root", required=True)
    parser.add_argument("--output-md", required=True)
    parser.add_argument("--output-docx", required=True)
    parser.add_argument("--reference-docx", default="")
    parser.add_argument("--journal-style", choices=tuple(PROFILE_PRESETS.keys()), default="journal-manuscript")
    args = parser.parse_args()

    project_root = Path(args.project_root)
    output_md = Path(args.output_md)
    output_docx = Path(args.output_docx)
    response_md = project_root / "response_to_reviewers.md"
    response_docx = project_root / "response_to_reviewers.docx"

    markdown_to_docx(
        output_md,
        output_docx,
        args.reference_docx,
        page_break_before_comment=False,
        header_text="Revised Manuscript",
        profile_name=args.journal_style,
        doc_kind="manuscript",
    )
    markdown_to_docx(
        response_md,
        response_docx,
        args.reference_docx,
        page_break_before_comment=True,
        header_text="Response to Reviewers",
        include_toc=True,
        profile_name="review-response",
        doc_kind="response",
    )

    state = read_json(project_root / "project_state.json", {})
    state.setdefault("outputs", {})
    state["outputs"]["response_md"] = str(response_md.resolve())
    state["outputs"]["response_docx"] = str(response_docx.resolve())
    state["outputs"]["output_md"] = str(output_md.resolve())
    state["outputs"]["output_docx"] = str(output_docx.resolve())
    state.setdefault("inputs", {})
    state["inputs"]["journal_style"] = args.journal_style
    (project_root / "project_state.json").write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")

    print(json.dumps({"ok": True, "response_docx": str(response_docx.resolve()), "output_docx": str(output_docx.resolve()), "journal_style": args.journal_style}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
